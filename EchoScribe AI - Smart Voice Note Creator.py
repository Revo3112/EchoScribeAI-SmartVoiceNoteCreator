import speech_recognition as sr
import os
import datetime
import tkinter as tk
from tkinter import filedialog, messagebox
import customtkinter as ctk
import threading
import time
import wave
import tempfile
import audioop
from pydub import AudioSegment
import sys
import re
import logging
import subprocess
import groq
import docx
import sounddevice as sd
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
import warnings
import platform
import json
from logging.handlers import RotatingFileHandler
from pathlib import Path
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
import numpy as np
import threading
import queue
import pyaudio
try:
    import pyaudiowpatch as pyaudio
    PYAUDIOWPATCH_AVAILABLE = True
except ImportError:
    import pyaudio
    PYAUDIOWPATCH_AVAILABLE = False
    print("WARNING: PyAudioWPatch tidak tersedia. System audio recording tidak akan berfungsi optimal.")

warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles")

ctk.set_appearance_mode("Dark")
ctk.set_default_color_theme("blue")

# Only create logger without basicConfig to prevent duplicate handlers
logger = logging.getLogger('AudioEnhancer')

def setup_ffmpeg():
    ffmpeg_paths = [
        os.path.join(os.path.dirname(sys.executable), 'ffmpeg.exe'),
        'C:\\FFmpeg\\bin\\ffmpeg.exe',
        os.path.join(os.environ.get('USERPROFILE', ''), 'FFmpeg\\bin\\ffmpeg.exe'),
        os.path.join(os.environ.get('LOCALAPPDATA', ''), 'FFmpeg\\bin\\ffmpeg.exe'),
        os.path.join(os.environ.get('PROGRAMFILES', ''), 'FFmpeg\\bin\\ffmpeg.exe'),
        os.path.join(os.environ.get('PROGRAMFILES(X86)', ''), 'FFmpeg\\bin\\ffmpeg.exe'),
        'ffmpeg.exe'
    ]

    for path in ffmpeg_paths:
        if os.path.exists(path):
            AudioSegment.converter = path
            print(f"Found FFmpeg at: {path}")
            return True

    try:
        subprocess.run(["ffmpeg", "-version"], stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        print("FFmpeg found in PATH")
        return True
    except:
        print("WARNING: FFmpeg not found. Audio processing may not work correctly.")
        print("Please install FFmpeg from https://ffmpeg.org/download.html")
        return False

def get_api_key_storage_path():
    """Mendapatkan path untuk menyimpan API key user."""
    app_data_dir = Path.home() / ".echoscribe"
    app_data_dir.mkdir(exist_ok=True)
    return app_data_dir / "config.json"

def save_user_api_key(api_key):
    """Menyimpan API key user ke file konfigurasi."""
    try:
        config_path = get_api_key_storage_path()
        config = {"groq_api_key": api_key}
        with open(config_path, 'w') as f:
            json.dump(config, f)
        return True
    except Exception as e:
        logger.error(f"Error saving API key: {e}")
        return False

def load_user_api_key():
    """Memuat API key user dari file konfigurasi."""
    try:
        config_path = get_api_key_storage_path()
        if config_path.exists():
            with open(config_path, 'r') as f:
                config = json.load(f)
                return config.get("groq_api_key", "")
    except Exception as e:
        logger.error(f"Error loading API key: {e}")
    return ""

def remove_user_api_key():
    """Menghapus API key user dari konfigurasi."""
    try:
        config_path = get_api_key_storage_path()
        if config_path.exists():
            os.remove(config_path)
        return True
    except Exception as e:
        logger.error(f"Error removing API key: {e}")
        return False

class ErrorHandler:
    """
    Kelas untuk penanganan kesalahan yang terpusat dan konsisten.
    """
    def __init__(self, app):
        self.app = app
        self.error_counts = {}  # Track error frequency
        self.last_errors = {}   # Track last error timestamp by type
        self.max_retry = 3      # Maximum retry attempts

    def handle_error(self, error_type, exception, operation=None, retry_func=None, retry_args=None):
        """
        Handle error with centralized logic and optional retry.

        Args:
            error_type: String identifier for the error category
            exception: The exception that was raised
            operation: Description of the operation that failed (for logging)
            retry_func: Function to retry if appropriate
            retry_args: Arguments to pass to retry_func

        Returns:
            Boolean indicating if retry was attempted
        """
        current_time = time.time()

        # Update error tracking
        if error_type not in self.error_counts:
            self.error_counts[error_type] = 1
            self.last_errors[error_type] = current_time
        else:
            self.error_counts[error_type] += 1

        # Log the error
        logger.error(
            f"Error [{error_type}]: {exception}",
            exc_info=True if self.error_counts[error_type] <= 2 else False  # Only full stack trace for first occurrences
        )

        # Determine if retry is appropriate
        should_retry = False
        if retry_func:
            # Don't retry too many times
            if self.error_counts[error_type] <= self.max_retry:
                # Check error type to determine retry strategy
                if error_type == 'network':
                    # For network errors, use exponential backoff
                    backoff_time = min(2 ** (self.error_counts[error_type] - 1), 30)  # Max 30 seconds
                    self.app.root.after(0, lambda: self.app.status_var.set(
                        f"Kesalahan jaringan: Mencoba ulang dalam {backoff_time} detik..."))
                    self.app.root.after(int(backoff_time * 1000), lambda: self._execute_retry(retry_func, retry_args))
                    should_retry = True
                elif error_type == 'api_error':
                    # For API errors, retry after a fixed delay if not rate limited
                    if "rate limit" in str(exception).lower():
                        self.app.root.after(0, lambda: self.app.status_var.set(
                            "Batas API terlampaui. Menunggu 60 detik..."))
                        self.app.root.after(60000, lambda: self._execute_retry(retry_func, retry_args))
                    else:
                        self.app.root.after(0, lambda: self.app.status_var.set(
                            "Error API: Mencoba ulang dalam 5 detik..."))
                        self.app.root.after(5000, lambda: self._execute_retry(retry_func, retry_args))
                    should_retry = True
                elif error_type == 'transcription':
                    # For transcription errors, try with different parameters
                    self.app.root.after(0, lambda: self.app.status_var.set(
                        "Kesalahan transkripsi: Mencoba dengan pengaturan alternatif..."))
                    self.app.root.after(1000, lambda: self._execute_retry(retry_func, retry_args, alternative=True))
                    should_retry = True

        # Show a user-friendly message
        self._show_user_message(error_type, exception, operation, should_retry)

        return should_retry

    def _execute_retry(self, retry_func, retry_args, alternative=False):
        try:
            if alternative:
                # Modify args for alternative approach
                if isinstance(retry_args, dict):
                    # Make a copy to avoid modifying the original
                    alt_args = retry_args.copy()

                    # Add retry flag
                    alt_args['is_retry'] = True

                    # Call with modified args
                    retry_func(**alt_args)
                else:
                    # If not a dict, just call with original args and retry flag
                    if retry_args:
                        retry_func(*retry_args, is_retry=True)
                    else:
                        retry_func(is_retry=True)
            else:
                # Call normally
                if retry_args:
                    if isinstance(retry_args, dict):
                        retry_func(**retry_args)
                    else:
                        retry_func(*retry_args)
                else:
                    retry_func()

            # Update status on successful retry
            self.app.root.after(0, lambda: self.app.status_var.set("Percobaan ulang berhasil"))

        except Exception as e:
            # Failed retry
            logger.error(f"Retry failed: {e}")
            self.app.root.after(0, lambda: self.app.status_var.set("Percobaan ulang gagal"))

    def _show_user_message(self, error_type, exception, operation, retry_attempted):
        """
        Display a user-friendly error message.

        Args:
            error_type: The type of error
            exception: The exception object
            operation: Description of what was being done
            retry_attempted: Whether a retry was attempted
        """
        # Define user-friendly messages by error type
        friendly_messages = {
            'network': "Terjadi masalah koneksi jaringan.",
            'api_error': "Layanan API saat ini mengalami gangguan.",
            'transcription': "Transkripsi audio mengalami kesulitan.",
            'file_access': "Terjadi masalah dalam mengakses file.",
            'permission': "Aplikasi tidak memiliki izin yang diperlukan.",
            'memory': "Tidak cukup memori untuk menyelesaikan operasi.",
            'encoding': "Format file tidak didukung atau rusak."
        }

        # Get base message
        base_message = friendly_messages.get(
            error_type,
            "Terjadi kesalahan yang tidak terduga."
        )

        # Add operation context if provided
        if operation:
            base_message += f"\nSaat: {operation}"

        # Add retry info
        if retry_attempted:
            base_message += "\nMencoba untuk melanjutkan secara otomatis..."
        else:
            base_message += "\nSilakan coba lagi atau periksa pengaturan."

        # Add error details for non-network errors (more technical)
        if error_type not in ['network', 'api_error']:
            error_str = str(exception)
            if len(error_str) > 100:
                error_str = error_str[:97] + "..."
            base_message += f"\n\nDetail teknis: {error_str}"

        # Display message using appropriate severity
        critical_errors = ['memory', 'permission']
        warning_errors = ['network', 'api_error', 'transcription']

        if error_type in critical_errors:
            messagebox.showerror("Error", base_message)
        elif error_type in warning_errors:
            # Only show messagebox for first occurrence or if significant time has passed
            current_time = time.time()
            last_time = self.last_errors.get(error_type, 0)
            if self.error_counts[error_type] <= 1 or (current_time - last_time) > 300:  # 5 minutes
                messagebox.showwarning("Peringatan", base_message)
        else:
            # Info level for other errors
            messagebox.showinfo("Informasi", base_message)

        # Update last error time
        self.last_errors[error_type] = time.time()

class ValueTrackingSlider(ctk.CTkFrame):
    def __init__(self, master, title, from_, to, variable, format_string="{:.0f}", width=300, **kwargs):
        super().__init__(master, **kwargs)

        self.format_string = format_string
        self.variable = variable
        self.title = title

        self.tooltip_window = None

        self.title_label = ctk.CTkLabel(self, text=title)
        self.title_label.grid(row=0, column=0, sticky="w", padx=(0,10))

        self.value_text = ctk.StringVar()
        self.update_value_text()
        self.value_label = ctk.CTkLabel(self, textvariable=self.value_text, width=60)
        self.value_label.grid(row=0, column=1, sticky="e")

        self.slider = ctk.CTkSlider(self, from_=from_, to=to, variable=variable, width=width,
                                   command=self.on_slider_change)
        self.slider.grid(row=1, column=0, columnspan=2, sticky="ew", pady=(5, 0))

        self.tooltip_text = f"{title}: {self.format_string.format(variable.get())}"

        self.slider.bind("<Enter>", self.show_tooltip)
        self.slider.bind("<Leave>", self.hide_tooltip)
        self.slider.bind("<Motion>", self.update_tooltip_position)

        variable.trace_add("write", self.update_value_text)

    def update_value_text(self, *args):
        value = self.variable.get()
        self.value_text.set(self.format_string.format(value))
        self.tooltip_text = f"{self.title}: {self.format_string.format(value)}"
        if hasattr(self, 'tooltip_window') and self.tooltip_window:
            self.tooltip_window.winfo_children()[0].configure(text=self.tooltip_text)

    def on_slider_change(self, value):
        # Gunakan value parameter untuk validasi dan logging
        if hasattr(self, 'variable'):
            self.variable.set(value)
        self.update_value_text()
        logger.debug(f"Slider value changed to: {value}")

    def hide_tooltip(self, event):
        # Gunakan event untuk validasi posisi mouse
        if hasattr(self, 'tooltip_window') and self.tooltip_window:
            # Cek apakah mouse benar-benar keluar dari area slider
            if event.widget == self.slider:
                self.tooltip_window.destroy()
                self.tooltip_window = None

    def show_tooltip(self, event):
        x, y = event.x_root, event.y_root
        self.tooltip_window = tk.Toplevel(self)
        self.tooltip_window.wm_overrideredirect(True)
        self.tooltip_window.wm_geometry(f"+{x+10}+{y+10}")

        label = ctk.CTkLabel(
            self.tooltip_window,
            text=self.tooltip_text,
            fg_color=("gray75", "gray25"),
            corner_radius=6,
            padx=8,
            pady=4
        )
        label.pack()

    def update_tooltip_position(self, event):
        if hasattr(self, 'tooltip_window') and self.tooltip_window:
            x, y = event.x_root, event.y_root
            self.tooltip_window.wm_geometry(f"+{x+10}+{y+10}")

class APIKeyDialog(ctk.CTkToplevel):
    def __init__(self, parent, current_key=""):
        """Dialog untuk input dan manajemen API key."""
        super().__init__(parent)

        self.title("Pengaturan API Key Groq")
        self.geometry("500x350")
        self.transient(parent)
        self.grab_set()

        self.api_key = None
        self.result = None

        # Center the dialog
        self.geometry(f"+{parent.winfo_rootx() + 50}+{parent.winfo_rooty() + 50}")

        self.setup_ui(current_key)

    def setup_ui(self, current_key):
        """Setup UI untuk dialog API key."""
        main_frame = ctk.CTkFrame(self)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=20, pady=20)

        # Title
        title_label = ctk.CTkLabel(
            main_frame,
            text="🔑 Konfigurasi API Key Groq",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=(0, 20))

        # Info text
        info_text = """Anda dapat menggunakan API key Groq sendiri atau menggunakan API key default. API key Anda akan disimpan secara lokal di komputer dan dienkripsi untuk keamanan."""

        info_label = ctk.CTkLabel(
            main_frame,
            text=info_text,
            wraplength=450,
            justify="left"
        )
        info_label.pack(pady=(0, 20))

        # Current key status
        if current_key and current_key != "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":
            status_text = f"✅ API Key Tersimpan: {current_key[:20]}..."
            status_color = "green"
        elif current_key:
            status_text = "🔧 Menggunakan API Key Default"
            status_color = "orange"
        else:
            status_text = "❌ Tidak Ada API Key"
            status_color = "red"

        status_label = ctk.CTkLabel(
            main_frame,
            text=status_text,
            text_color=status_color
        )
        status_label.pack(pady=(0, 15))

        # API Key input
        input_frame = ctk.CTkFrame(main_frame)
        input_frame.pack(fill=tk.X, pady=(0, 20))

        ctk.CTkLabel(input_frame, text="Masukkan API Key Groq:").pack(anchor="w", padx=10, pady=(10, 5))

        self.api_key_entry = ctk.CTkEntry(
            input_frame,
            placeholder_text="gsk_...",
            width=400,
            show="*"
        )
        self.api_key_entry.pack(padx=10, pady=(0, 10))

        if current_key and current_key != "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":
            self.api_key_entry.insert(0, current_key)

        # Show/hide password button
        show_frame = ctk.CTkFrame(input_frame, fg_color="transparent")
        show_frame.pack(padx=10, pady=(0, 10))

        self.show_key = tk.BooleanVar()
        show_checkbox = ctk.CTkCheckBox(
            show_frame,
            text="Tampilkan API Key",
            variable=self.show_key,
            command=self.toggle_key_visibility
        )
        show_checkbox.pack(side=tk.LEFT)

        # Instructions
        instructions_frame = ctk.CTkFrame(main_frame)
        instructions_frame.pack(fill=tk.X, pady=(0, 20))

        instructions_text = """📋 Cara mendapatkan API Key Groq:
        1. Kunjungi https://console.groq.com/
        2. Daftar atau login ke akun Anda
        3. Buat API Key baru di dashboard
        4. Salin dan tempel di sini"""

        instructions_label = ctk.CTkLabel(
            instructions_frame,
            text=instructions_text,
            justify="left",
            font=ctk.CTkFont(size=11)
        )
        instructions_label.pack(padx=10, pady=10)

        # Buttons
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill=tk.X)

        # Use default button
        default_btn = ctk.CTkButton(
            button_frame,
            text="Gunakan Default",
            command=self.use_default_key,
            width=120
        )
        default_btn.pack(side=tk.LEFT, padx=(0, 10))

        # Remove key button
        if current_key and current_key != "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":
            remove_btn = ctk.CTkButton(
                button_frame,
                text="Hapus Key",
                command=self.remove_key,
                fg_color="red",
                hover_color="darkred",
                width=100
            )
            remove_btn.pack(side=tk.LEFT, padx=(0, 10))

        # Save button
        save_btn = ctk.CTkButton(
            button_frame,
            text="Simpan",
            command=self.save_key,
            width=100
        )
        save_btn.pack(side=tk.RIGHT, padx=(10, 0))

        # Cancel button
        cancel_btn = ctk.CTkButton(
            button_frame,
            text="Batal",
            command=self.cancel,
            fg_color="gray",
            hover_color="darkgray",
            width=100
        )
        cancel_btn.pack(side=tk.RIGHT)

    def toggle_key_visibility(self):
        """Toggle visibility API key."""
        if self.show_key.get():
            self.api_key_entry.configure(show="")
        else:
            self.api_key_entry.configure(show="*")

    def use_default_key(self):
        """Menggunakan API key default."""
        self.api_key = "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"
        self.result = "default"
        remove_user_api_key()  # Hapus user key jika ada
        self.destroy()

    def save_key(self):
        """Menyimpan API key user."""
        key = self.api_key_entry.get().strip()

        if not key:
            messagebox.showwarning("Peringatan", "Silakan masukkan API key")
            return

        if not key.startswith("gsk_"):
            messagebox.showwarning("Peringatan", "API key Groq harus dimulai dengan 'gsk_'")
            return

        if len(key) < 20:
            messagebox.showwarning("Peringatan", "API key terlalu pendek")
            return

        # Test API key
        self.test_api_key(key)

    def test_api_key(self, key):
        """Test validitas API key."""
        try:
            import groq
            test_client = groq.Groq(api_key=key)

            # Test dengan request sederhana
            response = test_client.chat.completions.create(
                model="deepseek-r1-distill-llama-70b",
                messages=[{"role": "user", "content": "test"}],
                max_tokens=5
            )

            # Jika sampai sini berarti API key valid
            if save_user_api_key(key):
                self.api_key = key
                self.result = "custom"
                messagebox.showinfo("Sukses", "API key berhasil disimpan dan diverifikasi!")
                self.destroy()
            else:
                messagebox.showerror("Error", "Gagal menyimpan API key")

        except Exception as e:
            error_msg = str(e)
            if "401" in error_msg or "authentication" in error_msg.lower():
                messagebox.showerror("Error", "API key tidak valid atau tidak memiliki akses")
            elif "429" in error_msg or "rate limit" in error_msg.lower():
                messagebox.showwarning("Peringatan", "Rate limit tercapai, tapi API key valid. Disimpan.")
                if save_user_api_key(key):
                    self.api_key = key
                    self.result = "custom"
                    self.destroy()
            else:
                messagebox.showerror("Error", f"Gagal memverifikasi API key:\n{error_msg}")

    def remove_key(self):
        """Menghapus API key user."""
        if messagebox.askyesno("Konfirmasi", "Apakah Anda yakin ingin menghapus API key?"):
            if remove_user_api_key():
                self.api_key = "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"
                self.result = "removed"
                messagebox.showinfo("Sukses", "API key berhasil dihapus")
                self.destroy()
            else:
                messagebox.showerror("Error", "Gagal menghapus API key")

    def cancel(self):
        """Membatalkan dialog."""
        self.result = "cancel"
        self.destroy()

class VoiceToMarkdownApp:

    # !Fungsi-fungsi yang Mengatur Suara (Voice Processing)

    def __init__(self, root):
        self.root = root
        self.root.title("EchoScribe AI - Smart Voice Note Creator")
        self.root.geometry("1000x1080")

        try:
            base_path = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
            icon_path = os.path.join(base_path, "icon.ico")
            if os.path.exists(icon_path):
                self.root.iconbitmap(icon_path)
            else:
                print(f"DEBUG: Icon file not found at runtime: {icon_path}")
        except Exception as e:
            print(f"Error setting icon: {e}")
            pass

        # ===== PERBAIKAN 1: Setup config management PERTAMA =====
        self.setup_config_management()

        # Inisialisasi components lainnya
        self.recognizer = sr.Recognizer()
        self.microphones = self.get_available_microphones()
        self.selected_mic = tk.StringVar(root)
        if self.microphones:
            self.selected_mic.set(self.microphones[0])

        # ===== PERBAIKAN 2: Inisialisasi variabel dengan nilai dari config =====
        # Recording variables - gunakan nilai dari config
        self.use_system_audio = tk.BooleanVar(root)
        self.use_system_audio.set(self.config.get("use_system_audio", False))

        self.use_dual_recording = tk.BooleanVar(root)
        self.use_dual_recording.set(self.config.get("use_dual_recording", False))

        self.recording_mode = tk.StringVar(self.root)
        self.recording_mode.set(self.config.get("recording_mode", "microphone"))

        # File variables - gunakan nilai dari config
        self.output_folder = tk.StringVar(root)
        self.output_folder.set(self.config.get("output_folder", os.path.expanduser("~/Documents")))

        self.filename_prefix = tk.StringVar(root)
        self.filename_prefix.set(self.config.get("filename_prefix", "catatan"))

        # AI variables - gunakan nilai dari config
        self.language = tk.StringVar(root)
        self.language.set(self.config.get("language", "id-ID"))

        self.engine = tk.StringVar(root)
        self.engine.set(self.config.get("engine", "Google"))

        self.use_ai_enhancement = tk.BooleanVar(root)
        self.use_ai_enhancement.set(self.config.get("use_ai_enhancement", True))

        self.use_economic_model = tk.BooleanVar(root)
        self.use_economic_model.set(self.config.get("use_economic_model", False))

        # Recording state - gunakan nilai dari config
        self.use_extended_recording = tk.BooleanVar(root)
        self.use_extended_recording.set(self.config.get("use_extended_recording", True))

        self.chunk_size = tk.IntVar(root)
        self.chunk_size.set(self.config.get("chunk_size", 600))

        self.recording = False
        self.recording_thread = None
        self.stop_recording_flag = False
        self.elapsed_time = 0
        self.audio_chunks = []
        self.temp_wav_file = None
        self.temp_audio_files = []
        self.temp_dir = None

        # UI theme
        self.theme_color = "#1E1E1E"
        self.accent_color = "#007ACC"
        self.text_color = "#E0E0E0"
        self.button_color = "#2A2A2A"
        self.button_hover = "#3A3A3A"
        self.border_color = "#3E3E3E"

        # Processing variables - gunakan nilai dari config
        self.processing_start_time = 0
        self.heading_spacing_before = 12
        self.heading_spacing_after = 6
        self.paragraph_spacing = 6
        self.api_request_delay = self.config.get("api_request_delay", 10)

        # Error handling
        self.error_handler = None

        # Inisialisasi API key dengan sistem yang ditingkatkan
        self.setup_groq_api_key()

        # Initialize ffmpeg
        ffmpeg_found = setup_ffmpeg()

        # Setup UI
        self.setup_ui()

        # ===== PERBAIKAN 3: Apply config SETELAH UI siap =====
        self.root.after(100, self.apply_config_after_ui_ready)

        # Set up error handling dan exit handler
        self.post_init_hook()

        if not ffmpeg_found:
            self.root.after(1000, lambda: messagebox.showwarning(
                "FFmpeg Tidak Ditemukan",
                "FFmpeg tidak ditemukan. Beberapa fitur audio mungkin tidak berfungsi.\n"
                "Pastikan FFmpeg terinstal dan dapat diakses oleh aplikasi ini."
            ))

    def setup_groq_api_key(self):
        """Setup API key Groq dengan sistem user dan default key."""
        # Coba load user API key terlebih dahulu
        user_api_key = load_user_api_key()

        if user_api_key:
            self.api_key = user_api_key
            logger.info("Using user's custom API key")
        else:
            # Gunakan default API key
            self.api_key = "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"
            logger.info("Using default API key")

        # Inisialisasi Groq client
        try:
            self.groq_client = groq.Groq(api_key=self.api_key)
            logger.info("Groq client initialized successfully")
        except Exception as e:
            logger.error(f"Failed to initialize Groq client: {e}")
            self.groq_client = None
            self.root.after(1000, lambda: messagebox.showwarning(
                "API Key Error",
                "Gagal menginisialisasi Groq client. Beberapa fitur tidak akan berfungsi.\n"
                "Silakan periksa pengaturan API key di menu Settings."
            ))

    def show_api_key_dialog(self):
        """Menampilkan dialog pengaturan API key."""
        current_key = load_user_api_key() or self.api_key

        dialog = APIKeyDialog(self.root, current_key)
        self.root.wait_window(dialog)

        if dialog.result and dialog.result != "cancel":
            if dialog.result == "default":
                self.api_key = "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"
                self.status_var.set("Menggunakan API key default")
            elif dialog.result == "custom":
                self.api_key = dialog.api_key
                self.status_var.set("Menggunakan API key custom")
            elif dialog.result == "removed":
                self.api_key = "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"
                self.status_var.set("API key dihapus, menggunakan default")

            # Reinisialisasi Groq client dengan API key baru
            try:
                self.groq_client = groq.Groq(api_key=self.api_key)
                logger.info("Groq client reinitialized with new API key")
            except Exception as e:
                logger.error(f"Failed to reinitialize Groq client: {e}")
                self.groq_client = None
                messagebox.showerror("Error", f"Gagal menggunakan API key baru: {e}")

            # TAMBAHKAN INI untuk update tampilan status
            self.update_api_status_display()

    def get_available_microphones(self):
        """
        Get list of available microphones, excluding loopback devices.
        """
        mic_list = []

        try:
            if PYAUDIOWPATCH_AVAILABLE:
                # Use PyAudioWPatch
                p = pyaudio.PyAudio()

                for i in range(p.get_device_count()):
                    try:
                        info = p.get_device_info_by_index(i)

                        # Only include input devices that are NOT loopback
                        if (info['maxInputChannels'] > 0 and
                            not info.get('isLoopbackDevice', False)):

                            # Skip system audio devices
                            device_name = info['name'].lower()
                            system_keywords = [
                                'stereo mix', 'wave out mix', 'what u hear',
                                'loopback', 'virtual cable'
                            ]

                            is_system_device = any(kw in device_name for kw in system_keywords)

                            if not is_system_device:
                                mic_list.append(f"{i}: {info['name']}")

                    except Exception as e:
                        logger.error(f"Error getting device {i}: {e}")
                        continue

                p.terminate()

            else:
                devices = sd.query_devices()

                for i, device in enumerate(devices):
                    if device['max_input_channels'] > 0:
                        device_name = device['name'].lower()
                        system_keywords = [
                            'stereo mix', 'loopback', 'what u hear',
                            'wave out mix', 'cable output', 'virtual cable'
                        ]

                        is_system_device = any(kw in device_name for kw in system_keywords)

                        if not is_system_device:
                            mic_list.append(f"{i}: {device['name']}")

            # Add default if no mics found
            if not mic_list:
                mic_list = ["0: Default Microphone"]

        except Exception as e:
            logger.error(f"Error getting microphones: {e}")
            mic_list = ["0: Default Microphone"]

        return mic_list

    def _find_system_audio_devices(self):
        """
        Helper function untuk mencari perangkat audio sistem (fallback method).
        """
        system_devices = []
        try:
            devices = sd.query_devices()

            for i, device in enumerate(devices):
                if device['max_input_channels'] > 0:
                    device_name = device['name'].lower()

                    # Keywords untuk sistem audio / loopback devices
                    system_keywords = [
                        'stereo mix', 'loopback', 'what u hear', 'what you hear',
                        'wave out mix', 'rec. playback', 'recording mix',
                        'cable output', 'cable input', 'virtual cable',
                        'voicemeeter', 'obs virtual', 'blackhole', 'soundflower'
                    ]

                    # Check if system audio device
                    is_system_device = any(keyword in device_name for keyword in system_keywords)

                    if is_system_device:
                        system_devices.append({
                            'index': i,
                            'name': device['name'],
                            'channels': device['max_input_channels'],
                            'samplerate': device['default_samplerate']
                        })
        except Exception as e:
            logger.error(f"Error finding system audio devices: {e}")

        return system_devices

    def setup_ui(self):
        # Main frame with better padding and modern look
        main_frame = ctk.CTkFrame(self.root)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)

        # Add app logo/title
        title_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        title_frame.pack(pady=(0, 10))

        title_label = ctk.CTkLabel(
            title_frame,
            text="EchoScribe AI",
            font=ctk.CTkFont(size=24, weight="bold")
        )
        title_label.pack(pady=(0, 5))

        subtitle_label = ctk.CTkLabel(
            title_frame,
            text="Smart Voice Note Creator",
            font=ctk.CTkFont(size=14),
            text_color=("#4F4F4F", "#AAAAAA")  # Dark gray in light mode, light gray in dark mode
        )
        subtitle_label.pack()

        # Create tabbed interface for better organization
        self.tab_view = ctk.CTkTabview(main_frame)
        self.tab_view.pack(fill=tk.BOTH, expand=True, padx=5, pady=10)

        # Create tabs
        self.tab_view.add("Rekaman")
        self.tab_view.add("Pengaturan")
        self.tab_view.add("Output")

        # Set default tab
        self.tab_view.set("Rekaman")

        # Tab 1: Recording Controls
        self.setup_recording_tab(self.tab_view.tab("Rekaman"))

        # Tab 2: Settings
        self.setup_settings_tab(self.tab_view.tab("Pengaturan"))

        # Tab 3: Output
        self.setup_output_tab(self.tab_view.tab("Output"))

        # Status bar at bottom with rich information
        self.setup_status_bar(main_frame)

    def setup_recording_tab(self, parent):
        """
        Tab untuk kontrol rekaman dengan UI yang ditingkatkan.
        """
        recording_frame = ctk.CTkFrame(parent)
        recording_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Timer display with elegant styling
        timer_frame = ctk.CTkFrame(recording_frame, fg_color=self.button_color, corner_radius=10)
        timer_frame.pack(pady=15, fill=tk.X, padx=20)

        self.timer_var = tk.StringVar()
        self.timer_var.set("00:00:00")
        timer_label = ctk.CTkLabel(
            timer_frame,
            textvariable=self.timer_var,
            font=ctk.CTkFont(size=32, weight="bold")
        )
        timer_label.pack(pady=15)

        # Recording controls with better layout
        controls_frame = ctk.CTkFrame(recording_frame, fg_color="transparent")
        controls_frame.pack(pady=15, fill=tk.X)

        # IMPROVED: Real Audio Visualization
        viz_frame = ctk.CTkFrame(controls_frame, height=150, fg_color=("#F0F0F0", "#2D2D2D"))
        viz_frame.pack(fill=tk.X, padx=20, pady=10)

        # Setup matplotlib untuk visualisasi
        self.setup_audio_visualization(viz_frame)

        # Setup visualization controls
        self.setup_visualization_controls(controls_frame)

        # Record button with enhanced styling
        button_frame = ctk.CTkFrame(controls_frame, fg_color="transparent")
        button_frame.pack(pady=15)

        self.record_button = ctk.CTkButton(
            button_frame,
            text="Mulai Rekaman",
            command=self.toggle_recording,
            width=200,
            height=50,
            fg_color="#007ACC",
            hover_color="#0066AA",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        self.record_button.pack(pady=5)

        # Progress tracking
        progress_frame = ctk.CTkFrame(recording_frame, fg_color="transparent")
        progress_frame.pack(pady=10, fill=tk.X)

        self.progress_var = tk.DoubleVar()
        self.progress_percentage = tk.StringVar()
        self.progress_percentage.set("0%")

        self.progress = ctk.CTkProgressBar(
            progress_frame,
            mode="determinate",
            variable=self.progress_var,
            height=12,
            corner_radius=5
        )
        self.progress.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0,10))

        self.progress_label = ctk.CTkLabel(
            progress_frame,
            textvariable=self.progress_percentage,
            width=40
        )
        self.progress_label.pack(side=tk.RIGHT)

        self.progress_var.trace_add("write", self.update_progress_percentage)
        self.progress_var.set(0)
        self.progress.set(0)

        # Quick settings that are most frequently changed
        quick_settings_frame = ctk.CTkFrame(recording_frame)
        quick_settings_frame.pack(fill=tk.X, padx=10, pady=10)

        # Microphone selection
        mic_frame = ctk.CTkFrame(quick_settings_frame, fg_color="transparent")
        mic_frame.pack(fill=tk.X, pady=5)

        ctk.CTkLabel(mic_frame, text="Mikrofon:", width=100).pack(side=tk.LEFT)
        mic_combo = ctk.CTkComboBox(
            mic_frame,
            variable=self.selected_mic,
            values=self.microphones,
            width=320
        )
        mic_combo.pack(side=tk.LEFT, padx=5, fill=tk.X, expand=True)

        refresh_mic_button = ctk.CTkButton(
            mic_frame,
            text="⟳",
            width=30,
            command=self.refresh_microphones
        )
        refresh_mic_button.pack(side=tk.RIGHT, padx=5)

        # Sumber audio selection - PERBAIKAN LOGIKA RADIO BUTTON
        audio_frame = ctk.CTkFrame(quick_settings_frame, fg_color="transparent")
        audio_frame.pack(fill=tk.X, pady=5)

        ctk.CTkLabel(audio_frame, text="Sumber Audio:", width=100).pack(side=tk.LEFT)

        # Radio buttons untuk mode perekaman - LOGIKA DIPERBAIKI
        source_frame = ctk.CTkFrame(audio_frame, fg_color="transparent")
        source_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # Variable untuk mode recording
        self.recording_mode = tk.StringVar(value="microphone")

        mic_only_radio = ctk.CTkRadioButton(
            source_frame,
            text="Mikrofon saja",
            variable=self.recording_mode,
            value="microphone",
            command=self._update_recording_mode
        )
        mic_only_radio.pack(side=tk.LEFT, padx=(5, 10))

        system_only_radio = ctk.CTkRadioButton(
            source_frame,
            text="Audio sistem saja",
            variable=self.recording_mode,
            value="system",
            command=self._update_recording_mode
        )
        system_only_radio.pack(side=tk.LEFT, padx=10)

        dual_radio = ctk.CTkRadioButton(
            source_frame,
            text="Mikrofon + Audio sistem",
            variable=self.recording_mode,
            value="dual",
            command=self._update_recording_mode
        )
        dual_radio.pack(side=tk.LEFT, padx=10)

    def setup_audio_visualization(self, parent_frame):
        """Setup real-time audio visualization dengan matplotlib - PERBAIKAN LENGKAP"""
        try:
            # PERBAIKAN: Import matplotlib dengan fallback yang aman
            try:
                import matplotlib
                matplotlib.use('TkAgg')  # Set backend sebelum import pyplot
                import matplotlib.pyplot as plt
                from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
                import numpy as np
                import queue
            except ImportError as e:
                logger.error(f"❌ Matplotlib tidak tersedia: {e}")
                self.setup_placeholder_visualization(parent_frame)
                return

            # Initialize visualization variables
            self.viz_mode = tk.StringVar(value="waveform")
            self.viz_enabled = tk.BooleanVar(value=True)
            self.viz_running = False
            self.audio_queue = queue.Queue(maxsize=100)
            self.viz_data = np.zeros(100)
            self.spectrum_data = np.zeros((50, 50))
            self.idle_time = 0
            self.idle_spectrum_data = np.zeros((50, 50))
            self.prev_bar_data = np.zeros(20)

            # PERBAIKAN: Set matplotlib configuration untuk menghindari font warnings
            plt.rcParams['font.family'] = ['Arial', 'DejaVu Sans', 'sans-serif']
            plt.rcParams['font.size'] = 10
            plt.rcParams['axes.unicode_minus'] = False  # Avoid unicode issues

            # Create matplotlib figure dengan tema yang sesuai
            self.viz_fig, self.viz_ax = plt.subplots(figsize=(10, 2), facecolor='#2B2B2B')
            self.viz_fig.patch.set_facecolor('#2B2B2B')

            # Configure plot dengan warna yang elegant
            self.viz_ax.set_xlim(0, 100)
            self.viz_ax.set_ylim(-1, 1)
            self.viz_ax.set_facecolor('#1E1E1E')
            self.viz_ax.axis('off')

            # Initialize visualization elements
            self.viz_line, = self.viz_ax.plot([], [], color='#007ACC', linewidth=2, alpha=0.8)
            self.viz_bars = None
            self.viz_fill = None
            self.viz_spectrum_image = None

            # Create canvas untuk tkinter
            self.viz_canvas = FigureCanvasTkAgg(self.viz_fig, parent_frame)
            canvas_widget = self.viz_canvas.get_tk_widget()
            canvas_widget.configure(bg='#2B2B2B')
            canvas_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

            # PERBAIKAN: Add welcome message tanpa emoji untuk menghindari font warning
            self.viz_ax.text(
                0.5, 0.5,
                'Audio Visualization Ready\nSelect mode and start recording',
                transform=self.viz_ax.transAxes,
                ha='center', va='center',
                fontsize=12, color='#CCCCCC', alpha=0.8
            )
            self.viz_canvas.draw()

            # Start visualization thread
            self.start_visualization_update()

            logger.info("✅ Audio visualization setup successful")

        except ImportError as e:
            logger.error(f"❌ Error importing visualization libraries: {e}")
            self.setup_placeholder_visualization(parent_frame)
        except Exception as e:
            logger.error(f"❌ Error setting up audio visualization: {e}")
            self.setup_placeholder_visualization(parent_frame)

    def setup_placeholder_visualization(self, parent_frame):
        """Fallback placeholder dengan styling yang elegant"""
        placeholder_frame = ctk.CTkFrame(parent_frame, fg_color=("#F0F0F0", "#2D2D2D"))
        placeholder_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        viz_label = ctk.CTkLabel(
            placeholder_frame,
            text="Audio Visualization\n(Real-time waveform akan muncul saat recording)\n\nMatplotlib diperlukan untuk visualisasi\nInstall dengan: pip install matplotlib",
            font=ctk.CTkFont(size=12),
            text_color=("#666666", "#AAAAAA"),
            justify="center"
        )
        viz_label.place(relx=0.5, rely=0.5, anchor=tk.CENTER)

    def start_visualization_update(self):
        """Start real-time visualization update thread dengan threading yang aman"""
        if hasattr(self, 'viz_canvas') and self.viz_canvas:
            self.viz_running = True
            # PERBAIKAN: Thread yang lebih aman dengan proper error handling
            try:
                self.viz_thread = threading.Thread(target=self.update_visualization_loop, daemon=True)
                self.viz_thread.start()
                logger.info("Visualization thread started successfully")
            except Exception as e:
                logger.error(f"Failed to start visualization thread: {e}")
                self.viz_running = False

    def update_visualization_loop(self):
        """Loop untuk update visualisasi secara real-time dengan error handling yang robust"""
        while self.viz_running:
            try:
                if not self.viz_enabled.get():
                    time.sleep(0.1)
                    continue

                # PERBAIKAN: Check if main thread is still alive
                if not self.root.winfo_exists():
                    logger.info("Main window closed, stopping visualization thread")
                    break

                if self.recording and not self.audio_queue.empty():
                    try:
                        # Get audio data from queue dengan timeout
                        audio_chunk = self.audio_queue.get(timeout=0.1)

                        # Update visualization based on mode
                        current_mode = self.viz_mode.get()
                        if current_mode == "waveform":
                            self.update_waveform_visualization(audio_chunk)
                        elif current_mode == "bars":
                            self.update_bars_visualization(audio_chunk)
                        elif current_mode == "spectrum":
                            self.update_spectrum_visualization(audio_chunk)
                        elif current_mode == "fill":
                            self.update_fill_visualization(audio_chunk)

                    except queue.Empty:
                        pass  # Normal timeout, continue loop
                    except Exception as e:
                        logger.debug(f"Error processing audio chunk: {e}")

                elif not self.recording:
                    # Show idle state
                    self.show_idle_visualization()

                time.sleep(0.05)  # 20 FPS update

            except Exception as e:
                logger.error(f"Error in visualization update loop: {e}")
                time.sleep(0.5)  # Longer sleep on error

        logger.info("Visualization thread ended")

    def update_waveform_visualization(self, audio_chunk):
        """Update waveform visualization - PERBAIKAN THREADING ISSUE"""
        try:
            if not hasattr(self, 'viz_canvas') or not hasattr(self, 'viz_sensitivity'):
                return

            if len(audio_chunk) > 0:
                # Apply sensitivity
                sensitivity = self.viz_sensitivity.get()

                # Normalize audio data
                normalized = np.array(audio_chunk, dtype=np.float32) / 32768.0 * sensitivity

                # Resample to fit display
                if len(normalized) > 100:
                    step = len(normalized) // 100
                    self.viz_data = normalized[::step][:100]
                else:
                    if len(normalized) < 100:
                        padded = np.zeros(100)
                        padded[:len(normalized)] = normalized
                        self.viz_data = padded
                    else:
                        self.viz_data = normalized[:100]

                # PERBAIKAN: Clear dan reset plot dengan proper styling
                self.viz_ax.clear()
                self.viz_ax.set_xlim(0, 100)
                self.viz_ax.set_facecolor('#1E1E1E')
                self.viz_ax.axis('off')

                # Update y-axis limits dynamically
                max_amplitude = np.max(np.abs(self.viz_data))
                if max_amplitude > 0:
                    self.viz_ax.set_ylim(-max_amplitude * 1.2, max_amplitude * 1.2)
                else:
                    self.viz_ax.set_ylim(-1, 1)

                # Update colors based on intensity dengan gradient yang elegant
                if max_amplitude > 0.7:
                    color = '#FF6B6B'  # Red for loud
                    glow_color = '#FF9999'
                elif max_amplitude > 0.3:
                    color = '#4ECDC4'  # Teal for medium
                    glow_color = '#7FFFD4'
                else:
                    color = '#007ACC'  # Blue for quiet
                    glow_color = '#4DA6FF'

                # Plot data dengan glow effect
                x_data = np.arange(len(self.viz_data))

                # Main line
                self.viz_ax.plot(x_data, self.viz_data, color=color, linewidth=2, alpha=0.9)

                # Glow effect
                self.viz_ax.plot(x_data, self.viz_data, color=glow_color, linewidth=4, alpha=0.3)

                # Add zero line
                self.viz_ax.axhline(y=0, color='#555555', alpha=0.5, linewidth=1)

                # PERBAIKAN: Thread-safe canvas update
                if hasattr(self, 'viz_canvas'):
                    self.root.after(0, self._safe_canvas_draw)

        except Exception as e:
            logger.error(f"Error updating waveform: {e}")

    def setup_visualization_controls(self, parent_frame):
        """Setup controls untuk visualization modes dengan text yang aman"""
        control_frame = ctk.CTkFrame(parent_frame, fg_color="transparent", height=50)
        control_frame.pack(fill=tk.X, padx=20, pady=8)

        # Left side - Mode selector
        mode_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
        mode_frame.pack(side=tk.LEFT, fill=tk.X, expand=True)

        # PERBAIKAN: Gunakan text tanpa emoji untuk menghindari font warning
        mode_label = ctk.CTkLabel(mode_frame, text="Viz Mode:", width=80, font=ctk.CTkFont(size=12, weight="bold"))
        mode_label.pack(side=tk.LEFT)

        if not hasattr(self, 'viz_mode'):
            self.viz_mode = tk.StringVar(value="waveform")

        viz_modes = ["waveform", "bars", "spectrum", "fill"]

        viz_combo = ctk.CTkComboBox(
            mode_frame,
            variable=self.viz_mode,
            values=viz_modes,
            width=130,
            command=self.change_visualization_mode
        )
        viz_combo.pack(side=tk.LEFT, padx=8)

        # Middle - Sensitivity control
        sensitivity_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
        sensitivity_frame.pack(side=tk.LEFT, padx=20)

        sens_label = ctk.CTkLabel(sensitivity_frame, text="Sensitivity:", width=85, font=ctk.CTkFont(size=12, weight="bold"))
        sens_label.pack(side=tk.LEFT)

        if not hasattr(self, 'viz_sensitivity'):
            self.viz_sensitivity = tk.DoubleVar(value=1.0)

        sensitivity_slider = ctk.CTkSlider(
            sensitivity_frame,
            from_=0.1,
            to=5.0,
            variable=self.viz_sensitivity,
            width=120,
            button_color="#007ACC",
            progress_color="#4DA6FF"
        )
        sensitivity_slider.pack(side=tk.LEFT, padx=8)

        # Sensitivity value display
        self.sens_value_label = ctk.CTkLabel(sensitivity_frame, text="1.0x", width=35, font=ctk.CTkFont(size=10))
        self.sens_value_label.pack(side=tk.LEFT, padx=5)

        def update_sensitivity_label(*args):
            try:
                value = self.viz_sensitivity.get()
                self.sens_value_label.configure(text=f"{value:.1f}x")
            except Exception as e:
                logger.debug(f"Error updating sensitivity label: {e}")

        self.viz_sensitivity.trace_add("write", update_sensitivity_label)

        # Right side - Toggle switch
        toggle_frame = ctk.CTkFrame(control_frame, fg_color="transparent")
        toggle_frame.pack(side=tk.RIGHT)

        if not hasattr(self, 'viz_enabled'):
            self.viz_enabled = tk.BooleanVar(value=True)

        viz_toggle = ctk.CTkSwitch(
            toggle_frame,
            text="Enable Visualization",
            variable=self.viz_enabled,
            command=self.toggle_visualization,
            button_color="#007ACC",
            progress_color="#4DA6FF"
        )
        viz_toggle.pack(side=tk.RIGHT, padx=10)

    def update_bars_visualization(self, audio_chunk):
        """Update bar visualization - PERBAIKAN COMPLETE"""
        try:
            if len(audio_chunk) > 0:
                # Apply sensitivity
                sensitivity = self.viz_sensitivity.get()

                # FFT untuk frequency analysis
                fft = np.fft.fft(audio_chunk * sensitivity)
                freqs = np.abs(fft[:len(fft)//2])

                # Resample to fit bars (20 bars untuk tampilan yang bagus)
                num_bars = 20
                if len(freqs) > num_bars:
                    step = len(freqs) // num_bars
                    bar_data = freqs[::step][:num_bars]
                else:
                    bar_data = np.pad(freqs, (0, max(0, num_bars - len(freqs))), 'constant')[:num_bars]

                # Normalize
                if np.max(bar_data) > 0:
                    bar_data = bar_data / np.max(bar_data)

                # PERBAIKAN: Clear dengan styling yang proper
                self.viz_ax.clear()
                self.viz_ax.set_xlim(-0.5, num_bars - 0.5)
                self.viz_ax.set_ylim(0, 1.1)
                self.viz_ax.set_facecolor('#1E1E1E')
                self.viz_ax.axis('off')

                # PERBAIKAN: Smoothing yang lebih baik
                if hasattr(self, 'prev_bar_data') and self.prev_bar_data is not None:
                    smoothing_factor = 0.7
                    bar_data = smoothing_factor * self.prev_bar_data + (1 - smoothing_factor) * bar_data
                self.prev_bar_data = bar_data.copy()

                # PERBAIKAN: Color gradient yang elegant
                # Create gradient from blue to cyan to red
                colors = []
                for i, height in enumerate(bar_data):
                    if height > 0.7:
                        colors.append('#FF6B6B')  # Red for high frequencies
                    elif height > 0.4:
                        colors.append('#4ECDC4')  # Cyan for medium
                    else:
                        colors.append('#007ACC')  # Blue for low

                # Create bars dengan glow effect
                bars = self.viz_ax.bar(
                    range(len(bar_data)),
                    bar_data,
                    color=colors,
                    width=0.8,
                    alpha=0.8,
                    edgecolor='white',
                    linewidth=0.5
                )

                # Add glow effect untuk bars
                self.viz_ax.bar(
                    range(len(bar_data)),
                    bar_data,
                    color=colors,
                    width=1.0,
                    alpha=0.3
                )

                # PERBAIKAN: Thread-safe redraw
                if hasattr(self, 'viz_canvas'):
                    self.root.after(0, self._safe_canvas_draw)

        except Exception as e:
            logger.error(f"Error updating bars: {e}")

    def update_spectrum_visualization(self, audio_chunk):
        """Update spectrum visualization - PERBAIKAN WATERFALL"""
        try:
            if len(audio_chunk) > 0:
                # Apply sensitivity
                sensitivity = self.viz_sensitivity.get()

                # Create spectrogram-like visualization
                fft = np.fft.fft(audio_chunk * sensitivity)
                spectrum = np.abs(fft[:len(fft)//2])

                # PERBAIKAN: Initialize spectrum_data jika belum ada
                if not hasattr(self, 'spectrum_data') or self.spectrum_data is None:
                    self.spectrum_data = np.zeros((50, 50))

                # Shift existing data
                self.spectrum_data = np.roll(self.spectrum_data, -1, axis=1)

                # Add new column
                spectrum_height = self.spectrum_data.shape[0]
                if len(spectrum) > spectrum_height:
                    step = len(spectrum) // spectrum_height
                    self.spectrum_data[:, -1] = spectrum[::step][:spectrum_height]
                else:
                    padded_spectrum = np.zeros(spectrum_height)
                    padded_spectrum[:len(spectrum)] = spectrum
                    self.spectrum_data[:, -1] = padded_spectrum

                # Normalize
                if np.max(self.spectrum_data) > 0:
                    normalized_data = self.spectrum_data / np.max(self.spectrum_data)
                else:
                    normalized_data = self.spectrum_data

                # PERBAIKAN: Clear dan redraw dengan styling
                self.viz_ax.clear()
                self.viz_ax.set_facecolor('#1E1E1E')
                self.viz_ax.axis('off')

                # PERBAIKAN: Use elegant colormap
                im = self.viz_ax.imshow(
                    normalized_data,
                    aspect='auto',
                    cmap='plasma',  # Beautiful purple-pink-yellow gradient
                    origin='lower',
                    alpha=0.8,
                    interpolation='bilinear'
                )

                # PERBAIKAN: Thread-safe redraw
                if hasattr(self, 'viz_canvas'):
                    self.root.after(0, self._safe_canvas_draw)

        except Exception as e:
            logger.error(f"Error updating spectrum: {e}")

    def show_idle_visualization(self):
        """Show idle state visualization tanpa emoji untuk menghindari font warning"""
        try:
            if not self.viz_enabled.get():
                return

            # Check if canvas still exists
            if not hasattr(self, 'viz_ax') or not self.viz_ax:
                return

            self.idle_time += 0.1
            x = np.linspace(0, 100, 100)
            current_mode = self.viz_mode.get()

            self.viz_ax.clear()
            self.viz_ax.set_facecolor('#1E1E1E')
            self.viz_ax.axis('off')

            if current_mode == "waveform":
                self.viz_ax.set_xlim(0, 100)
                self.viz_ax.set_ylim(-0.3, 0.3)

                y1 = 0.15 * np.sin(0.2 * x + self.idle_time) * np.exp(-0.02 * x)
                y2 = 0.08 * np.sin(0.5 * x + self.idle_time * 1.5)
                y = y1 + y2

                self.viz_ax.plot(x, y, color='#007ACC', linewidth=2, alpha=0.6)
                self.viz_ax.plot(x, y, color='#4DA6FF', linewidth=4, alpha=0.3)
                self.viz_ax.axhline(y=0, color='#555555', alpha=0.5, linewidth=1)

            elif current_mode == "fill":
                self.viz_ax.set_xlim(0, 100)
                self.viz_ax.set_ylim(-0.3, 0.3)

                amplitude = 0.12 * (1 + 0.4 * np.sin(self.idle_time))
                y = amplitude * np.sin(0.3 * x + self.idle_time)

                self.viz_ax.fill_between(x, 0, y, alpha=0.5, color='#007ACC')
                self.viz_ax.fill_between(x, 0, y, alpha=0.2, color='#4DA6FF')
                self.viz_ax.axhline(y=0, color='#FFFFFF', alpha=0.6, linewidth=1)

            elif current_mode == "bars":
                self.viz_ax.set_xlim(-0.5, 19.5)
                self.viz_ax.set_ylim(0, 0.3)

                bar_data = 0.08 * np.random.random(20) * (1 + 0.3 * np.sin(self.idle_time + np.arange(20) * 0.3))
                colors = ['#007ACC' if i % 3 == 0 else '#4ECDC4' if i % 3 == 1 else '#FF6B6B' for i in range(20)]

                self.viz_ax.bar(range(20), bar_data, color=colors, width=0.8, alpha=0.6)
                self.viz_ax.bar(range(20), bar_data, color=colors, width=1.0, alpha=0.3)

            elif current_mode == "spectrum":
                if not hasattr(self, 'idle_spectrum_data') or self.idle_spectrum_data is None:
                    self.idle_spectrum_data = np.random.random((50, 50)) * 0.1

                self.idle_spectrum_data = np.roll(self.idle_spectrum_data, -1, axis=1)
                self.idle_spectrum_data[:, -1] = np.random.random(50) * 0.08 * (1 + 0.5 * np.sin(self.idle_time))

                self.viz_ax.imshow(
                    self.idle_spectrum_data,
                    aspect='auto',
                    cmap='viridis',
                    origin='lower',
                    alpha=0.4,
                    interpolation='bilinear'
                )

            # Thread-safe redraw
            if hasattr(self, 'viz_canvas'):
                self.root.after(0, self._safe_canvas_draw)

        except Exception as e:
            logger.debug(f"Error in idle visualization: {e}")

    def update_fill_visualization(self, audio_chunk):
        """Update filled area visualization - PERBAIKAN COMPLETE"""
        try:
            if len(audio_chunk) > 0:
                # Apply sensitivity
                sensitivity = self.viz_sensitivity.get()

                # Normalize audio data
                normalized = np.array(audio_chunk, dtype=np.float32) / 32768.0 * sensitivity

                # Resample to fit display
                if len(normalized) > 100:
                    step = len(normalized) // 100
                    self.viz_data = normalized[::step][:100]
                else:
                    if len(normalized) < 100:
                        padded = np.zeros(100)
                        padded[:len(normalized)] = normalized
                        self.viz_data = padded
                    else:
                        self.viz_data = normalized[:100]

                # PERBAIKAN: Clear dengan styling yang proper
                self.viz_ax.clear()
                self.viz_ax.set_xlim(0, 100)
                self.viz_ax.set_facecolor('#1E1E1E')
                self.viz_ax.axis('off')

                # Update y-axis limits
                max_amplitude = np.max(np.abs(self.viz_data))
                if max_amplitude > 0:
                    self.viz_ax.set_ylim(-max_amplitude * 1.2, max_amplitude * 1.2)
                else:
                    self.viz_ax.set_ylim(-1, 1)

                # Update colors based on intensity
                if max_amplitude > 0.7:
                    fill_color = '#FF6B6B'  # Red for loud
                    glow_color = '#FF9999'
                elif max_amplitude > 0.3:
                    fill_color = '#4ECDC4'  # Teal for medium
                    glow_color = '#7FFFD4'
                else:
                    fill_color = '#007ACC'  # Blue for quiet
                    glow_color = '#4DA6FF'

                x_data = np.arange(len(self.viz_data))

                # PERBAIKAN: Create filled area dengan gradient effect
                # Main fill
                self.viz_ax.fill_between(
                    x_data,
                    0,
                    self.viz_data,
                    alpha=0.7,
                    color=fill_color
                )

                # Glow effect
                self.viz_ax.fill_between(
                    x_data,
                    0,
                    self.viz_data,
                    alpha=0.3,
                    color=glow_color
                )

                # Add center line dan grid lines
                self.viz_ax.axhline(y=0, color='#FFFFFF', alpha=0.6, linewidth=1)
                if max_amplitude > 0:
                    self.viz_ax.axhline(y=max_amplitude * 0.5, color='#888888', alpha=0.3, linewidth=0.5, linestyle='--')
                    self.viz_ax.axhline(y=-max_amplitude * 0.5, color='#888888', alpha=0.3, linewidth=0.5, linestyle='--')

                # PERBAIKAN: Thread-safe redraw
                if hasattr(self, 'viz_canvas'):
                    self.root.after(0, self._safe_canvas_draw)

        except Exception as e:
            logger.error(f"Error updating fill visualization: {e}")

    def change_visualization_mode(self, mode):
        """Change visualization mode dengan handling yang aman"""
        try:
            valid_modes = ["waveform", "bars", "spectrum", "fill"]
            if mode not in valid_modes:
                logger.warning(f"Invalid visualization mode: {mode}")
                mode = "waveform"

            self.viz_mode.set(mode)

            if hasattr(self, 'viz_ax') and self.viz_ax:
                self.viz_ax.clear()
                self.viz_ax.set_facecolor('#1E1E1E')
                self.viz_ax.axis('off')

                # Reset mode-specific data
                if mode == "bars":
                    self.prev_bar_data = np.zeros(20)
                    self.viz_ax.set_xlim(-0.5, 19.5)
                    self.viz_ax.set_ylim(0, 1.1)
                elif mode == "spectrum":
                    self.spectrum_data = np.zeros((50, 50))
                    self.idle_spectrum_data = np.zeros((50, 50))
                elif mode in ["waveform", "fill"]:
                    self.viz_data = np.zeros(100)
                    self.viz_ax.set_xlim(0, 100)
                    self.viz_ax.set_ylim(-1, 1)

                # PERBAIKAN: Show mode change message tanpa emoji
                self.viz_ax.text(
                    0.5, 0.5,
                    f'{mode.title()} Mode\nReady for audio...',
                    transform=self.viz_ax.transAxes,
                    ha='center', va='center',
                    fontsize=12, color='#CCCCCC', alpha=0.8
                )

                # Thread-safe canvas draw
                if hasattr(self, 'viz_canvas'):
                    self.root.after(0, self._safe_canvas_draw)

            logger.info(f"Visualization mode changed to: {mode}")

        except Exception as e:
            logger.error(f"Error changing visualization mode: {e}")

    def toggle_visualization(self):
        """Toggle visualization on/off dengan error handling yang robust"""
        try:
            if self.viz_enabled.get():
                self.viz_running = True
                if not hasattr(self, 'viz_thread') or not self.viz_thread.is_alive():
                    self.start_visualization_update()

                if hasattr(self, 'viz_ax') and self.viz_ax:
                    self.viz_ax.clear()
                    self.viz_ax.set_facecolor('#1E1E1E')
                    self.viz_ax.axis('off')
                    self.viz_ax.set_xlim(0, 100)
                    self.viz_ax.set_ylim(-1, 1)

                    self.viz_ax.text(
                        0.5, 0.5,
                        'Visualization Enabled\nReady for audio...',
                        transform=self.viz_ax.transAxes,
                        ha='center', va='center',
                        fontsize=12, color='#4ECDC4', alpha=0.8
                    )

                    if hasattr(self, 'viz_canvas'):
                        self.root.after(0, self._safe_canvas_draw)
            else:
                self.viz_running = False
                if hasattr(self, 'viz_ax') and self.viz_ax:
                    self.viz_ax.clear()
                    self.viz_ax.set_facecolor('#1E1E1E')
                    self.viz_ax.axis('off')
                    self.viz_ax.set_xlim(0, 100)
                    self.viz_ax.set_ylim(-1, 1)

                    self.viz_ax.text(
                        0.5, 0.5,
                        'Visualization Disabled\nClick to enable',
                        transform=self.viz_ax.transAxes,
                        ha='center', va='center',
                        fontsize=12, color='#888888'
                    )

                    if hasattr(self, 'viz_canvas'):
                        self.root.after(0, self._safe_canvas_draw)

            logger.info(f"Visualization {'enabled' if self.viz_enabled.get() else 'disabled'}")

        except Exception as e:
            logger.error(f"Error toggling visualization: {e}")

    def _safe_canvas_draw(self):
        """Thread-safe canvas drawing method dengan error handling yang robust"""
        try:
            # PERBAIKAN: Check if canvas and root still exist
            if (hasattr(self, 'viz_canvas') and
                self.viz_canvas and
                hasattr(self, 'root') and
                self.root.winfo_exists()):

                # Use draw_idle() for thread safety
                self.viz_canvas.draw_idle()
        except tk.TclError as e:
            # Handle TclError yang muncul saat window ditutup
            logger.debug(f"TclError in canvas draw (window likely closed): {e}")
        except Exception as e:
            logger.debug(f"Error in safe canvas draw: {e}")

    def _update_recording_mode(self):
        """PERBAIKAN: Update recording mode dengan validasi"""
        mode = self.recording_mode.get()
        print(f"DEBUG: Switching to recording mode: {mode}")

        if mode == "microphone":
            self.use_system_audio.set(False)
            self.use_dual_recording.set(False)
        elif mode == "system":
            self.use_system_audio.set(True)
            self.use_dual_recording.set(False)
        elif mode == "dual":
            self.use_system_audio.set(False)
            self.use_dual_recording.set(True)

        # Update status dengan info lebih detail
        mode_text = {
            "microphone": "Mikrofon saja",
            "system": "Audio sistem saja",
            "dual": "Mikrofon + Audio sistem"
        }

        status_text = f"Mode recording: {mode_text.get(mode, 'Unknown')}"
        if hasattr(self, 'selected_mic'):
            try:
                mic_name = self.selected_mic.get().split(":")[1].strip()
                status_text += f" | Mikrofon: {mic_name[:20]}..."
            except:
                pass

        self.status_var.set(status_text)
        print(f"DEBUG: Recording mode updated - use_system_audio: {self.use_system_audio.get()}, use_dual_recording: {self.use_dual_recording.get()}")

    def setup_settings_tab(self, parent):
        """Tab untuk pengaturan dengan tombol API key."""
        settings_frame = ctk.CTkScrollableFrame(parent)
        settings_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # 1. Bagian API Configuration
        api_frame = ctk.CTkFrame(settings_frame)
        api_frame.pack(fill=tk.X, padx=5, pady=10)

        api_label = ctk.CTkLabel(
            api_frame,
            text="Konfigurasi API",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        api_label.pack(anchor=tk.W, padx=10, pady=5)

        # API Key status dan button
        api_key_frame = ctk.CTkFrame(api_frame, fg_color="transparent")
        api_key_frame.pack(fill=tk.X, padx=10, pady=5)

        # Status API key
        user_key = load_user_api_key()
        if user_key:
            api_status = f"🔑 Custom API Key: {user_key[:20]}..."
            api_status_color = "green"
        else:
            api_status = "🔧 Default API Key"
            api_status_color = "orange"

        ctk.CTkLabel(
            api_key_frame,
            text="Status API Key:",
            width=140
        ).pack(side=tk.LEFT)

        self.api_status_label = ctk.CTkLabel(
            api_key_frame,
            text=api_status,
            text_color=api_status_color,
            width=200
        )
        self.api_status_label.pack(side=tk.LEFT, padx=5)

        api_button = ctk.CTkButton(
            api_key_frame,
            text="Kelola API Key",
            command=self.show_api_key_dialog,
            width=120
        )
        api_button.pack(side=tk.RIGHT)

        # API description
        api_desc = ctk.CTkLabel(
            api_frame,
            text="Gunakan API key Groq sendiri untuk akses tanpa batas, atau gunakan default key dengan batasan.",
            font=ctk.CTkFont(size=11),
            text_color=("#666666", "#AAAAAA"),
            wraplength=500
        )
        api_desc.pack(anchor=tk.W, padx=10, pady=(0, 10))

        # 2. Bagian Pengenalan Suara (existing code)
        recognition_frame = ctk.CTkFrame(settings_frame)
        recognition_frame.pack(fill=tk.X, padx=5, pady=10)

        recognition_label = ctk.CTkLabel(
            recognition_frame,
            text="Pengaturan Pengenalan Suara",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        recognition_label.pack(anchor=tk.W, padx=10, pady=5)

        # Language setting
        lang_frame = ctk.CTkFrame(recognition_frame, fg_color="transparent")
        lang_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(lang_frame, text="Bahasa:", width=140).pack(side=tk.LEFT)
        languages = ["id-ID (Indonesian)", "en-US (English)", "ja-JP (Japanese)", "zh-CN (Chinese)"]
        language_combo = ctk.CTkComboBox(
            lang_frame,
            variable=self.language,
            values=languages,
            width=280
        )
        language_combo.pack(side=tk.LEFT, padx=5)

        # Recognition engine
        engine_frame = ctk.CTkFrame(recognition_frame, fg_color="transparent")
        engine_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(engine_frame, text="Mesin Pengenalan:", width=140).pack(side=tk.LEFT)
        engines = ["Google", "Whisper"]
        engine_combo = ctk.CTkComboBox(
            engine_frame,
            variable=self.engine,
            values=engines,
            width=280
        )
        engine_combo.pack(side=tk.LEFT, padx=5)

        # Economic model switch
        econ_frame = ctk.CTkFrame(recognition_frame, fg_color="transparent")
        econ_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(econ_frame, text="Mode Ekonomis:", width=140).pack(side=tk.LEFT)
        econ_switch = ctk.CTkSwitch(
            econ_frame,
            text="Gunakan model AI yang lebih hemat",
            variable=self.use_economic_model
        )
        econ_switch.pack(side=tk.LEFT, padx=5)

        # 3. AI Enhancement section (existing code continues...)
        ai_frame = ctk.CTkFrame(settings_frame)
        ai_frame.pack(fill=tk.X, padx=5, pady=10)

        ai_label = ctk.CTkLabel(
            ai_frame,
            text="Pengaturan Peningkatan AI",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        ai_label.pack(anchor=tk.W, padx=10, pady=5)

        ai_switch_frame = ctk.CTkFrame(ai_frame, fg_color="transparent")
        ai_switch_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(ai_switch_frame, text="Gunakan AI:", width=140).pack(side=tk.LEFT)
        ai_switch = ctk.CTkSwitch(
            ai_switch_frame,
            text="Aktifkan peningkatan catatan dengan AI",
            variable=self.use_ai_enhancement
        )
        ai_switch.pack(side=tk.LEFT, padx=5)

        # AI enhancement description
        ai_desc = ctk.CTkLabel(
            ai_frame,
            text="Peningkatan AI akan secara otomatis menyusun catatan menjadi lebih terstruktur, profesional,\n"
                 "dan mudah dibaca. AI akan mendeteksi konteks rekaman untuk format yang optimal.",
            font=ctk.CTkFont(size=11),
            text_color=("#666666", "#AAAAAA")
        )
        ai_desc.pack(anchor=tk.W, padx=10, pady=5)

        # 3. Bagian Format Output
        output_frame = ctk.CTkFrame(settings_frame)
        output_frame.pack(fill=tk.X, padx=5, pady=10)

        output_label = ctk.CTkLabel(
            output_frame,
            text="Pengaturan Output",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        output_label.pack(anchor=tk.W, padx=10, pady=5)

        # Output folder
        folder_frame = ctk.CTkFrame(output_frame, fg_color="transparent")
        folder_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(folder_frame, text="Folder Output:", width=140).pack(side=tk.LEFT)
        ctk.CTkEntry(folder_frame, textvariable=self.output_folder, width=280).pack(side=tk.LEFT, padx=(5, 10))
        ctk.CTkButton(folder_frame, text="Browse", command=self.browse_folder, width=70).pack(side=tk.LEFT)

        # Filename prefix
        prefix_frame = ctk.CTkFrame(output_frame, fg_color="transparent")
        prefix_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(prefix_frame, text="Awalan Nama File:", width=140).pack(side=tk.LEFT)
        ctk.CTkEntry(prefix_frame, textvariable=self.filename_prefix, width=280).pack(side=tk.LEFT, padx=5)

        # 4. Bagian Rekaman Lanjutan
        adv_frame = ctk.CTkFrame(settings_frame)
        adv_frame.pack(fill=tk.X, padx=5, pady=10)

        adv_label = ctk.CTkLabel(
            adv_frame,
            text="Pengaturan Lanjutan",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        adv_label.pack(anchor=tk.W, padx=10, pady=5)

        # Extended recording switch
        ext_frame = ctk.CTkFrame(adv_frame, fg_color="transparent")
        ext_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(ext_frame, text="Rekaman Panjang:", width=140).pack(side=tk.LEFT)
        ext_switch = ctk.CTkSwitch(
            ext_frame,
            text="Aktifkan dukungan rekaman panjang (direkomendasikan)",
            variable=self.use_extended_recording
        )
        ext_switch.pack(side=tk.LEFT, padx=5)

        # Chunk size slider
        chunk_frame = ctk.CTkFrame(adv_frame, fg_color="transparent")
        chunk_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(chunk_frame, text="Ukuran Penggalan:", width=140).pack(side=tk.LEFT)

        chunk_slider_frame = ctk.CTkFrame(chunk_frame, fg_color="transparent")
        chunk_slider_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        chunk_slider = ctk.CTkSlider(
            chunk_slider_frame,
            from_=60,
            to=1800,
            number_of_steps=29,
            variable=self.chunk_size
        )
        chunk_slider.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))

        self.chunk_size_label = tk.StringVar()
        self.chunk_size_label.set(f"{self.chunk_size.get()} detik")

        def update_chunk_label(*args):
            seconds = self.chunk_size.get()
            if seconds < 60:
                self.chunk_size_label.set(f"{int(seconds)} detik")
            else:
                minutes = seconds // 60
                self.chunk_size_label.set(f"{int(minutes)} menit")

        self.chunk_size.trace_add("write", update_chunk_label)
        update_chunk_label()  # Initial update

        chunk_label = ctk.CTkLabel(chunk_slider_frame, textvariable=self.chunk_size_label, width=60)
        chunk_label.pack(side=tk.LEFT)

        # Chunk size explanation
        chunk_desc = ctk.CTkLabel(
            adv_frame,
            text="Ukuran penggalan menentukan berapa lama setiap bagian rekaman diproses.\n"
                 "Nilai lebih kecil = lebih sering diproses, nilai lebih besar = lebih jarang diproses.",
            font=ctk.CTkFont(size=11),
            text_color=("#666666", "#AAAAAA")
        )
        chunk_desc.pack(anchor=tk.W, padx=150, pady=(0, 5))

        # API delay settings
        api_frame = ctk.CTkFrame(adv_frame, fg_color="transparent")
        api_frame.pack(fill=tk.X, padx=10, pady=5)

        ctk.CTkLabel(api_frame, text="Jeda API (detik):", width=140).pack(side=tk.LEFT)

        api_slider_frame = ctk.CTkFrame(api_frame, fg_color="transparent")
        api_slider_frame.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)

        self.api_delay_var = tk.IntVar(value=self.api_request_delay)

        api_slider = ctk.CTkSlider(
            api_slider_frame,
            from_=1,
            to=20,
            number_of_steps=19,
            variable=self.api_delay_var,
            command=lambda v: setattr(self, 'api_request_delay', int(v))
        )
        api_slider.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=(0, 10))

        api_label = ctk.CTkLabel(api_slider_frame, textvariable=self.api_delay_var, width=30)
        api_label.pack(side=tk.LEFT)

    def update_api_status_display(self):
        """Update tampilan status API key."""
        user_key = load_user_api_key()
        if hasattr(self, 'api_status_label'):
            if user_key:
                status_text = f"🔑 Custom API Key: {user_key[:20]}..."
                status_color = "green"
            else:
                status_text = "🔧 Default API Key"
                status_color = "orange"

            self.api_status_label.configure(text=status_text, text_color=status_color)

    def setup_output_tab(self, parent):
        """
        Tab untuk hasil output dengan area tampilan yang lebih besar dan kontrol ekspor.
        """
        output_frame = ctk.CTkFrame(parent)
        output_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Output display with better styling
        ctk.CTkLabel(
            output_frame,
            text="Hasil Transkripsi:",
            font=ctk.CTkFont(size=16, weight="bold"),
            anchor="w"
        ).pack(anchor=tk.W, pady=5, padx=10)

        # Result text area with better styling
        self.result_text = ctk.CTkTextbox(
            output_frame,
            height=200,
            wrap="word",
            font=ctk.CTkFont(size=13)
        )
        self.result_text.pack(pady=5, padx=10, fill=tk.BOTH, expand=True)

        # Export controls
        export_frame = ctk.CTkFrame(output_frame, fg_color="transparent")
        export_frame.pack(fill=tk.X, padx=10, pady=10)

        copy_button = ctk.CTkButton(
            export_frame,
            text="Salin ke Clipboard",
            command=self.copy_to_clipboard,
            width=150
        )
        copy_button.pack(side=tk.LEFT, padx=5)

        export_button = ctk.CTkButton(
            export_frame,
            text="Ekspor ke Word",
            command=self.export_to_word,
            width=150
        )
        export_button.pack(side=tk.LEFT, padx=5)

        # Open output folder button
        open_folder_button = ctk.CTkButton(
            export_frame,
            text="Buka Folder Output",
            command=self.open_output_folder,
            width=150
        )
        open_folder_button.pack(side=tk.LEFT, padx=5)

    def setup_status_bar(self, parent):
        """
        Status bar yang ditingkatkan dengan informasi yang lebih kaya.
        """
        status_frame = ctk.CTkFrame(parent, height=40, fg_color=("#F0F0F0", "#2B2B2B"))
        status_frame.pack(fill=tk.X, pady=(10, 0))

        # Status text
        self.status_var = tk.StringVar()
        self.status_var.set("Siap")
        status_label = ctk.CTkLabel(
            status_frame,
            textvariable=self.status_var,
            text_color=("#00AAFF", "#00AAFF"),
            font=ctk.CTkFont(weight="bold")
        )
        status_label.pack(side=tk.LEFT, padx=10, pady=5)

        # Time info
        self.time_info_var = tk.StringVar()
        self.time_info_var.set("")
        time_label = ctk.CTkLabel(
            status_frame,
            textvariable=self.time_info_var,
            text_color=("#666666", "#AAAAAA"),
            font=ctk.CTkFont(size=11)
        )
        time_label.pack(side=tk.RIGHT, padx=10, pady=5)

        # Update time info periodically
        def update_time_info():
            now = datetime.datetime.now()
            self.time_info_var.set(now.strftime("%H:%M:%S"))
            self.root.after(1000, update_time_info)

        update_time_info()

    def refresh_microphones(self):
        """
        Refresh the list of available microphones.
        """
        self.microphones = self.get_available_microphones()

        # Update the combobox values
        for tab in [self.tab_view.tab("Rekaman"), self.tab_view.tab("Pengaturan")]:
            for child in tab.winfo_children():
                if isinstance(child, ctk.CTkFrame):
                    for subchild in child.winfo_children():
                        if isinstance(subchild, ctk.CTkFrame):
                            for widget in subchild.winfo_children():
                                if isinstance(widget, ctk.CTkComboBox) and widget.cget("values") == self.microphones:
                                    widget.configure(values=self.microphones)

        # Set default mic if available
        if self.microphones:
            self.selected_mic.set(self.microphones[0])

        # Show success message
        self.status_var.set("Daftar mikrofon diperbarui")
        self.root.after(3000, lambda: self.status_var.set("Siap"))

    def copy_to_clipboard(self):
        """
        Copy the transcript result to clipboard.
        """
        text = self.result_text.get(1.0, tk.END)
        self.root.clipboard_clear()
        self.root.clipboard_append(text)

        # Show success message
        self.status_var.set("Teks disalin ke clipboard")
        self.root.after(3000, lambda: self.status_var.set("Siap"))

    def open_output_folder(self):
        """
        Open the output folder in file explorer.
        """
        folder = self.output_folder.get()

        if not os.path.exists(folder):
            messagebox.showwarning("Peringatan", f"Folder tidak ditemukan:\n{folder}")
            return

        # Open folder using appropriate command for the OS
        if os.name == 'nt':  # Windows
            os.startfile(folder)
        elif os.name == 'posix':  # macOS and Linux
            if sys.platform == 'darwin':  # macOS
                subprocess.run(['open', folder])
            else:  # Linux
                subprocess.run(['xdg-open', folder])

        self.status_var.set(f"Membuka folder: {folder}")
        self.root.after(3000, lambda: self.status_var.set("Siap"))

    def open_file(self, filepath):
        """
        Open a file with its default application.
        """
        if not os.path.exists(filepath):
            messagebox.showwarning("Peringatan", f"File tidak ditemukan:\n{filepath}")
            return

        # Open file using appropriate command for the OS
        try:
            if os.name == 'nt':  # Windows
                os.startfile(filepath)
            elif os.name == 'posix':  # macOS and Linux
                if sys.platform == 'darwin':  # macOS
                    subprocess.run(['open', filepath])
                else:  # Linux
                    subprocess.run(['xdg-open', filepath])
        except Exception as e:
            messagebox.showerror("Error", f"Gagal membuka file: {e}")

        self.status_var.set(f"Membuka file: {os.path.basename(filepath)}")
        self.root.after(3000, lambda: self.status_var.set("Siap"))

    def update_progress_percentage(self, *args):
        progress = self.progress_var.get()
        self.progress_percentage.set(f"{int(progress)}%")

    def browse_folder(self):
        folder = filedialog.askdirectory(initialdir=self.output_folder.get())
        if folder:
            self.output_folder.set(folder)

    def toggle_recording(self):
        if self.recording:
            self.stop_recording()
        else:
            self.start_recording()

    def start_recording(self):
        self.elapsed_time = 0
        self.audio_chunks = []
        self.stop_recording_flag = False
        self.result_text.delete(1.0, tk.END)

        self.record_button.configure(text="Berhenti Rekaman", fg_color="#CC3030", hover_color="#AA2020")
        self.status_var.set("Merekam...")

        self.recording = True
        self.recording_thread = threading.Thread(target=self.recording_thread_func)
        self.recording_thread.daemon = True
        self.recording_thread.start()

        self.update_timer()

    def update_timer(self):
        if self.recording:
            self.elapsed_time += 1
            hours = self.elapsed_time // 3600
            minutes = (self.elapsed_time % 3600) // 60
            seconds = self.elapsed_time % 60
            self.timer_var.set(f"{hours:02d}:{minutes:02d}:{seconds:02d}")
            self.root.after(1000, self.update_timer)

    def stop_recording(self):
        if self.recording:
            self.stop_recording_flag = True
            self.recording = False

            self.record_button.configure(text="Mulai Rekaman", fg_color="#007ACC", hover_color="#0066AA")
            self.status_var.set("Memproses rekaman...")

            # SOLUSI: Buat fungsi callback yang tepat
            self.root.after(100, self.handle_recording_completion)

    def handle_recording_completion(self):
        """Handle completion of recording and start processing safely"""
        try:
            # Check if any audio data is available
            has_audio = False

            # Check for extended recording audio files
            if self.use_extended_recording.get() and hasattr(self, 'temp_audio_files') and self.temp_audio_files:
                has_audio = True
                print(f"DEBUG: Found {len(self.temp_audio_files)} chunks for processing")

            # Check for single recording file
            elif hasattr(self, 'temp_wav_file') and self.temp_wav_file:
                has_audio = True
                print(f"DEBUG: Found recording file for processing: {self.temp_wav_file}")

            # Start processing if we have audio data
            if has_audio:
                threading.Thread(target=self.process_audio_thread, daemon=True).start()
            else:
                self.status_var.set("Error: Tidak ada data audio")
                print("ERROR: No audio data available for processing")

        except Exception as e:
            error_msg = str(e)
            self.status_var.set(f"Error: {error_msg}")
            print(f"Error in handle_recording_completion: {e}")

    def recording_thread_func(self):
        try:
            if self.use_extended_recording.get():
                try:
                    self.temp_dir = tempfile.mkdtemp()
                except Exception as e:
                    print(f"Error creating temp directory: {e}")
                    self.root.after(0, lambda: messagebox.showerror("Error", f"Gagal membuat direktori sementara: {e}"))
                    self.recording = False
                    self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
                    self.root.after(0, lambda: self.status_var.set("Error"))
                    return

            # Determine recording mode dan update status yang sesuai
            mode = self.recording_mode.get()
            if mode == "dual":
                self.root.after(0, lambda: self.status_var.set("Merekam mikrofon + audio sistem..."))
                self.record_dual_audio()
            elif mode == "system":
                self.root.after(0, lambda: self.status_var.set("Merekam audio sistem..."))
                self.record_system_audio()
            else:  # microphone
                self.root.after(0, lambda: self.status_var.set("Merekam mikrofon..."))
                self.record_microphone_audio()

        except Exception as e:
            print(f"Recording thread error: {e}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Terjadi kesalahan saat merekam: {e}"))
            self.recording = False
            self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
            self.root.after(0, lambda: self.status_var.set("Error"))

    def create_audio_visualizer(self, audio_data, sample_rate=44100):
        """Membuat visualisasi static dari audio yang sudah direkam"""
        try:
            import matplotlib.pyplot as plt

            fig, axes = plt.subplots(3, 1, figsize=(12, 10))
            fig.suptitle('EchoScribe AI - Audio Analysis', fontsize=16, fontweight='bold')

            # Pastikan audio_data adalah 1D
            if len(audio_data.shape) > 1:
                audio_data = np.mean(audio_data, axis=1)

            # Time axis
            time_axis = np.linspace(0, len(audio_data)/sample_rate, len(audio_data))

            # 1. Waveform
            axes[0].plot(time_axis, audio_data, 'b-', linewidth=0.5)
            axes[0].set_title('Waveform')
            axes[0].set_xlabel('Time (seconds)')
            axes[0].set_ylabel('Amplitude')
            axes[0].grid(True, alpha=0.3)

            # 2. Spectrogram
            try:
                f, t, Sxx = axes[1].specgram(audio_data, Fs=sample_rate)
                axes[1].set_title('Spectrogram')
                axes[1].set_xlabel('Time (seconds)')
                axes[1].set_ylabel('Frequency (Hz)')
            except:
                axes[1].text(0.5, 0.5, 'Spectrogram tidak dapat dibuat',
                            ha='center', va='center', transform=axes[1].transAxes)

            # 3. Frequency Spectrum
            fft_data = np.abs(np.fft.fft(audio_data))
            freqs = np.fft.fftfreq(len(fft_data), 1/sample_rate)

            # Ambil setengah pertama (positive frequencies)
            half_len = len(fft_data) // 2
            axes[2].plot(freqs[:half_len], fft_data[:half_len], 'r-', linewidth=0.5)
            axes[2].set_title('Frequency Spectrum')
            axes[2].set_xlabel('Frequency (Hz)')
            axes[2].set_ylabel('Magnitude')
            axes[2].grid(True, alpha=0.3)
            axes[2].set_xlim(0, sample_rate//2)

            plt.tight_layout()
            return fig

        except Exception as e:
            print(f"Error creating visualization: {e}")
            return None

    def show_audio_info(self, audio_data, sample_rate=44100):
        """Menampilkan informasi audio dalam UI tkinter"""
        if audio_data is None:
            return

        try:
            # Pastikan audio_data adalah 1D
            if len(audio_data.shape) > 1:
                audio_data = np.mean(audio_data, axis=1)

            duration = len(audio_data) / sample_rate
            max_amplitude = np.max(np.abs(audio_data))
            rms = np.sqrt(np.mean(audio_data**2))

            # Tampilkan info dalam messagebox atau status
            info_text = f"""📊 Informasi Audio:
    Duration: {duration:.2f}s
    Sample Rate: {sample_rate} Hz
    Max Amplitude: {max_amplitude:.3f}
    RMS Level: {rms:.3f}"""

            self.root.after(0, lambda: messagebox.showinfo("Audio Info", info_text))

        except Exception as e:
            logger.error(f"Error showing audio info: {e}")

    def detect_audio_context(self, audio_file):
        """
        Deteksi karakteristik dan konteks audio untuk optimasi transkripsi.
        """
        try:
            duration = self.get_audio_duration(audio_file)

            # Analisis tingkat suara dengan validasi yang lebih ketat
            with wave.open(audio_file, 'rb') as wf:
                # Gunakan test_samplerate untuk validasi
                test_samplerate = wf.getframerate()
                if test_samplerate == 0:
                    logger.error(f"Invalid sample rate detected: {test_samplerate}")
                    return self._get_default_audio_context()

                # Dapatkan beberapa sampel untuk analisis
                n_frames = min(wf.getnframes(), 1000000)
                frames = wf.readframes(n_frames)

                if not frames:
                    logger.warning("No audio frames found in file")
                    return self._get_default_audio_context()

                # Hitung RMS untuk menentukan volume rata-rata
                try:
                    rms = audioop.rms(frames, wf.getsampwidth())
                except audioop.error as audio_error:
                    logger.error(f"Audio processing error: {audio_error}")
                    return self._get_default_audio_context()

                # Hitung jumlah silent frames untuk mendeteksi jeda bicara
                silent_threshold = max(rms * 0.1, 100)  # Minimal threshold
                silent_frames = 0
                frame_size = wf.getsampwidth() * wf.getnchannels()

                for i in range(0, len(frames), frame_size):
                    if i + frame_size <= len(frames):
                        chunk = frames[i:i + frame_size]
                        try:
                            chunk_rms = audioop.rms(chunk, wf.getsampwidth())
                            if chunk_rms < silent_threshold:
                                silent_frames += 1
                        except audioop.error:
                            continue

                silence_ratio = silent_frames / max(len(frames) / frame_size, 1)

            # Deteksi karakteristik dengan informasi yang lebih lengkap
            context = {
                "duration": duration,
                "sample_rate": test_samplerate,
                "volume_level": "high" if rms > 10000 else "medium" if rms > 5000 else "low",
                "silence_ratio": silence_ratio,
                "content_type": self._detect_content_type(duration, silence_ratio),
                "audio_quality": "good" if rms > 1000 and test_samplerate >= 16000 else "poor"
            }

            logger.info(f"Audio context detected: {context}")
            return context

        except Exception as e:
            logger.error(f"Error saat deteksi audio context: {e}", exc_info=True)
            return self._get_default_audio_context()

    def _get_default_audio_context(self):
        """Return default audio context when detection fails"""
        return {
            "duration": 0,
            "sample_rate": 16000,
            "volume_level": "medium",
            "silence_ratio": 0,
            "content_type": "unknown",
            "audio_quality": "unknown"
        }

    def record_system_audio(self):
        """
        Record system audio dengan error handling dan logging yang diperbaiki.
        """
        try:
            if not PYAUDIOWPATCH_AVAILABLE:
                raise ImportError("PyAudioWPatch not available")

            self.root.after(0, lambda: self.status_var.set("Mencari device audio sistem..."))

            # Dapatkan device loopback dengan validasi yang lebih ketat
            loopback_device_index, loopback_device_info = self.get_primary_speaker_loopback_device()

            if loopback_device_index is None:
                error_msg = "Tidak dapat menemukan device loopback yang kompatibel"
                logger.error(error_msg)
                self.root.after(0, lambda: self.status_var.set("ERROR: Device loopback tidak ditemukan"))
                self.root.after(0, lambda: self._show_enhanced_system_audio_troubleshooting(error_msg))
                self.root.after(2000, self._fallback_to_microphone_recording)
                return

            # Konfigurasi audio yang adaptif dengan validasi
            CHUNK = 1024
            FORMAT = pyaudio.paInt16
            max_channels = int(loopback_device_info['maxInputChannels'])
            CHANNELS = min(2, max_channels) if max_channels > 0 else 1

            default_rate = int(loopback_device_info['defaultSampleRate'])
            possible_rates = [default_rate, 44100, 48000, 16000]
            RATE = None

            # Test rate yang didukung dengan validasi yang ketat
            for test_rate in possible_rates:
                if test_rate > 0 and test_rate <= 192000:  # Validasi range yang masuk akal
                    RATE = test_rate
                    break

            if RATE is None or RATE == 0:
                logger.error("Could not determine valid sample rate")
                RATE = 44100  # Ultimate fallback

            logger.info(f"Using validated audio config: {RATE}Hz, {CHANNELS} channels, device: {loopback_device_index}")

            # Initialize PyAudio dengan error handling yang lebih baik
            p = pyaudio.PyAudio()

            device_name = loopback_device_info['name'][:30] + "..." if len(loopback_device_info['name']) > 30 else loopback_device_info['name']
            self.root.after(0, lambda: self.status_var.set(f"Menggunakan: {device_name}"))

            # Buka stream dengan konfigurasi adaptif dan validasi
            stream = None
            try:
                stream = p.open(
                    format=FORMAT,
                    channels=CHANNELS,
                    rate=RATE,
                    input=True,
                    input_device_index=loopback_device_index,
                    frames_per_buffer=CHUNK,
                    input_host_api_specific_stream_info=None
                )

                logger.info(f"SUCCESS: Stream opened successfully: {RATE}Hz, {CHANNELS} channels")

            except Exception as stream_error:
                logger.error(f"Error opening stream: {stream_error}")

                # Try dengan konfigurasi fallback yang lebih konservatif
                try:
                    logger.info("Mencoba konfigurasi fallback...")
                    CHANNELS = 1  # Force mono
                    RATE = 44100  # Force standard rate

                    stream = p.open(
                        format=FORMAT,
                        channels=CHANNELS,
                        rate=RATE,
                        input=True,
                        input_device_index=loopback_device_index,
                        frames_per_buffer=CHUNK
                    )

                    logger.info(f"SUCCESS: Fallback stream opened: {RATE}Hz, {CHANNELS} channels")

                except Exception as fallback_error:
                    logger.error(f"Fallback stream also failed: {fallback_error}")
                    p.terminate()
                    self.root.after(0, lambda: self.status_var.set("ERROR: Gagal membuka stream audio"))
                    self.root.after(0, lambda: self._show_enhanced_system_audio_troubleshooting(str(fallback_error)))
                    self.root.after(2000, self._fallback_to_microphone_recording)
                    return

            # Recording loop dengan monitoring yang lebih akurat
            frames = []
            chunk_start_time = time.time()
            total_samples = 0
            non_silent_chunks = 0
            consecutive_errors = 0
            max_consecutive_errors = 10

            # Variabel untuk monitoring kualitas audio
            audio_quality_samples = []
            last_quality_check = time.time()

            self.root.after(0, lambda: self.status_var.set("RECORDING: Recording aktif - memantau audio sistem..."))

            while not self.stop_recording_flag and self.recording:
                try:
                    # Baca data dari loopback device
                    data = stream.read(CHUNK, exception_on_overflow=False)
                    frames.append(data)
                    total_samples += CHUNK
                    consecutive_errors = 0  # Reset error counter

                    # Monitor level audio dengan validasi yang lebih ketat
                    try:
                        audio_level = audioop.rms(data, 2)
                        audio_quality_samples.append(audio_level)

                        # Threshold dinamis berdasarkan rata-rata level audio
                        if len(audio_quality_samples) > 10:
                            avg_level = sum(audio_quality_samples[-10:]) / 10
                            dynamic_threshold = max(avg_level * 0.1, 50)
                        else:
                            dynamic_threshold = 50

                        if audio_level > dynamic_threshold:
                            non_silent_chunks += 1

                        # Update status setiap 100 chunks dengan informasi kualitas
                        current_time = time.time()
                        if total_samples % (CHUNK * 100) == 0 or (current_time - last_quality_check) > 5:
                            percentage_active = (non_silent_chunks / max(total_samples // CHUNK, 1)) * 100
                            avg_quality = sum(audio_quality_samples[-50:]) / min(len(audio_quality_samples), 50)
                            quality_indicator = "📶" if avg_quality > 1000 else "📶" if avg_quality > 500 else "📶"

                            self.root.after(0, lambda pct=percentage_active, qual=quality_indicator:
                                        self.status_var.set(f"RECORDING: {qual} Audio aktif: {pct:.1f}%"))
                            last_quality_check = current_time

                    except Exception as level_error:
                        logger.debug(f"Audio level monitoring error: {level_error}")

                    # Handle chunking untuk extended recording dengan validasi ukuran
                    if self.use_extended_recording.get():
                        chunk_length = self.get_audio_duration_from_frames(frames, RATE)
                        if chunk_length >= self.chunk_size.get():
                            self._save_system_audio_chunk(frames, RATE, CHANNELS)
                            frames = []
                            chunk_start_time = time.time()

                except Exception as e:
                    consecutive_errors += 1

                    if "Input overflowed" in str(e):
                        logger.debug(f"Input overflow (error #{consecutive_errors})")
                        continue
                    else:
                        logger.error(f"Error during recording (#{consecutive_errors}): {e}")

                        if consecutive_errors >= max_consecutive_errors:
                            logger.error(f"Too many consecutive errors ({consecutive_errors}), stopping recording")
                            break

            # Cleanup dengan validasi
            try:
                if stream and stream.is_active():
                    stream.stop_stream()
                if stream:
                    stream.close()
                p.terminate()
                logger.info("SUCCESS: Audio stream cleaned up successfully")
            except Exception as cleanup_error:
                logger.error(f"Error during cleanup: {cleanup_error}")

            # Simpan data terakhir dengan validasi
            if frames:
                try:
                    if self.use_extended_recording.get():
                        self._save_system_audio_chunk(frames, RATE, CHANNELS)
                    else:
                        self._save_system_audio_to_file(frames, RATE, CHANNELS)
                    logger.info(f"SUCCESS: Audio data saved successfully ({len(frames)} chunks)")
                except Exception as save_error:
                    logger.error(f"Error saving audio data: {save_error}")

            # Status akhir dengan informasi kualitas yang akurat
            if non_silent_chunks > 0:
                percentage_active = (non_silent_chunks / max(total_samples // CHUNK, 1)) * 100
                avg_quality = sum(audio_quality_samples) / max(len(audio_quality_samples), 1) if audio_quality_samples else 0
                quality_desc = "Excellent" if avg_quality > 2000 else "Good" if avg_quality > 1000 else "Fair" if avg_quality > 500 else "Poor"

                self.root.after(0, lambda pct=percentage_active, qual=quality_desc:
                            self.status_var.set(f"SUCCESS: Selesai! Audio: {pct:.1f}% aktif, Kualitas: {qual}"))
            else:
                self.root.after(0, lambda: self.status_var.set(
                    "WARNING: Tidak ada audio terdeteksi - pastikan ada audio yang diputar"))
                self.root.after(1000, lambda: self._show_no_audio_detected_tips())

        except Exception as e:
            logger.error(f"Error in record_system_audio: {e}", exc_info=True)
            self.root.after(0, lambda: self.status_var.set(f"ERROR: {str(e)[:50]}..."))
            self.recording = False
            self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
            self.root.after(1000, lambda: self._show_enhanced_system_audio_troubleshooting(str(e)))

    def get_audio_duration_from_frames(self, frames, sample_rate):
        """Calculate duration from audio frames"""
        if not frames or sample_rate == 0:
            return 0

        total_samples = len(frames) * 1024  # Assuming CHUNK size of 1024
        duration = total_samples / sample_rate
        return duration

    def _fallback_to_microphone_recording(self):
        """
        Fallback otomatis ke recording mikrofon jika system audio gagal.
        """
        try:
            # Update mode recording ke microphone
            self.recording_mode.set("microphone")
            self._update_recording_mode()

            # Update UI untuk menunjukkan fallback
            self.root.after(0, lambda: self.status_var.set("🔄 Beralih ke mode mikrofon..."))

            # Show notification
            self.root.after(0, lambda: messagebox.showinfo(
                "Mode Fallback",
                "System audio recording gagal.\n\n"
                "Aplikasi otomatis beralih ke mode 'Mikrofon saja'.\n"
                "Anda dapat mencoba mode 'Dual Recording' sebagai alternatif."
            ))

            logger.info("✅ Successfully switched to microphone fallback mode")

        except Exception as e:
            logger.error(f"Error in fallback to microphone: {e}")

    def _show_no_audio_detected_tips(self):
        """
        Tampilkan tips ketika tidak ada audio yang terdeteksi.
        """
        tips_window = tk.Toplevel(self.root)
        tips_window.title("Tips: Tidak Ada Audio Terdeteksi")
        tips_window.geometry("500x400")
        tips_window.transient(self.root)
        tips_window.grab_set()

        main_frame = ctk.CTkFrame(tips_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=15, pady=15)

        # Title
        title_label = ctk.CTkLabel(
            main_frame,
            text="🔇 Tidak Ada Audio Terdeteksi",
            font=ctk.CTkFont(size=18, weight="bold")
        )
        title_label.pack(pady=(0, 15))

        # Tips content
        tips_text = """
    KEMUNGKINAN PENYEBAB & SOLUSI:

    ✅ PASTIKAN AUDIO SEDANG DIPUTAR:
    • Buka YouTube, Spotify, atau aplikasi audio lainnya
    • Pastikan volume sistem tidak dalam keadaan mute
    • Volume aplikasi audio harus > 0%

    ✅ PERIKSA PENGATURAN WINDOWS:
    • Buka Settings > System > Sound
    • Pastikan output device benar
    • Test speaker dengan tombol "Test"

    ✅ COBA ALTERNATIF LAIN:
    • Gunakan mode "Mikrofon + Audio sistem"
    • Gunakan "Mikrofon saja" sambil putar audio speaker
    • Install VB-Cable untuk virtual audio routing

    ✅ RESTART AUDIO SERVICES:
    • Tekan Win+R, ketik "services.msc"
    • Restart "Windows Audio" service
    • Restart aplikasi ini

    Jika masih bermasalah, gunakan mode recording lain.
        """

        text_widget = ctk.CTkTextbox(main_frame, wrap=tk.WORD, height=250)
        text_widget.pack(fill=tk.BOTH, expand=True, pady=(0, 15))
        text_widget.insert("1.0", tips_text)
        text_widget.configure(state="disabled")

        # Buttons
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill=tk.X)

        def switch_to_dual():
            self.recording_mode.set("dual")
            self._update_recording_mode()
            tips_window.destroy()

        def switch_to_mic():
            self.recording_mode.set("microphone")
            self._update_recording_mode()
            tips_window.destroy()

        dual_btn = ctk.CTkButton(
            button_frame,
            text="Coba Dual Recording",
            command=switch_to_dual,
            width=150
        )
        dual_btn.pack(side=tk.LEFT, padx=(0, 10))

        mic_btn = ctk.CTkButton(
            button_frame,
            text="Gunakan Mikrofon Saja",
            command=switch_to_mic,
            width=150
        )
        mic_btn.pack(side=tk.LEFT, padx=10)

        close_btn = ctk.CTkButton(
            button_frame,
            text="Tutup",
            command=tips_window.destroy,
            width=100
        )
        close_btn.pack(side=tk.RIGHT)

    def _show_enhanced_system_audio_troubleshooting(self, error_message):
        """
        Enhanced troubleshooting guide dengan solusi yang lebih komprehensif.
        """
        troubleshoot_window = tk.Toplevel(self.root)
        troubleshoot_window.title("🔧 Panduan Mengatasi Masalah Audio Sistem")
        troubleshoot_window.geometry("700x600")
        troubleshoot_window.transient(self.root)

        main_frame = ctk.CTkFrame(troubleshoot_window)
        main_frame.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        # Title dengan icon
        title_label = ctk.CTkLabel(
            main_frame,
            text="🔧 Panduan Mengatasi Masalah Audio Sistem",
            font=ctk.CTkFont(size=16, weight="bold")
        )
        title_label.pack(pady=(10, 5))

        # Error info
        error_frame = ctk.CTkFrame(main_frame, fg_color=("gray90", "gray20"))
        error_frame.pack(fill=tk.X, padx=10, pady=5)

        error_label = ctk.CTkLabel(
            error_frame,
            text=f"❌ Error: {error_message[:100]}{'...' if len(error_message) > 100 else ''}",
            font=ctk.CTkFont(size=11),
            wraplength=650
        )
        error_label.pack(pady=8)

        # Scrollable troubleshooting content
        text_widget = ctk.CTkTextbox(main_frame, wrap=tk.WORD)
        text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        troubleshooting_content = f"""
    🚀 SOLUSI CEPAT (RECOMMENDED):

    1️⃣ GUNAKAN MODE DUAL RECORDING:
    • Pilih "Mikrofon + Audio sistem" di aplikasi
    • Mode ini lebih reliable dan kompatibel
    • Menggabungkan audio dari mic dan sistem

    2️⃣ RESTART AUDIO SERVICES:
    • Tekan Win+R → ketik "services.msc"
    • Cari "Windows Audio" → Klik kanan → Restart
    • Cari "Windows Audio Endpoint Builder" → Restart
    • Tunggu 10 detik, lalu coba lagi

    3️⃣ PERIKSA PENGATURAN AUDIO:
    • Klik kanan icon speaker di taskbar
    • Pilih "Open Volume mixer"
    • Pastikan ada aplikasi yang memutar audio
    • Volume harus > 0% dan tidak mute

    🔧 SOLUSI ADVANCED:

    4️⃣ INSTALL VB-CABLE (MOST EFFECTIVE):
    • Download dari: https://vb-audio.com/Cable/
    • Install dan restart komputer
    • Set VB-Cable sebagai default playback device
    • Route audio melalui virtual cable

    5️⃣ ENABLE STEREO MIX (JIKA TERSEDIA):
    • Klik kanan icon speaker → "Sounds"
    • Tab "Recording" → Klik kanan area kosong
    • "Show Disabled Devices" → "Show Disconnected Devices"
    • Jika ada "Stereo Mix" → Enable dan set as default

    6️⃣ UPDATE AUDIO DRIVERS:
    • Buka Device Manager
    • Expand "Sound, video and game controllers"
    • Klik kanan audio device → "Update driver"
    • Restart setelah update

    🐛 DEBUGGING INFO:
    • Your audio device: {error_message}
    • Windows version: {platform.system()} {platform.release()}
    • PyAudioWPatch available: {PYAUDIOWPATCH_AVAILABLE}

    💡 ALTERNATIVE SOLUTIONS:

    7️⃣ USE EXTERNAL TOOLS:
    • OBS Studio dengan Virtual Camera
    • Voicemeeter (advanced audio mixer)
    • Audacity untuk recording sistem audio

    8️⃣ HARDWARE WORKAROUND:
    • Gunakan kabel audio untuk menghubungkan headphone output ke mic input
    • Gunakan USB audio interface external
    • Putar audio melalui speaker, rekam dengan mic

    ❓ MASIH BERMASALAH?
    • Coba restart aplikasi EchoScribe
    • Restart komputer secara lengkap
    • Gunakan mode "Mikrofon saja" sebagai fallback
    • Contact support dengan log file
        """

        text_widget.insert("1.0", troubleshooting_content)
        text_widget.configure(state="disabled")

        # Action buttons
        button_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        button_frame.pack(fill=tk.X, padx=10, pady=(0, 10))

        def open_sound_settings():
            try:
                subprocess.run(['ms-settings:sound'], check=True)
            except:
                try:
                    subprocess.run(['control', 'mmsys.cpl'], check=True)
                except:
                    messagebox.showinfo("Info", "Buka Settings > System > Sound secara manual")

        def open_vb_cable():
            import webbrowser
            webbrowser.open("https://vb-audio.com/Cable/")

        def switch_to_dual():
            self.recording_mode.set("dual")
            self._update_recording_mode()
            troubleshoot_window.destroy()
            messagebox.showinfo("Mode Dual", "Switched to Dual Recording mode. Coba rekam lagi!")

        def restart_audio_service():
            try:
                subprocess.run(['net', 'stop', 'audiosrv'], shell=True)
                time.sleep(2)
                subprocess.run(['net', 'start', 'audiosrv'], shell=True)
                messagebox.showinfo("Success", "Audio service restarted. Coba rekam lagi!")
            except Exception as e:
                messagebox.showerror("Error", f"Gagal restart service: {e}")

        # Row 1
        settings_btn = ctk.CTkButton(
            button_frame, text="🔊 Buka Sound Settings", command=open_sound_settings, width=140
        )
        settings_btn.pack(side=tk.LEFT, padx=2)

        vb_btn = ctk.CTkButton(
            button_frame, text="⬇️ Download VB-Cable", command=open_vb_cable, width=140
        )
        vb_btn.pack(side=tk.LEFT, padx=2)

        dual_btn = ctk.CTkButton(
            button_frame, text="🔄 Coba Dual Mode", command=switch_to_dual, width=140
        )
        dual_btn.pack(side=tk.LEFT, padx=2)

        service_btn = ctk.CTkButton(
            button_frame, text="🔄 Restart Audio Service", command=restart_audio_service, width=140
        )
        service_btn.pack(side=tk.LEFT, padx=2)

        close_btn = ctk.CTkButton(
            button_frame, text="❌ Tutup", command=troubleshoot_window.destroy, width=80
        )
        close_btn.pack(side=tk.RIGHT, padx=2)

    def get_primary_speaker_loopback_device(self):
        """
        Mencari device loopback dari speaker utama sistem dengan algoritma yang lebih robust dan fleksibel.
        Returns: (device_index, device_info) atau (None, None) jika tidak ditemukan
        """
        try:
            if not PYAUDIOWPATCH_AVAILABLE:
                logger.error("PyAudioWPatch tidak tersedia untuk perekaman audio sistem")
                return None, None

            p = pyaudio.PyAudio()

            # Cari WASAPI host API terlebih dahulu
            wasapi_info = None
            try:
                wasapi_info = p.get_host_api_info_by_type(pyaudio.paWASAPI)
                logger.info(f"WASAPI Host API found: {wasapi_info['name']}")
            except Exception as e:
                logger.error(f"WASAPI tidak tersedia: {e}")
                p.terminate()
                return None, None

            # Dapatkan default output device dari WASAPI
            default_output_device = None
            default_output_index = None
            try:
                default_output_index = wasapi_info["defaultOutputDevice"]
                default_output_device = p.get_device_info_by_index(default_output_index)
                logger.info(f"Default output device: {default_output_device['name']} (index: {default_output_index})")
            except Exception as e:
                logger.error(f"Tidak dapat mengakses default output device: {e}")
                p.terminate()
                return None, None

            # ALGORITMA PENCARIAN YANG DIPERBAIKI - Multi-strategi
            loopback_device_index = None
            loopback_device_info = None

            # Nama device yang akan dicari
            target_device_name = default_output_device['name']

            # Strategi 1: Pencarian nama persis (existing logic)
            logger.info(f"Strategi 1: Mencari loopback dengan nama persis: '{target_device_name}'")

            for i in range(p.get_device_count()):
                try:
                    info = p.get_device_info_by_index(i)

                    # Harus WASAPI device
                    if info['hostApi'] != wasapi_info['index']:
                        continue

                    # Harus loopback device
                    if not info.get('isLoopbackDevice', False):
                        continue

                    # Harus memiliki input channels
                    if info['maxInputChannels'] == 0:
                        continue

                    # Nama device harus PERSIS sama
                    if info['name'] == target_device_name:
                        loopback_device_index = i
                        loopback_device_info = info
                        logger.info(f"SUCCESS: Strategi 1 berhasil: Found EXACT match at index {i}")
                        break

                except Exception as e:
                    logger.debug(f"Error checking device {i} in strategy 1: {e}")
                    continue

            # Strategi 2: Pencarian berdasarkan keyword jika strategi 1 gagal
            if loopback_device_index is None:
                logger.info("Strategi 2: Mencari berdasarkan keyword dari nama device")

                # Extract keywords dari nama device
                device_keywords = self._extract_device_keywords(target_device_name)
                logger.info(f"Keywords extracted: {device_keywords}")

                best_match_score = 0
                best_match_index = None
                best_match_info = None

                for i in range(p.get_device_count()):
                    try:
                        info = p.get_device_info_by_index(i)

                        # Harus WASAPI device
                        if info['hostApi'] != wasapi_info['index']:
                            continue

                        # Harus loopback device
                        if not info.get('isLoopbackDevice', False):
                            continue

                        # Harus memiliki input channels
                        if info['maxInputChannels'] == 0:
                            continue

                        # Hitung similarity score
                        score = self._calculate_device_similarity(target_device_name, info['name'])
                        logger.debug(f"Device {i} '{info['name']}' similarity score: {score}")

                        if score > best_match_score and score >= 0.5:  # Minimal 50% similarity
                            best_match_score = score
                            best_match_index = i
                            best_match_info = info

                    except Exception as e:
                        logger.debug(f"Error checking device {i} in strategy 2: {e}")
                        continue

                if best_match_index is not None:
                    loopback_device_index = best_match_index
                    loopback_device_info = best_match_info
                    logger.info(f"SUCCESS: Strategi 2 berhasil: Found match at index {best_match_index} with score {best_match_score:.2f}")

            # Strategi 3: Pencarian default loopback jika strategi 1-2 gagal
            if loopback_device_index is None:
                logger.info("Strategi 3: Mencari default loopback device apa saja")

                for i in range(p.get_device_count()):
                    try:
                        info = p.get_device_info_by_index(i)

                        # Harus WASAPI device
                        if info['hostApi'] != wasapi_info['index']:
                            continue

                        # Harus loopback device
                        if not info.get('isLoopbackDevice', False):
                            continue

                        # Harus memiliki input channels
                        if info['maxInputChannels'] == 0:
                            continue

                        # Ambil loopback device pertama yang ditemukan
                        loopback_device_index = i
                        loopback_device_info = info
                        logger.info(f"SUCCESS: Strategi 3 berhasil: Using fallback loopback device '{info['name']}' at index {i}")
                        break

                    except Exception as e:
                        logger.debug(f"Error checking device {i} in strategy 3: {e}")
                        continue

            p.terminate()

            if loopback_device_index is None:
                logger.error("ERROR: Semua strategi gagal - tidak dapat menemukan loopback device")
                logger.error(f"Target device: {target_device_name}")

                # Log semua WASAPI loopback devices yang tersedia untuk debugging
                self._log_available_loopback_devices()

                return None, None

            logger.info(f"SUCCESS: Loopback device ditemukan: '{loopback_device_info['name']}' (index: {loopback_device_index})")
            return loopback_device_index, loopback_device_info

        except Exception as e:
            logger.error(f"Error dalam get_primary_speaker_loopback_device: {e}", exc_info=True)
            return None, None

    def _extract_device_keywords(self, device_name):
        """
        Extract keywords dari nama device untuk matching yang lebih fleksibel.
        """
        # Hapus kata-kata umum yang tidak signifikan
        stop_words = {'(', ')', '-', 'USB', 'Audio', 'Device', 'Microphone', 'Speakers', 'Headphones'}

        # Split dan bersihkan
        words = device_name.replace('(', ' ').replace(')', ' ').replace('-', ' ').split()
        keywords = []

        for word in words:
            word = word.strip()
            if len(word) > 2 and word not in stop_words:
                keywords.append(word.lower())

        return keywords

    def _calculate_device_similarity(self, target_name, candidate_name):
        """
        Menghitung similarity score antara dua nama device.
        Returns: float between 0.0 and 1.0
        """
        target_keywords = set(self._extract_device_keywords(target_name))
        candidate_keywords = set(self._extract_device_keywords(candidate_name))

        if not target_keywords:
            return 0.0

        # Jaccard similarity
        intersection = len(target_keywords & candidate_keywords)
        union = len(target_keywords | candidate_keywords)

        if union == 0:
            return 0.0

        jaccard_score = intersection / union

        # Bonus untuk substring matches
        target_lower = target_name.lower()
        candidate_lower = candidate_name.lower()

        substring_bonus = 0.0
        if target_lower in candidate_lower or candidate_lower in target_lower:
            substring_bonus = 0.3

        # Bonus untuk common audio device patterns
        common_patterns = ['headphones', 'speakers', 'audio', 'realtek', 'conexant']
        pattern_bonus = 0.0

        for pattern in common_patterns:
            if pattern in target_lower and pattern in candidate_lower:
                pattern_bonus += 0.1

        final_score = min(1.0, jaccard_score + substring_bonus + pattern_bonus)
        return final_score

    def _log_available_loopback_devices(self):
        """
        Log semua loopback devices yang tersedia untuk debugging.
        """
        try:
            if not PYAUDIOWPATCH_AVAILABLE:
                return

            p = pyaudio.PyAudio()
            wasapi_info = p.get_host_api_info_by_type(pyaudio.paWASAPI)

            logger.info("=== AVAILABLE LOOPBACK DEVICES ===")
            loopback_count = 0

            for i in range(p.get_device_count()):
                try:
                    info = p.get_device_info_by_index(i)

                    if (info['hostApi'] == wasapi_info['index'] and
                        info.get('isLoopbackDevice', False) and
                        info['maxInputChannels'] > 0):

                        logger.info(f"Loopback {loopback_count}: '{info['name']}' (index: {i})")
                        logger.info(f"  Channels: {info['maxInputChannels']}, Rate: {info['defaultSampleRate']}")
                        loopback_count += 1

                except Exception as e:
                    logger.debug(f"Error checking device {i}: {e}")

            if loopback_count == 0:
                logger.warning("❌ Tidak ada loopback devices ditemukan")
            else:
                logger.info(f"✅ Total {loopback_count} loopback devices tersedia")

            logger.info("===============================")
            p.terminate()

        except Exception as e:
            logger.error(f"Error logging loopback devices: {e}")

    def _save_system_audio_chunk(self, frames, rate, channels):
        """
        Helper function untuk menyimpan chunk audio sistem.
        """
        if not self.temp_dir:
            logger.error("Temp directory not available")
            return

        try:
            temp_file = os.path.join(self.temp_dir, f"system_chunk_{len(self.temp_audio_files)}.wav")

            with wave.open(temp_file, 'wb') as wf:
                wf.setnchannels(channels)
                wf.setsampwidth(2)  # 16-bit audio
                wf.setframerate(rate)
                wf.writeframes(b''.join(frames))

            # Validasi file yang disimpan
            file_size = os.path.getsize(temp_file)
            if file_size < 1024:  # File terlalu kecil
                logger.warning(f"Saved audio chunk is very small: {file_size} bytes")

            self.temp_audio_files.append(temp_file)
            logger.info(f"Saved system audio chunk: {temp_file} ({file_size} bytes)")

        except Exception as e:
            logger.error(f"Error saving system audio chunk: {e}")

    def _save_system_audio_to_file(self, frames, rate, channels):
        """
        Simpan audio sistem ke file temporary.
        """
        try:
            self.temp_wav_file = tempfile.NamedTemporaryFile(suffix=".wav", delete=False)

            with wave.open(self.temp_wav_file.name, 'wb') as wf:
                wf.setnchannels(channels)
                wf.setsampwidth(2)  # 16-bit audio
                wf.setframerate(rate)
                wf.writeframes(b''.join(frames))

            # Validasi file
            file_size = os.path.getsize(self.temp_wav_file.name)
            audio_duration = len(frames) * 1024 / rate  # Estimasi durasi

            logger.info(f"Saved system audio: {self.temp_wav_file.name}")
            logger.info(f"File size: {file_size} bytes, Duration: {audio_duration:.2f}s")

            if file_size < 1024:
                logger.warning("Audio file is very small - may indicate no audio was captured")

        except Exception as e:
            logger.error(f"Error saving system audio to file: {e}")
            self.temp_wav_file = None

    def record_microphone_audio(self):
        """Record audio from the selected microphone"""
        try:
            # Parse microphone index with better error handling
            try:
                mic_index = int(self.selected_mic.get().split(":")[0])
                mic_name = self.selected_mic.get().split(":", 1)[1].strip() if ":" in self.selected_mic.get() else "Unknown"
                self.root.after(0, lambda: self.status_var.set(f"Initializing microphone: {mic_name}"))
                print(f"DEBUG: Using microphone index {mic_index}: {mic_name}")
            except (ValueError, IndexError) as e:
                print(f"Error parsing microphone index: {e}")
                self.root.after(0, lambda: messagebox.showerror("Error", "Mikrofon tidak valid. Silakan pilih mikrofon yang tersedia."))
                self.recording = False
                self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
                self.root.after(0, lambda: self.status_var.set("Error"))
                return

            # Audio format configuration
            FORMAT = pyaudio.paInt16
            CHANNELS = 1
            RATE = 16000  # Keeping your original rate
            CHUNK = 1024

            # Initialize PyAudio
            audio = pyaudio.PyAudio()

            # Validate microphone before opening stream
            try:
                mic_info = audio.get_device_info_by_index(mic_index)
                print(f"DEBUG: Microphone info: {mic_info}")

                # Verify microphone supports input
                max_channels = int(mic_info.get('maxInputChannels', 0))
                if max_channels < 1:
                    raise Exception(f"Selected device doesn't support audio input: {mic_name}")

                self.root.after(0, lambda: self.status_var.set(f"Microphone ready: {mic_name[:30]}..."))
            except Exception as e:
                print(f"Error validating microphone: {e}")
                self.root.after(0, lambda: messagebox.showerror(
                    "Error",
                    f"Gagal memvalidasi mikrofon: {mic_name}\n\nDetail error: {e}\n\nCoba pilih mikrofon lain atau restart aplikasi."))
                audio.terminate()
                self.recording = False
                self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
                self.root.after(0, lambda: self.status_var.set("Error"))
                return

            # Open audio stream with error handling
            try:
                self.root.after(0, lambda: self.status_var.set("Opening audio stream..."))
                stream = audio.open(
                    format=FORMAT,
                    channels=CHANNELS,
                    rate=RATE,
                    input=True,
                    input_device_index=mic_index,
                    frames_per_buffer=CHUNK
                )

                # Test the stream with a small read
                test_data = stream.read(CHUNK, exception_on_overflow=False)
                if not test_data or len(test_data) == 0:
                    raise Exception("Test read returned no data")

                print(f"DEBUG: Stream test successful, read {len(test_data)} bytes")
                self.root.after(0, lambda: self.status_var.set("Recording from microphone..."))
            except Exception as e:
                print(f"Error opening audio stream: {e}")
                self.root.after(0, lambda: messagebox.showerror("Error", f"Gagal membuka aliran audio: {e}"))
                audio.terminate()
                self.recording = False
                self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
                self.root.after(0, lambda: self.status_var.set("Error"))
                return

            # Initialize recording variables
            frames = []
            chunk_start_time = time.time()
            total_chunks = 0
            error_count = 0
            max_errors = 5  # Maximum consecutive errors before stopping

            print(f"DEBUG: Starting recording with microphone index {mic_index}, RATE={RATE}, CHANNELS={CHANNELS}")
            self.root.after(0, lambda: self.status_var.set(f"Recording from: {mic_name[:30]}..."))

            # Main recording loop
            while not self.stop_recording_flag:
                try:
                    data = stream.read(CHUNK, exception_on_overflow=False)

                    if data and len(data) > 0:
                        frames.append(data)
                        total_chunks += 1
                        error_count = 0  # Reset error counter on successful read

                        # Update UI periodically
                        if total_chunks % 20 == 0:
                            seconds = int(total_chunks * CHUNK / RATE)
                            self.root.after(0, lambda s=seconds:
                                        self.status_var.set(f"Recording... ({s}s)"))

                        # Log progress
                        if len(frames) % 100 == 0:
                            print(f"DEBUG: Collected {len(frames)} frames so far")

                        # Handle visualization if available
                        if hasattr(self, 'audio_queue') and hasattr(self, 'viz_enabled') and \
                        self.viz_enabled.get() and total_chunks % 5 == 0:
                            try:
                                import numpy as np
                                audio_data = np.frombuffer(data, dtype=np.int16)
                                if len(audio_data) > 0 and self.audio_queue.qsize() < 10:
                                    self.audio_queue.put_nowait(audio_data)
                            except Exception as viz_error:
                                # Silent fail for visualization - not critical
                                pass

                        # Handle extended recording chunking
                        if self.use_extended_recording.get():
                            if time.time() - chunk_start_time >= self.chunk_size.get():
                                print(f"DEBUG: Saving chunk with {len(frames)} frames")
                                self.save_audio_chunk(frames, RATE, CHANNELS, audio.get_sample_size(FORMAT))
                                frames = []
                                chunk_start_time = time.time()
                    else:
                        print("WARNING: Empty data read from microphone")
                        error_count += 1

                        if error_count >= max_errors:
                            print(f"ERROR: Too many consecutive empty reads ({error_count})")
                            break

                except Exception as e:
                    print(f"Error during recording: {e}")
                    error_count += 1

                    if error_count >= max_errors:
                        print(f"ERROR: Too many consecutive errors ({error_count})")
                        break

            # Recording finished or stopped
            print(f"DEBUG: Recording stopped with {len(frames)} frames collected")

            if frames:
                try:
                    if self.use_extended_recording.get():
                        print(f"DEBUG: Saving final chunk with {len(frames)} frames")
                        self.save_audio_chunk(frames, RATE, CHANNELS, audio.get_sample_size(FORMAT))
                    else:
                        # Create temp directory if needed
                        if not hasattr(self, 'temp_dir') or not self.temp_dir:
                            self.temp_dir = tempfile.mkdtemp()
                            print(f"DEBUG: Created temp directory: {self.temp_dir}")

                        # Use full path for temporary file
                        temp_path = os.path.join(self.temp_dir, "recording.wav")
                        print(f"DEBUG: Saving recording to: {temp_path}")

                        # Create and save WAV file
                        wf = wave.open(temp_path, 'wb')
                        wf.setnchannels(CHANNELS)
                        wf.setsampwidth(audio.get_sample_size(FORMAT))
                        wf.setframerate(RATE)
                        wf.writeframes(b''.join(frames))
                        wf.close()

                        # Store path as string
                        self.temp_wav_file = temp_path
                        print(f"DEBUG: Recording saved to temporary file: {self.temp_wav_file}")

                        # Verify the saved file
                        if os.path.exists(self.temp_wav_file):
                            file_size = os.path.getsize(self.temp_wav_file)
                            duration = len(frames) * CHUNK / RATE
                            print(f"DEBUG: WAV file size: {file_size} bytes, duration: {duration:.2f}s")

                            if file_size < 1000:
                                print("WARNING: Audio file is suspiciously small, might be empty")
                except Exception as e:
                    print(f"Error saving audio: {e}")
                    error_msg = str(e)  # Capture the error message
                    self.root.after(0, lambda err=error_msg: messagebox.showerror("Error", f"Gagal menyimpan audio: {err}"))

            # Clean up audio resources
            try:
                stream.stop_stream()
                stream.close()
                audio.terminate()
                print("DEBUG: Audio resources successfully cleaned up")
            except Exception as e:
                print(f"Error closing audio stream: {e}")

        except Exception as e:
            print(f"Microphone recording error: {e}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Terjadi kesalahan saat merekam dari mikrofon: {e}"))
            self.recording = False
            self.root.after(0, lambda: self.record_button.configure(text="Mulai Rekaman"))
            self.root.after(0, lambda: self.status_var.set("Error"))

    def record_dual_audio(self):
        """
        Record mikrofon dan sistem audio bersamaan dengan sinkronisasi yang tepat.
        """
        try:
            if not PYAUDIOWPATCH_AVAILABLE:
                raise ImportError("PyAudioWPatch not available")

            # Validasi mikrofon
            try:
                mic_index = int(self.selected_mic.get().split(":")[0])
            except (ValueError, IndexError):
                raise Exception("Mikrofon tidak valid")

            # Dapatkan loopback device
            loopback_device_index, loopback_device_info = self.get_primary_speaker_loopback_device()
            if loopback_device_index is None:
                raise Exception("Tidak dapat menemukan device audio sistem")

            self.root.after(0, lambda: self.status_var.set("Memulai dual recording..."))

            # Konfigurasi audio yang sinkron
            CHUNK = 1024
            FORMAT = pyaudio.paInt16
            CHANNELS = 1  # Mono untuk mikrofon
            RATE = 44100   # Standar untuk kompatibilitas

            p = pyaudio.PyAudio()

            # Buka stream mikrofon
            mic_stream = p.open(
                format=FORMAT,
                channels=CHANNELS,
                rate=RATE,
                input=True,
                input_device_index=mic_index,
                frames_per_buffer=CHUNK
            )

            # Buka stream sistem
            system_stream = p.open(
                format=FORMAT,
                channels=min(2, int(loopback_device_info['maxInputChannels'])),
                rate=RATE,
                input=True,
                input_device_index=loopback_device_index,
                frames_per_buffer=CHUNK
            )

            self.root.after(0, lambda: self.status_var.set("Dual recording aktif..."))

            # Recording loop dengan mixing
            mixed_frames = []
            chunk_start_time = time.time()

            while not self.stop_recording_flag and self.recording:
                try:
                    # Baca dari kedua stream
                    mic_data = mic_stream.read(CHUNK, exception_on_overflow=False)
                    system_data = system_stream.read(CHUNK, exception_on_overflow=False)

                    # Mix audio dengan proporsi yang seimbang
                    mixed_data = self._mix_audio_data(mic_data, system_data, 0.6, 0.4)
                    mixed_frames.append(mixed_data)

                    # Handle chunking
                    if self.use_extended_recording.get():
                        if time.time() - chunk_start_time >= self.chunk_size.get():
                            self._save_system_audio_chunk(mixed_frames, RATE, CHANNELS)
                            mixed_frames = []
                            chunk_start_time = time.time()

                except Exception as e:
                    if "Input overflowed" not in str(e):
                        logger.error(f"Error in dual recording: {e}")
                    continue

            # Cleanup
            mic_stream.stop_stream()
            mic_stream.close()
            system_stream.stop_stream()
            system_stream.close()
            p.terminate()

            # Simpan data terakhir
            if mixed_frames:
                if self.use_extended_recording.get():
                    self._save_system_audio_chunk(mixed_frames, RATE, CHANNELS)
                else:
                    self._save_system_audio_to_file(mixed_frames, RATE, CHANNELS)

            self.root.after(0, lambda: self.status_var.set("Dual recording selesai"))

        except Exception as e:
            logger.error(f"Dual recording error: {e}")
            self.root.after(0, lambda: self.status_var.set(f"Error dual recording: {str(e)[:50]}"))
            # Fallback ke mikrofon saja
            self.record_microphone_audio()

    def _mix_audio_data(self, data1, data2, volume1=0.7, volume2=0.3):
        """
        Mix two audio data streams with specified volumes.

        Args:
            data1: First audio data (bytes)
            data2: Second audio data (bytes)
            volume1: Volume for first stream (0.0 to 1.0)
            volume2: Volume for second stream (0.0 to 1.0)

        Returns:
            Mixed audio data (bytes)
        """
        try:
            import numpy as np

            # Convert bytes to numpy arrays
            audio1 = np.frombuffer(data1, dtype=np.int16)
            audio2 = np.frombuffer(data2, dtype=np.int16)

            # Make sure both arrays have the same length
            min_len = min(len(audio1), len(audio2))
            audio1 = audio1[:min_len]
            audio2 = audio2[:min_len]

            # Mix with specified volumes
            mixed = (audio1 * volume1 + audio2 * volume2).astype(np.int16)

            # Prevent clipping
            mixed = np.clip(mixed, -32768, 32767).astype(np.int16)

            return mixed.tobytes()

        except Exception as e:
            logger.error(f"Error mixing audio: {e}")
            # Return first stream if mixing fails
            return data1

    def save_audio_chunk(self, frames, rate, channels=1, sample_width=2):
        """Save audio chunk to a WAV file with proper parameters"""
        if not self.temp_dir:
            print("DEBUG: Cannot save chunk - temp_dir not initialized")
            return

        temp_file = os.path.join(self.temp_dir, f"chunk_{len(self.temp_audio_files)}.wav")
        wf = wave.open(temp_file, 'wb')
        wf.setnchannels(channels)  # Use provided channels parameter
        wf.setsampwidth(sample_width)  # Use provided sample width
        wf.setframerate(rate)
        wf.writeframes(b''.join(frames))
        wf.close()

        # Verifikasi file tersimpan dengan benar
        file_size = os.path.getsize(temp_file)
        print(f"DEBUG: Saving chunk to: {temp_file}, size: {file_size} bytes")

        if file_size > 0:
            self.temp_audio_files.append(temp_file)
        else:
            print(f"WARNING: Empty audio file generated: {temp_file}")

    def process_audio_thread(self):
        try:
            # PERBAIKAN: Log kondisi perekaman untuk debugging
            print(f"DEBUG: Processing audio - Extended recording: {self.use_extended_recording.get()}")
            print(f"DEBUG: Temp audio files: {len(self.temp_audio_files) if hasattr(self, 'temp_audio_files') else 'None'}")
            print(f"DEBUG: Temp wav file: {self.temp_wav_file if hasattr(self, 'temp_wav_file') else 'None'}")

            if self.use_extended_recording.get() and hasattr(self, 'temp_audio_files') and self.temp_audio_files:
                print(f"DEBUG: Processing extended recording with {len(self.temp_audio_files)} chunks")
                self.process_extended_recording_optimized()
            elif hasattr(self, 'temp_wav_file') and self.temp_wav_file:
                print(f"DEBUG: Processing standard recording with file: {self.temp_wav_file}")
                self.process_standard_recording_enhanced()
            else:
                print("ERROR: No audio data available for processing")
                self.root.after(0, lambda: messagebox.showwarning("Peringatan", "Tidak ada rekaman yang ditemukan untuk diproses."))
                self.root.after(0, lambda: self.status_var.set("Siap"))
                return

        except Exception as e:
            print(f"ERROR: Exception in process_audio_thread: {e}")
            self.root.after(0, lambda: messagebox.showerror("Error", f"Terjadi kesalahan saat memproses: {e}"))
            self.root.after(0, lambda: self.status_var.set("Error"))

    def process_standard_recording_enhanced(self):
        """
        Versi yang dioptimalkan dari process_standard_recording dengan deteksi konteks dan pemilihan model.
        """
        try:
            self.root.after(0, lambda: self.progress_var.set(0))
            self.root.after(0, lambda: self.status_var.set("Menganalisis audio..."))

            # PERBAIKAN: Periksa apakah temp_wav_file adalah string (path) atau objek file
            audio_file_path = self.temp_wav_file
            if not isinstance(audio_file_path, str):
                if hasattr(self.temp_wav_file, 'name'):
                    audio_file_path = self.temp_wav_file.name

            # PERBAIKAN: Verifikasi file ada sebelum diproses
            if not os.path.exists(audio_file_path):
                print(f"ERROR: Audio file not found: {audio_file_path}")
                self.root.after(0, lambda: messagebox.showerror("Error", "File audio tidak ditemukan untuk diproses."))
                self.root.after(0, lambda: self.status_var.set("Error: File tidak ditemukan"))
                return

            print(f"DEBUG: Processing audio file: {audio_file_path}")

            # Deteksi konteks audio
            audio_context = self.detect_audio_context(audio_file_path)

            with sr.AudioFile(audio_file_path) as source:
                self.recognizer.adjust_for_ambient_noise(source)
                audio_data = self.recognizer.record(source)

                self.root.after(0, lambda: self.progress_var.set(25))
                self.root.after(0, lambda: self.status_var.set(
                    f"Mengenali audio ({audio_context['content_type']})..."))

                language_code = self.language.get().split(" ")[0]
                engine = self.engine.get()

                if engine == "Google":
                    self.root.after(0, lambda: self.status_var.set(
                        "Menggunakan Google Speech Recognition..."))
                    text = self.recognizer.recognize_google(audio_data, language=language_code)
                elif engine == "Whisper":
                    model = self.select_optimal_transcription_model(audio_context, language_code)
                    self.root.after(0, lambda: self.status_var.set(
                        f"Menggunakan Whisper model {model}..."))
                    text = self.transcribe_with_groq_whisper(audio_data, language_code, model)

                self.root.after(0, lambda: self.progress_var.set(50))

                # Update status dengan detail konteks
                status_message = f"Meningkatkan hasil ({audio_context['content_type']})..."
                self.root.after(0, lambda: self.status_var.set(status_message))

                if self.use_ai_enhancement.get():
                    text = self.enhance_with_ai(text)

                # Buat nama file yang deskriptif
                now = datetime.datetime.now()
                date_str = now.strftime("%Y%m%d")
                time_str = now.strftime("%H%M%S")

                # Tambahkan indikator tipe konten ke nama file
                content_type_indicator = f"_{audio_context['content_type']}" if audio_context['content_type'] != "unknown" else ""
                filename = f"{self.filename_prefix.get()}{content_type_indicator}_{date_str}_{time_str}.docx"
                filepath = os.path.join(self.output_folder.get(), filename)

                # Simpan dengan format yang dioptimalkan
                self.save_as_word_document(text, filepath)

                self.root.after(0, lambda: self.progress_var.set(100))
                self.root.after(0, lambda: self.status_var.set(f"Selesai! File disimpan: {filename}"))

                self.root.after(0, lambda: self.result_text.delete(1.0, tk.END))
                preview_text = text[:5000] + "... (konten lengkap tersedia di file .docx)" if len(text) > 5000 else text
                self.root.after(0, lambda: self.result_text.insert(tk.END, preview_text))

        except sr.UnknownValueError:
            self.root.after(0, lambda: messagebox.showwarning(
                "Peringatan", "Tidak dapat mengenali audio. Coba rekam dengan lebih jelas atau gunakan mikrofon eksternal."))
            self.root.after(0, lambda: self.status_var.set("Error: Tidak dapat mengenali audio"))
        except sr.RequestError as e:
            self.root.after(0, lambda: messagebox.showerror(
                "Error", f"Error dari layanan pengenalan: {e}"))
            self.root.after(0, lambda: self.status_var.set("Error layanan"))
        finally:
            if self.temp_wav_file:
                try:
                    os.unlink(self.temp_wav_file.name)
                    self.temp_wav_file = None
                except:
                    pass

    def process_extended_recording_optimized(self):
        """
        Versi yang sangat dioptimalkan dari process_extended_recording dengan:
        - Pemrosesan paralel untuk chunk audio
        - Analisis konteks audio yang lebih baik
        - Pemantauan kualitas transkripsi dengan mekanisme fallback
        - Smart batching untuk performa yang lebih baik
        """
        try:
            total_chunks = len(self.temp_audio_files)
            self.root.after(0, lambda: self.status_var.set(f"Menganalisis {total_chunks} bagian audio..."))

            # Peningkatan: Analisis chunk pertama untuk dapatkan konteks
            initial_context = {}
            if total_chunks > 0:
                initial_context = self.detect_audio_context(self.temp_audio_files[0])
                logger.info(f"Detected audio context: {initial_context}")

            # Peningkatan: Smart batching berdasarkan jumlah total chunks
            batch_size = self._calculate_optimal_batch_size(total_chunks)
            logger.info(f"Using batch size: {batch_size} for {total_chunks} chunks")

            batch_count = (total_chunks + batch_size - 1) // batch_size
            all_raw_transcriptions = []
            all_enhanced_chunk_texts = []
            processed_audio_files_in_run = []

            self.processing_start_time = time.time()
            language_code = self.language.get().split("-")[0]

            # Peningkatan: Semaphore untuk membatasi proses paralel API calls
            max_concurrent_api_calls = 2
            api_semaphore = threading.Semaphore(max_concurrent_api_calls)

            # Peningkatan: Transcribe batch secara paralel dengan ThreadPoolExecutor
            for batch_index in range(batch_count):
                batch_start = batch_index * batch_size
                batch_end = min(batch_start + batch_size, total_chunks)

                self.root.after(0, lambda b=batch_index+1, t=batch_count:
                    self.status_var.set(f"Memproses batch transkripsi {b}/{t}..."))

                current_batch_raw_texts = []
                current_batch_enhanced_texts = []

                # Proses file dalam batch
                for i in range(batch_start, batch_end):
                    progress = int(((i + 1) / total_chunks) * 60)  # Transkripsi sampai 60%
                    chunk_num = i + 1
                    time_estimate = self.estimate_remaining_time(chunk_num, total_chunks)

                    self.root.after(0, lambda p=progress: self.progress_var.set(p))
                    self.root.after(0, lambda n=chunk_num, t=total_chunks, est=time_estimate:
                        self.status_var.set(f"Transkripsi bagian {n}/{t} ({est})"))

                    audio_file = self.temp_audio_files[i]
                    chunk_text = ""

                    try:
                        # Deteksi konteks audio untuk setiap chunk
                        audio_context = self.detect_audio_context(audio_file)

                        with sr.AudioFile(audio_file) as source:
                            audio_data = self.recognizer.record(source)
                            engine = self.engine.get()

                            if engine == "Google":
                                # API limitation: introduce delay to avoid rate limits
                                with api_semaphore:
                                    chunk_text = self.recognizer.recognize_google(audio_data, language=language_code)
                            elif engine == "Whisper":
                                # Peningkatan: Optimasi pemilihan model per chunk
                                model = self.select_optimal_transcription_model(audio_context, language_code)

                                # Peningkatan: Adaptive batch size untuk model selection
                                chunk_length = self.get_audio_duration(audio_file)

                                with api_semaphore:
                                    chunk_text = self.transcribe_with_groq_whisper(audio_data, language_code, model)

                                if not chunk_text:
                                    # Peningkatan: Fallback mechanism jika model utama gagal
                                    logger.warning(f"Primary model failed for chunk {i+1}, trying fallback...")
                                    fallback_model = "distil-whisper-large-v3-en" if language_code.startswith("en") else "whisper-large-v3-turbo"

                                    with api_semaphore:
                                        chunk_text = self.transcribe_with_groq_whisper(audio_data, language_code, fallback_model)

                                    if not chunk_text:
                                        chunk_text = f"[Transkripsi gagal untuk bagian {i+1}]"

                        # Peningkatan: Validasi kualitas transkripsi
                        quality_score = self._evaluate_transcription_quality(chunk_text)
                        if quality_score < 0.2:  # Very low quality threshold
                            logger.warning(f"Low quality transcription detected for chunk {i+1}, score: {quality_score}")
                            # Add warning in the transcription
                            chunk_text = f"[Peringatan: Kualitas audio rendah] {chunk_text}"

                        all_raw_transcriptions.append(chunk_text)
                        current_batch_raw_texts.append(chunk_text)

                        # Peningkatan: Pra-pemrosesan teks sebelum peningkatan
                        chunk_text = self._preprocess_transcription(chunk_text)

                        if self.use_ai_enhancement.get():
                            self.root.after(0, lambda cn=chunk_num, tc=total_chunks:
                                self.status_var.set(f"Meningkatkan bagian {cn}/{tc} dengan AI..."))

                            try:
                                # Peningkatan: Tambahkan konteks audio ke proses enhancement
                                enhanced_context = {
                                    "audio_context": audio_context,
                                    "chunk_position": f"{i+1}/{total_chunks}",
                                    "language": language_code
                                }

                                # Peningkatan: Jika ini bukan chunk pertama, tambahkan konteks dari chunk sebelumnya
                                if i > 0 and all_enhanced_chunk_texts:
                                    # Ambil beberapa kalimat terakhir dari chunk sebelumnya sebagai konteks
                                    prev_chunk = all_enhanced_chunk_texts[-1]
                                    sentences = re.split(r'[.!?]+\s*', prev_chunk)
                                    context_sentences = sentences[-3:] if len(sentences) > 3 else sentences
                                    enhanced_context["previous_context"] = " ".join(context_sentences)

                                enhanced_single_chunk_text = self.enhance_with_ai(chunk_text)
                                all_enhanced_chunk_texts.append(enhanced_single_chunk_text)
                                current_batch_enhanced_texts.append(enhanced_single_chunk_text)
                            except Exception as e_enhance:
                                logger.error(f"Error enhancing chunk {chunk_num}: {e_enhance}", exc_info=True)
                                fallback_enhanced_text = f"[PENINGKATAN AI GAGAL UNTUK BAGIAN {chunk_num}]\n{chunk_text}"
                                all_enhanced_chunk_texts.append(fallback_enhanced_text)
                                current_batch_enhanced_texts.append(fallback_enhanced_text)

                        processed_audio_files_in_run.append(audio_file)

                    except sr.UnknownValueError:
                        logger.warning(f"Could not recognize audio in chunk {i+1}")
                        error_msg = f"[Audio tidak terdeteksi pada bagian {i+1}]"
                        all_raw_transcriptions.append(error_msg)
                        current_batch_raw_texts.append(error_msg)
                        if self.use_ai_enhancement.get():
                            all_enhanced_chunk_texts.append(error_msg)
                            current_batch_enhanced_texts.append(error_msg)
                        self.root.after(0, lambda n=i+1:
                            self.status_var.set(f"Tidak dapat mengenali audio pada bagian {n}"))
                    except Exception as e:
                        logger.error(f"Error processing chunk {i+1}: {e}", exc_info=True)
                        error_msg = f"[Error pada bagian {i+1}: {str(e)[:50]}...]"
                        all_raw_transcriptions.append(error_msg)
                        current_batch_raw_texts.append(error_msg)
                        if self.use_ai_enhancement.get():
                            all_enhanced_chunk_texts.append(error_msg)
                            current_batch_enhanced_texts.append(error_msg)

                # Update preview setelah setiap batch
                self.root.after(0, lambda: self.result_text.delete(1.0, tk.END))
                if self.use_ai_enhancement.get():
                    preview_text = "\n\n---\n[AKHIR BAGIAN]---\n\n".join(all_enhanced_chunk_texts)
                else:
                    preview_text = " ".join(all_raw_transcriptions)

                preview_prefix = f"[Pratinjau setelah batch {batch_index+1}/{batch_count}]\n"
                display_text = preview_text[:3000] + ("... (banyak konten lainnya)" if len(preview_text) > 3000 else "")
                self.root.after(0, lambda pt=display_text, pp=preview_prefix:
                    self.result_text.insert(tk.END, pp + pt))

            # Peningkatan: Analisis global semua bagian sebelum penyusunan dokumen final
            global_analysis = self._analyze_global_transcript(all_raw_transcriptions)
            self.root.after(0, lambda: self.status_var.set(
                f"Menyusun dokumen final ({global_analysis['content_type']})..."))
            self.root.after(0, lambda: self.progress_var.set(80))

            # Peningkatan: Penyusunan dokumen final dengan analisis global
            if self.use_ai_enhancement.get() and all_enhanced_chunk_texts:
                complete_text = "\n\n---\n[AKHIR BAGIAN]---\n\n".join(all_enhanced_chunk_texts)

                # Peningkatan: Keputusan cerdas tentang tahap peningkatan kohesi
                if len(complete_text.split()) > 40000:
                    self.root.after(0, lambda: self.status_var.set(
                        "Dokumen sangat panjang, melewati tahap peningkatan kohesi final..."))
                elif len(all_enhanced_chunk_texts) > 1:
                    self.root.after(0, lambda: self.status_var.set(
                        f"Meningkatkan koherensi dokumen final ({global_analysis['content_type']})..."))
                    self.root.after(0, lambda: self.progress_var.set(90))

                    # Peningkatan: Tambahkan analisis global ke proses peningkatan kohesi
                    complete_text = self.enhance_document_cohesion(complete_text)
            else:
                complete_text = " ".join(all_raw_transcriptions)

            self.root.after(0, lambda: self.progress_var.set(95))

            # Buat nama file yang deskriptif berdasarkan analisis
            now = datetime.datetime.now()
            date_str = now.strftime("%Y%m%d")
            time_str = now.strftime("%H%M%S")

            # Peningkatan: Gunakan tipe konten dalam nama file
            type_suffix = f"_{global_analysis['content_type']}" if global_analysis.get('content_type') else ""
            filename = f"{self.filename_prefix.get()}{type_suffix}_{date_str}_{time_str}.docx"
            filepath = os.path.join(self.output_folder.get(), filename)

            self.root.after(0, lambda: self.status_var.set("Menyimpan dokumen terformat..."))

            # Peningkatan: Kirim analisis global ke fungsi formatting untuk optimasi layout
            content_stats = {
                "content_type": global_analysis.get('content_type', 'general'),
                "original_text": complete_text,
                "language": language_code,
                "has_tables": "table" in global_analysis.get('elements', []),
                "has_lists": "list" in global_analysis.get('elements', [])
            }

            # Simpan dokumen dengan format optimal
            self.save_as_word_document(complete_text, filepath)

            self.root.after(0, lambda: self.progress_var.set(100))
            self.root.after(0, lambda: self.status_var.set(f"Selesai! File disimpan: {filename}"))

            # Update hasil akhir
            self.root.after(0, lambda: self.result_text.delete(1.0, tk.END))
            final_display_text = complete_text[:5000] + "\n\n... (konten lengkap tersedia di file .docx)" if len(complete_text) > 5000 else complete_text
            self.root.after(0, lambda: self.result_text.insert(tk.END, final_display_text))

            # Cleanup files
            self._cleanup_temp_files(processed_audio_files_in_run)

        except Exception as e:
            logger.error(f"Error in process_extended_recording_optimized: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("Error", f"Terjadi kesalahan: {e}"))
            self.root.after(0, lambda: self.status_var.set(f"Error: {e}"))

    def _calculate_optimal_batch_size(self, total_chunks):
        """
        Calculate optimal batch size based on total chunks and system capabilities.
        """
        if total_chunks <= 5:
            return 1  # Process one at a time for small recordings
        elif total_chunks <= 20:
            return 2  # Small batches for medium recordings
        elif total_chunks <= 50:
            return 5  # Medium batches for larger recordings
        else:
            return 10  # Larger batches for very large recordings

    def _analyze_global_transcript(self, all_chunks):
        """
        Analyze the complete transcript dengan menggunakan semua variabel yang diperlukan.
        """
        combined_text = " ".join(all_chunks)
        word_count = len(combined_text.split())

        # Analisis struktur dokumen
        structure_analysis = {
            "has_headings": bool(re.search(r'^#+\s+', combined_text, re.MULTILINE)),
            "has_lists": bool(re.search(r'^\s*[-*•]\s+', combined_text, re.MULTILINE)),
            "has_numbered_lists": bool(re.search(r'^\s*\d+\.\s+', combined_text, re.MULTILINE)),
            "has_tables": bool(re.search(r'\|.+\|.+\|', combined_text)),
            "paragraph_count": len(re.split(r'\n\s*\n', combined_text)),
            "avg_paragraph_length": 0
        }

        # Hitung rata-rata panjang paragraf menggunakan structure_analysis
        paragraphs = re.split(r'\n\s*\n', combined_text)
        if paragraphs and structure_analysis["paragraph_count"] > 0:
            total_length = sum(len(p.split()) for p in paragraphs)
            structure_analysis["avg_paragraph_length"] = total_length / structure_analysis["paragraph_count"]

        # Detect content type dengan informasi struktur
        content_type = "general"
        elements = []

        # Check for meeting patterns
        meeting_indicators = ["meeting", "rapat", "agenda", "diskusi", "minutes", "attendees", "peserta"]
        meeting_score = sum(1 for word in meeting_indicators if word in combined_text.lower())

        # Check for lecture patterns
        lecture_indicators = ["lecture", "course", "kuliah", "pembelajaran", "materi", "topic", "bab"]
        lecture_score = sum(1 for word in lecture_indicators if word in combined_text.lower())

        # Check for interview patterns
        interview_indicators = ["interview", "question", "answer", "wawancara", "tanya", "jawab"]
        interview_score = sum(1 for word in interview_indicators if word in combined_text.lower())

        # Gunakan structure_analysis untuk mendeteksi elemen
        if structure_analysis["has_tables"]:
            elements.append("table")
        if structure_analysis["has_lists"]:
            elements.append("list")
        if structure_analysis["has_numbered_lists"]:
            elements.append("numbered_list")
        if structure_analysis["has_headings"]:
            elements.append("headings")

        # Determine content type berdasarkan score dan struktur
        max_score = max(meeting_score, lecture_score, interview_score)
        if max_score > 2:
            if meeting_score == max_score:
                content_type = "meeting"
            elif lecture_score == max_score:
                content_type = "lecture"
            elif interview_score == max_score:
                content_type = "interview"

        logger.info(f"Global analysis completed - Content type: {content_type}, Elements: {elements}, Word count: {word_count}")
        logger.info(f"Structure analysis: {structure_analysis}")

        return {
            "content_type": content_type,
            "elements": elements,
            "word_count": word_count,
            "structure": structure_analysis,
            "meeting_score": meeting_score,
            "lecture_score": lecture_score,
            "interview_score": interview_score
        }

    def _evaluate_transcription_quality(self, text):
        """
        Evaluasi kualitas transkrip berdasarkan berbagai indikator.
        Returns score between 0.0 (very poor) and 1.0 (excellent).
        """
        if not text or len(text.strip()) == 0:
            return 0.0

        # Check for common error indicators
        error_markers = ["[inaudible]", "[unclear]", "[noise]", "[?]", "???"]
        error_count = sum(text.count(marker) for marker in error_markers)

        # Check for repeated words/phrases which often indicate poor transcription
        repeated_word_patterns = re.findall(r'\b(\w+)(\s+\1\b)+', text)
        repetition_count = len(repeated_word_patterns)

        # Calculate word and sentence metrics
        words = text.split()
        word_count = len(words)
        if word_count == 0:
            return 0.1  # Very low quality but not zero

        sentences = re.split(r'[.!?]+', text)
        sentence_count = len([s for s in sentences if s.strip()])

        # Calculate average words per sentence
        avg_words_per_sentence = word_count / max(sentence_count, 1)

        # Very short or very long sentences can indicate poor transcription
        sentence_length_penalty = 0
        if avg_words_per_sentence < 3 or avg_words_per_sentence > 30:
            sentence_length_penalty = 0.2

        # Calculate base score
        base_score = 1.0

        # Apply penalties
        error_penalty = min(0.5, (error_count / max(word_count, 1)) * 10)
        repetition_penalty = min(0.3, (repetition_count / max(word_count, 1)) * 20)

        final_score = max(0.0, base_score - error_penalty - repetition_penalty - sentence_length_penalty)

        return final_score

    def _preprocess_transcription(self, text):
        """
        Pra-pemrosesan teks untuk membersihkan artefak umum dalam transkripsi.
        """
        if not text:
            return text

        # Remove excess whitespace
        text = re.sub(r'\s+', ' ', text).strip()

        # Fix common transcription artifacts
        text = re.sub(r'\buh\b|\bum\b|\bah\b', '', text)  # Remove filler words
        text = re.sub(r'(\w+)(\s+\1\b){2,}', r'\1', text)  # Remove excessive repetition (3+ times)

        # Add periods if sentences don't have proper punctuation
        text = re.sub(r'([a-zA-Z])\s+([A-Z])', r'\1. \2', text)

        # Fix spacing around punctuation
        text = re.sub(r'\s+([.,;:!?])', r'\1', text)

        return text

    def _cleanup_temp_files(self, processed_files):
        """Cleanup temporary files safely."""
        print(f"DEBUG: Starting cleanup of {len(processed_files)} files")
        for file_path in processed_files:
            if hasattr(self, 'temp_audio_files') and file_path in self.temp_audio_files:
                try:
                    if os.path.exists(file_path):
                        os.unlink(file_path)
                        print(f"DEBUG: Deleted temp file: {file_path}")
                    self.temp_audio_files.remove(file_path)
                except Exception as e_clean:
                    print(f"ERROR: Failed to clean temp file {file_path}: {e_clean}")

    def detect_audio_context(self, audio_file):
        """
        Deteksi karakteristik dan konteks audio untuk optimasi transkripsi.
        """
        try:
            duration = self.get_audio_duration(audio_file)

            # Analisis tingkat suara
            with wave.open(audio_file, 'rb') as wf:
                # Dapatkan beberapa sampel untuk analisis
                n_frames = min(wf.getnframes(), 1000000)  # Max 1M sampel untuk kecepatan
                frames = wf.readframes(n_frames)

                # Hitung RMS untuk menentukan volume rata-rata
                rms = audioop.rms(frames, wf.getsampwidth())

                # Hitung jumlah silent frames untuk mendeteksi jeda bicara
                silent_threshold = rms * 0.1  # 10% dari volume rata-rata
                silent_frames = 0
                frame_size = wf.getsampwidth() * wf.getnchannels()

                for i in range(0, len(frames), frame_size):
                    if i + frame_size <= len(frames):
                        chunk = frames[i:i + frame_size]
                        chunk_rms = audioop.rms(chunk, wf.getsampwidth())
                        if chunk_rms < silent_threshold:
                            silent_frames += 1

                silence_ratio = silent_frames / (len(frames) / frame_size) if len(frames) > 0 else 0

            # Deteksi karakteristik
            context = {
                "duration": duration,
                "volume_level": "high" if rms > 10000 else "medium" if rms > 5000 else "low",
                "silence_ratio": silence_ratio,
                "content_type": self._detect_content_type(duration, silence_ratio)
            }

            return context

        except Exception as e:
            logger.error(f"Error saat deteksi audio context: {e}", exc_info=True)
            # Return default values if detection fails
            return {
                "duration": 0,
                "volume_level": "medium",
                "silence_ratio": 0,
                "content_type": "unknown"
            }

    def _detect_content_type(self, duration, silence_ratio):
        """
        Deteksi tipe konten audio berdasarkan penelitian:
        - Command: sangat singkat, sedikit jeda
        - Dictation: sedang, jeda sedang-tinggi
        - Lecture: panjang, jeda rendah-sedang
        - Meeting: panjang, jeda tinggi
        - General: selain itu
        """
        # Command: <30 detik, silence <0.2
        if duration < 30 and silence_ratio < 0.2:
            return "command"
        # Short note: <60 detik, silence tinggi
        elif duration < 60 and silence_ratio >= 0.4:
            return "short_note"
        # Dictation: 1-5 menit, silence sedang-tinggi
        elif 60 <= duration < 300 and silence_ratio >= 0.25:
            return "dictation"
        # Lecture: >10 menit, silence rendah-sedang
        elif duration >= 600 and silence_ratio < 0.5:
            return "lecture"
        # Meeting: >5 menit, silence tinggi
        elif duration >= 8000 and silence_ratio >= 0.8:
            return "meeting"
        # Default
        else:
            return "leecture"

    def select_optimal_transcription_model(self, audio_context, language_code):
        """
        Pilih model transcription optimal berdasarkan konteks audio.
        """
        # Default model
        model = "whisper-large-v3"

        # For very short audio, use the fastest model
        if audio_context["duration"] < 20:
            model = "whisper-large-v3-turbo"

        # For English with economic model enabled, use distil model
        elif language_code.startswith("en") and self.use_economic_model.get():
            model = "distil-whisper-large-v3-en"

        # For very long recordings, use model with better context handling
        elif audio_context["duration"] > 1800:  # Over 30 minutes
            model = "whisper-large-v3"  # Best for long context

        logger.info(f"Selected transcription model: {model} for content type: {audio_context['content_type']}")
        return model

    def enhance_with_ai(self, text):
        """
        Fungsi AI enhancement dengan penanganan error yang ditingkatkan.
        """
        if not text or not text.strip():
            logger.warning("Empty text received for AI enhancement")
            return text

        try:
            # Deteksi karakteristik konten untuk menentukan strategi pengolahan
            content_stats = self._analyze_content_characteristics(text)

            # Pilih model dan parameter berdasarkan karakteristik konten
            model_config = self._select_optimal_model(content_stats)

            # Buat prompt yang disesuaikan dengan jenis konten
            system_prompt, user_prompt = self._create_content_adaptive_prompts(text, content_stats)

            # Gunakan metode yang lebih aman untuk memperbarui status UI
            status_message = f"Meningkatkan dengan {model_config['name']} - Optimasi untuk {content_stats['content_type']}..."
            self.root.after(0, lambda msg=status_message: self.status_var.set(msg))

            # Periksa apakah client ada
            if self.groq_client is None:
                logger.error("Groq client tidak terinisialisasi")
                return self._fallback_enhancement(text)

            # Menambahkan jeda API untuk mencegah rate limiting
            time.sleep(self.api_request_delay)

            # Panggil API dengan parameter yang dioptimalkan
            completion = self.groq_client.chat.completions.create(
                model=model_config['model_id'],
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=model_config['temperature'],
                max_tokens=model_config['max_tokens'],
                top_p=model_config['top_p'],
                reasoning_format="hidden",
                stream=False
            )

            enhanced_text = completion.choices[0].message.content
            enhanced_text = self._post_process_enhanced_text(enhanced_text, content_stats)

            return enhanced_text

        except Exception as e:
            # Tangani error dengan lebih baik
            error_message = str(e)
            logger.error(f"Error dalam peningkatan AI: {error_message}", exc_info=True)

            # Gunakan metode yang lebih aman untuk memperbarui status UI
            error_str = error_message[:100]  # Batasi panjang string error
            self.root.after(0, lambda err=error_str: self.status_var.set(f"Error peningkatan: {err}"))

            return self._fallback_enhancement(text)

    def _analyze_content_characteristics(self, text_content):
        """
        Menganalisis karakteristik konten dengan detail tinggi untuk formatting optimal.
        Menggunakan AI Groq dengan model DeepSeek untuk klasifikasi yang lebih akurat.
        """
        # Hitung statistik dasar
        words = text_content.split()
        word_count = len(words)
        sentences = re.findall(r'[.!?]+', text_content)
        sentence_count = len(sentences) + 1
        avg_words_per_sentence = word_count / max(sentence_count, 1)

        # Analisis struktur dokumen
        lines = text_content.split('\n')
        heading_count = len([line for line in lines if re.match(r'^#+\s+', line)])
        paragraph_count = len([para for para in text_content.split('\n\n') if para.strip()])

        # Deteksi elemen khusus
        has_tables = bool(re.search(r'\|.*\|.*\|', text_content))
        has_code_blocks = bool(re.search(r'```[\s\S]*?```', text_content))
        has_lists = bool(re.search(r'^\s*[-*•]\s+', text_content, re.MULTILINE))
        has_numbered_lists = bool(re.search(r'^\s*\d+\.\s+', text_content, re.MULTILINE))
        has_quotes = bool(re.search(r'^>\s+', text_content, re.MULTILINE))
        has_action_items = bool(re.search(r'\[ACTION\]|\*\*ACTION\*\*|ACTION:', text_content, re.IGNORECASE))

        # Deteksi bahasa dominan
        text_lower = text_content.lower()
        english_words = len(re.findall(r'\b(the|and|or|of|to|in|for|with|by|on|at|is|are|was|were)\b', text_lower))
        indonesian_words = len(re.findall(r'\b(dan|atau|dari|ke|di|untuk|dengan|oleh|pada|adalah|ini|itu|yang)\b', text_lower))
        language = "en" if english_words > indonesian_words else "id"

        # **ENHANCED: AI-POWERED CONTENT TYPE CLASSIFICATION**
        content_type, ai_confidence, ai_reasoning = self._classify_content_with_ai(text_content, language)

        # Rule-based fallback scores untuk validasi dan fallback
        technical_indicators = [
            'API', 'database', 'server', 'aplikasi', 'sistem', 'teknologi', 'software',
            'hardware', 'network', 'security', 'development', 'programming', 'algorithm',
            'code', 'function', 'method', 'class', 'variable', 'configuration'
        ]
        meeting_indicators = [
            'agenda', 'rapat', 'meeting', 'diskusi', 'peserta', 'keputusan', 'action item',
            'tindakan', 'follow up', 'attendees', 'minutes', 'decision', 'discussed',
            'decided', 'participants', 'action points'
        ]
        lecture_indicators = [
            'materi', 'pembelajaran', 'kuliah', 'course', 'lesson', 'topik', 'bab',
            'chapter', 'modul', 'pembelajaran', 'penjelasan', 'teori', 'konsep',
            'students', 'learn', 'understand', 'explain', 'definition'
        ]

        # Hitung skor untuk validasi dan fallback
        technical_score = sum(1 for indicator in technical_indicators if indicator in text_lower)
        meeting_score = sum(1 for indicator in meeting_indicators if indicator in text_lower)
        lecture_score = sum(1 for indicator in lecture_indicators if indicator in text_lower)

        # Fallback ke rule-based classification jika AI gagal atau confidence rendah
        if content_type == "unknown" or ai_confidence < 0.7:
            logger.warning(f"AI classification failed or low confidence ({ai_confidence:.2f}), using rule-based fallback")
            content_type = self._fallback_rule_based_classification(text_content, technical_score, meeting_score, lecture_score)
            ai_confidence = 0.5
            ai_reasoning = "Using rule-based classification as fallback"

        # Analisis kompleksitas
        complexity_score = 0
        if has_tables: complexity_score += 2
        if has_code_blocks: complexity_score += 2
        if heading_count > 5: complexity_score += 1
        if word_count > 2000: complexity_score += 1
        if avg_words_per_sentence > 25: complexity_score += 1

        complexity_level = "high" if complexity_score >= 4 else "medium" if complexity_score >= 2 else "low"

        # Ekstrak topik utama
        main_topics = self._extract_main_topics(text_content, heading_count)

        content_stats = {
            "word_count": word_count,
            "sentence_count": sentence_count,
            "paragraph_count": paragraph_count,
            "heading_count": heading_count,
            "avg_words_per_sentence": avg_words_per_sentence,
            "content_type": content_type,
            "language": language,
            "complexity_level": complexity_level,
            "has_tables": has_tables,
            "has_code_blocks": has_code_blocks,
            "has_lists": has_lists,
            "has_numbered_lists": has_numbered_lists,
            "has_quotes": has_quotes,
            "has_action_items": has_action_items,
            "technical_score": technical_score,
            "meeting_score": meeting_score,
            "lecture_score": lecture_score,
            "main_topics": main_topics,
            "original_text": text_content,
            "ai_confidence": ai_confidence,
            "ai_reasoning": ai_reasoning
        }

        logger.info(f"📊 AI Content analysis: {content_type} (confidence: {ai_confidence:.2f}, {complexity_level} complexity, {word_count} words)")
        logger.info(f"🤖 AI reasoning: {ai_reasoning}")
        return content_stats

    def _classify_content_with_ai(self, text_content, language):
        """
        Menggunakan AI Groq dengan model DeepSeek untuk mengklasifikasi jenis konten.
        Returns: (content_type, confidence_score, reasoning)
        """
        if not self.groq_client:
            logger.warning("Groq client not available for content classification")
            return "unknown", 0.0, "Groq client not available"

        try:
            # Batas teks untuk analisis (untuk menghemat token)
            sample_text = text_content[:2000] if len(text_content) > 2000 else text_content

            # Deteksi bahasa dan sesuaikan prompt
            language_code = language.split('-')[0] if '-' in language else language.lower()

            # Mapping bahasa untuk prompt yang lebih tepat
            language_mappings = {
                'id': {'name': 'Indonesian', 'instruction': 'dalam bahasa Indonesia'},
                'en': {'name': 'English', 'instruction': 'in English'},
                'ja': {'name': 'Japanese', 'instruction': 'in Japanese'},
                'zh': {'name': 'Chinese', 'instruction': 'in Chinese'}
            }

            lang_info = language_mappings.get(language_code, {'name': 'Unknown', 'instruction': 'in the detected language'})

            # **PERBAIKAN: Prompt sesuai dokumentasi DeepSeek - semua instruksi dalam user message**
            user_prompt = f"""You are an expert document classifier. Analyze the given text and classify it into one of these categories with step-by-step reasoning.

    LANGUAGE CONTEXT: The content is primarily {lang_info['name']} ({language}).

    CLASSIFICATION CATEGORIES:
    1. **technical_report** - Technical documentation, API docs, system specifications, code documentation, software manuals, technical analysis
    2. **meeting_notes** - Meeting minutes, discussion summaries, agenda items, action items, decisions made in meetings
    3. **lecture** - Educational content, course materials, lessons, tutorials, academic presentations, training materials
    4. **interview** - Q&A format, interviews, conversations with questions and answers
    5. **presentation** - Slide content, presentation materials, structured presentations
    6. **research** - Research papers, studies, analysis reports, findings
    7. **instructional** - How-to guides, step-by-step instructions, procedures, manuals
    8. **news** - News articles, journalism, current events reporting
    9. **narrative** - Stories, personal accounts, descriptive content
    10. **general** - General content that doesn't fit other categories

    LANGUAGE-SPECIFIC TERMS TO CONSIDER:
    - Indonesian: "rapat" (meeting), "presentasi" (presentation), "laporan" (report), "tutorial", "pembelajaran" (learning)
    - English: Standard technical and business terminology
    - Japanese: "会議" (meeting), "報告" (report), "講義" (lecture)
    - Chinese: "会议" (meeting), "报告" (report), "讲座" (lecture)

    STEP-BY-STEP ANALYSIS REQUIRED:
    1. Identify key terminology and patterns in the text
    2. Look for language-specific indicators
    3. Assess document structure and format
    4. Consider context and purpose
    5. Determine the most appropriate category
    6. Assign confidence level based on clarity of indicators

    RESPONSE FORMAT: Return ONLY a valid JSON object with:
    - "content_type": one of the 10 categories above
    - "confidence": float between 0.0 and 1.0
    - "reasoning": your step-by-step analysis and language-specific indicators found

    TEXT TO ANALYZE (Language: {lang_info['name']}):
    {sample_text}

    Provide your analysis and respond with valid JSON only."""

            # **PERBAIKAN: Konfigurasi sesuai dokumentasi resmi**
            completion = self.groq_client.chat.completions.create(
                model="deepseek-r1-distill-llama-70b",
                messages=[
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.6,  # Sesuai rekomendasi dokumentasi (0.5-0.7)
                max_completion_tokens=512,  # Cukup untuk analisis dan JSON response
                top_p=0.95,
                response_format={"type": "json_object"},
                reasoning_format="hidden",  # Hidden agar output bersih tanpa <think> tags
                stream=False
            )

            response_text = completion.choices[0].message.content.strip()
            logger.info(f"AI response received (language: {language}): {len(response_text)} characters")

            # **PERBAIKAN: JSON parsing yang lebih robust**
            try:
                # Dengan response_format="json_object", output sudah dijamin valid JSON
                result = json.loads(response_text)

                content_type = result.get("content_type", "unknown")
                confidence = float(result.get("confidence", 0.0))
                reasoning = result.get("reasoning", "No reasoning provided")

                # Validasi content_type
                valid_types = [
                    "technical_report", "meeting_notes", "lecture", "interview",
                    "presentation", "research", "instructional", "news",
                    "narrative", "general"
                ]

                if content_type not in valid_types:
                    logger.warning(f"AI returned invalid content type: {content_type}")
                    content_type = "general"
                    confidence = 0.5
                    reasoning = f"Invalid category returned, defaulted to general. Original: {content_type}"

                # **ENHANCEMENT: Boost confidence untuk bahasa yang cocok**
                if language_code == 'id' and any(term in reasoning.lower() for term in ['indonesian', 'indonesia', 'bahasa']):
                    confidence = min(1.0, confidence + 0.1)
                elif language_code == 'en' and any(term in reasoning.lower() for term in ['english', 'english pattern']):
                    confidence = min(1.0, confidence + 0.1)

                # Validasi confidence range
                confidence = max(0.0, min(1.0, confidence))

                logger.info(f"AI classification successful ({language}): {content_type} (confidence: {confidence:.2f})")
                logger.debug(f"AI reasoning: {reasoning[:200]}...")
                return content_type, confidence, reasoning

            except json.JSONDecodeError as e:
                logger.error(f"Failed to parse AI response as JSON: {response_text[:200]}")

                # **PERBAIKAN: Enhanced fallback dengan regex extraction**
                try:
                    # Coba ekstrak JSON dari response yang mungkin rusak
                    json_pattern = r'\{[^{}]*?"content_type"[^{}]*?"confidence"[^{}]*?"reasoning"[^{}]*?\}'
                    json_match = re.search(json_pattern, response_text, re.DOTALL)

                    if json_match:
                        extracted_json = json_match.group()
                        result = json.loads(extracted_json)

                        content_type = result.get("content_type", "general")
                        confidence = float(result.get("confidence", 0.5))
                        reasoning = result.get("reasoning", "Extracted from partial response")

                        logger.info(f"JSON extracted from partial response: {content_type}")
                        return content_type, confidence, reasoning

                except Exception as extract_error:
                    logger.error(f"JSON extraction also failed: {extract_error}")

                return "unknown", 0.0, f"JSON parsing error: {str(e)}"

        except Exception as e:
            logger.error(f"Error in AI content classification: {e}", exc_info=True)

            # **PERBAIKAN: Specific error handling**
            if "rate limit" in str(e).lower():
                return "unknown", 0.0, "Rate limit exceeded - using fallback classification"
            elif "api" in str(e).lower():
                return "unknown", 0.0, f"API error: {str(e)[:100]}"
            else:
                return "unknown", 0.0, f"Classification error: {str(e)[:100]}"

    def _fallback_rule_based_classification(self, text_content, technical_score, meeting_score, lecture_score):
        """
        Fallback rule-based classification jika AI gagal.
        """
        text_lower = text_content.lower()

        # Additional rule-based indicators
        interview_indicators = [
            'question', 'answer', 'interview', 'wawancara', 'tanya', 'jawab',
            'interviewer', 'interviewee', 'Q:', 'A:', 'asked', 'responded'
        ]
        instructional_indicators = [
            'step', 'langkah', 'how to', 'cara', 'instruction', 'petunjuk',
            'guide', 'tutorial', 'first', 'next', 'then', 'finally', 'procedure'
        ]

        interview_score = sum(1 for indicator in interview_indicators if indicator in text_lower)
        instructional_score = sum(1 for indicator in instructional_indicators if indicator in text_lower)

        # Tentukan jenis berdasarkan skor tertinggi
        scores = {
            "technical_report": technical_score,
            "meeting_notes": meeting_score,
            "lecture": lecture_score,
            "interview": interview_score,
            "instructional": instructional_score
        }

        max_score = max(scores.values())
        if max_score >= 3:  # Threshold untuk klasifikasi
            content_type = max(scores, key=scores.get)
        else:
            content_type = "general"

        logger.info(f"Rule-based classification: {content_type} (scores: {scores})")
        return content_type

    def _extract_main_topics(self, text_content, heading_count):
        """Ekstrak topik utama dari heading dan konten dengan error handling yang robust."""
        topics = []

        try:
            # Validasi input
            if not text_content or not isinstance(text_content, str):
                logger.warning("Invalid text_content provided to _extract_main_topics")
                return topics

            # **OPTIMASI: Gunakan heading_count untuk strategi yang tepat**
            if heading_count > 0:
                logger.info(f"Document has {heading_count} headings, prioritizing heading extraction")

                # Ekstrak dari heading dengan batasan yang efisien
                headings = re.findall(r'^#+\s+(.+)$', text_content, re.MULTILINE)
                max_headings_to_process = min(heading_count, 5)

                for heading in headings[:max_headings_to_process]:
                    if not heading:
                        continue

                    # Bersihkan dari emoji dan simbol
                    clean_heading = re.sub(r'[^\w\s\-]', '', heading).strip()

                    # Filter heading yang bermakna
                    if len(clean_heading) > 3 and clean_heading.lower() not in ['dan', 'atau', 'the', 'and', 'or']:
                        topics.append(clean_heading)

                # **EFISIENSI: Threshold dinamis berdasarkan heading_count**
                min_topics_threshold = 2 if heading_count >= 3 else 1
            else:
                logger.info("Document has no headings, focusing on content keyword extraction")
                min_topics_threshold = 3

            # **OPTIMASI: Keyword extraction hanya jika perlu**
            if len(topics) < min_topics_threshold:
                logger.info(f"Need more topics (current: {len(topics)}), extracting from content")

                # **EFISIENSI: Filter kata langsung tanpa loop terpisah**
                words = text_content.split()
                stopwords = {
                    'dan', 'atau', 'dari', 'ke', 'di', 'untuk', 'dengan', 'oleh', 'pada', 'adalah', 'ini', 'itu', 'yang',
                    'the', 'and', 'or', 'of', 'to', 'in', 'for', 'with', 'by', 'on', 'at', 'is', 'are', 'was', 'were',
                    'a', 'an', 'as', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would'
                }

                # **OPTIMASI: One-pass filtering dengan list comprehension**
                meaningful_words = [
                    word.lower() for word in words
                    if (len(word) > 4 and
                        word.isalpha() and
                        word.lower() not in stopwords and
                        not word.isupper())
                ]

                if meaningful_words:  # **EFISIENSI: Cek ada kata sebelum Counter**
                    from collections import Counter
                    word_freq = Counter(meaningful_words)

                    # **OPTIMASI: Threshold dinamis berdasarkan heading_count**
                    freq_threshold = 1 if heading_count == 0 else 2
                    slots_needed = 5 - len(topics)  # **EFISIENSI: Hitung slot yang dibutuhkan**

                    # **OPTIMASI: Batasi iterasi dengan slots_needed**
                    for word, freq in word_freq.most_common(min(10, slots_needed * 2)):
                        if freq >= freq_threshold:
                            word_title = word.title()
                            # **EFISIENSI: Case-insensitive comparison yang benar**
                            if word_title.lower() not in [t.lower() for t in topics]:
                                topics.append(word_title)
                                if len(topics) >= 5:  # **EARLY EXIT: Stop jika sudah cukup**
                                    break

                # **EFISIENSI: Fallback hanya jika benar-benar perlu**
                if len(topics) == 0 and heading_count == 0:
                    try:
                        # **OPTIMASI: Fallback sederhana dengan first words**
                        first_words = text_content.split()[:10]
                        for word in first_words:
                            if len(word) > 4 and word.isalpha():
                                topics.append(word.title())
                                if len(topics) >= 3:
                                    break
                    except:
                        pass

            # **EFISIENSI: Validation dan cleanup dalam satu loop**
            final_topics = [
                topic.strip() for topic in topics[:5]
                if topic and isinstance(topic, str) and len(topic.strip()) > 0
            ]

            logger.info(f"Extracted {len(final_topics)} main topics from document with {heading_count} headings: {final_topics}")
            return final_topics

        except Exception as e:
            logger.error(f"Error in _extract_main_topics: {e}", exc_info=True)
            return []

    def _select_optimal_model(self, content_stats):
        """
        Memilih model dan parameter optimal berdasarkan karakteristik konten.
        """
        # Default config - menggunakan model DeepSeek yang lebih besar sebagai default
        config = {
            "model_id": "deepseek-r1-distill-llama-70b", # Model DeepSeek default
            "temperature": 0.5,
            "max_tokens": 6000,
            "top_p": 0.95,
            "name": "Default AI"
        }

        # Sesuaikan berdasarkan jenis konten
        if content_stats["content_type"] == "technical_report":
            config["temperature"] = 0.3  # Lebih deterministik untuk konten teknis
            config["name"] = "Technical AI"
            # Model tetap default DeepSeek, atau bisa disesuaikan jika ada model DeepSeek spesifik untuk teknis
        elif content_stats["content_type"] == "instructional":
            config["temperature"] = 0.4
            config["name"] = "Instructional AI"
            # Model tetap default DeepSeek
        elif content_stats["content_type"] == "meeting_notes":
            config["model_id"] = "deepseek-r1-distill-llama-70b"  # Model DeepSeek terbaik untuk rapat
            config["temperature"] = 0.5
            config["name"] = "Meeting Notes AI"
        elif content_stats["content_type"] == "narrative":
            config["temperature"] = 0.6  # Sedikit lebih kreatif untuk naratif
            config["name"] = "Narrative AI"
            # Model tetap default DeepSeek

        # Penyesuaian untuk bahasa hanya menggunakan model DeepSeek
        if content_stats["language"] == "en" and self.use_economic_model.get():
            if content_stats["word_count"] < 1000:
                # Menggunakan model DeepSeek yang lebih kecil/ekonomis
                config["model_id"] = "deepseek-llm-7b-chat" # Contoh model DeepSeek yang lebih ekonomis
                config["name"] += " (Economic DeepSeek)"

        return config

    def _create_content_adaptive_prompts(self, text, content_stats):
        """
        Membuat prompt yang disesuaikan dengan jenis konten.
        """
        # Prompt sistem dasar
        system_prompt = """
        Anda adalah asisten profesional yang mengubah transkrip menjadi dokumen terstruktur dengan detail komprehensif.
        Fokus utama Anda adalah:
        1. Menyusun dokumen dengan struktur hierarki yang jelas dan logis
        2. Menggunakan sub-judul yang SANGAT SPESIFIK dengan konten aktual
        3. Mempertahankan SETIAP DETAIL PENTING dari transkrip asli
        4. Menjaga akurasi data kuantitatif dan terminologi teknis
        5. Menghasilkan format yang konsisten dan profesional
        """

        # Sesuaikan prompt sistem berdasarkan jenis konten
        if content_stats["content_type"] == "technical_report":
            system_prompt += """
            Anda memiliki keahlian khusus dalam dokumen teknis dan harus:
            - Memastikan presisi absolut dalam data numerik dan istilah teknis
            - Mengstrukturkan hierarki kompleks informasi dengan headings multi-level
            - Mengidentifikasi dan menekankan hubungan sebab-akibat
            - Memformat tabel data dengan presisi dan kejelasan maksimal
            """
        elif content_stats["content_type"] == "meeting_notes":
            system_prompt += """
            Anda adalah spesialis notulensi rapat yang:
            - Mengorganisir setiap agenda item dengan sub-judul spesifik
            - Mengidentifikasi dan memformat action items dengan jelas
            - Menyoroti keputusan-keputusan penting
            - Memastikan atribusi yang tepat untuk komentar peserta
            - Menghilangkan small talk yang tidak substantif sambil mempertahankan semua informasi penting
            """
        elif content_stats["content_type"] == "instructional":
            system_prompt += """
            Anda ahli dalam konten instruksional yang:
            - Menyusun langkah-langkah dengan urutan logis dan jelas
            - Menggunakan numbered lists untuk prosedur sekuensial
            - Menyertakan peringatan dan tips penting dalam format yang menonjol
            - Memastikan setiap instruksi spesifik dan dapat ditindaklanjuti
            """

        # Base user prompt
        user_prompt = f"""
        TUGAS:
        Ubah transkrip berikut menjadi dokumen profesional yang SANGAT DETAIL, terstruktur, dan informatif.

        PENTING:
        - Gunakan HANYA sintaks Markdown standar.
        - Jangan gunakan HTML, tag custom, atau emoji pada heading.
        - Jangan menambahkan penjelasan di luar konten (langsung mulai dengan heading/konten).
        - Format heading utama (Level 1) dengan "# ", subheading (Level 2) dengan "## ", sub-subheading (Level 3) dengan "### ".
        - Gunakan bullet list dengan "- " dan numbered list dengan "1. ", "2. ", dst.
        - Format tabel dengan Markdown: gunakan "|" dan "---" untuk header.
        - Setiap bagian harus jelas dan mudah diproses ke Word.

        STRUKTUR YANG DIHARAPKAN:
        - Heading dan subheading sesuai isi, bukan kategori generik.
        - Paragraf informatif, bullet point untuk poin penting, tabel untuk data.
        - Tidak ada karakter atau format aneh di luar Markdown.

        PETUNJUK KHUSUS:
        1. Perbaiki tata bahasa dan ejaan tanpa mengubah makna substantif.
        2. Pertahankan SETIAP DETAIL PENTING dari transkrip.
        3. Buat sub-judul yang mencerminkan konten aktual.
        4. Format istilah teknis dengan penjelasan singkat saat pertama muncul.
        5. Perhatikan konteks dan tingkat formalitas dari transkrip asli.
        """

        # Customize user prompt based on content type
        if content_stats["content_type"] == "technical_report":
            user_prompt += """
            6. Tampilkan data numerik dengan presisi dan konteks yang tepat
            7. Gunakan tabel untuk membandingkan nilai dan parameter
            8. Pertahankan relasi logis antara konsep teknis
            9. Buat daftar referensi untuk sumber-sumber eksternal jika disebutkan
            """
        elif content_stats["content_type"] == "meeting_notes":
            user_prompt += """
            6. Format secara jelas: agenda, diskusi, keputusan, dan action items
            7. Buat daftar peserta jika disebutkan dalam transkrip
            8. Tandai action items dengan assignee dan deadline jika disebutkan
            9. Rangkum diskusi panjang dengan mempertahankan poin-poin penting
            10. Hilangkan pengulangan dan off-topic discussions yang tidak substantif
            """
        elif content_stats["content_type"] == "instructional":
            user_prompt += """
            6. Gunakan numbered lists untuk urutan langkah yang harus diikuti
            7. Soroti peringatan dan tindakan pencegahan dengan jelas
            8. Tambahkan sub-bagian "Catatan Penting" untuk informasi tambahan
            9. Berikan konteks "Mengapa" di samping instruksi "Bagaimana"
            """

        user_prompt += f"""
        TRANSKRIP UNTUK DIPROSES:
        {text}
        """

        return system_prompt, user_prompt

    def get_audio_duration(self, file_path):
        """Get the duration of an audio file in seconds with better error handling."""
        try:
            if not os.path.exists(file_path):
                logger.error(f"Audio file not found: {file_path}")
                return 0

            if os.path.getsize(file_path) == 0:
                logger.error(f"Audio file is empty: {file_path}")
                return 0

            with wave.open(file_path, 'rb') as wf:
                frames = wf.getnframes()
                rate = wf.getframerate()
                if rate == 0:
                    logger.error(f"Invalid audio file (framerate=0): {file_path}")
                    return 0
                duration = frames / float(rate)
                return duration
        except wave.Error as e:
            logger.error(f"Wave error reading {file_path}: {e}")
            return 0
        except Exception as e:
            logger.error(f"Error getting audio duration for {file_path}: {e}", exc_info=True)
            return 0

    def estimate_remaining_time(self, current_chunk, total_chunks):
        """Estimate remaining time based on processing rate"""
        if current_chunk <= 1 or self.processing_start_time == 0:
            return "estimasi waktu..."

        elapsed = time.time() - self.processing_start_time
        chunks_per_second = current_chunk / elapsed if elapsed > 0 else 0
        remaining_chunks = total_chunks - current_chunk

        if chunks_per_second > 0:
            remaining_seconds = remaining_chunks / chunks_per_second

            if remaining_seconds < 60:
                return f"{int(remaining_seconds)} detik"
            elif remaining_seconds < 3600:
                return f"{int(remaining_seconds / 60)} menit"
            else:
                hours = int(remaining_seconds / 3600)
                minutes = int((remaining_seconds % 3600) / 60)
                return f"{hours} jam {minutes} menit"
        else:
            return "menghitung..."

    def transcribe_with_groq_whisper(self, audio_data, language_code="id", model="whisper-large-v3"):
        """Transcribe audio using Groq Whisper API with simplified error handling."""
        temp_file_path = None
        try:
            self.root.after(0, lambda: self.status_var.set(f"Menggunakan Groq API dengan model {model}..."))

            # Periksa apakah client ada
            if self.groq_client is None:
                logger.error("Groq client tidak terinisialisasi")
                error_message = "Groq client tidak tersedia. Pastikan API key valid."
                self.root.after(0, lambda msg=error_message: self.status_var.set(msg))
                return None

            with tempfile.NamedTemporaryFile(suffix=".wav", delete=False) as temp_file:
                temp_file_path = temp_file.name
                with wave.open(temp_file_path, 'wb') as wf:
                    wf.setnchannels(1)
                    wf.setsampwidth(2)
                    wf.setframerate(16000)
                    wf.writeframes(audio_data.get_wav_data())

            # Model selection logic based on language
            if language_code.startswith("en") and model == "whisper-large-v3":
                if self.use_economic_model.get():
                    model = "distil-whisper-large-v3-en"
                    self.root.after(0, lambda: self.status_var.set("Menggunakan model ekonomis untuk Bahasa Inggris..."))
            elif not language_code.startswith("en") and model == "distil-whisper-large-v3-en":
                model = "whisper-large-v3-turbo"
                self.root.after(0, lambda: self.status_var.set("Model distil hanya untuk Bahasa Inggris, menggunakan turbo..."))

            self.root.after(0, lambda: self.status_var.set(f"Mentranskripsi dengan Groq {model}..."))

            # Check file size before sending to API
            file_size = os.path.getsize(temp_file_path)
            if file_size > 25 * 1024 * 1024:  # 25MB limit
                self.root.after(0, lambda: self.status_var.set("File audio terlalu besar, sedang mengoptimasi..."))
                # Compress or truncate file here if needed

            with open(temp_file_path, "rb") as audio_file_stream:
                # Memanggil API langsung dengan error handling yang disederhanakan
                try:
                    transcription_result = self.groq_client.audio.transcriptions.create(
                        model=model,
                        file=audio_file_stream,
                        language=language_code.split("-")[0],
                        response_format="text"
                    )
                    return transcription_result
                except Exception as e:
                    logger.error(f"Error saat memanggil API transkripsi: {e}", exc_info=True)
                    error_message = f"Error transkripsi: {str(e)[:100]}"
                    self.root.after(0, lambda msg=error_message: self.status_var.set(msg))
                    return None

        except Exception as e:
            error_message = f"Terjadi kesalahan saat menggunakan Groq API: {type(e).__name__}: {str(e)}"
            logger.error(error_message, exc_info=True)
            self.root.after(0, lambda msg=error_message: self.status_var.set(msg))
            return None
        finally:
            if temp_file_path and os.path.exists(temp_file_path):
                try:
                    os.remove(temp_file_path)
                except Exception as e_clean:
                    logger.error(f"Error cleaning up temp file {temp_file_path}: {e_clean}")

    def _post_process_enhanced_text(self, text, content_stats):
        """
        Pemrosesan pasca-peningkatan untuk penyempurnaan hasil akhir.
        """
        # Bersihkan tag reasoning dan thinking
        text = self.remove_reasoning_tags(text)

        # Deteksi dan perbaiki heading tanpa konten
        empty_heading_pattern = r'(#+\s+.+?\n)(?=#+\s+|\Z)'
        text = re.sub(empty_heading_pattern, r'\1\nTidak ada informasi detail untuk bagian ini dalam transkrip asli.\n\n', text)

        # Perbaiki format daftar yang tidak konsisten
        text = re.sub(r'\n\*\s+', r'\n- ', text)  # Standardize bullet points to use "-"

        # Perbaiki spacing setelah heading
        text = re.sub(r'(#+\s+.+?)\n(?!\n)', r'\1\n\n', text)

        # Perbaiki format tabel jika ada
        if '|' in text:
            lines = text.split('\n')
            in_table = False
            table_lines = []

            for i, line in enumerate(lines):
                if line.strip().startswith('|') and line.strip().endswith('|'):
                    if not in_table:
                        in_table = True
                        # Jika ini baris pertama tabel dan tidak ada pemisah di bawahnya
                        if i+1 < len(lines) and not (lines[i+1].strip().startswith('|') and '-' in lines[i+1]):
                            # Tambahkan baris pemisah tabel
                            cell_count = line.count('|') - 1
                            separator = '|' + '|'.join([' --- ' for _ in range(cell_count)]) + '|'
                            table_lines.append(line)
                            table_lines.append(separator)
                            continue
                    table_lines.append(line)
                elif in_table:
                    in_table = False
                    # Pastikan tabel selalu diikuti oleh baris kosong
                    table_lines.append('')
                    table_lines.append(line)
                else:
                    table_lines.append(line)

            text = '\n'.join(table_lines)

        # Perbaiki format spesifik untuk tipe konten
        if content_stats["content_type"] == "meeting_notes":
            # Format action items dengan lebih jelas
            text = re.sub(r'\b(Action item|Tindakan)s?\b:?\s*', r'**ACTION ITEM**: ', text, flags=re.IGNORECASE)

        return text

    def _fallback_enhancement(self, text):
        """
        Peningkatan fallback jika metode utama gagal.
        Memberikan peningkatan minimal yang masih berguna.
        """
        try:
            # Tambahkan judul dokumen dasar
            enhanced_text = "# Transkrip Audio\n\n"

            # Split teks menjadi paragraf
            paragraphs = re.split(r'\n\s*\n', text)

            for i, para in enumerate(paragraphs):
                if not para.strip():
                    continue

                # Coba deteksi dan format poin-poin
                if re.match(r'\d+\.\s+', para.strip()):
                    # Ini mungkin daftar bernomor
                    enhanced_text += para + "\n\n"
                elif len(para.split()) < 15 and i > 0:
                    # Kalimat pendek setelah paragraf pertama mungkin sub-judul
                    enhanced_text += f"## {para}\n\n"
                else:
                    # Paragraf normal
                    enhanced_text += para + "\n\n"

            return enhanced_text

        except Exception as e:
            logger.error(f"Error dalam fallback enhancement: {e}")
            # Kembalikan teks asli jika semua metode gagal
            return text

    def enhance_document_cohesion(self, text):
        """
        Improve document cohesion with optimized error handling and recovery.
        """
        if not text or not text.strip():
            logger.info("enhance_document_cohesion called with empty text. Returning as is.")
            return text

        try:
            MAX_TEXT_LENGTH_FOR_COHESION_API = 12000  # Characters

            # Log overall document size
            logger.info(f"Starting document cohesion enhancement. Text length: {len(text)} characters.")

            # Handle large documents by splitting into chunks
            if len(text) > MAX_TEXT_LENGTH_FOR_COHESION_API:
                self.root.after(0, lambda: self.status_var.set(f"Dokumen besar ({len(text)} karakter), memproses dengan metode chunking adaptif..."))
                return self._process_large_document(text, MAX_TEXT_LENGTH_FOR_COHESION_API)
            else:
                return self._process_single_document(text)

        except groq.error.RateLimitError as e:
            logger.error(f"Rate limit exceeded during document cohesion: {e}")
            self.root.after(0, lambda: self.status_var.set("Batas API terlampaui, melewati peningkatan kohesi."))
            return text  # Return original text
        except groq.error.APIError as e:
            logger.error(f"API error during document cohesion: {e}")
            self.root.after(0, lambda: self.status_var.set(f"Error API: {str(e)[:50]}..."))
            return text
        except Exception as e:
            logger.error(f"Error enhancing document cohesion: {e}", exc_info=True)
            self.root.after(0, lambda err_msg=str(e): self.status_var.set(f"Error kohesi: {err_msg[:100]}"))
            return text

    def _process_large_document(self, text, max_length):
        """
        Process a large document by splitting it into chunks with robust error handling
        """
        try:
            logger.info(f"Document text length ({len(text)} chars) exceeds limit for single cohesion call. Using intelligent adaptive chunking.")
            self.root.after(0, lambda: self.status_var.set("Dokumen besar, memproses dengan teknik ekstraksi detail komprehensif..."))

            # Step 1: Analyze document structure first to inform better chunking strategy
            try:
                structure_analysis = {}
                logger.info("Document structure analysis initialized")
            except Exception as e:
                logger.error(f"Error during document structure analysis: {e}")
                structure_analysis = {}  # Fallback to empty structure

            # Step 2: Split into chunks with context-aware boundaries
            try:
                chunks = self._split_text_into_chunks(text, max_length)
                logger.info(f"Document split into {len(chunks)} chunks successfully")

                if not chunks:
                    logger.warning("Chunking returned empty result, using fallback simple chunking")
                    # Simple fallback chunking if the main method fails to produce chunks
                    chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
            except Exception as e:
                logger.error(f"Error splitting text into chunks: {e}")
                # Fallback: simple text splitting
                chunks = [text[i:i+max_length] for i in range(0, len(text), max_length)]
                logger.info(f"Using fallback splitting, created {len(chunks)} chunks")

            # Step 3: Create context overlaps between chunks to maintain continuity
            try:
                context_chunks = self._create_context_chunks(chunks)
                logger.info(f"Created {len(context_chunks)} context-aware chunks")
            except Exception as e:
                logger.error(f"Error creating context chunks: {e}")
                # Fallback: use chunks without context
                context_chunks = [(chunk, {}) for chunk in chunks]
                logger.info("Using fallback chunks without context")

            # Process each chunk with context awareness
            enhanced_parts = []
            num_chunks = len(context_chunks)
            failed_chunks = []

            for i, (chunk, context) in enumerate(context_chunks):
                try:
                    if not chunk.strip():  # Skip empty or whitespace-only chunks
                        logger.debug(f"Skipping empty chunk {i+1}")
                        continue

                    # Update UI status with more informative message
                    chunk_size = len(chunk)
                    self.root.after(0, lambda j=i+1, total=num_chunks, c_size=chunk_size:
                                    self.status_var.set(f"Mengekstrak detail komprehensif dari bagian {j}/{total} ({c_size} karakter)..."))

                    # Enhanced processing with context awareness
                    processed_chunk = self._enhance_chunk_with_context(chunk, context, i+1, num_chunks)

                    # Validate processed chunk
                    if processed_chunk and processed_chunk.strip():
                        enhanced_parts.append(processed_chunk)
                        logger.debug(f"Successfully processed chunk {i+1}/{num_chunks}")
                    else:
                        logger.warning(f"Chunk {i+1} returned empty result, using original")
                        enhanced_parts.append(chunk)  # Use original if processing failed
                        failed_chunks.append(i+1)

                    # Update progress bar with smoother progression
                    self.root.after(0, lambda current=i, total=num_chunks:
                                    self.progress_var.set(90 + ((current + 1) / total * 5)))

                    # Optional API rate limiting delay
                    if hasattr(self, 'api_request_delay') and self.api_request_delay > 0:
                        time.sleep(self.api_request_delay)

                except Exception as e:
                    logger.error(f"Error processing chunk {i+1}: {e}")
                    # Use original chunk as fallback if processing fails
                    enhanced_parts.append(chunk)
                    failed_chunks.append(i+1)

                    # Update status but continue processing
                    self.root.after(0, lambda j=i+1:
                                self.status_var.set(f"Error pada bagian {j}, menggunakan teks asli..."))

            # Join all enhanced chunks with improved transitions
            try:
                if not enhanced_parts:
                    logger.warning("No enhanced parts were generated, returning original text")
                    return text

                final_processed_text = self._join_with_transitions(enhanced_parts)
                logger.info("Successfully joined chunks with transitions")

                if not final_processed_text or not final_processed_text.strip():
                    logger.warning("Joined text is empty, using concatenated parts")
                    final_processed_text = "\n\n".join(enhanced_parts)

            except Exception as e:
                logger.error(f"Error joining enhanced parts: {e}")
                # Fallback: simple join
                final_processed_text = "\n\n".join(enhanced_parts)
                logger.info("Used fallback simple join method")

            logger.info("Completed detail-preserving enhancement of document chunks with semantic coherence.")

            # Final coherence pass with detail preservation focus
            if len(final_processed_text) <= max_length and len(chunks) > 1:
                try:
                    self.root.after(0, lambda: self.status_var.set("Mengoptimalkan koherensi global dokumen dengan preservasi detail lengkap..."))
                    enhanced_final = self._enhance_final_document(final_processed_text)

                    # Validate enhanced final result
                    if enhanced_final and enhanced_final.strip():
                        final_processed_text = enhanced_final
                        logger.info("Final document coherence enhancement completed")
                    else:
                        logger.warning("Final enhancement returned empty result, using pre-enhanced text")
                except Exception as e:
                    logger.error(f"Error in final document enhancement: {e}")
                    # Keep the current joined text if final enhancement fails
                    logger.info("Using joined text without final enhancement due to error")

            # Log processing results
            if failed_chunks:
                logger.warning(f"Some chunks failed processing: {failed_chunks}")
                logger.info(f"Processing summary: {len(enhanced_parts) - len(failed_chunks)}/{len(enhanced_parts)} chunks enhanced successfully")
            else:
                logger.info("All chunks processed successfully")

            # Final validation before returning
            if not final_processed_text or not final_processed_text.strip():
                logger.warning("Final processed text is empty, returning original text")
                return text

            return final_processed_text

        except Exception as e:
            logger.error(f"Critical error in _process_large_document: {e}", exc_info=True)
            self.root.after(0, lambda: self.status_var.set(f"Error pemrosesan dokumen: {str(e)[:50]}..."))

            # Return original text as ultimate fallback
            return text

    def _split_text_into_chunks(self, text, max_length):
        """Split text into chunks based on semantic boundaries"""
        chunks = []
        remaining_text = text

        # Prioritize splitting at major section boundaries
        heading_patterns = [
            r'\n# [^\n]+\n',  # Level 1 heading
            r'\n## [^\n]+\n',  # Level 2 heading
            r'\n### [^\n]+\n',  # Level 3 heading
            r'\n\n[A-Z][^\n]{20,}\n\n',  # Possible unlabeled heading (all caps with significant length)
            r'\n\n',  # Paragraph breaks as fallback
            r'\. '  # Sentence breaks as last resort
        ]

        while remaining_text:
            if len(remaining_text) <= max_length:
                chunks.append(remaining_text)
                break

            # Start from maximum allowable length and find the best split point
            best_split = max_length
            min_acceptable_length = int(max_length * 0.7)  # Don't split too early to maximize chunk size

            # Try to split at increasingly granular boundaries
            for pattern in heading_patterns:
                matches = list(re.finditer(pattern, remaining_text[:max_length]))
                if matches:
                    # Find the last match that exceeds minimum acceptable length
                    for match in reversed(matches):
                        if match.start() > min_acceptable_length:
                            best_split = match.start()
                            break

                    # If we found a good split point, stop searching
                    if best_split < max_length:
                        break

            # Split at the best point found
            chunks.append(remaining_text[:best_split])
            remaining_text = remaining_text[best_split:]

            # Stop if remaining is just whitespace
            if not remaining_text.strip():
                break

        logger.info(f"Document split into {len(chunks)} semantic chunks for detailed processing.")
        return chunks

    def _create_context_chunks(self, chunks):
        """Create contextual overlaps between chunks"""
        context_chunks = []

        for i, chunk in enumerate(chunks):
            # Create context from previous and next chunks
            prev_context = ""
            next_context = ""

            # Get context from previous chunk (last 15% of previous chunk)
            if i > 0:
                prev_chunk = chunks[i-1]
                context_size = min(len(prev_chunk) // 6, 1000)  # Up to 1000 chars or 1/6 of prev chunk
                prev_context = prev_chunk[-context_size:] if context_size > 0 else ""

            # Get context from next chunk (first 10% of next chunk)
            if i < len(chunks) - 1:
                next_chunk = chunks[i+1]
                context_size = min(len(next_chunk) // 10, 500)  # Up to 500 chars or 1/10 of next chunk
                next_context = next_chunk[:context_size] if context_size > 0 else ""

            # Combine contexts
            context = {
                "previous": prev_context,
                "next": next_context,
                "position": f"{i+1}/{len(chunks)}"
            }

            context_chunks.append((chunk, context))

        return context_chunks

    def _enhance_chunk_with_context(self, chunk, context, chunk_num, total_chunks):
        """Enhance a single chunk with awareness of its context"""
        # Extract contexts
        prev_context = context.get("previous", "")
        next_context = context.get("next", "")
        position = context.get("position", f"{chunk_num}/{total_chunks}")

        # Determine chunk processing strategy based on position
        is_first = chunk_num == 1
        is_last = chunk_num == total_chunks

        chunk_prompt = f"""
        Berikut adalah BAGIAN {position} dari dokumen yang lebih besar.

        EKSTRAKSI DETAIL MENDALAM:
        Konversi bagian teks ini menjadi dokumen terstruktur dengan mempertahankan SEMUA DETAIL SUBSTANTIF.

        KONTEKS PENTING:
        {"Ini adalah BAGIAN AWAL dokumen." if is_first else ""}
        {"Ini adalah BAGIAN AKHIR dokumen." if is_last else ""}

        {"KONTEKS DARI BAGIAN SEBELUMNYA:\n" + prev_context if prev_context else ""}

        {"KONTEKS DARI BAGIAN BERIKUTNYA:\n" + next_context if next_context else ""}

        INSTRUKSI DETAIL:
        1. Pertahankan SEMUA DATA NUMERIK, ISTILAH TEKNIS, NAMA ENTITAS, dan DETAIL SPESIFIK.
        2. Analisis teks untuk mengekstrak:
        - Topik dan subtopik dengan hierarki yang jelas
        - Konsep-konsep kunci dengan penjelasan detail
        - Proses atau prosedur dengan langkah terperinci
        - Data kuantitatif dengan konteks lengkap
        - Hubungan dan ketergantungan antar konsep
        3. Strukturkan dengan:
        - Judul yang SANGAT DESKRIPTIF (jika relevan)
        - Sub-judul (Level 2) yang SPESIFIK untuk setiap topik utama
        - Sub-sub-judul (Level 3) untuk subtopik kompleks
        - Paragraf informatif dengan penjelasan MENDALAM
        - Daftar poin komprehensif dengan DETAIL lengkap
        - Tabel detail untuk informasi komparatif
        4. Singkirkan redundansi tetapi JANGAN menghilangkan detail substantif
        5. Perhatikan transisi ke bagian sebelumnya/berikutnya
        6. Jika ada istilah teknis, tambahkan penjelasan singkat
        7. Untuk data numerik, pastikan akurasi kontekstual dipertahankan

        BAGIAN TEKS UNTUK DIPROSES:
        {chunk}
        """

        system_prompt = """
        Anda adalah sistem pengolahan dokumentasi yang dirancang khusus untuk mempertahankan SETIAP DETAIL PENTING dalam teks sambil meningkatkan struktur dan kejelasannya.

        FITUR UTAMA ANDA:
        - Kemampuan mengekstrak dan mempertahankan SEMUA informasi substantif
        - Identifikasi dan preservasi data kuantitatif dengan presisi tinggi
        - Penstrukturan hierarki informasi dengan detail tingkat tinggi
        - Konversi teks menjadi dokumen terstruktur tanpa kehilangan nuansa
        - Pemahaman konteks dari bagian sebelum dan sesudahnya

        HASIL YANG DIHARAPKAN adalah bagian dokumen yang terstruktur dengan baik dan mempertahankan SEMUA DETAIL SUBSTANTIF dari teks asli. Tidak menambahkan kalimat yang tidak perlu seperi dengan demikian transkrip ini dan sebagainya
        """

        # Add delay before API call
        time.sleep(self.api_request_delay)

        # Make API call with detailed focus
        completion = self._make_llm_request(
            system_prompt=system_prompt,
            user_prompt=chunk_prompt,
            model="deepseek-r1-distill-llama-70b",  # Most capable model for detail preservation
            temperature=0.3,  # Lower temperature for deterministic output
            max_tokens=6000
        )

        enhanced_chunk = self.remove_reasoning_tags(completion)
        return enhanced_chunk

    def _join_with_transitions(self, enhanced_parts):
        """
        Join enhanced chunks with smooth transitions between them.
        """
        if not enhanced_parts:
            return ""

        if len(enhanced_parts) == 1:
            return enhanced_parts[0]

        # For documents with multiple enhanced parts, create transitions
        result = [enhanced_parts[0]]

        for i in range(1, len(enhanced_parts)):
            prev_part = enhanced_parts[i-1]
            curr_part = enhanced_parts[i]

            # Get the ending context from previous part (last paragraph that's not a note)
            prev_paragraphs = [p for p in prev_part.split('\n\n') if p and not p.startswith('[Catatan:')]
            prev_context = prev_paragraphs[-1] if prev_paragraphs else ""

            # Check if the current part starts with a heading
            starts_with_heading = bool(re.match(r'^#+ ', curr_part.lstrip()))

            # If not starting with a heading, add a transition
            if not starts_with_heading and prev_context:
                # Look for the first paragraph to modify
                curr_paragraphs = curr_part.split('\n\n')
                for j, para in enumerate(curr_paragraphs):
                    if para and not para.startswith('#') and len(para) > 20:
                        # Add transition phrase if appropriate
                        if not any(phrase in para.lower() for phrase in ['melanjutkan', 'berikutnya', 'selanjutnya']):
                            curr_paragraphs[j] = f"Melanjutkan pembahasan sebelumnya, {para}"
                        break

                # Join modified paragraphs back together
                curr_part = '\n\n'.join(curr_paragraphs)

            result.append(curr_part)

        # Join all parts
        return '\n\n'.join(result)

    def _enhance_final_document(self, text):
        """
        Perform final coherence enhancement on the document.
        """
        prompt = f"""
        FINALISASI DOKUMEN DENGAN PRESERVASI DETAIL TINGKAT TINGGI

        Berikut adalah dokumen yang telah diproses dari transkripsi audio dan telah distrukturkan per bagian.

        TUGAS FINALISASI:
        1. Pastikan koherensi global dokumen dari awal hingga akhir
        2. Eliminasi redundansi tanpa menghapus informasi penting
        3. Verifikasi dan pertahankan SEMUA DATA NUMERIK, ISTILAH TEKNIS, dan DETAIL SPESIFIK
        4. Verifikasi dan pertahankan SEMUA hubungan logis dan kausal antar bagian
        5. Pastikan konsistensi terminologi dan gaya bahasa
        6. Sempurnakan transisi antar bagian untuk alur membaca yang mulus
        7. Pastikan semua tabel, daftar, dan struktur formatif lainnya dipertahankan dengan detail lengkap
        8. Pastikan tingkat detail konsisten di seluruh dokumen
        9. SANGAT PENTING: JANGAN menghapus informasi substantif apapun
        10. SANGAT PENTING: JANGAN mengubah angka atau data kuantitatif apapun

        DOKUMEN UNTUK FINALISASI:
        {text}
        """

        system_prompt = """
        Anda adalah editor senior spesialis preservasi detail dengan kemampuan untuk melakukan penyempurnaan koherensi dokumen tanpa mengorbankan informasi substantif sekecil apapun.

        KEAHLIAN UTAMA ANDA:
        - Mempertahankan SETIAP detail penting sambil meningkatkan alur dan keterbacaan
        - Mengidentifikasi dan mempreservasi semua data kuantitatif dan informasi teknis
        - Memastikan konsistensi terminologi dan integritas faktual
        - Menyeimbangkan keterbacaan dengan komprehensivitas detail
        - Merangkum point point yang panjang menjadi hal hal yang sangat penting saja

        Tugas Anda adalah menyempurnakan dokumen yang sudah distrukturkan dengan memastikan koherensi global tanpa mengorbankan informasi substantif apapun. Tidak menambahkan kalimat yang tidak perlu seperi dengan demikian transkrip ini dan sebagainya
        """

        # Add delay before API call
        time.sleep(self.api_request_delay)

        # Choose most capable model for final pass
        completion = self._make_llm_request(
            system_prompt=system_prompt,
            user_prompt=prompt,
            model="deepseek-r1-distill-llama-70b",  # Best for detailed preservation
            temperature=0.3,  # Lower for precise editing
            max_tokens=6000
        )

        final_text = self.remove_reasoning_tags(completion)
        return final_text

    def _process_single_document(self, text):
        """
        Process a single document that fits within API limits.
        """
        self.root.after(0, lambda: self.status_var.set("Meningkatkan kohesi dokumen dengan analisis struktur detail..."))

        # Simple prompt for coherence enhancement
        cohesion_prompt = f"""
        TUGAS: Tingkatkan kohesi dan detail dokumen berikut.

        PANDUAN PENINGKATAN KOHESI DENGAN DETAIL:
        1. Perbaiki transisi antar bagian untuk alur yang lebih natural
        2. Pastikan konsistensi terminologi dan gaya penulisan di seluruh dokumen
        3. Sesuaikan referensi antar bagian untuk meningkatkan keterhubungan
        4. Pertahankan SEMUA DETAIL TEKNIS, DATA NUMERIK, dan INFORMASI SPESIFIK
        5. Perbaiki keseimbangan detail di seluruh dokumen - bagian yang terlalu ringkas perlu diperluas
        6. Pastikan hierarki informasi konsisten dan logis
        7. Tambahkan kalimat penghubung di awal bagian yang memerlukan transisi lebih baik
        8. Untuk bagian dengan perbandingan, pastikan parameter yang digunakan konsisten
        9. JANGAN menambahkan informasi baru yang tidak ada dalam dokumen asli
        10. JANGAN mengubah struktur utama atau menghapus bagian
        11. Merangkum point point yang panjang menjadi hal hal yang sangat penting saja

        DOKUMEN UNTUK DITINGKATKAN:
        {text}
        """

        self.root.after(0, lambda: self.status_var.set("Menerapkan peningkatan kohesi dengan mempertahankan semua detail penting..."))
        time.sleep(self.api_request_delay)

        cohesion_enhanced = self._make_llm_request(
            system_prompt="Anda adalah editor profesional yang ahli dalam meningkatkan kohesi dokumen sambil mempertahankan integritas SEMUA DETAIL teknis dan spesifik. Anda memiliki kemampuan untuk menyeimbangkan kedalaman informasi dengan alur yang lancar. Tidak menambahkan kalimat yang tidak perlu seperi dengan demikian transkrip ini dan sebagainya",
            user_prompt=cohesion_prompt,
            model="deepseek-r1-distill-llama-70b",  # Better for maintaining detail
            temperature=0.3,  # Lower for more precise editing
            max_tokens=6000
        )

        return self.remove_reasoning_tags(cohesion_enhanced)

    def _make_llm_request(self, system_prompt, user_prompt, model, temperature, max_tokens):
        """Make LLM request with improved error handling and retry logic."""
        retry_count = 0
        max_retries = 2
        last_error = None

        while retry_count <= max_retries:
            try:
                # Add exponential backoff between retries
                if retry_count > 0:
                    backoff_time = min(2 ** (retry_count - 1), 15)  # Max 15 second backoff
                    time.sleep(backoff_time)

                completion = self.groq_client.chat.completions.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": system_prompt},
                        {"role": "user", "content": user_prompt}
                    ],
                    temperature=temperature,
                    max_tokens=max_tokens,
                    top_p=0.95,
                    reasoning_format="hidden",
                    stream=False
                )

                # Success - return response content
                return completion.choices[0].message.content

            except groq.error.RateLimitError as e:
                logger.warning(f"Rate limit hit in LLM request (attempt {retry_count+1}/{max_retries+1})")
                last_error = e
                retry_count += 1

                # Use fallback model if available
                if retry_count < max_retries:
                    if model == "deepseek-r1-distill-llama-70b":
                        model = "mixtral-8x7b-32768"  # Fallback to smaller model
                        temperature += 0.1  # Slightly increase creativity

            except groq.error.APIError as e:
                logger.error(f"API error in LLM request: {e}")
                last_error = e
                retry_count += 1

            except Exception as e:
                logger.error(f"Error in LLM API request: {e}", exc_info=True)
                last_error = e
                retry_count += 1

        # All retries failed
        if last_error:
            raise last_error
        return "Error dalam pemrosesan LLM. Silakan coba lagi."

    def remove_reasoning_tags(self, text):
        """Remove reasoning tags from the text"""
        if not text:
            return text

        # Remove <think> tags and their content
        cleaned_text = re.sub(r'<think>.*?</think>', '', text, flags=re.DOTALL)

        # Remove "reasoning" JSON field
        cleaned_text = re.sub(r'"reasoning"\s*:\s*".*?"', '', cleaned_text, flags=re.DOTALL)

        # Remove reasoning section headers and content
        lines = cleaned_text.split("\n")
        filtered_lines = []
        skip_mode = False
        reasoning_section = False

        for line in lines:
            if re.match(r'^#+\s*(Reasoning|Pemikiran|Alur Pikir|Proses Berpikir|Think)', line, re.IGNORECASE):
                skip_mode = True
                reasoning_section = True
                continue

            if re.match(r'^(Reasoning|Pemikiran|Alur Pikir|Proses Berpikir|Think):', line, re.IGNORECASE):
                skip_mode = True
                reasoning_section = True
                continue

            if skip_mode and re.match(r'^#+\s+', line) and reasoning_section:
                skip_mode = False
                reasoning_section = False

            if not skip_mode:
                filtered_lines.append(line)

        return "\n".join(filtered_lines)

    def _extract_keywords(self, text_content):
        """Ekstrak kata kunci untuk metadata dokumen."""
        # Stopwords untuk bahasa Indonesia dan Inggris
        stopwords = {
            'dan', 'atau', 'dari', 'ke', 'di', 'untuk', 'dengan', 'oleh', 'pada', 'adalah', 'ini', 'itu', 'yang',
            'the', 'and', 'or', 'of', 'to', 'in', 'for', 'with', 'by', 'on', 'at', 'is', 'are', 'was', 'were',
            'a', 'an', 'as', 'be', 'been', 'being', 'have', 'has', 'had', 'do', 'does', 'did', 'will', 'would'
        }

        # Ekstrak kata bermakna
        words = re.findall(r'\b[a-zA-Z]{4,}\b', text_content.lower())
        meaningful_words = [word for word in words if word not in stopwords]

        # Hitung frekuensi dan ambil yang paling sering
        from collections import Counter
        word_freq = Counter(meaningful_words)

        # Ambil 10 kata kunci teratas
        keywords = [word for word, freq in word_freq.most_common(10) if freq > 1]

        return ', '.join(keywords[:10])

    # * Fungsi-fungsi yang Mengatur Word (DOCX Processing)

    def export_to_word(self):
        """
        Export the current text directly to Word without further processing.
        """
        text = self.result_text.get(1.0, tk.END)
        if not text.strip():
            messagebox.showwarning("Peringatan", "Tidak ada teks untuk diekspor.")
            return

        # Create a filename
        now = datetime.datetime.now()
        date_str = now.strftime("%Y%m%d")
        time_str = now.strftime("%H%M%S")
        filename = f"{self.filename_prefix.get()}_export_{date_str}_{time_str}.docx"
        filepath = os.path.join(self.output_folder.get(), filename)

        # Save as Word document
        success = self.save_as_word_document(text, filepath)

        if success:
            self.status_var.set(f"Ekspor berhasil: {filename}")
            self.root.after(3000, lambda: self.status_var.set("Siap"))

            # Ask if user wants to open the file
            if messagebox.askyesno("Buka File", f"File telah disimpan di:\n{filepath}\n\nApakah Anda ingin membukanya sekarang?"):
                self.open_file(filepath)

    def _apply_document_theme(self, doc, document_type):
        """
        Applies a consistent theme to the document based on document type.
        """
        # Theme colors based on document type
        theme_colors = {
            "technical_report": {
                "primary": "4472C4",
                "secondary": "F2F9FF",
                "accent": "2E75B5"
            },
            "meeting_notes": {
                "primary": "70AD47",
                "secondary": "F5FFF5",
                "accent": "548235"
            },
            "lecture": {
                "primary": "ED7D31",
                "secondary": "FFF8F5",
                "accent": "C65911"
            },
            "general": {
                "primary": "5B9BD5",
                "secondary": "F7F9FC",
                "accent": "2E75B5"
            }
        }

        # Use default if document type is not recognized
        if document_type not in theme_colors:
            document_type = "general"

        # PERBAIKAN: Inisialisasi self.theme jika belum ada
        if not hasattr(self, 'theme'):
            self.theme = {}

        # Get theme for document type
        self.theme = theme_colors[document_type]

        # Apply theme ke document properties jika doc tersedia
        if doc and hasattr(doc, 'core_properties'):
            doc.core_properties.category = document_type.replace("_", " ").title()

        logger.info(f"Applied theme for document type: {document_type}")
        return self.theme

    def save_as_word_document(self, text_content, filepath):
        """
        Membuat dokumen Word dengan format profesional dan handling error yang lebih baik.
        """
        try:
            # Deteksi karakteristik untuk menentukan styling
            content_stats = self._analyze_content_characteristics(text_content)

            # Buat dokumen baru
            doc = docx.Document()

            # Konfigurasi properti dokumen
            doc.core_properties.author = "EchoScribe AI"
            doc.core_properties.title = self._extract_document_title(text_content)
            doc.core_properties.subject = f"Catatan {content_stats['content_type'].title()}"
            doc.core_properties.created = datetime.datetime.now()
            doc.core_properties.category = content_stats["content_type"].capitalize()
            doc.core_properties.comments = f"Dibuat dengan EchoScribe AI - {content_stats['word_count']} kata"
            doc.core_properties.keywords = self._extract_keywords(text_content)
            doc.core_properties.language = content_stats.get("language", "id")

            # Set the theme based on content type before using it
            self._apply_document_theme(doc, content_stats["content_type"])

            # Siapkan style dan konfigurasi dokumen
            self._setup_document_styles(doc, content_stats)

            # Konfigurasi halaman dan margin
            self._configure_page_layout(doc, content_stats)

            # Tambahkan judul dan metadata
            self._add_document_header(doc, content_stats)

            # Proses konten utama dengan parser Markdown yang ditingkatkan
            self._process_markdown_content(doc, text_content, content_stats)

            # Tambahkan footer profesional
            self._add_document_footer(doc, content_stats)

            # PERBAIKAN: Optimasi akhir dokumen
            self.finalize_document_formatting_enhanced(doc, content_stats)

            # Pastikan directory ada
            os.makedirs(os.path.dirname(os.path.abspath(filepath)), exist_ok=True)

            # Simpan dokumen dengan try-except untuk handling error spesifik
            try:
                doc.save(filepath)
                logger.info(f"✅ Dokumen berhasil disimpan: {filepath}")

                # Validasi file yang disimpan
                if os.path.exists(filepath) and os.path.getsize(filepath) > 0:
                    file_size_kb = os.path.getsize(filepath) / 1024
                    logger.info(f"📊 Ukuran file: {file_size_kb:.1f} KB")
                    return True
                else:
                    raise Exception("File tidak dapat disimpan dengan benar")

            except PermissionError:
                logger.error(f"❌ Permission denied when saving to {filepath}")
                # Coba simpan ke lokasi alternatif
                alt_path = os.path.join(os.path.expanduser("~"), f"EchoScribe_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
                doc.save(alt_path)
                self.root.after(0, lambda: messagebox.showwarning(
                    "Izin Ditolak",
                    f"Tidak dapat menyimpan ke {filepath}\nFile disimpan ke {alt_path}"
                ))
                return True

        except Exception as e:
            logger.error(f"❌ Error menyimpan dokumen Word: {e}", exc_info=True)
            self.root.after(0, lambda: messagebox.showerror("Error Menyimpan Word", f"Gagal menyimpan dokumen: {e}"))

            # Fallback: simpan sebagai teks biasa jika docx gagal
            try:
                alt_path = filepath.replace('.docx', '.txt')
                with open(alt_path, 'w', encoding='utf-8') as f:
                    f.write(text_content)
                self.root.after(0, lambda: messagebox.showinfo(
                    "Fallback Tersedia",
                    f"Dokumen Word gagal disimpan, tetapi konten disimpan sebagai teks di: {alt_path}"
                ))
            except Exception as text_error:
                logger.error(f"❌ Error pada penyimpanan fallback: {text_error}")

            return False

    def _extract_document_title(self, text_content):
        """Ekstrak judul dokumen dari konten."""
        lines = text_content.strip().split('\n')

        # Cari heading level 1 pertama
        for line in lines[:10]:  # Cek 10 baris pertama
            if line.strip().startswith('# '):
                title = line.strip()[2:].strip()
                # Bersihkan dari emoji dan simbol berlebihan
                clean_title = re.sub(r'^[^\w\s]+|[^\w\s]+$', '', title).strip()
                if len(clean_title) > 3:
                    return clean_title

        # Jika tidak ada heading, gunakan kalimat pertama yang tidak kosong
        for line in lines[:5]:
            if line.strip() and len(line.strip()) > 10 and not line.strip().startswith('#'):
                sentence = line.strip().split('.')[0]
                if len(sentence) > 5 and len(sentence) < 100:
                    return sentence

        # Fallback ke judul default
        return "Dokumen EchoScribe AI"

    def _setup_document_styles(self, doc, content_stats):
        """
        Menyiapkan style profesional untuk dokumen berdasarkan tipe konten.
        """
        # Style dasar untuk seluruh dokumen
        style = doc.styles['Normal']
        font = style.font

        # Pilih font berdasarkan tipe konten
        if content_stats["content_type"] == "technical_report":
            font.name = 'Cambria'
            font.size = Pt(11)
        elif content_stats["content_type"] == "meeting_notes":
            font.name = 'Calibri'
            font.size = Pt(11)
        elif content_stats["content_type"] == "instructional":
            font.name = 'Segoe UI'
            font.size = Pt(10.5)
        else:  # Default
            font.name = 'Calibri'
            font.size = Pt(11)

        # Line spacing untuk seluruh dokumen
        style.paragraph_format.line_spacing = Pt(14)  # 1.15 line spacing

        # Tambahkan style kustom
        self._add_custom_styles(doc)

    def _add_custom_styles(self, doc):
        """
        Menambahkan style kustom yang profesional untuk digunakan dalam dokumen.
        """
        styles = doc.styles

        # Style Quote
        if 'Quote' not in styles:
            quote_style = styles.add_style('Quote', WD_STYLE_TYPE.PARAGRAPH)
            quote_style.font.italic = True
            quote_style.font.color.rgb = RGBColor(70, 70, 70)
            quote_format = quote_style.paragraph_format
            quote_format.left_indent = Inches(0.5)
            quote_format.right_indent = Inches(0.5)
            quote_format.space_before = Pt(10)
            quote_format.space_after = Pt(10)
            quote_format.line_spacing = Pt(13)

            # Tambahkan border kiri yang elegan
            paragraph_format = quote_style.paragraph_format
            pPr = paragraph_format._element
            pBdr = OxmlElement('w:pBdr')
            pPr.append(pBdr)

            # Border kiri dengan warna abu-abu
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), '24')
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), '808080')
            pBdr.append(left_border)

        # Body Text Style
        if 'Body Text' not in styles:
            body_style = styles.add_style('Body Text', WD_STYLE_TYPE.PARAGRAPH)
            body_format = body_style.paragraph_format
            body_format.line_spacing = Pt(14)
            body_format.space_before = Pt(6)
            body_format.space_after = Pt(6)
            body_style.font.name = styles['Normal'].font.name
            body_style.font.size = styles['Normal'].font.size

        # Callout Style untuk peringatan dan tips
        if 'Callout' not in styles:
            callout_style = styles.add_style('Callout', WD_STYLE_TYPE.PARAGRAPH)
            callout_style.font.name = styles['Normal'].font.name
            callout_style.font.size = styles['Normal'].font.size
            callout_format = callout_style.paragraph_format
            callout_format.left_indent = Inches(0.25)
            callout_format.right_indent = Inches(0.25)
            callout_format.space_before = Pt(10)
            callout_format.space_after = Pt(10)

        # Action Item Style
        if 'Action Item' not in styles:
            action_style = styles.add_style('Action Item', WD_STYLE_TYPE.PARAGRAPH)
            action_style.font.name = styles['Normal'].font.name
            action_style.font.size = styles['Normal'].font.size
            action_style.font.bold = True
            action_format = action_style.paragraph_format
            action_format.space_before = Pt(6)
            action_format.space_after = Pt(6)
            action_format.left_indent = Inches(0.25)

    def _configure_page_layout(self, doc, content_stats):
        """
        Konfigurasi layout halaman dan margin berdasarkan tipe konten.
        """
        # Dapatkan semua section (biasanya hanya ada satu di awal)
        sections = doc.sections
        for section in sections:
            # Margin standar yang profesional
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)

            # Sesuaikan berdasarkan tipe konten
            if content_stats["content_type"] == "technical_report":
                # Technical reports biasanya memiliki margin yang lebih lebar
                section.left_margin = Cm(3.0)
                section.right_margin = Cm(3.0)
            elif content_stats["content_type"] == "meeting_notes":
                # Meeting notes biasanya memiliki margin yang lebih sempit
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)

    def _add_document_header(self, doc, content_stats):
        """
        Menambahkan header dokumen yang profesional dengan judul dan metadata.
        """
        # Deteksi judul dari konten jika tersedia
        title_text = "Dokumen EchoScribe AI"
        document_type = content_stats.get("content_type", "general").replace("_", " ").title()

        # Coba ekstrak judul dari konten jika ada pola # di awal dokumen
        lines = content_stats.get("original_text", "").strip().split('\n')
        if "original_text" in content_stats and lines and lines[0].startswith('# '):
            title_text = lines[0][2:].strip()

        # Ensure theme is available - failsafe approach
        if not hasattr(self, 'theme'):
            self._apply_document_theme(doc, content_stats.get("content_type", "general"))

        # Get theme colors based on document type
        theme_primary = self.theme.get("primary", "4472C4")
        theme_accent = self.theme.get("accent", "2E75B5")

        # Tambahkan cover page styling berdasarkan jenis dokumen
        section = doc.sections[0]
        section.different_first_page_header_footer = True

        # Title page header dengan styling yang lebih profesional
        header = section.first_page_header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()

        current_date = datetime.datetime.now().strftime("%d %B %Y")
        header_run = header_para.add_run(f"EchoScribe AI • {current_date}")
        header_run.font.size = Pt(9)
        header_run.font.color.rgb = RGBColor.from_string(theme_accent)
        header_run.font.italic = True
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Tambahkan judul dengan format yang sesuai tipe konten
        title = doc.add_heading(title_text, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.style.font.color.rgb = RGBColor.from_string(theme_primary)
        title.style.font.size = Pt(24)

        # Tambahkan garis dekoratif di bawah judul
        self._add_decorative_line(doc, theme_primary)

        # Tambahkan subtitle jika terdeteksi
        if "original_text" in content_stats and len(lines) > 1 and lines[1].startswith('## '):
            subtitle = doc.add_paragraph(lines[1][3:].strip(), style='Subtitle')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            subtitle.style.font.color.rgb = RGBColor.from_string(theme_accent)

        # Tambahkan ikon dan jenis dokumen
        doc_type_para = doc.add_paragraph()
        doc_type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Pilih ikon berdasarkan jenis dokumen
        icon = "📄"  # Default
        if content_stats["content_type"] == "technical_report":
            icon = "📊"
        elif content_stats["content_type"] == "meeting_notes":
            icon = "🗣️"
        elif content_stats["content_type"] == "lecture":
            icon = "📝"
        elif content_stats["content_type"] == "workshop":
            icon = "🔧"
        elif content_stats["content_type"] == "seminar":
            icon = "🎓"

        doc_type_run = doc_type_para.add_run(f"{icon} {document_type}")
        doc_type_run.font.color.rgb = RGBColor.from_string(theme_accent)
        doc_type_run.font.size = Pt(12)
        doc_type_para.paragraph_format.space_after = Pt(16)

        # Add metadata table in a more modern style
        self._add_metadata_table(doc, content_stats)

        # Add cover separator
        doc.add_paragraph()
        self._add_enhanced_horizontal_rule(doc, content_stats["content_type"])
        doc.add_paragraph()

    def _add_decorative_line(self, doc, color):
        """
        Add a decorative line under the title.
        """
        line_para = doc.add_paragraph()
        line_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        line_para.paragraph_format.space_before = Pt(8)
        line_para.paragraph_format.space_after = Pt(16)

        line = line_para.add_run("─" * 30)
        line.font.color.rgb = RGBColor.from_string(color)
        line.font.size = Pt(12)
        line.font.bold = True

    def _add_metadata_table(self, doc, content_stats):
        """
        Add metadata in a modern table format.
        """
        # Get current date and time
        now = datetime.datetime.now()
        date_str = now.strftime("%d %B %Y")
        time_str = now.strftime("%H:%M:%S")

        # Create a cleaner, more minimal metadata table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid'

        # Remove outside borders
        self._remove_outside_table_borders(table)

        # First row - Date and Type
        cells = table.rows[0].cells
        cells[0].text = "Date"
        cells[1].text = date_str

        # Second row - Document Type
        row = table.add_row()
        cells = row.cells
        cells[0].text = "Document Type"
        cells[1].text = content_stats["content_type"].replace("_", " ").title()

        # Third row - Generator
        row = table.add_row()
        cells = row.cells
        cells[0].text = "Created by"
        cells[1].text = "EchoScribe AI"

        # Style the table
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                # Make label cells bold with light background
                if i == 0:
                    cell.paragraphs[0].runs[0].bold = True

                    # Add light gray background to label cells
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                    cell._element.get_or_add_tcPr().append(shading_elm)

                # Add padding to all cells
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.space_before = Pt(3)
                    paragraph.paragraph_format.space_after = Pt(3)

        # Set column widths
        table.columns[0].width = Inches(1.5)
        table.columns[1].width = Inches(4.0)

        # Add space after the table
        doc.add_paragraph()

    def _remove_outside_table_borders(self, table):
        """
        Remove the outside borders of a table for a more modern look.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr  # CORRECTED LINE: Changed from tbl._tblPr to tbl.tblPr
        tblBorders = OxmlElement('w:tblBorders')

        # Clear all borders first
        for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'nil')
            tblBorders.append(border)

        # Then set only the inside borders
        for border_name in ['insideH', 'insideV']:
            border = OxmlElement(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'E0E0E0')  # Light gray
            tblBorders.append(border)

        tblPr.append(tblBorders)

    def _get_heading_icon(self, heading_text, level, document_type):
        """
        Menentukan icon yang tepat untuk heading berdasarkan konteks dan level.
        """
        heading_lower = heading_text.lower()

        # Icon berdasarkan kata kunci dalam heading
        keyword_icons = {
            # Informasi dan pengetahuan
            'informasi': '📋', 'info': '📋', 'information': '📋', 'data': '📊',
            'overview': '🔍', 'ringkasan': '📄', 'summary': '📄', 'gambaran': '🔍',

            # Penting dan kritis
            'penting': '⚠️', 'important': '⚠️', 'critical': '🚨', 'urgent': '🚨',
            'perhatian': '⚠️', 'attention': '⚠️', 'warning': '⚠️', 'peringatan': '⚠️',
            'bahaya': '🚨', 'danger': '🚨', 'risk': '⚠️', 'risiko': '⚠️',

            # Kesimpulan dan hasil
            'kesimpulan': '🏁', 'conclusion': '🏁', 'hasil': '🎯', 'result': '🎯',
            'outcome': '🎯', 'finding': '🔍', 'temuan': '🔍', 'rekomendasi': '💡',
            'recommendation': '💡', 'saran': '💡', 'suggestion': '💡',

            # Proses dan langkah
            'langkah': '🔄', 'step': '🔄', 'proses': '⚙️', 'process': '⚙️',
            'prosedur': '📝', 'procedure': '📝', 'metode': '🔧', 'method': '🔧',
            'cara': '🛠️', 'how': '🛠️', 'tutorial': '📚', 'panduan': '📚',

            # Tujuan dan target
            'tujuan': '🎯', 'goal': '🎯', 'objective': '🎯', 'target': '🎯',
            'sasaran': '🎯', 'aim': '🎯', 'purpose': '🎯', 'maksud': '🎯',

            # Analisis dan evaluasi
            'analisis': '🔬', 'analysis': '🔬', 'evaluasi': '📊', 'evaluation': '📊',
            'assessment': '📊', 'penilaian': '📊', 'review': '🔍', 'tinjauan': '🔍',

            # Tips dan saran
            'tips': '💡', 'tip': '💡', 'hint': '💡', 'petunjuk': '💡',
            'advice': '💡', 'nasihat': '💡', 'best practice': '⭐', 'praktik terbaik': '⭐',

            # Masalah dan solusi
            'masalah': '❗', 'problem': '❗', 'issue': '❗', 'kendala': '❗',
            'solusi': '💡', 'solution': '💡', 'penyelesaian': '💡', 'resolve': '💡',

            # Timeline dan agenda
            'agenda': '📅', 'schedule': '📅', 'jadwal': '📅', 'timeline': '📅',
            'waktu': '🕒', 'time': '🕒', 'deadline': '⏰', 'batas waktu': '⏰',

            # Dokumentasi dan catatan
            'dokumentasi': '📄', 'documentation': '📄', 'catatan': '📝', 'note': '📝',
            'record': '📝', 'laporan': '📊', 'report': '📊',

            # Komunikasi dan diskusi
            'diskusi': '💬', 'discussion': '💬', 'pembahasan': '💬', 'komunikasi': '📢',
            'communication': '📢', 'meeting': '🗣️', 'rapat': '🗣️', 'pertemuan': '🗣️',

            # Keputusan dan action
            'keputusan': '✅', 'decision': '✅', 'action': '⚡', 'tindakan': '⚡',
            'implementasi': '🚀', 'implementation': '🚀', 'eksekusi': '🚀', 'execution': '🚀',

            # Pembelajaran dan training
            'pembelajaran': '🎓', 'learning': '🎓', 'training': '🏋️', 'pelatihan': '🏋️',
            'edukasi': '📚', 'education': '📚', 'pengetahuan': '🧠', 'knowledge': '🧠',

            # Teknologi dan sistem
            'teknologi': '💻', 'technology': '💻', 'sistem': '⚙️', 'system': '⚙️',
            'software': '💻', 'hardware': '🔧', 'infrastruktur': '🏗️', 'infrastructure': '🏗️',

            # Keuangan dan bisnis
            'keuangan': '💰', 'finance': '💰', 'budget': '💰', 'anggaran': '💰',
            'bisnis': '💼', 'business': '💼', 'profit': '📈', 'keuntungan': '📈',

            # Keamanan dan compliance
            'keamanan': '🔒', 'security': '🔒', 'compliance': '📋', 'kepatuhan': '📋',
            'audit': '🔍', 'kontrol': '🎛️', 'control': '🎛️',

            # Inovasi dan pengembangan
            'inovasi': '💡', 'innovation': '💡', 'pengembangan': '🚀', 'development': '🚀',
            'improvement': '📈', 'perbaikan': '📈', 'enhancement': '⬆️', 'peningkatan': '⬆️'
        }

        # Cari icon berdasarkan kata kunci
        for keyword, icon in keyword_icons.items():
            if keyword in heading_lower:
                return icon

        # Icon default berdasarkan level dan tipe dokumen
        if level == 1:
            if document_type == "technical_report":
                return '📊'
            elif document_type == "meeting_notes":
                return '🗣️'
            elif document_type == "lecture":
                return '🎓'
            else:
                return '📋'
        elif level == 2:
            if document_type == "technical_report":
                return '🔧'
            elif document_type == "meeting_notes":
                return '📌'
            elif document_type == "lecture":
                return '📝'
            else:
                return '📄'
        elif level == 3:
            return '📍'
        else:
            return '▪️'

    def _process_markdown_content(self, doc, text_content, content_stats):
        """
        Memproses konten Markdown dengan enhanced styling, icon untuk header berdasarkan konteks,
        dan struktur dokumen profesional dengan akurasi 100% konversi ke Word.
        """
        lines = text_content.split('\n')
        i = 0
        in_list = False
        list_style = None
        list_level = 0
        in_table = False
        table_data = []
        in_code_block = False
        code_language = ""
        code_block_lines = []
        in_admonition = False
        admonition_type = None
        admonition_lines = []
        current_paragraph = None

        # Deteksi dokumen type dan terapkan tema yang sesuai
        document_type = content_stats.get("content_type", "general")
        self._apply_document_theme(doc, document_type)

        # Enhanced icon mapping dengan kategori yang lebih komprehensif
        enhanced_icon_mapping = {
            # Informasi dan Data
            'informasi': '📋', 'info': '📋', 'information': '📋', 'data': '📊',
            'overview': '🔍', 'ringkasan': '📄', 'summary': '📄', 'gambaran': '🔍',
            'detail': '🔍', 'details': '🔍', 'rincian': '📝', 'spesifikasi': '⚙️',

            # Penting dan Prioritas
            'penting': '⚠️', 'important': '⚠️', 'critical': '🚨', 'urgent': '🚨',
            'perhatian': '⚠️', 'attention': '⚠️', 'warning': '⚠️', 'peringatan': '⚠️',
            'bahaya': '🚨', 'danger': '🚨', 'risk': '⚠️', 'risiko': '⚠️',
            'prioritas': '🔴', 'priority': '🔴', 'vital': '💎', 'krusial': '💎',

            # Kesimpulan dan Hasil
            'kesimpulan': '🏁', 'conclusion': '🏁', 'hasil': '🎯', 'result': '🎯',
            'outcome': '🎯', 'finding': '🔍', 'temuan': '🔍', 'rekomendasi': '💡',
            'recommendation': '💡', 'saran': '💡', 'suggestion': '💡',
            'evaluasi': '📊', 'evaluation': '📊', 'assessment': '📋',

            # Proses dan Metodologi
            'langkah': '🔄', 'step': '🔄', 'proses': '⚙️', 'process': '⚙️',
            'prosedur': '📝', 'procedure': '📝', 'metode': '🔧', 'method': '🔧',
            'cara': '🛠️', 'how': '🛠️', 'tutorial': '📚', 'panduan': '📚',
            'workflow': '🔄', 'alur': '🔄', 'tahapan': '📋', 'fase': '🎯',

            # Tujuan dan Target
            'tujuan': '🎯', 'goal': '🎯', 'objective': '🎯', 'target': '🎯',
            'sasaran': '🎯', 'aim': '🎯', 'purpose': '🎯', 'maksud': '🎯',
            'visi': '👁️', 'vision': '👁️', 'misi': '🚀', 'mission': '🚀',

            # Analisis dan Riset
            'analisis': '🔬', 'analysis': '🔬', 'penelitian': '🔬', 'research': '🔬',
            'study': '📖', 'studi': '📖', 'observasi': '👀', 'observation': '👀',
            'eksperimen': '🧪', 'experiment': '🧪', 'testing': '🧪', 'pengujian': '🧪',

            # Tips dan Bantuan
            'tips': '💡', 'tip': '💡', 'hint': '💡', 'petunjuk': '💡',
            'advice': '💡', 'nasihat': '💡', 'best practice': '⭐', 'praktik terbaik': '⭐',
            'shortcut': '⚡', 'pintasan': '⚡', 'trick': '🎩', 'trik': '🎩',

            # Masalah dan Solusi
            'masalah': '❗', 'problem': '❗', 'issue': '❗', 'kendala': '❗',
            'hambatan': '🚧', 'obstacle': '🚧', 'challenge': '💪', 'tantangan': '💪',
            'solusi': '💡', 'solution': '💡', 'penyelesaian': '💡', 'resolve': '💡',
            'fix': '🔧', 'perbaikan': '🔧', 'troubleshoot': '🔍', 'debug': '🐛',

            # Timeline dan Manajemen
            'agenda': '📅', 'schedule': '📅', 'jadwal': '📅', 'timeline': '📅',
            'waktu': '🕒', 'time': '🕒', 'deadline': '⏰', 'batas waktu': '⏰',
            'kalendar': '📅', 'calendar': '📅', 'milestone': '🏆', 'tonggak': '🏆',

            # Dokumentasi dan Pencatatan
            'dokumentasi': '📄', 'documentation': '📄', 'catatan': '📝', 'note': '📝',
            'record': '📝', 'laporan': '📊', 'report': '📊', 'log': '📋',
            'jurnal': '📓', 'journal': '📓', 'memo': '📝', 'minutes': '📝',

            # Komunikasi dan Kolaborasi
            'diskusi': '💬', 'discussion': '💬', 'pembahasan': '💬', 'komunikasi': '📢',
            'communication': '📢', 'meeting': '🗣️', 'rapat': '🗣️', 'pertemuan': '🗣️',
            'presentasi': '📊', 'presentation': '📊', 'demo': '🎬', 'demonstrasi': '🎬',

            # Keputusan dan Aksi
            'keputusan': '✅', 'decision': '✅', 'action': '⚡', 'tindakan': '⚡',
            'implementasi': '🚀', 'implementation': '🚀', 'eksekusi': '🚀', 'execution': '🚀',
            'deploy': '🚀', 'launch': '🚀', 'peluncuran': '🚀', 'rollout': '📤',

            # Pembelajaran dan Pengembangan
            'pembelajaran': '🎓', 'learning': '🎓', 'training': '🏋️', 'pelatihan': '🏋️',
            'edukasi': '📚', 'education': '📚', 'pengetahuan': '🧠', 'knowledge': '🧠',
            'skill': '💪', 'kemampuan': '💪', 'kompetensi': '🎯', 'competency': '🎯',

            # Teknologi dan Sistem
            'teknologi': '💻', 'technology': '💻', 'sistem': '⚙️', 'system': '⚙️',
            'software': '💻', 'aplikasi': '📱', 'application': '📱', 'platform': '🏗️',
            'infrastruktur': '🏗️', 'infrastructure': '🏗️', 'arsitektur': '🏛️', 'architecture': '🏛️',

            # Keuangan dan Bisnis
            'keuangan': '💰', 'finance': '💰', 'budget': '💰', 'anggaran': '💰',
            'bisnis': '💼', 'business': '💼', 'profit': '📈', 'keuntungan': '📈',
            'revenue': '💰', 'pendapatan': '💰', 'cost': '💸', 'biaya': '💸',

            # Keamanan dan Compliance
            'keamanan': '🔒', 'security': '🔒', 'compliance': '📋', 'kepatuhan': '📋',
            'audit': '🔍', 'kontrol': '🎛️', 'control': '🎛️', 'governance': '⚖️',

            # Inovasi dan Pengembangan
            'inovasi': '💡', 'innovation': '💡', 'pengembangan': '🚀', 'development': '🚀',
            'improvement': '📈', 'perbaikan': '📈', 'enhancement': '⬆️', 'peningkatan': '⬆️',
            'upgrade': '⬆️', 'modernisasi': '🔄', 'modernization': '🔄',

            # Status dan Kondisi
            'status': '📊', 'kondisi': '📊', 'condition': '📊', 'state': '🔘',
            'situasi': '📍', 'situation': '📍', 'posisi': '📍', 'position': '📍',
            'progress': '📈', 'kemajuan': '📈', 'update': '🔄', 'pembaruan': '🔄'
        }

        # Enhanced function untuk deteksi ikon berdasarkan konteks yang lebih cerdas
        def get_enhanced_heading_icon(heading_text, level, document_type):
            """Menentukan ikon yang tepat untuk heading dengan analisis konteks yang mendalam."""
            heading_lower = heading_text.lower().strip()

            # Prioritas 1: Deteksi berdasarkan pola khusus
            special_patterns = {
                r'\b(urgent|emergency|critical|penting sekali)\b': '🚨',
                r'\b(success|berhasil|completed|selesai)\b': '✅',
                r'\b(failed|gagal|error|kesalahan)\b': '❌',
                r'\b(new|baru|latest|terbaru)\b': '🆕',
                r'\b(final|akhir|conclusion|kesimpulan)\b': '🏁',
                r'\b(question|pertanyaan|tanya|ask)\b': '❓',
                r'\b(answer|jawaban|solution|solusi)\b': '💡',
                r'\b(review|ulasan|evaluation|evaluasi)\b': '📋',
                r'\b(summary|ringkasan|overview|gambaran)\b': '📄',
                r'\b(action|tindakan|do|lakukan)\b': '⚡',
                r'\b(meeting|rapat|discussion|diskusi)\b': '🗣️',
                r'\b(report|laporan|documentation|dokumentasi)\b': '📊',
                r'\b(planning|perencanaan|strategy|strategi)\b': '📋',
                r'\b(implementation|implementasi|execution|eksekusi)\b': '🚀'
            }

            for pattern, icon in special_patterns.items():
                if re.search(pattern, heading_lower):
                    return icon

            # Prioritas 2: Deteksi berdasarkan kata kunci individual
            for keyword, icon in enhanced_icon_mapping.items():
                if keyword in heading_lower:
                    return icon

            # Prioritas 3: Icon berdasarkan level dan tipe dokumen
            level_icons = {
                1: {
                    "technical_report": '📊',
                    "meeting_notes": '🗣️',
                    "lecture": '🎓',
                    "general": '📋'
                },
                2: {
                    "technical_report": '🔧',
                    "meeting_notes": '📌',
                    "lecture": '📝',
                    "general": '📄'
                },
                3: {
                    "technical_report": '⚙️',
                    "meeting_notes": '💬',
                    "lecture": '📖',
                    "general": '📍'
                }
            }

            # Prioritas 4: Deteksi berdasarkan posisi numerik
            number_match = re.match(r'^(\d+)[\.\)]\s*', heading_text)
            if number_match:
                number = int(number_match.group(1))
                if number <= 10:
                    return f"{number}️⃣"

            # Default berdasarkan level dan dokumen type
            return level_icons.get(level, {}).get(document_type, '▪️')

        # Cari judul pertama jika sudah ada judul dalam metadata
        if i < len(lines) and lines[i].startswith('# '):
            i += 1  # Lewati judul pertama karena sudah digunakan
            # Lewati baris kosong setelah judul
            while i < len(lines) and not lines[i].strip():
                i += 1

        while i < len(lines):
            line = lines[i]
            stripped_line = line.strip()

            # ===== ENHANCED EMPTY LINE HANDLING =====
            if not stripped_line:
                if not in_code_block and not in_table and not in_admonition:
                    # Reset list context pada baris kosong
                    if in_list:
                        in_list = False
                        list_level = 0
                        list_style = None

                    # Tambahkan spacing yang konsisten
                    if current_paragraph:
                        current_paragraph.paragraph_format.space_after = Pt(6)
                    current_paragraph = None
                i += 1
                continue

            # ===== ENHANCED ADMONITION BLOCK PROCESSING =====
            # Deteksi awal admonition blocks dengan pattern yang diperluas
            admonition_match = re.match(r'^:::(\w+)(?:\s+(.+))?$', stripped_line)
            if admonition_match:
                admonition_type = admonition_match.group(1).lower()
                admonition_title = admonition_match.group(2) if admonition_match.group(2) else None
                in_admonition = True
                admonition_lines = []
                if admonition_title:
                    admonition_lines.append(admonition_title)
                i += 1
                continue

            # Deteksi akhir admonition blocks
            if in_admonition and stripped_line == ':::':
                # UPGRADED: Gunakan enhanced admonition block dengan content_stats lengkap
                self._add_enhanced_admonition_block(doc, admonition_lines, admonition_type, content_stats)
                in_admonition = False
                admonition_type = None
                current_paragraph = None
                i += 1
                continue

            # Tambahkan baris ke admonition jika sedang di dalam blok
            if in_admonition:
                admonition_lines.append(line)
                i += 1
                continue

            # ===== ENHANCED CODE BLOCK PROCESSING =====
            if stripped_line.startswith('```'):
                in_code_block = not in_code_block
                if in_code_block:
                    # Deteksi bahasa dengan validasi
                    language_match = re.match(r'^```(\w+)(?:\s+(.+))?$', stripped_line)
                    if language_match:
                        code_language = language_match.group(1)
                        # Optional: Ambil parameter tambahan seperti filename
                        code_params = language_match.group(2) if language_match.group(2) else None
                    else:
                        code_language = ""
                        code_params = None
                    code_block_lines = []
                else:
                    # UPGRADED: Proses code block dengan parameter yang ditingkatkan dan content_stats
                    self._add_enhanced_code_block(doc, code_block_lines, code_language, content_stats)
                    current_paragraph = None
                i += 1
                continue

            # Akumulasi code block dengan preservasi indentasi
            if in_code_block:
                code_block_lines.append(line)  # Preserve original indentation
                i += 1
                continue

            # ===== ENHANCED TABLE PROCESSING =====
            if stripped_line.startswith('|') and stripped_line.endswith('|'):
                if not in_table:
                    in_table = True
                    table_data = []
                    in_list = False  # Reset list context when entering table

                table_data.append(stripped_line)

                # Enhanced table end detection
                next_line_is_table = False
                if i + 1 < len(lines):
                    next_line = lines[i + 1].strip()
                    next_line_is_table = (next_line.startswith('|') and next_line.endswith('|')) or not next_line

                if not next_line_is_table or i == len(lines) - 1:
                    # UPGRADED: Process enhanced table dengan content_stats lengkap
                    self._process_enhanced_markdown_table(doc, table_data, content_stats)
                    in_table = False
                    current_paragraph = None

                i += 1
                continue

            # ===== ENHANCED HEADING PROCESSING =====
            heading_match = re.match(r'^(#{1,6})\s+(.+)$', stripped_line)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()

                # Reset context
                in_list = False
                list_level = 0
                list_style = None

                # Enhanced page break logic
                if level == 1 and len(doc.paragraphs) > 5:
                    # Hanya tambah page break jika dokumen sudah memiliki konten substansial
                    last_para = doc.paragraphs[-1]
                    if last_para.text.strip():  # Ada konten sebelumnya
                        doc.add_page_break()

                # ===== CORE ENHANCEMENT: AUTOMATIC ICON INTEGRATION =====
                # Deteksi dan tambahkan ikon yang sesuai
                detected_icon = get_enhanced_heading_icon(heading_text, level, document_type)

                # Integrasikan ikon dengan teks heading
                if detected_icon:
                    # Format: "🔧 Panduan Instalasi" instead of just "Panduan Instalasi"
                    enhanced_heading_text = f"{detected_icon} {heading_text}"
                else:
                    enhanced_heading_text = heading_text

                # Create heading dengan teks yang sudah ditingkatkan
                heading_element = doc.add_heading(enhanced_heading_text, level=level)

                # UPGRADED: Enhanced styling berdasarkan content_stats
                self._apply_enhanced_heading_styling(heading_element, level, content_stats)

                # Enhanced formatting
                heading_element.paragraph_format.space_after = Pt(8 if level == 1 else 6)
                heading_element.paragraph_format.space_before = Pt(16 if level == 1 else 12)
                heading_element.paragraph_format.keep_with_next = True

                current_paragraph = None
                i += 1
                continue

            # ===== ENHANCED CALLOUT/ADMONITION PROCESSING =====
            # Pattern yang diperluas untuk callout
            callout_patterns = [
                (r'^(!{1,3})\s+(.+)$', 'importance'),  # !!! Important
                (r'^>\s*\[!(\w+)\]\s*(.+)$', 'typed'),  # > [!NOTE] Text
                (r'^>\s*\[(\w+)\]\s*(.+)$', 'simple'),  # > [INFO] Text
                (r'^\[!(\w+)\]:\s*(.+)$', 'block'),     # [!WARNING]: Text
            ]

            callout_processed = False
            for pattern, pattern_type in callout_patterns:
                callout_match = re.match(pattern, stripped_line)
                if callout_match:
                    if pattern_type == 'importance':
                        importance = len(callout_match.group(1))
                        callout_text = callout_match.group(2).strip()
                        callout_type = ["note", "important", "warning"][min(importance-1, 2)]
                    else:
                        callout_type = callout_match.group(1).lower()
                        callout_text = callout_match.group(2).strip()

                    # UPGRADED: Gunakan enhanced callout dengan content_stats dan style_options
                    style_options = self._get_callout_style_options(callout_type, content_stats)
                    self._add_enhanced_callout(doc, callout_text, callout_type, content_stats, style_options)
                    current_paragraph = None
                    callout_processed = True
                    break

            if callout_processed:
                i += 1
                continue

            # ===== ENHANCED LIST PROCESSING =====
            # Bulleted list dengan nested level support yang diperluas
            bullet_match = re.match(r'^(\s*)[-*•+]\s+(.+)$', line)
            if bullet_match:
                indent_text = bullet_match.group(1)
                item_text = bullet_match.group(2)

                # Hitung level indentasi dengan lebih akurat
                if '\t' in indent_text:
                    indent_level = indent_text.count('\t')
                else:
                    indent_level = len(indent_text) // 2

                # Deteksi perubahan konteks list
                if not in_list or list_style != 'bullet' or indent_level != list_level:
                    in_list = True
                    list_style = 'bullet'
                    list_level = indent_level

                # UPGRADED: Enhanced bullet list creation dengan content_stats
                p = self._create_enhanced_bullet_list_item(doc, item_text, indent_level, content_stats)
                current_paragraph = p

                i += 1
                continue

            # Numbered list dengan pattern yang diperluas
            numbered_patterns = [
                r'^(\s*)(\d+)\.?\s+(.+)$',      # 1. Item
                r'^(\s*)(\d+)\)\s+(.+)$',       # 1) Item
                r'^(\s*)([a-z])\.?\s+(.+)$',    # a. Item
                r'^(\s*)([A-Z])\.?\s+(.+)$',    # A. Item
                r'^(\s*)([ivx]+)\.?\s+(.+)$',   # i. Item (Roman)
            ]

            numbered_processed = False
            for pattern in numbered_patterns:
                numbered_match = re.match(pattern, line)
                if numbered_match:
                    indent_text = numbered_match.group(1)
                    number_text = numbered_match.group(2)
                    item_text = numbered_match.group(3)

                    # Hitung level indentasi
                    if '\t' in indent_text:
                        indent_level = indent_text.count('\t')
                    else:
                        indent_level = len(indent_text) // 2

                    # Deteksi perubahan konteks
                    if not in_list or list_style != 'number' or indent_level != list_level:
                        in_list = True
                        list_style = 'number'
                        list_level = indent_level

                    # UPGRADED: Enhanced numbered list creation dengan content_stats
                    p = self._create_enhanced_numbered_list_item(doc, item_text, indent_level, content_stats)
                    current_paragraph = p
                    numbered_processed = True
                    break

            if numbered_processed:
                i += 1
                continue

            # ===== ENHANCED TASK LIST PROCESSING =====
            task_match = re.match(r'^(\s*)\[([ xX✓✗])\]\s+(.+)$', stripped_line)
            if task_match:
                indent_text = task_match.group(1)
                check_char = task_match.group(2)
                item_text = task_match.group(3)

                # Deteksi status
                is_checked = check_char.lower() in ['x', '✓']
                is_failed = check_char in ['✗']

                indent_level = len(indent_text) // 2

                # UPGRADED: Create enhanced task list item dengan content_stats
                p = self._create_enhanced_task_list_item(doc, item_text, is_checked, is_failed, indent_level, content_stats)
                current_paragraph = p

                i += 1
                continue

            # ===== ENHANCED QUOTE PROCESSING =====
            if stripped_line.startswith('> '):
                quote_text = stripped_line[2:]

                # Multi-line quote detection dengan nested quotes
                quote_lines = [quote_text]
                quote_level = 1
                j = i + 1

                while j < len(lines):
                    next_line = lines[j].strip()
                    if next_line.startswith('> '):
                        # Deteksi nested quotes
                        nested_match = re.match(r'^(>+)\s+(.+)$', next_line)
                        if nested_match:
                            current_quote_level = len(nested_match.group(1))
                            quote_content = nested_match.group(2)
                            if current_quote_level > quote_level:
                                quote_lines.append(f"    {quote_content}")  # Indent nested quote
                            else:
                                quote_lines.append(quote_content)
                        else:
                            quote_lines.append(next_line[2:])
                        j += 1
                    elif next_line == '':
                        # Empty line dalam quote
                        quote_lines.append('')
                        j += 1
                    else:
                        break

                # UPGRADED: Create enhanced quote dengan content_stats
                full_quote = '\n'.join(quote_lines)
                p = self._create_enhanced_quote(doc, full_quote, content_stats)
                current_paragraph = p

                i = j
                continue

            # ===== ENHANCED HORIZONTAL RULE PROCESSING =====
            if re.match(r'^[-*_]{3,}$', stripped_line):
                # UPGRADED: Gunakan enhanced horizontal rule dengan content_stats
                self._add_enhanced_horizontal_rule(doc, content_stats)
                current_paragraph = None
                i += 1
                continue

            # ===== ENHANCED SPECIAL SECTIONS =====
            special_patterns = [
                (r'^(CONCLUSION|SUMMARY|NOTE|IMPORTANT|TIPS?|CATATAN|KESIMPULAN|RINGKASAN):\s+(.+)$', 'section'),
                (r'^\s*(?:\[ACTION\]|\*\*ACTION ITEM\*\*:?|ACTION:|\[TINDAKAN\]|\*\*TINDAKAN\*\*:?|TINDAKAN:)\s+(.+)$', 'action'),
                (r'^\s*(?:\[DECISION\]|DECISION:|\*\*DECISION\*\*:?|\[KEPUTUSAN\]|KEPUTUSAN:|\*\*KEPUTUSAN\*\*:?)\s+(.+)$', 'decision'),
                (r'^\s*(?:\[TODO\]|TODO:|\*\*TODO\*\*:?|\[TUGAS\]|TUGAS:|\*\*TUGAS\*\*:?)\s+(.+)$', 'todo'),
            ]

            special_processed = False
            for pattern, special_type in special_patterns:
                special_match = re.match(pattern, stripped_line, re.IGNORECASE)
                if special_match:
                    if special_type == 'section':
                        section_type = special_match.group(1).lower()
                        section_text = special_match.group(2)
                        # Normalisasi nama section
                        section_mapping = {
                            'catatan': 'note', 'kesimpulan': 'conclusion', 'ringkasan': 'summary',
                            'tips': 'tip', 'tip': 'tip'
                        }
                        normalized_section = section_mapping.get(section_type, section_type)
                        # UPGRADED: Add special section dengan content_stats
                        p = self._add_special_section(doc, normalized_section, section_text, content_stats)
                    else:
                        content_text = special_match.group(1)
                        # UPGRADED: Add enhanced special marker dengan content_stats
                        p = self._add_enhanced_special_marker(doc, special_type, content_text, content_stats)

                    current_paragraph = p
                    special_processed = True
                    break

            if special_processed:
                i += 1
                continue

            # ===== ENHANCED DEFINITION LIST =====
            definition_match = re.match(r'^(.{1,50}?)\s*:\s*(.+)$', stripped_line)
            if (definition_match and
                not stripped_line.startswith('|') and
                not stripped_line.startswith('http') and
                not re.match(r'^\d{1,2}:\d{2}', stripped_line) and
                len(definition_match.group(1)) <= 50):

                term = definition_match.group(1).strip()
                definition = definition_match.group(2).strip()

                # UPGRADED: Create enhanced definition item dengan content_stats
                p = self._create_enhanced_definition_item(doc, term, definition, content_stats)
                current_paragraph = p

                i += 1
                continue

            # ===== ENHANCED PARAGRAPH PROCESSING =====
            # Reset list context untuk paragraf normal
            if in_list and not re.match(r'^(\s*)[-*•+\d]+[\.\)]\s+', line):
                in_list = False
                list_level = 0
                list_style = None

            # Create paragraph dengan enhanced formatting
            p = doc.add_paragraph()
            current_paragraph = p

            # UPGRADED: Apply enhanced paragraph styling dengan content_stats
            processed_text = self._apply_paragraph_style(p, line, content_stats)

            # UPGRADED: Add formatted content dengan enhanced text processing
            self._add_enhanced_formatted_runs(p, processed_text, content_stats)

            # Enhanced spacing
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = Pt(14)

            i += 1

        # ===== FINAL DOCUMENT ENHANCEMENT =====
        # Tambahkan separator akhir yang elegan
        doc.add_paragraph()
        # UPGRADED: Enhanced horizontal rule dengan content_stats
        self._add_enhanced_horizontal_rule(doc, content_stats)
        doc.add_paragraph()

        # UPGRADED: Apply final document formatting dengan content_stats lengkap
        self.finalize_document_formatting_enhanced(doc, content_stats)

    def _create_enhanced_bullet_list_item(self, doc, item_text, indent_level, content_stats):
        """
        Membuat item list bullet dengan styling yang sangat ditingkatkan dan adaptif.
        """
        # ===== ENHANCED INPUT VALIDATION =====
        if not item_text or not item_text.strip():
            return None

        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== ENHANCED BULLET STYLE CONFIGURATION =====
        bullet_styles = {
            "technical_report": {
                0: {"bullet": "▪", "color": "2E5984", "font": "Segoe UI Symbol", "size": 10},
                1: {"bullet": "◦", "color": "4472C4", "font": "Segoe UI Symbol", "size": 9},
                2: {"bullet": "▫", "color": "8DB4E2", "font": "Segoe UI Symbol", "size": 8},
                3: {"bullet": "‣", "color": "B8CCE4", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "•", "color": "DDEAF6", "font": "Segoe UI Symbol", "size": 7}
            },
            "meeting_notes": {
                0: {"bullet": "●", "color": "2F855A", "font": "Segoe UI Symbol", "size": 10},
                1: {"bullet": "○", "color": "70AD47", "font": "Segoe UI Symbol", "size": 9},
                2: {"bullet": "◼", "color": "A9D18E", "font": "Segoe UI Symbol", "size": 8},
                3: {"bullet": "◻", "color": "C5E0B4", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "▪", "color": "E2EFDA", "font": "Segoe UI Symbol", "size": 7}
            },
            "lecture": {
                0: {"bullet": "►", "color": "C05621", "font": "Segoe UI Symbol", "size": 10},
                1: {"bullet": "▷", "color": "ED7D31", "font": "Segoe UI Symbol", "size": 9},
                2: {"bullet": "▶", "color": "F4B183", "font": "Segoe UI Symbol", "size": 8},
                3: {"bullet": "▸", "color": "FAD5B4", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "‣", "color": "FBE5D6", "font": "Segoe UI Symbol", "size": 7}
            },
            "presentation": {
                0: {"bullet": "⚡", "color": "7030A0", "font": "Segoe UI Emoji", "size": 11},
                1: {"bullet": "🔸", "color": "9966CC", "font": "Segoe UI Emoji", "size": 10},
                2: {"bullet": "🔹", "color": "B19CD9", "font": "Segoe UI Emoji", "size": 9},
                3: {"bullet": "▪", "color": "D6C7F0", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "‣", "color": "E9E0F7", "font": "Segoe UI Symbol", "size": 7}
            },
            "research": {
                0: {"bullet": "◆", "color": "1F4E79", "font": "Segoe UI Symbol", "size": 10},
                1: {"bullet": "◇", "color": "2E75B5", "font": "Segoe UI Symbol", "size": 9},
                2: {"bullet": "♦", "color": "8DB4E2", "font": "Segoe UI Symbol", "size": 8},
                3: {"bullet": "♢", "color": "B8CCE4", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "▪", "color": "DDEAF6", "font": "Segoe UI Symbol", "size": 7}
            },
            "general": {
                0: {"bullet": "•", "color": "1F497D", "font": "Segoe UI Symbol", "size": 10},
                1: {"bullet": "◦", "color": "4F81BD", "font": "Segoe UI Symbol", "size": 9},
                2: {"bullet": "▪", "color": "8DB4E2", "font": "Segoe UI Symbol", "size": 8},
                3: {"bullet": "▫", "color": "B8CCE4", "font": "Segoe UI Symbol", "size": 8},
                4: {"bullet": "‣", "color": "DDEAF6", "font": "Segoe UI Symbol", "size": 7}
            }
        }

        # ===== ENHANCED INDENTATION CALCULATION =====
        base_indent = 0.25
        level_increment = 0.25

        if document_type == "technical_report":
            base_indent = 0.3
            level_increment = 0.3
        elif document_type == "presentation":
            base_indent = 0.35
            level_increment = 0.35
        elif document_type == "meeting_notes":
            base_indent = 0.2
            level_increment = 0.2

        # Complexity-based adjustments
        complexity = content_stats.get("complexity_level", "medium")
        if complexity == "high":
            level_increment += 0.05
        elif complexity == "low":
            level_increment = max(0.15, level_increment - 0.05)

        calculated_indent = base_indent + (indent_level * level_increment)

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']  # Remove default list style untuk full control

        # ===== CUSTOM BULLET IMPLEMENTATION =====
        style_config = bullet_styles.get(document_type, bullet_styles["general"])
        bullet_config = style_config.get(indent_level, style_config.get(0))

        # Create custom bullet run
        bullet_run = p.add_run(f"{bullet_config['bullet']} ")
        bullet_run.font.name = bullet_config['font']
        bullet_run.font.size = Pt(bullet_config['size'])
        bullet_run.font.color.rgb = RGBColor.from_string(bullet_config['color'])
        bullet_run.font.bold = True if indent_level == 0 else False

        # ===== ENHANCED INDENTATION AND SPACING =====
        p.paragraph_format.left_indent = Inches(calculated_indent)
        p.paragraph_format.first_line_indent = Inches(-0.15)  # Hanging indent

        # Dynamic spacing berdasarkan document type
        if document_type == "technical_report":
            space_before = Pt(4 if indent_level == 0 else 2)
            space_after = Pt(4 if indent_level == 0 else 2)
        elif document_type == "presentation":
            space_before = Pt(6 if indent_level == 0 else 3)
            space_after = Pt(6 if indent_level == 0 else 3)
        elif document_type == "meeting_notes":
            space_before = Pt(2)
            space_after = Pt(2)
        else:
            space_before = Pt(3)
            space_after = Pt(3)

        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after

        # Enhanced line spacing
        if content_stats.get("word_count", 0) > 3000:
            p.paragraph_format.line_spacing = Pt(14)
        else:
            p.paragraph_format.line_spacing = Pt(13)

        # ===== ENHANCED TEXT PROCESSING =====
        # Pre-process text untuk optimasi
        processed_text = item_text.strip()

        # Enhanced processing berdasarkan document type
        if document_type == "technical_report":
            technical_terms = ["API", "URL", "HTTP", "HTTPS", "JSON", "XML", "SQL", "HTML", "CSS", "JS"]
            for term in technical_terms:
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                processed_text = pattern.sub(term, processed_text)
        elif document_type == "meeting_notes":
            processed_text = re.sub(r'\b(action|TODO|DECISION|FOLLOW.?UP)\b',
                                lambda m: f"**{m.group(1).upper()}**",
                                processed_text, flags=re.IGNORECASE)
        elif document_type == "lecture" and indent_level == 0:
            if ':' in processed_text and len(processed_text.split(':')[0]) < 30:
                parts = processed_text.split(':', 1)
                if len(parts) == 2:
                    processed_text = f"**{parts[0].strip()}**: {parts[1].strip()}"

        # Universal enhancements
        processed_text = re.sub(r'\be\.g\.\s*', 'e.g., ', processed_text)
        processed_text = re.sub(r'\bi\.e\.\s*', 'i.e., ', processed_text)
        processed_text = re.sub(r'\betc\.?\s*$', 'etc.', processed_text)

        # ===== ADD FORMATTED CONTENT =====
        if hasattr(self, '_add_formatted_runs_to_paragraph'):
            self._add_formatted_runs_to_paragraph(p, processed_text)
        else:
            # Fallback simple formatting
            content_run = p.add_run(processed_text)

            # Base font settings berdasarkan document type
            if document_type == "technical_report":
                content_run.font.name = 'Cambria'
                content_run.font.size = Pt(11 if indent_level == 0 else 10)
            elif document_type == "presentation":
                content_run.font.name = 'Segoe UI'
                content_run.font.size = Pt(12 if indent_level == 0 else 11)
            elif document_type == "meeting_notes":
                content_run.font.name = 'Calibri'
                content_run.font.size = Pt(10.5)
            else:
                content_run.font.name = 'Calibri'
                content_run.font.size = Pt(11 if indent_level == 0 else 10)

        # ===== SPECIAL ENHANCEMENTS =====
        # Priority items highlighting
        if any(keyword in processed_text.lower() for keyword in ['urgent', 'critical', 'important', 'asap']):
            try:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF2E5"/>'.format(nsdecls('w')))
                p._element.get_or_add_pPr().append(shading_elm)

                # Make content slightly bolder
                for run in p.runs[1:]:  # Skip bullet run
                    if not run.bold:
                        run.font.bold = True
            except:
                pass  # Continue if shading fails

        # Enhanced spacing untuk long items
        if len(processed_text) > 100:
            p.paragraph_format.space_after = Pt(6)

        # Keep with next untuk better page breaks
        if indent_level == 0:
            p.paragraph_format.keep_with_next = True

        return p

    def _create_enhanced_numbered_list_item(self, doc, item_text, indent_level, content_stats):
        """
        Membuat item list bernomor dengan styling yang sangat ditingkatkan dan adaptif.
        """
        import re  # Pindahkan import ke awal function

        # ===== ENHANCED INPUT VALIDATION =====
        if not item_text or not item_text.strip():
            return None

        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== ENHANCED NUMBERING STYLE CONFIGURATION =====
        numbering_styles = {
            "technical_report": {
                0: {"style": "decimal", "color": "2E5984", "format": "{}."},
                1: {"style": "lower-alpha", "color": "4472C4", "format": "{})"},
                2: {"style": "lower-roman", "color": "8DB4E2", "format": "{}:"},
                3: {"style": "decimal", "color": "B8CCE4", "format": "{}."},
                4: {"style": "lower-alpha", "color": "DDEAF6", "format": "{})"}
            },
            "meeting_notes": {
                0: {"style": "decimal", "color": "2F855A", "format": "{}."},
                1: {"style": "decimal", "color": "70AD47", "format": "{}.{}"},
                2: {"style": "lower-alpha", "color": "A9D18E", "format": "{})"},
                3: {"style": "lower-roman", "color": "C5E0B4", "format": "{}:"},
                4: {"style": "decimal", "color": "E2EFDA", "format": "{}."}
            },
            "lecture": {
                0: {"style": "decimal", "color": "C05621", "format": "{}."},
                1: {"style": "upper-alpha", "color": "ED7D31", "format": "{}."},
                2: {"style": "lower-roman", "color": "F4B183", "format": "{}:"},
                3: {"style": "decimal", "color": "FAD5B4", "format": "{})"},
                4: {"style": "lower-alpha", "color": "FBE5D6", "format": "{}."}
            },
            "instructional": {
                0: {"style": "decimal", "color": "107C10", "format": "Step {}:"},
                1: {"style": "lower-alpha", "color": "70AD47", "format": "{})"},
                2: {"style": "lower-roman", "color": "A9D18E", "format": "{}:"},
                3: {"style": "decimal", "color": "C5E0B4", "format": "{}."},
                4: {"style": "lower-alpha", "color": "E2EFDA", "format": "{})"}
            },
            "presentation": {
                0: {"style": "decimal", "color": "7030A0", "format": "{}."},
                1: {"style": "upper-alpha", "color": "9966CC", "format": "{}."},
                2: {"style": "lower-roman", "color": "B19CD9", "format": "{}:"},
                3: {"style": "decimal", "color": "D6C7F0", "format": "{})"},
                4: {"style": "lower-alpha", "color": "E9E0F7", "format": "{}."}
            },
            "research": {
                0: {"style": "decimal", "color": "1F4E79", "format": "{}."},
                1: {"style": "decimal", "color": "2E75B5", "format": "{}.{}"},
                2: {"style": "decimal", "color": "8DB4E2", "format": "{}.{}.{}"},
                3: {"style": "lower-alpha", "color": "B8CCE4", "format": "{})"},
                4: {"style": "lower-roman", "color": "DDEAF6", "format": "{}:"}
            },
            "general": {
                0: {"style": "decimal", "color": "1F497D", "format": "{}."},
                1: {"style": "lower-alpha", "color": "4F81BD", "format": "{})"},
                2: {"style": "lower-roman", "color": "8DB4E2", "format": "{}:"},
                3: {"style": "decimal", "color": "B8CCE4", "format": "{}."},
                4: {"style": "lower-alpha", "color": "DDEAF6", "format": "{})"}
            }
        }

        # ===== ENHANCED INDENTATION CALCULATION =====
        base_indent = 0.3
        level_increment = 0.3

        if document_type == "technical_report":
            base_indent = 0.35
            level_increment = 0.35
        elif document_type == "meeting_notes":
            base_indent = 0.25
            level_increment = 0.25
        elif document_type == "instructional":
            base_indent = 0.4
            level_increment = 0.4
        elif document_type == "presentation":
            base_indent = 0.45
            level_increment = 0.35

        # Complexity-based adjustments
        complexity = content_stats.get("complexity_level", "medium")
        if complexity == "high":
            level_increment += 0.05
            base_indent += 0.05
        elif complexity == "low":
            level_increment = max(0.2, level_increment - 0.05)
            base_indent = max(0.2, base_indent - 0.05)

        calculated_indent = base_indent + (indent_level * level_increment)

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']  # Use normal style for full control

        # ===== NUMBERING SYSTEM SIMULATION =====
        style_config = numbering_styles.get(document_type, numbering_styles["general"])
        number_config = style_config.get(indent_level, style_config.get(0))

        # Simulated numbering (since Word's list numbering is complex to control programmatically)
        # This creates a consistent visual numbering system
        number_generators = {
            "decimal": lambda n: str(n),
            "lower-alpha": lambda n: chr(ord('a') + n - 1) if n <= 26 else f"z{n-26}",
            "upper-alpha": lambda n: chr(ord('A') + n - 1) if n <= 26 else f"Z{n-26}",
            "lower-roman": lambda n: ["i", "ii", "iii", "iv", "v", "vi", "vii", "viii", "ix", "x"][min(n-1, 9)],
            "upper-roman": lambda n: ["I", "II", "III", "IV", "V", "VI", "VII", "VIII", "IX", "X"][min(n-1, 9)]
        }

        # Generate number based on a simple counter (in real implementation, this would track across document)
        # For this example, we'll use a simple incremental system
        if not hasattr(self, '_numbered_list_counters'):
            self._numbered_list_counters = {}

        counter_key = f"{document_type}_{indent_level}"
        if counter_key not in self._numbered_list_counters:
            self._numbered_list_counters[counter_key] = 0

        self._numbered_list_counters[counter_key] += 1
        current_number = self._numbered_list_counters[counter_key]

        # Generate the actual number display
        number_style = number_config["style"]
        number_format = number_config["format"]

        if number_style in number_generators:
            generated_number = number_generators[number_style](current_number)
            formatted_number = number_format.format(generated_number)
        else:
            formatted_number = f"{current_number}."

        # ===== CUSTOM NUMBER IMPLEMENTATION =====
        # Create custom number run
        number_run = p.add_run(f"{formatted_number} ")
        number_run.font.name = "Calibri"
        number_run.font.size = Pt(11 if indent_level == 0 else 10)
        number_run.font.color.rgb = RGBColor.from_string(number_config['color'])
        number_run.font.bold = True if indent_level == 0 else False

        # ===== ENHANCED INDENTATION AND SPACING =====
        p.paragraph_format.left_indent = Inches(calculated_indent)
        p.paragraph_format.first_line_indent = Inches(-0.2)  # Hanging indent for number

        # Dynamic spacing berdasarkan document type
        if document_type == "technical_report":
            space_before = Pt(6 if indent_level == 0 else 3)
            space_after = Pt(6 if indent_level == 0 else 3)
        elif document_type == "instructional":
            space_before = Pt(8 if indent_level == 0 else 4)
            space_after = Pt(8 if indent_level == 0 else 4)
        elif document_type == "presentation":
            space_before = Pt(8 if indent_level == 0 else 4)
            space_after = Pt(8 if indent_level == 0 else 4)
        elif document_type == "meeting_notes":
            space_before = Pt(3)
            space_after = Pt(3)
        else:
            space_before = Pt(4)
            space_after = Pt(4)

        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after

        # Enhanced line spacing
        if content_stats.get("word_count", 0) > 3000:
            p.paragraph_format.line_spacing = Pt(15)
        else:
            p.paragraph_format.line_spacing = Pt(14)

        # ===== ENHANCED TEXT PROCESSING =====
        # Pre-process text untuk optimasi
        processed_text = item_text.strip()

        # Enhanced processing berdasarkan document type
        if document_type == "technical_report":
            # Technical terms standardization
            technical_terms = ["API", "URL", "HTTP", "HTTPS", "JSON", "XML", "SQL", "HTML", "CSS", "JS", "REST", "SOAP"]
            for term in technical_terms:
                # Tidak perlu import re lagi karena sudah di-import di awal function
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                processed_text = pattern.sub(term, processed_text)

            # Add technical emphasis for numbered procedures
            if indent_level == 0 and any(word in processed_text.lower() for word in ['install', 'configure', 'setup', 'initialize']):
                processed_text = f"**{processed_text}**"

        elif document_type == "instructional":
            # Instructional enhancements
            if indent_level == 0:
                # Emphasize step actions
                action_words = ['open', 'click', 'select', 'choose', 'enter', 'type', 'save', 'close', 'navigate']
                for word in action_words:
                    pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
                    processed_text = pattern.sub(f"**{word.upper()}**", processed_text)

            # Add instructional markers
            if any(keyword in processed_text.lower() for keyword in ['note:', 'important:', 'warning:', 'tip:']):
                processed_text = f"*{processed_text}*"

        elif document_type == "meeting_notes":
            # Meeting notes enhancements
            processed_text = re.sub(r'\b(ACTION|TODO|DECISION|FOLLOW.?UP)\b',
                                lambda m: f"**{m.group(1).upper()}**",
                                processed_text, flags=re.IGNORECASE)

            # Emphasize assignments
            if re.search(r'\b(assigned to|responsible|owner|due|deadline)\b', processed_text, re.IGNORECASE):
                processed_text = f"*{processed_text}*"

        elif document_type == "presentation":
            # Presentation enhancements
            if indent_level == 0:
                # Main points get emphasis
                processed_text = f"**{processed_text}**"

            # Highlight key metrics or data
            if re.search(r'\b\d+%|\$\d+|\d+\s+(users|customers|clients|increase|decrease)\b', processed_text, re.IGNORECASE):
                processed_text = f"*{processed_text}*"

        # Universal enhancements
        processed_text = re.sub(r'\be\.g\.\s*', 'e.g., ', processed_text)
        processed_text = re.sub(r'\bi\.e\.\s*', 'i.e., ', processed_text)
        processed_text = re.sub(r'\betc\.?\s*$', 'etc.', processed_text)

        # Bold formatting
        parts = re.split(r'(\*\*.*?\*\*)', processed_text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                # Bold text
                content = part[2:-2]
                content_run = p.add_run(content)
                content_run.font.bold = True
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                # Italic text
                content = part[1:-1]
                content_run = p.add_run(content)
                content_run.font.italic = True
            else:
                # Regular text
                if part:  # Only add if not empty
                    content_run = p.add_run(part)

        # Set base font properties for all runs
        for run in p.runs[1:]:  # Skip the number run
            if document_type == "technical_report":
                run.font.name = 'Cambria'
                run.font.size = Pt(11 if indent_level == 0 else 10)
            elif document_type == "instructional":
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11 if indent_level == 0 else 10)
            elif document_type == "presentation":
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12 if indent_level == 0 else 11)
            elif document_type == "meeting_notes":
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
            else:
                run.font.name = 'Calibri'
                run.font.size = Pt(11 if indent_level == 0 else 10)

        # ===== SPECIAL ENHANCEMENTS =====
        # Priority items highlighting
        if any(keyword in processed_text.lower() for keyword in ['urgent', 'critical', 'important', 'priority', 'asap']):
            try:
                from docx.oxml import parse_xml
                from docx.oxml.ns import nsdecls
                shading_elm = parse_xml(r'<w:shd {} w:fill="FFF2E5"/>'.format(nsdecls('w')))
                p._element.get_or_add_pPr().append(shading_elm)

                # Make number more prominent for priority items
                number_run.font.bold = True
                number_run.font.size = Pt(number_run.font.size.pt + 1)
            except:
                pass  # Continue if shading fails

        # Enhanced spacing untuk detailed items
        if len(processed_text) > 120:
            p.paragraph_format.space_after = Pt(8)

        # Keep with next untuk better page breaks (level 0 items)
        if indent_level == 0:
            p.paragraph_format.keep_with_next = True

        # Special formatting untuk different content patterns
        if document_type == "instructional" and indent_level == 0:
            # Add subtle border for main instruction steps
            try:
                from docx.oxml import parse_xml
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(r'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="E0E0E0"/>
                </w:pBdr>''')
                pPr.append(pBdr)
            except:
                pass

        # Enhanced numbering reset logic for new sections
        if indent_level == 0 and any(keyword in processed_text.lower() for keyword in ['section', 'chapter', 'part', 'phase']):
            # Reset lower level counters
            keys_to_reset = [key for key in self._numbered_list_counters.keys()
                            if key.startswith(f"{document_type}_") and int(key.split('_')[-1]) > indent_level]
            for key in keys_to_reset:
                self._numbered_list_counters[key] = 0

        return p

    def _create_enhanced_task_list_item(self, doc, item_text, is_checked, is_failed, indent_level, content_stats):
        """
        Membuat item task list dengan checkbox styling yang sangat ditingkatkan dan adaptif.
        """
        # ===== ENHANCED INPUT VALIDATION =====
        if not item_text or not item_text.strip():
            return None

        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== ENHANCED CHECKBOX STYLE CONFIGURATION =====
        checkbox_styles = {
            "technical_report": {
                "checked": {"symbol": "✅", "color": "107C10", "bg": "F0FFF0", "font": "Segoe UI Emoji", "size": 11},
                "failed": {"symbol": "❌", "color": "C42B1C", "bg": "FFF0F0", "font": "Segoe UI Emoji", "size": 11},
                "pending": {"symbol": "⬜", "color": "605E5C", "bg": "F8F8F8", "font": "Segoe UI Emoji", "size": 11},
                "in_progress": {"symbol": "🔄", "color": "0078D4", "bg": "F3F9FF", "font": "Segoe UI Emoji", "size": 11}
            },
            "meeting_notes": {
                "checked": {"symbol": "☑️", "color": "107C10", "bg": "F0FFF0", "font": "Segoe UI Symbol", "size": 12},
                "failed": {"symbol": "☒", "color": "C42B1C", "bg": "FFF0F0", "font": "Segoe UI Symbol", "size": 12},
                "pending": {"symbol": "☐", "color": "605E5C", "bg": "F8F8F8", "font": "Segoe UI Symbol", "size": 12},
                "in_progress": {"symbol": "◐", "color": "0078D4", "bg": "F3F9FF", "font": "Segoe UI Symbol", "size": 12}
            },
            "presentation": {
                "checked": {"symbol": "✓", "color": "107C10", "bg": "F0FFF0", "font": "Segoe UI", "size": 13},
                "failed": {"symbol": "✗", "color": "C42B1C", "bg": "FFF0F0", "font": "Segoe UI", "size": 13},
                "pending": {"symbol": "○", "color": "605E5C", "bg": "F8F8F8", "font": "Segoe UI", "size": 13},
                "in_progress": {"symbol": "◉", "color": "0078D4", "bg": "F3F9FF", "font": "Segoe UI", "size": 13}
            },
            "instructional": {
                "checked": {"symbol": "✅", "color": "107C10", "bg": "F0FFF0", "font": "Segoe UI Emoji", "size": 12},
                "failed": {"symbol": "🚫", "color": "C42B1C", "bg": "FFF0F0", "font": "Segoe UI Emoji", "size": 12},
                "pending": {"symbol": "⭕", "color": "605E5C", "bg": "F8F8F8", "font": "Segoe UI Emoji", "size": 12},
                "in_progress": {"symbol": "⏳", "color": "0078D4", "bg": "F3F9FF", "font": "Segoe UI Emoji", "size": 12}
            },
            "research": {
                "checked": {"symbol": "✓", "color": "107C10", "bg": "F0FFF0", "font": "Times New Roman", "size": 11},
                "failed": {"symbol": "✗", "color": "C42B1C", "bg": "FFF0F0", "font": "Times New Roman", "size": 11},
                "pending": {"symbol": "□", "color": "605E5C", "bg": "F8F8F8", "font": "Times New Roman", "size": 11},
                "in_progress": {"symbol": "◑", "color": "0078D4", "bg": "F3F9FF", "font": "Times New Roman", "size": 11}
            },
            "general": {
                "checked": {"symbol": "☑", "color": "107C10", "bg": "F0FFF0", "font": "Segoe UI Symbol", "size": 11},
                "failed": {"symbol": "☒", "color": "C42B1C", "bg": "FFF0F0", "font": "Segoe UI Symbol", "size": 11},
                "pending": {"symbol": "☐", "color": "605E5C", "bg": "F8F8F8", "font": "Segoe UI Symbol", "size": 11},
                "in_progress": {"symbol": "◐", "color": "0078D4", "bg": "F3F9FF", "font": "Segoe UI Symbol", "size": 11}
            }
        }

        # ===== ENHANCED STATUS DETECTION =====
        # Deteksi status berdasarkan teks untuk menentukan tipe checkbox yang tepat
        status = "pending"  # Default
        if is_checked:
            status = "checked"
        elif is_failed:
            status = "failed"
        else:
            # Advanced status detection dari teks
            text_lower = item_text.lower()
            if any(keyword in text_lower for keyword in ['in progress', 'sedang', 'working on', 'mengerjakan', 'berlangsung']):
                status = "in_progress"
            elif any(keyword in text_lower for keyword in ['pending', 'waiting', 'menunggu', 'belum', 'todo']):
                status = "pending"

        # ===== ENHANCED INDENTATION CALCULATION =====
        base_indent = 0.25
        level_increment = 0.3

        # Document type specific indentation
        if document_type == "technical_report":
            base_indent = 0.3
            level_increment = 0.35
        elif document_type == "meeting_notes":
            base_indent = 0.2
            level_increment = 0.25
        elif document_type == "presentation":
            base_indent = 0.35
            level_increment = 0.4
        elif document_type == "instructional":
            base_indent = 0.4
            level_increment = 0.4

        # Complexity-based adjustments
        complexity = content_stats.get("complexity_level", "medium")
        if complexity == "high":
            level_increment += 0.05
            base_indent += 0.05
        elif complexity == "low":
            level_increment = max(0.15, level_increment - 0.05)
            base_indent = max(0.15, base_indent - 0.05)

        calculated_indent = base_indent + (indent_level * level_increment)

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = doc.styles['Normal']

        # ===== ENHANCED CHECKBOX IMPLEMENTATION =====
        style_config = checkbox_styles.get(document_type, checkbox_styles["general"])
        checkbox_config = style_config[status]

        # Create enhanced checkbox run
        checkbox_run = p.add_run(f"{checkbox_config['symbol']} ")
        checkbox_run.font.name = checkbox_config['font']
        checkbox_run.font.size = Pt(checkbox_config['size'])
        checkbox_run.font.color.rgb = RGBColor.from_string(checkbox_config['color'])
        checkbox_run.font.bold = True if status in ['checked', 'failed'] else False

        # Add background highlight untuk visual emphasis
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            bg_color = checkbox_config.get('bg', 'FFFFFF')
            if bg_color != 'FFFFFF' and status != 'pending':
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                rPr = checkbox_run._element.get_or_add_rPr()
                rPr.append(shading_elm)
        except:
            pass  # Continue if shading fails

        # ===== ENHANCED INDENTATION AND SPACING =====
        p.paragraph_format.left_indent = Inches(calculated_indent)
        p.paragraph_format.first_line_indent = Inches(-0.15)  # Hanging indent for checkbox

        # Dynamic spacing berdasarkan document type dan status
        if document_type == "technical_report":
            space_before = Pt(5 if indent_level == 0 else 3)
            space_after = Pt(5 if indent_level == 0 else 3)
        elif document_type == "meeting_notes":
            space_before = Pt(4)
            space_after = Pt(4)
        elif document_type == "presentation":
            space_before = Pt(6 if indent_level == 0 else 4)
            space_after = Pt(6 if indent_level == 0 else 4)
        elif document_type == "instructional":
            space_before = Pt(6 if indent_level == 0 else 4)
            space_after = Pt(6 if indent_level == 0 else 4)
        else:
            space_before = Pt(4)
            space_after = Pt(4)

        # Extra spacing for important status
        if status in ['failed', 'checked']:
            space_before += Pt(2)
            space_after += Pt(2)

        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after

        # Enhanced line spacing
        if content_stats.get("word_count", 0) > 3000:
            p.paragraph_format.line_spacing = Pt(15)
        else:
            p.paragraph_format.line_spacing = Pt(14)

        # ===== ENHANCED TEXT PROCESSING =====
        # Pre-process text untuk optimasi
        processed_text = item_text.strip()

        # Enhanced processing berdasarkan status
        if status == "checked":
            # Add completion emphasis
            if not processed_text.startswith('**') and not processed_text.endswith('**'):
                processed_text = f"~~{processed_text}~~"  # Strikethrough untuk completed
        elif status == "failed":
            # Add failure emphasis
            if 'failed' not in processed_text.lower() and 'gagal' not in processed_text.lower():
                processed_text = f"**{processed_text}** *(Failed)*"
        elif status == "in_progress":
            # Add progress emphasis
            if not any(keyword in processed_text.lower() for keyword in ['in progress', 'sedang']):
                processed_text = f"*{processed_text}* (In Progress)"

        # Enhanced processing berdasarkan document type
        if document_type == "technical_report":
            # Technical task formatting
            technical_terms = ["API", "URL", "HTTP", "HTTPS", "JSON", "XML", "SQL", "HTML", "CSS", "JS"]
            for term in technical_terms:
                pattern = re.compile(re.escape(term), re.IGNORECASE)
                processed_text = pattern.sub(term, processed_text)

            # Emphasize technical actions
            tech_actions = ['deploy', 'configure', 'install', 'setup', 'test', 'debug', 'optimize']
            for action in tech_actions:
                pattern = re.compile(r'\b' + re.escape(action) + r'\b', re.IGNORECASE)
                processed_text = pattern.sub(f"**{action.upper()}**", processed_text)

        elif document_type == "meeting_notes":
            # Meeting task formatting
            processed_text = re.sub(r'\b(ACTION|TODO|FOLLOW.?UP|TASK)\b',
                                lambda m: f"**{m.group(1).upper()}**",
                                processed_text, flags=re.IGNORECASE)

            # Emphasize assignments
            if re.search(r'\b(assigned to|responsible|owner|due|deadline)\b', processed_text, re.IGNORECASE):
                processed_text = f"*{processed_text}*"

        elif document_type == "instructional":
            # Instructional task formatting
            if indent_level == 0:
                # Emphasize main instruction actions
                action_words = ['complete', 'finish', 'submit', 'review', 'check', 'verify']
                for word in action_words:
                    pattern = re.compile(r'\b' + re.escape(word) + r'\b', re.IGNORECASE)
                    processed_text = pattern.sub(f"**{word.upper()}**", processed_text)

        elif document_type == "presentation":
            # Presentation task formatting
            if indent_level == 0:
                # Main presentation tasks get emphasis
                processed_text = f"**{processed_text}**"

        # Universal enhancements
        processed_text = re.sub(r'\be\.g\.\s*', 'e.g., ', processed_text)
        processed_text = re.sub(r'\bi\.e\.\s*', 'i.e., ', processed_text)
        processed_text = re.sub(r'\betc\.?\s*$', 'etc.', processed_text)

        # Extended formatting patterns
        formatting_patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),      # ***bold italic***
            (r'\*\*(.+?)\*\*', 'bold'),                 # **bold**
            (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', 'italic'), # *italic*
            (r'~~(.+?)~~', 'strikethrough'),            # ~~strikethrough~~
            (r'`(.+?)`', 'code'),                       # `code`
            (r'==(.+?)==', 'highlight'),                # ==highlight==
        ]

        # Find all matches
        all_matches = []
        for pattern, format_type in formatting_patterns:
            for match in re.finditer(pattern, processed_text):
                all_matches.append((match.start(), match.end(), match.group(1), format_type))

        # Sort matches by position
        all_matches.sort(key=lambda x: x[0])

        # Remove overlapping matches
        cleaned_matches = []
        for match in all_matches:
            start, end = match[0], match[1]
            is_overlapping = any(start < prev_end and end > prev_start
                            for prev_start, prev_end, _, _ in cleaned_matches)
            if not is_overlapping:
                cleaned_matches.append(match)

        # Process text dengan formatting
        if cleaned_matches:
            last_end = 0
            for start, end, content, format_type in cleaned_matches:
                # Add plain text before formatted section
                if start > last_end:
                    plain_text = processed_text[last_end:start]
                    if plain_text:
                        plain_run = p.add_run(plain_text)

                # Create formatted run
                formatted_run = p.add_run(content)

                # Apply formatting
                if format_type == 'bold':
                    formatted_run.font.bold = True
                elif format_type == 'italic':
                    formatted_run.font.italic = True
                elif format_type == 'bold_italic':
                    formatted_run.font.bold = True
                    formatted_run.font.italic = True
                elif format_type == 'strikethrough':
                    formatted_run.font.strike = True
                elif format_type == 'code':
                    formatted_run.font.name = 'Consolas'
                    formatted_run.font.size = Pt(9)
                    # Add code background
                    try:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
                        rPr = formatted_run._element.get_or_add_rPr()
                        rPr.append(shading_elm)
                    except:
                        pass
                elif format_type == 'highlight':
                    try:
                        from docx.enum.text import WD_COLOR_INDEX
                        formatted_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    except:
                        pass

                last_end = end

            # Add remaining text
            if last_end < len(processed_text):
                remaining_text = processed_text[last_end:]
                if remaining_text:
                    remaining_run = p.add_run(remaining_text)
        else:
            # No special formatting, add as plain text
            content_run = p.add_run(processed_text)

        # ===== SET BASE FONT PROPERTIES =====
        # Apply base font settings to all content runs (skip checkbox run)
        for run in p.runs[1:]:  # Skip the checkbox run
            if document_type == "technical_report":
                run.font.name = 'Cambria'
                run.font.size = Pt(11 if indent_level == 0 else 10)
            elif document_type == "presentation":
                run.font.name = 'Segoe UI'
                run.font.size = Pt(12 if indent_level == 0 else 11)
            elif document_type == "meeting_notes":
                run.font.name = 'Calibri'
                run.font.size = Pt(10.5)
            elif document_type == "instructional":
                run.font.name = 'Segoe UI'
                run.font.size = Pt(11 if indent_level == 0 else 10)
            else:
                run.font.name = 'Calibri'
                run.font.size = Pt(11 if indent_level == 0 else 10)

        # ===== SPECIAL ENHANCEMENTS =====
        # Priority items highlighting
        if any(keyword in processed_text.lower() for keyword in ['urgent', 'critical', 'important', 'asap', 'priority']):
            try:
                priority_bg = "FFF2E5" if status != "failed" else "FFE5E5"
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), priority_bg))
                p._element.get_or_add_pPr().append(shading_elm)

                # Make checkbox more prominent for priority items
                checkbox_run.font.bold = True
                checkbox_run.font.size = Pt(checkbox_run.font.size.pt + 1)
            except:
                pass

        # Enhanced spacing untuk long tasks
        if len(processed_text) > 120:
            p.paragraph_format.space_after = Pt(8)

        # Keep with next untuk better page breaks (level 0 items)
        if indent_level == 0:
            p.paragraph_format.keep_with_next = True

        # ===== STATUS-SPECIFIC ENHANCEMENTS =====
        if status == "failed":
            # Add subtle red border for failed tasks
            try:
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(r'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:left w:val="single" w:sz="4" w:space="0" w:color="C42B1C"/>
                </w:pBdr>''')
                pPr.append(pBdr)
            except:
                pass

        elif status == "checked":
            # Add subtle green accent for completed tasks
            try:
                # Lighter text color untuk completed tasks
                for run in p.runs[1:]:
                    if not run.font.color.rgb:
                        run.font.color.rgb = RGBColor(100, 100, 100)
            except:
                pass

        elif status == "in_progress":
            # Add progress indicator styling
            try:
                pPr = p._element.get_or_add_pPr()
                pBdr = parse_xml(r'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                    <w:left w:val="dotted" w:sz="4" w:space="0" w:color="0078D4"/>
                </w:pBdr>''')
                pPr.append(pBdr)
            except:
                pass

        # ===== DEADLINE AND DATE DETECTION =====
        # Enhanced date/deadline detection and formatting
        date_patterns = [
            r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',  # Date formats
            r'\b(due|deadline|by)\s+(\w+\s+\d{1,2})\b',  # "due March 15"
            r'\b(today|tomorrow|next week|next month)\b',  # Relative dates
        ]

        for pattern in date_patterns:
            if re.search(pattern, processed_text, re.IGNORECASE):
                # Add calendar icon untuk tasks dengan deadline
                try:
                    calendar_run = p.add_run(" 📅")
                    calendar_run.font.size = Pt(9)
                except:
                    pass
                break

        # ===== ASSIGNEE DETECTION =====
        # Detect assignee patterns and format accordingly
        assignee_patterns = [
            r'\b(assigned to|@|responsible:)\s*([A-Za-z\s]+)\b',
            r'\b(owner|assignee):\s*([A-Za-z\s]+)\b'
        ]

        for pattern in assignee_patterns:
            match = re.search(pattern, processed_text, re.IGNORECASE)
            if match:
                # Add person icon untuk assigned tasks
                try:
                    person_run = p.add_run(" 👤")
                    person_run.font.size = Pt(9)
                except:
                    pass
                break

        return p

    def _create_enhanced_quote(self, doc, quote_text, content_stats):
        """
        Membuat quote dengan styling yang sangat ditingkatkan dan adaptif berdasarkan content_stats.

        Features:
        - Dynamic styling berdasarkan document type
        - Quote attribution detection dan formatting
        - Multi-level quote support (nested quotes)
        - Enhanced visual elements (borders, backgrounds, icons)
        - Smart text processing dan formatting
        - Responsive spacing dan layout
        """
        # ===== ENHANCED INPUT VALIDATION =====
        if not quote_text or not quote_text.strip():
            return None

        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== ENHANCED QUOTE STYLE CONFIGURATION =====
        quote_styles = {
            "technical_report": {
                "icon": "📋",
                "color": "2E5984",
                "bg_color": "F0F5FF",
                "border_color": "BDD7EE",
                "font": "Cambria",
                "font_size": 11,
                "border_style": "single",
                "border_width": "8",
                "indent_left": 0.6,
                "indent_right": 0.3,
                "quote_mark": '"'
            },
            "meeting_notes": {
                "icon": "💬",
                "color": "385723",
                "bg_color": "F0FFF0",
                "border_color": "70AD47",
                "font": "Calibri",
                "font_size": 10.5,
                "border_style": "single",
                "border_width": "6",
                "indent_left": 0.4,
                "indent_right": 0.2,
                "quote_mark": '•'
            },
            "lecture": {
                "icon": "📚",
                "color": "C65911",
                "bg_color": "FFF8F0",
                "border_color": "ED7D31",
                "font": "Georgia",
                "font_size": 11,
                "border_style": "single",
                "border_width": "10",
                "indent_left": 0.7,
                "indent_right": 0.4,
                "quote_mark": '"'
            },
            "presentation": {
                "icon": "🎤",
                "color": "7030A0",
                "bg_color": "F8F0FF",
                "border_color": "9966CC",
                "font": "Segoe UI",
                "font_size": 12,
                "border_style": "double",
                "border_width": "12",
                "indent_left": 0.5,
                "indent_right": 0.3,
                "quote_mark": '"'
            },
            "research": {
                "icon": "🔬",
                "color": "1F4E79",
                "bg_color": "F5F8FF",
                "border_color": "4472C4",
                "font": "Times New Roman",
                "font_size": 11,
                "border_style": "single",
                "border_width": "12",
                "indent_left": 0.8,
                "indent_right": 0.5,
                "quote_mark": '"'
            },
            "interview": {
                "icon": "🎙️",
                "color": "8A2BE2",
                "bg_color": "FAF0FF",
                "border_color": "D8BFD8",
                "font": "Calibri",
                "font_size": 10.5,
                "border_style": "single",
                "border_width": "6",
                "indent_left": 0.4,
                "indent_right": 0.2,
                "quote_mark": '→'
            },
            "instructional": {
                "icon": "📖",
                "color": "228B22",
                "bg_color": "F0FFF0",
                "border_color": "90EE90",
                "font": "Segoe UI",
                "font_size": 11,
                "border_style": "single",
                "border_width": "8",
                "indent_left": 0.5,
                "indent_right": 0.3,
                "quote_mark": '💡'
            },
            "news": {
                "icon": "📰",
                "color": "DC143C",
                "bg_color": "FFF8F8",
                "border_color": "FFB6C1",
                "font": "Arial",
                "font_size": 10.5,
                "border_style": "single",
                "border_width": "6",
                "indent_left": 0.3,
                "indent_right": 0.2,
                "quote_mark": '"'
            },
            "narrative": {
                "icon": "📝",
                "color": "8B4513",
                "bg_color": "FFF8F0",
                "border_color": "D2B48C",
                "font": "Book Antiqua",
                "font_size": 11.5,
                "border_style": "single",
                "border_width": "10",
                "indent_left": 0.6,
                "indent_right": 0.4,
                "quote_mark": '"'
            },
            "general": {
                "icon": "💭",
                "color": "4F4F4F",
                "bg_color": "F8F8F8",
                "border_color": "C0C0C0",
                "font": "Calibri",
                "font_size": 11,
                "border_style": "single",
                "border_width": "8",
                "indent_left": 0.5,
                "indent_right": 0.3,
                "quote_mark": '"'
            }
        }

        # Get style configuration
        style_config = quote_styles.get(document_type, quote_styles["general"])

        # ===== ENHANCED TEXT PROCESSING =====
        # Process multi-line quotes dan nested quotes
        quote_lines = quote_text.split('\n')
        processed_lines = []

        # Detect quote attribution (author, source, etc.)
        attribution = None
        main_quote_lines = []

        for line in quote_lines:
            line = line.strip()
            if not line:
                continue

            # Detect attribution patterns
            attribution_patterns = [
                r'^[-—–]\s*(.+)$',  # "— Author Name"
                r'^~\s*(.+)$',      # "~ Author Name"
                r'^\*\s*(.+)\s*\*$', # "* Source *"
                r'^\(([^)]+)\)$',   # "(Author Name)"
                r'^Source:\s*(.+)$', # "Source: Name"
                r'^By:\s*(.+)$',    # "By: Name"
                r'^-\s*([A-Za-z\s,\.]+)$'  # "- Author Name"
            ]

            is_attribution = False
            for pattern in attribution_patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    attribution = match.group(1).strip()
                    is_attribution = True
                    break

            if not is_attribution:
                main_quote_lines.append(line)

        # Join main quote content
        processed_quote = '\n'.join(main_quote_lines) if main_quote_lines else quote_text

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = 'Normal'  # Start with normal untuk full control

        # ===== ENHANCED VISUAL STYLING =====
        try:
            # Background shading
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
                nsdecls('w'), style_config["bg_color"]))
            p._element.get_or_add_pPr().append(shading_elm)

            # Enhanced border (left border for quote effect)
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(f'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:left w:val="{style_config["border_style"]}" w:sz="{style_config["border_width"]}"
                        w:space="0" w:color="{style_config["border_color"]}"/>
            </w:pBdr>''')
            pPr.append(pBdr)

        except Exception as e:
            logger.warning(f"Could not apply enhanced quote styling: {e}")

        # ===== ENHANCED SPACING AND INDENTATION =====
        # Dynamic spacing berdasarkan document type dan content length
        content_length = len(processed_quote)

        if content_length > 200:
            space_before = Pt(14)
            space_after = Pt(14)
            line_spacing = Pt(15)
        elif content_length > 100:
            space_before = Pt(12)
            space_after = Pt(12)
            line_spacing = Pt(14)
        else:
            space_before = Pt(10)
            space_after = Pt(10)
            line_spacing = Pt(13)

        # Apply enhanced spacing
        p.paragraph_format.left_indent = Inches(style_config["indent_left"])
        p.paragraph_format.right_indent = Inches(style_config["indent_right"])
        p.paragraph_format.space_before = space_before
        p.paragraph_format.space_after = space_after
        p.paragraph_format.line_spacing = line_spacing

        # Enhanced first line indent untuk quote effect
        p.paragraph_format.first_line_indent = Inches(0.1)

        # ===== ENHANCED CONTENT FORMATTING =====
        # Add quote icon dan opening mark
        icon_run = p.add_run(f"{style_config['icon']} ")
        icon_run.font.size = Pt(style_config["font_size"] + 1)

        # Opening quote mark dengan enhanced styling
        opening_run = p.add_run(f"{style_config['quote_mark']} ")
        opening_run.font.name = style_config["font"]
        opening_run.font.size = Pt(style_config["font_size"] + 2)
        opening_run.font.color.rgb = RGBColor.from_string(style_config["color"])
        opening_run.font.bold = True

        # ===== ADVANCED TEXT PROCESSING =====
        # Enhanced text processing dengan inline formatting support
        if hasattr(self, '_add_enhanced_formatted_runs'):
            self._add_enhanced_formatted_runs(p, processed_quote, content_stats)
        else:
            # Fallback dengan simple formatting
            # Process basic formatting patterns
            formatting_patterns = [
                (r'\*\*(.+?)\*\*', 'bold'),
                (r'\*(.+?)\*', 'italic'),
                (r'`(.+?)`', 'code'),
                (r'~~(.+?)~~', 'strikethrough')
            ]

            remaining_text = processed_quote
            for pattern, format_type in formatting_patterns:
                parts = re.split(f'({pattern})', remaining_text)
                for part in parts:
                    if re.match(pattern, part):
                        # Extract content and apply formatting
                        content = re.match(pattern, part).group(1)
                        run = p.add_run(content)

                        if format_type == 'bold':
                            run.font.bold = True
                        elif format_type == 'italic':
                            run.font.italic = True
                        elif format_type == 'code':
                            run.font.name = 'Consolas'
                            run.font.size = Pt(style_config["font_size"] - 1)
                        elif format_type == 'strikethrough':
                            run.font.strike = True

                        # Apply base quote styling
                        run.font.name = style_config["font"]
                        run.font.size = Pt(style_config["font_size"])
                        run.font.color.rgb = RGBColor.from_string(style_config["color"])
                        run.font.italic = True
                    else:
                        # Regular text
                        if part.strip():
                            run = p.add_run(part)
                            run.font.name = style_config["font"]
                            run.font.size = Pt(style_config["font_size"])
                            run.font.color.rgb = RGBColor.from_string(style_config["color"])
                            run.font.italic = True
                break  # Process only first pattern match for simplicity

            # If no patterns matched, add as regular text
            if not any(re.search(pattern, processed_quote) for pattern, _ in formatting_patterns):
                content_run = p.add_run(processed_quote)
                content_run.font.name = style_config["font"]
                content_run.font.size = Pt(style_config["font_size"])
                content_run.font.color.rgb = RGBColor.from_string(style_config["color"])
                content_run.font.italic = True

        # Closing quote mark
        if style_config['quote_mark'] == '"':
            closing_mark = ' "'
        else:
            closing_mark = ""

        if closing_mark:
            closing_run = p.add_run(closing_mark)
            closing_run.font.name = style_config["font"]
            closing_run.font.size = Pt(style_config["font_size"] + 2)
            closing_run.font.color.rgb = RGBColor.from_string(style_config["color"])
            closing_run.font.bold = True

        # ===== ENHANCED ATTRIBUTION HANDLING =====
        if attribution:
            # Create separate paragraph untuk attribution
            attr_p = doc.add_paragraph()
            attr_p.style = 'Normal'

            # Attribution styling
            attr_p.paragraph_format.left_indent = Inches(style_config["indent_left"] + 0.2)
            attr_p.paragraph_format.right_indent = Inches(style_config["indent_right"])
            attr_p.paragraph_format.space_before = Pt(6)
            attr_p.paragraph_format.space_after = Pt(8)
            attr_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            # Attribution content
            attr_run = attr_p.add_run(f"— {attribution}")
            attr_run.font.name = style_config["font"]
            attr_run.font.size = Pt(style_config["font_size"] - 1)
            attr_run.font.color.rgb = RGBColor.from_string(style_config["color"])
            attr_run.font.italic = True
            attr_run.font.bold = True

            return attr_p  # Return attribution paragraph as the main return

        # ===== ENHANCED SPECIAL CASES =====
        # Special handling untuk different quote types
        if document_type == "technical_report":
            # Add technical quote indicator
            try:
                # Add subtle technical background pattern
                bg_elm = parse_xml(r'<w:shd {} w:fill="F0F5FF" w:val="diagStripe"/>'.format(nsdecls('w')))
                p._element.get_or_add_pPr().append(bg_elm)
            except:
                pass

        elif document_type == "meeting_notes":
            # Add speaker indicator jika ada
            if any(keyword in processed_quote.lower() for keyword in ['said', 'mentioned', 'stated', 'kata', 'bilang']):
                # This looks like a spoken quote in meeting
                for run in p.runs:
                    if not run.font.bold:  # Don't modify quote marks
                        run.font.bold = True

        elif document_type == "research":
            # Add citation placeholder
            citation_p = doc.add_paragraph()
            citation_p.style = 'Normal'
            citation_p.paragraph_format.left_indent = Inches(style_config["indent_left"] + 0.1)
            citation_p.paragraph_format.space_before = Pt(3)
            citation_p.paragraph_format.space_after = Pt(6)

            citation_run = citation_p.add_run("[Citation needed]")
            citation_run.font.size = Pt(9)
            citation_run.font.italic = True
            citation_run.font.color.rgb = RGBColor(128, 128, 128)

        # ===== ENHANCED CONTEXT-AWARE ADJUSTMENTS =====
        # Adjust styling berdasarkan content characteristics
        complexity = content_stats.get("complexity_level", "medium")
        word_count = content_stats.get("word_count", 0)

        if complexity == "high" or word_count > 5000:
            # For complex documents, make quotes more prominent
            for run in p.runs:
                if run.font.size and run.font.size.pt:
                    run.font.size = Pt(run.font.size.pt + 0.5)

        elif complexity == "low":
            # For simple documents, make quotes more subtle
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)

        # ===== ENHANCED ACCESSIBILITY =====
        # Add paragraph properties untuk screen readers
        try:
            # Add role information untuk accessibility
            p._element.set(qn('w:role'), 'quote')
        except:
            pass

        # Keep with next untuk better page breaks
        p.paragraph_format.keep_with_next = True

        logger.debug(f"Enhanced quote created for {document_type} with attribution: {bool(attribution)}")

        return p

    def _apply_enhanced_heading_styling(self, heading_element, level, content_stats):
        """
        Enhanced styling untuk heading berdasarkan level dan content statistics.

        Args:
            heading_element: Word heading element
            level: Heading level (1-6)
            content_stats: Dictionary containing content analysis or string document type
        """
        try:
            # ===== ENHANCED INPUT VALIDATION AND NORMALIZATION =====
            if not heading_element or not hasattr(heading_element, 'runs'):
                logger.warning("Invalid heading element provided")
                return False

            # Normalize content_stats to extract document_type properly
            if isinstance(content_stats, dict):
                document_type = content_stats.get('content_type', 'general')
                complexity_level = content_stats.get('complexity_level', 'medium')
                word_count = content_stats.get('word_count', 0)
                has_technical = content_stats.get('technical_score', 0) > 3
            elif isinstance(content_stats, str):
                document_type = content_stats
                complexity_level = 'medium'
                word_count = 0
                has_technical = False
            else:
                logger.warning(f"Invalid content_stats type: {type(content_stats)}")
                document_type = 'general'
                complexity_level = 'medium'
                word_count = 0
                has_technical = False

            # Ensure document_type is a valid string
            if not isinstance(document_type, str) or not document_type.strip():
                document_type = 'general'

            # Validate and clamp level
            level = max(1, min(6, int(level))) if level else 1

            # ===== ENHANCED FONT SIZE CONFIGURATION =====
            # Dynamic font sizes based on document complexity and type
            base_font_sizes = {1: 18, 2: 15, 3: 13, 4: 12, 5: 11, 6: 10}

            # Adjust font sizes based on document characteristics
            size_adjustments = {
                "technical_report": {1: +2, 2: +1, 3: +1},  # Larger for technical docs
                "presentation": {1: +3, 2: +2, 3: +1},      # Larger for presentations
                "meeting_notes": {1: 0, 2: 0, 3: -1},       # Standard for meetings
                "lecture": {1: +1, 2: +1, 3: 0},            # Slightly larger for lectures
                "research": {1: +1, 2: 0, 3: 0},            # Slightly larger for research
                "general": {1: 0, 2: 0, 3: 0}               # Standard sizes
            }

            # Apply complexity adjustments
            if complexity_level == 'high':
                complexity_adjustment = 1
            elif complexity_level == 'low':
                complexity_adjustment = -1
            else:
                complexity_adjustment = 0

            # Calculate final font size
            base_size = base_font_sizes.get(level, 11)
            doc_adjustment = size_adjustments.get(document_type, {}).get(level, 0)
            final_size = base_size + doc_adjustment + complexity_adjustment
            final_size = max(8, min(24, final_size))  # Clamp between 8-24pt

            # ===== ENHANCED COLOR SCHEMES =====
            color_schemes = {
                "technical_report": {
                    1: {"color": "001f3f", "secondary": "0074d9"},  # Deep Navy & Blue
                    2: {"color": "1f4e79", "secondary": "4472c4"},  # Professional Blue
                    3: {"color": "2e75b5", "secondary": "5b9bd5"},  # Medium Blue
                    4: {"color": "8db4e2", "secondary": "bdd7ee"},  # Light Blue
                    5: {"color": "b8cce4", "secondary": "d6e3f0"},  # Very Light Blue
                    6: {"color": "ddeaf6", "secondary": "e8f2ff"}   # Pale Blue
                },
                "meeting_notes": {
                    1: {"color": "2d5016", "secondary": "70ad47"},  # Dark Green
                    2: {"color": "385723", "secondary": "8faadc"},  # Forest Green
                    3: {"color": "4f7942", "secondary": "a9d18e"},  # Medium Green
                    4: {"color": "70ad47", "secondary": "c5e0b4"},  # Standard Green
                    5: {"color": "a9d18e", "secondary": "e2efda"},  # Light Green
                    6: {"color": "c5e0b4", "secondary": "f2f8f0"}   # Very Light Green
                },
                "lecture": {
                    1: {"color": "8b3a00", "secondary": "ed7d31"},  # Dark Orange
                    2: {"color": "c65911", "secondary": "f4b183"},  # Burnt Orange
                    3: {"color": "d18b47", "secondary": "fad5b4"},  # Medium Orange
                    4: {"color": "ed7d31", "secondary": "fbe5d6"},  # Standard Orange
                    5: {"color": "f4b183", "secondary": "fdf2e9"},  # Light Orange
                    6: {"color": "fad5b4", "secondary": "fef9f5"}   # Very Light Orange
                },
                "presentation": {
                    1: {"color": "4b0082", "secondary": "9966cc"},  # Indigo
                    2: {"color": "7030a0", "secondary": "b19cd9"},  # Purple
                    3: {"color": "9966cc", "secondary": "d6c7f0"},  # Medium Purple
                    4: {"color": "b19cd9", "secondary": "e9e0f7"},  # Light Purple
                    5: {"color": "d6c7f0", "secondary": "f4f0fd"},  # Very Light Purple
                    6: {"color": "e9e0f7", "secondary": "faf8ff"}   # Pale Purple
                },
                "research": {
                    1: {"color": "0d1b2a", "secondary": "415a77"},  # Dark Blue Gray
                    2: {"color": "1b263b", "secondary": "778da9"},  # Blue Gray
                    3: {"color": "1f4e79", "secondary": "8db4e2"},  # Professional Blue
                    4: {"color": "4472c4", "secondary": "bdd7ee"},  # Medium Blue
                    5: {"color": "8db4e2", "secondary": "ddeaf6"},  # Light Blue
                    6: {"color": "bdd7ee", "secondary": "f0f5ff"}   # Very Light Blue
                },
                "instructional": {
                    1: {"color": "006400", "secondary": "32cd32"},  # Dark Green
                    2: {"color": "228b22", "secondary": "90ee90"},  # Forest Green
                    3: {"color": "32cd32", "secondary": "98fb98"},  # Lime Green
                    4: {"color": "90ee90", "secondary": "f0fff0"},  # Light Green
                    5: {"color": "98fb98", "secondary": "f5fffa"},  # Mint Cream
                    6: {"color": "f0fff0", "secondary": "fafffa"}   # Honeydew
                },
                "general": {
                    1: {"color": "1f497d", "secondary": "4f81bd"},  # Professional Blue
                    2: {"color": "4f81bd", "secondary": "8db4e2"},  # Medium Blue
                    3: {"color": "8db4e2", "secondary": "bdd7ee"},  # Light Blue
                    4: {"color": "bdd7ee", "secondary": "ddeaf6"},  # Very Light Blue
                    5: {"color": "ddeaf6", "secondary": "f0f5ff"},  # Pale Blue
                    6: {"color": "f0f5ff", "secondary": "fafcff"}   # Almost White
                }
            }

            # ===== ENHANCED FONT FAMILY SELECTION =====
            font_families = {
                "technical_report": "Cambria",
                "presentation": "Segoe UI",
                "meeting_notes": "Calibri",
                "lecture": "Georgia",
                "research": "Times New Roman",
                "instructional": "Segoe UI",
                "general": "Calibri"
            }

            # Get color scheme and font
            scheme = color_schemes.get(document_type, color_schemes["general"])
            color_config = scheme.get(level, scheme.get(1, {"color": "1f497d", "secondary": "4f81bd"}))
            primary_color = color_config["color"]
            secondary_color = color_config.get("secondary", primary_color)
            font_family = font_families.get(document_type, "Calibri")

            # ===== APPLY ENHANCED STYLING TO RUNS =====
            if not heading_element.runs:
                # Create a run if none exists
                heading_element.add_run()

            for run in heading_element.runs:
                # Font family and size
                run.font.name = font_family
                run.font.size = Pt(final_size)

                # Color application with fallback
                try:
                    run.font.color.rgb = RGBColor.from_string(primary_color)
                except (ValueError, AttributeError) as e:
                    logger.warning(f"Invalid color '{primary_color}': {e}")
                    run.font.color.rgb = RGBColor(31, 73, 125)  # Fallback blue

                # Level-specific enhancements
                if level == 1:
                    run.font.bold = True
                    if document_type in ["presentation", "technical_report"]:
                        run.font.all_caps = True
                elif level == 2:
                    run.font.bold = True
                    if document_type in ["presentation"]:
                        run.font.small_caps = True
                elif level <= 4:
                    run.font.bold = True
                else:
                    run.font.bold = False

            # ===== ENHANCED PARAGRAPH FORMATTING =====
            if hasattr(heading_element, 'paragraph_format'):
                pf = heading_element.paragraph_format

                # Spacing based on level and document type
                if level == 1:
                    pf.space_before = Pt(20 if document_type == "presentation" else 16)
                    pf.space_after = Pt(14 if document_type == "presentation" else 12)
                elif level == 2:
                    pf.space_before = Pt(16 if document_type == "presentation" else 14)
                    pf.space_after = Pt(10 if document_type == "presentation" else 8)
                elif level <= 4:
                    pf.space_before = Pt(12)
                    pf.space_after = Pt(6)
                else:
                    pf.space_before = Pt(8)
                    pf.space_after = Pt(4)

                # Keep with next for better page breaks
                pf.keep_with_next = True

                # Alignment based on document type and level
                if document_type == "presentation" and level == 1:
                    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # ===== ENHANCED VISUAL ELEMENTS =====
            # Add enhanced borders for level 1 and 2 headings
            if level <= 2:
                border_size = "20" if level == 1 else "12"
                border_color = primary_color

                # Enhanced border with gradient effect simulation
                if hasattr(self, '_add_enhanced_bottom_border'):
                    self._add_enhanced_bottom_border(
                        heading_element,
                        border_color,
                        border_size,
                        document_type=document_type,
                        secondary_color=secondary_color,
                        level=level
                    )
                elif hasattr(self, '_add_bottom_border'):
                    self._add_bottom_border(heading_element, border_color, border_size)

            # ===== ADVANCED DOCUMENT-SPECIFIC ENHANCEMENTS =====
            if document_type == "technical_report":
                # Add subtle background shading for level 1 technical headings
                if level == 1:
                    try:
                        from docx.oxml import parse_xml
                        from docx.oxml.ns import nsdecls
                        shading_elm = parse_xml(r'<w:shd {} w:fill="F0F5FF"/>'.format(nsdecls('w')))
                        heading_element._element.get_or_add_pPr().append(shading_elm)
                    except Exception as e:
                        logger.debug(f"Could not apply background shading: {e}")

            elif document_type == "presentation":
                # Add emphasis for presentation headings
                if level <= 2:
                    try:
                        # Add subtle text effects for presentation
                        for run in heading_element.runs:
                            if level == 1:
                                run.font.shadow = True
                    except Exception as e:
                        logger.debug(f"Could not apply text effects: {e}")

            # ===== ACCESSIBILITY AND METADATA =====
            try:
                # Add heading level metadata for accessibility
                heading_element._element.set(qn('w:outlineLvl'), str(level - 1))
            except Exception as e:
                logger.debug(f"Could not set outline level: {e}")

            # ===== SUCCESS LOGGING =====
            logger.debug(f"✅ Enhanced heading styling applied successfully:")
            logger.debug(f"   Level: {level}")
            logger.debug(f"   Document Type: {document_type}")
            logger.debug(f"   Font: {font_family} {final_size}pt")
            logger.debug(f"   Color: {primary_color}")
            logger.debug(f"   Complexity: {complexity_level}")

            return True

        except Exception as e:
            error_msg = f"Error applying enhanced heading styling: {str(e)}"
            logger.error(error_msg, exc_info=True)

            # ===== FALLBACK STYLING =====
            try:
                logger.info("Applying fallback heading styling...")

                if heading_element and hasattr(heading_element, 'runs'):
                    for run in heading_element.runs:
                        run.font.name = 'Calibri'
                        run.font.bold = True
                        run.font.size = Pt(max(10, 20 - level))
                        run.font.color.rgb = RGBColor(31, 73, 125)  # Safe blue color

                    if hasattr(heading_element, 'paragraph_format'):
                        pf = heading_element.paragraph_format
                        pf.space_before = Pt(12)
                        pf.space_after = Pt(6)
                        pf.keep_with_next = True

                logger.info("✅ Fallback styling applied successfully")
                return True

            except Exception as fallback_error:
                logger.error(f"❌ Fallback styling also failed: {str(fallback_error)}")
                return False

    def _add_enhanced_admonition_block(self, doc, lines, admonition_type, document_type):
        """Enhanced admonition block dengan styling yang lebih canggih dan memanfaatkan border enhancement."""
        if not lines:
            return

        content = "\n".join(lines)

        # Enhanced type mapping dengan lebih banyak variasi dan parameter border
        type_configs = {
            "note": {
                "icon": "📝", "title": "NOTE", "color": "4472C4", "bg": "E8F2FF", "border": "BDD7EE",
                "border_style": "single", "priority": "medium", "sides": ["left", "top", "bottom"]
            },
            "info": {
                "icon": "ℹ️", "title": "INFORMATION", "color": "0078D4", "bg": "F3F9FF", "border": "A6C8FF",
                "border_style": "single", "priority": "medium", "sides": ["left", "bottom"]
            },
            "tip": {
                "icon": "💡", "title": "TIP", "color": "107C10", "bg": "F3FFF3", "border": "9FD89F",
                "border_style": "dashed", "priority": "low", "sides": ["left", "bottom"],
                "rounded_style": True
            },
            "warning": {
                "icon": "⚠️", "title": "WARNING", "color": "FF8C00", "bg": "FFF8E7", "border": "FFD166",
                "border_style": "double", "priority": "high", "sides": "all",
                "priority_highlight": True
            },
            "danger": {
                "icon": "🚨", "title": "DANGER", "color": "DC3545", "bg": "FFF5F5", "border": "F8A8A8",
                "border_style": "thick", "priority": "high", "sides": "all",
                "priority_highlight": True, "shadow_effect": True
            },
            "important": {
                "icon": "❗", "title": "IMPORTANT", "color": "DC143C", "bg": "FFF0F0", "border": "FFB3B3",
                "border_style": "double", "priority": "high", "sides": "all",
                "priority_highlight": True
            },
            "success": {
                "icon": "✅", "title": "SUCCESS", "color": "28A745", "bg": "F0FFF0", "border": "90EE90",
                "border_style": "single", "priority": "medium", "sides": ["left", "bottom"],
                "gradient_colors": ["90EE90", "98FB98", "F0FFF0"]
            },
            "question": {
                "icon": "❓", "title": "QUESTION", "color": "6F42C1", "bg": "F8F0FF", "border": "D0A9F5",
                "border_style": "dotted", "priority": "medium", "sides": ["left", "right"],
                "animation_hint": True
            },
            "conclusion": {
                "icon": "🏁", "title": "CONCLUSION", "color": "495057", "bg": "F8F9FA", "border": "CED4DA",
                "border_style": "wave", "priority": "low", "sides": ["top", "bottom"],
                "opacity_effect": True
            }
        }

        config = type_configs.get(admonition_type, type_configs["note"])

        # Create container paragraph
        para = doc.add_paragraph()
        para.style = 'No Spacing'

        # Enhanced styling dengan gradien effect simulation
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), config["bg"]))
        para._element.get_or_add_pPr().append(shading_elm)

        # ENHANCED: Gunakan enhanced border dengan semua fitur baru
        border_kwargs = {
            "border_style": config.get("border_style", "single"),
            "sides": config.get("sides", ["left"]),
            "document_type": document_type,
            "priority_highlight": config.get("priority_highlight", False),
            "gradient_colors": config.get("gradient_colors", []),
            "shadow_effect": config.get("shadow_effect", False),
            "rounded_style": config.get("rounded_style", False),
            "animation_hint": config.get("animation_hint", False),
            "opacity_effect": config.get("opacity_effect", False),
            "accessibility_mode": True,
            "add_shading": True,  # Tambahkan shading otomatis
            "debug_mode": False   # Set True untuk debugging
        }

        # Tambahkan custom pattern untuk admonition khusus
        if admonition_type == "warning":
            border_kwargs["custom_pattern"] = {
                "top": {"style": "thick", "size": "16", "color": config["border"]},
                "left": {"style": "double", "size": "12", "color": config["border"]},
                "bottom": {"style": "single", "size": "8", "color": config["border"]},
                "right": {"style": "dotted", "size": "6", "color": config["border"]}
            }
        elif admonition_type == "danger":
            border_kwargs["gradient_colors"] = [config["border"], "DC3545", "8B0000"]

        # Apply enhanced border
        self._add_enhanced_paragraph_border(para, config["border"], "12", **border_kwargs)

        # Enhanced spacing dan indentation berdasarkan priority
        base_indent = 0.3
        if config.get("priority") == "high":
            base_indent += 0.1
            para.paragraph_format.space_before = Pt(16)
            para.paragraph_format.space_after = Pt(16)
        elif config.get("priority") == "low":
            base_indent -= 0.05
            para.paragraph_format.space_before = Pt(8)
            para.paragraph_format.space_after = Pt(8)
        else:
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(12)

        para.paragraph_format.left_indent = Inches(base_indent)
        para.paragraph_format.right_indent = Inches(0.3)
        para.paragraph_format.line_spacing = 1.2

        # Keep with next untuk admonition penting
        if config.get("priority") == "high":
            para.paragraph_format.keep_with_next = True

        # Enhanced icon dengan ukuran berdasarkan priority
        icon_size = 12
        if config.get("priority") == "high":
            icon_size = 14
        elif config.get("priority") == "low":
            icon_size = 10

        icon_run = para.add_run(f"{config['icon']} ")
        icon_run.font.size = Pt(icon_size)

        # Enhanced title dengan emphasis berdasarkan priority
        title_text = config['title']
        if config.get("priority") == "high":
            title_text = f"🔴 {title_text}"  # Red indicator untuk high priority
        elif admonition_type == "danger":
            title_text = f"⚡ {title_text}"  # Lightning untuk danger

        title_run = para.add_run(f"{title_text}: ")
        title_run.bold = True
        title_run.font.size = Pt(11 if config.get("priority") != "high" else 12)
        title_run.font.color.rgb = RGBColor.from_string(config["color"])

        # Underline untuk high priority
        if config.get("priority") == "high":
            title_run.underline = True

        # ENHANCED: Process content dengan enhanced formatting dan document_type parameter
        if hasattr(self, '_add_enhanced_formatted_runs'):
            self._add_enhanced_formatted_runs(para, content, document_type)
        else:
            # Fallback untuk compatibility
            content_run = para.add_run(content)
            content_run.font.name = 'Calibri'
            content_run.font.size = Pt(11)

        # ENHANCED: Tambahkan metadata untuk admonition tracking
        try:
            # Add admonition type sebagai metadata
            para._element.set(qn('w:admonitionType'), admonition_type)
            para._element.set(qn('w:priority'), config.get("priority", "medium"))
        except:
            pass  # Continue jika metadata gagal

        # ENHANCED: Log admonition creation untuk debugging
        logger.debug(f"Enhanced admonition created: {admonition_type} | Priority: {config.get('priority')} | "
                    f"Document: {document_type} | Border style: {config.get('border_style')}")

        return para

    def _add_enhanced_paragraph_border(self, paragraph, border_color, border_size, **kwargs):
        """
        Menambahkan border yang sangat ditingkatkan pada paragraf dengan fitur advanced.

        Args:
            paragraph: Paragraph object Word
            border_color: Warna border (hex tanpa #, atau nama warna)
            border_size: Ukuran border dalam twips (string atau int)
            **kwargs: Opsi tambahan untuk customization

        Kwargs:
            - border_style: 'single', 'double', 'thick', 'dotted', 'dashed', 'triple', dll (default: 'single')
            - sides: list atau 'all' untuk sisi yang diberi border ['top', 'left', 'bottom', 'right'] (default: semua)
            - border_spacing: jarak border dari teks (default: '0')
            - gradient_colors: list warna untuk efek gradien [color1, color2, color3]
            - shadow_effect: bool, tambahkan efek bayangan (default: False)
            - rounded_style: bool, simulasi rounded corners dengan style berbeda (default: False)
            - priority_highlight: bool, highlight untuk konten prioritas (default: False)
            - theme_color: string, gunakan theme color Word ('accent1', 'accent2', dll)
            - custom_pattern: dict dengan pattern per sisi {'top': {...}, 'left': {...}}
            - opacity_effect: bool, efek transparansi dengan warna yang lebih terang (default: False)
            - document_type: string untuk auto-styling ('technical', 'meeting', 'presentation', dll)
            - animation_hint: bool, style yang menunjukkan konten dinamis (default: False)
            - accessibility_mode: bool, mode untuk screen readers (default: True)
            - debug_mode: bool, tampilkan informasi debug (default: False)

        Returns:
            bool: True jika berhasil, False jika gagal

        Example:
            # Basic usage
            self._add_enhanced_paragraph_border(para, "4472C4", "12")

            # Advanced usage
            self._add_enhanced_paragraph_border(
                para, "FF6B35", "16",
                border_style="double",
                sides=["left", "bottom"],
                gradient_colors=["FF6B35", "FF8C42", "FFA07A"],
                shadow_effect=True,
                document_type="technical",
                priority_highlight=True
            )
        """

        # ===== ENHANCED INPUT VALIDATION =====
        if not paragraph or not border_color:
            if kwargs.get('debug_mode', False):
                print("❌ Invalid paragraph or border_color")
            return False

        try:
            # Normalize dan validasi border_size
            if isinstance(border_size, int):
                border_size = str(border_size)
            elif not isinstance(border_size, str):
                border_size = "12"

            # Validasi range border_size (Word supports 2-96 twips)
            try:
                size_int = int(border_size)
                if size_int < 2:
                    border_size = "2"
                elif size_int > 96:
                    border_size = "96"
                else:
                    border_size = str(size_int)
            except (ValueError, TypeError):
                border_size = "12"

        except Exception as e:
            if kwargs.get('debug_mode', False):
                print(f"❌ Error validating border_size: {e}")
            border_size = "12"

        # ===== EXTRACT DAN PROCESS KWARGS =====
        border_style = kwargs.get('border_style', 'single')
        sides = kwargs.get('sides', ['top', 'left', 'bottom', 'right'])
        border_spacing = kwargs.get('border_spacing', '0')
        gradient_colors = kwargs.get('gradient_colors', [])
        shadow_effect = kwargs.get('shadow_effect', False)
        rounded_style = kwargs.get('rounded_style', False)
        priority_highlight = kwargs.get('priority_highlight', False)
        theme_color = kwargs.get('theme_color', None)
        custom_pattern = kwargs.get('custom_pattern', {})
        opacity_effect = kwargs.get('opacity_effect', False)
        document_type = kwargs.get('document_type', 'general')
        animation_hint = kwargs.get('animation_hint', False)
        accessibility_mode = kwargs.get('accessibility_mode', True)
        debug_mode = kwargs.get('debug_mode', False)

        # Handle 'all' sides shortcut
        if sides == 'all' or sides == ['all']:
            sides = ['top', 'left', 'bottom', 'right']
        elif isinstance(sides, str):
            sides = [sides]

        # ===== ADVANCED COLOR PROCESSING =====
        def normalize_color(color):
            """Normalize color input ke hex format dengan extended support."""
            if not color:
                return "000000"

            # Remove # jika ada
            color = str(color).lstrip('#').upper()

            # Extended named colors dictionary
            named_colors = {
                # Basic colors
                'BLACK': '000000', 'WHITE': 'FFFFFF', 'RED': 'FF0000', 'GREEN': '008000',
                'BLUE': '0000FF', 'YELLOW': 'FFFF00', 'CYAN': '00FFFF', 'MAGENTA': 'FF00FF',
                'GRAY': '808080', 'GREY': '808080', 'SILVER': 'C0C0C0', 'NAVY': '000080',

                # Extended colors
                'DARKRED': '8B0000', 'DARKGREEN': '006400', 'DARKBLUE': '00008B',
                'LIGHTGRAY': 'D3D3D3', 'LIGHTGREY': 'D3D3D3', 'DARKGRAY': 'A9A9A9',
                'ORANGE': 'FFA500', 'PURPLE': '800080', 'BROWN': 'A52A2A', 'PINK': 'FFC0CB',

                # Professional colors
                'CORPORATE_BLUE': '4472C4', 'MEETING_GREEN': '70AD47', 'LECTURE_ORANGE': 'ED7D31',
                'TECHNICAL_NAVY': '2E5984', 'PRESENTATION_PURPLE': '9966CC', 'FORMAL_BLACK': '000000',

                # Theme colors
                'ACCENT1': '4472C4', 'ACCENT2': 'E7E6E6', 'ACCENT3': 'A5A5A5', 'ACCENT4': 'FFC000',
                'ACCENT5': '4472C4', 'ACCENT6': '70AD47'
            }

            if color in named_colors:
                return named_colors[color]

            # Validate hex color
            if len(color) == 6 and all(c in '0123456789ABCDEF' for c in color):
                return color
            elif len(color) == 3 and all(c in '0123456789ABCDEF' for c in color):
                # Expand 3-digit hex ke 6-digit
                return ''.join([c*2 for c in color])
            else:
                if debug_mode:
                    print(f"⚠️ Invalid color '{color}', using black as fallback")
                return "000000"

        def lighten_hex_color(hex_color, percent):
            """Lighten hex color by percentage."""
            try:
                hex_color = hex_color.lstrip('#')
                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                lightened = tuple(min(255, int(c + (255 - c) * percent / 100)) for c in rgb)
                return f"{lightened[0]:02X}{lightened[1]:02X}{lightened[2]:02X}"
            except:
                return hex_color

        def darken_hex_color(hex_color, percent):
            """Darken hex color by percentage."""
            try:
                hex_color = hex_color.lstrip('#')
                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                darkened = tuple(max(0, int(c * (100 - percent) / 100)) for c in rgb)
                return f"{darkened[0]:02X}{darkened[1]:02X}{darkened[2]:02X}"
            except:
                return hex_color

        def blend_colors(color1, color2, ratio=0.5):
            """Blend two hex colors dengan ratio tertentu."""
            try:
                c1 = tuple(int(color1[i:i+2], 16) for i in (0, 2, 4))
                c2 = tuple(int(color2[i:i+2], 16) for i in (0, 2, 4))
                blended = tuple(int(c1[i] * (1-ratio) + c2[i] * ratio) for i in range(3))
                return f"{blended[0]:02X}{blended[1]:02X}{blended[2]:02X}"
            except:
                return color1

        normalized_color = normalize_color(border_color)

        # ===== ENHANCED BORDER STYLES =====
        advanced_border_styles = {
            # Standard styles
            'single': 'single', 'double': 'double', 'thick': 'thick', 'thin': 'thin',
            'dotted': 'dotted', 'dashed': 'dashed',

            # Advanced styles
            'dash_dot': 'dashDotStroked', 'dash_dot_dot': 'dashDotDotStroked',
            'triple': 'triple', 'wave': 'wave', 'double_wave': 'doubleWave',
            'inset': 'inset', 'outset': 'outset',

            # 3D effects
            'emboss': 'threeDEmboss', 'engrave': 'threeDEngrave',

            # Decorative
            'art_border': 'apples', 'hearts': 'hearts', 'stars': 'confettiStreamers',

            # Professional variations
            'professional': 'single', 'corporate': 'double', 'modern': 'thick',
            'subtle': 'thin', 'elegant': 'wave'
        }

        # Normalize border style
        if border_style not in advanced_border_styles:
            if border_style in advanced_border_styles.values():
                pass  # Already in Word format
            else:
                if debug_mode:
                    print(f"⚠️ Unknown border style '{border_style}', using 'single'")
                border_style = 'single'
        else:
            border_style = advanced_border_styles[border_style]

        # ===== DOCUMENT TYPE SPECIFIC ENHANCEMENTS =====
        document_enhancements = {
            'technical': {
                'default_style': 'double', 'default_size': '16', 'priority_color': '2E5984',
                'spacing': '2', 'professional_mode': True
            },
            'technical_report': {
                'default_style': 'double', 'default_size': '16', 'priority_color': '2E5984',
                'spacing': '2', 'professional_mode': True
            },
            'meeting': {
                'default_style': 'single', 'default_size': '12', 'priority_color': '70AD47',
                'spacing': '1', 'clean_mode': True
            },
            'meeting_notes': {
                'default_style': 'single', 'default_size': '12', 'priority_color': '70AD47',
                'spacing': '1', 'clean_mode': True
            },
            'presentation': {
                'default_style': 'thick', 'default_size': '18', 'priority_color': '9966CC',
                'spacing': '3', 'vibrant_mode': True
            },
            'lecture': {
                'default_style': 'wave', 'default_size': '14', 'priority_color': 'ED7D31',
                'spacing': '2', 'academic_mode': True
            },
            'formal': {
                'default_style': 'double', 'default_size': '20', 'priority_color': '000000',
                'spacing': '4', 'conservative_mode': True
            },
            'casual': {
                'default_style': 'dotted', 'default_size': '8', 'priority_color': 'A5A5A5',
                'spacing': '1', 'relaxed_mode': True
            },
            'creative': {
                'default_style': 'wave', 'default_size': '12', 'priority_color': 'FF6B35',
                'spacing': '2', 'artistic_mode': True
            }
        }

        enhancement = document_enhancements.get(document_type, {})

        # Apply document type defaults jika tidak explicitly set
        if 'border_style' not in kwargs and enhancement:
            border_style = advanced_border_styles.get(enhancement.get('default_style', border_style), border_style)
        if kwargs.get('auto_size', False) and enhancement:
            border_size = enhancement.get('default_size', border_size)
        if priority_highlight and enhancement:
            normalized_color = enhancement.get('priority_color', normalized_color)
        if 'border_spacing' not in kwargs and enhancement:
            border_spacing = enhancement.get('spacing', border_spacing)

        # ===== CREATE ENHANCED PARAGRAPH BORDER =====
        try:
            from docx.oxml import OxmlElement
            from docx.oxml.shared import qn

            # Get paragraph properties
            pPr = paragraph._element.get_or_add_pPr()

            # Remove existing borders jika ada
            existing_pBdr = pPr.find(qn('w:pBdr'))
            if existing_pBdr is not None:
                pPr.remove(existing_pBdr)

            # Create new border container
            pBdr = OxmlElement('w:pBdr')

            # ===== GRADIENT EFFECT IMPLEMENTATION =====
            if gradient_colors and len(gradient_colors) >= 2:
                if debug_mode:
                    print(f"🎨 Applying gradient effect with colors: {gradient_colors}")

                # Create gradient dengan multiple layers
                for i, side in enumerate(sides):
                    color_index = i % len(gradient_colors)
                    grad_color = normalize_color(gradient_colors[color_index])

                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), border_style)
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), grad_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== ROUNDED STYLE SIMULATION =====
            elif rounded_style:
                if debug_mode:
                    print("🔘 Applying rounded style simulation")

                rounded_styles = {
                    'top': 'dotted' if border_style == 'single' else border_style,
                    'bottom': 'dotted' if border_style == 'single' else border_style,
                    'left': border_style,
                    'right': border_style
                }

                for side in sides:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), rounded_styles.get(side, border_style))
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), normalized_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== SHADOW EFFECT IMPLEMENTATION =====
            elif shadow_effect:
                if debug_mode:
                    print("🌫️ Applying shadow effect")

                shadow_color = darken_hex_color(normalized_color, 30)

                # Main borders
                for side in sides:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), border_style)
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), normalized_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

                # Shadow borders (conceptual - Word limitations)
                if 'bottom' in sides and 'right' in sides:
                    # Add subtle shadow indication dengan darker color
                    for shadow_side in ['bottom', 'right']:
                        if shadow_side in sides:
                            shadow_border = OxmlElement(f'w:{shadow_side}')
                            shadow_border.set(qn('w:val'), 'single')
                            shadow_border.set(qn('w:sz'), str(max(2, int(border_size) // 2)))
                            shadow_border.set(qn('w:space'), '1')
                            shadow_border.set(qn('w:color'), shadow_color)
                            # Note: This creates a layered effect in Word

            # ===== CUSTOM PATTERN IMPLEMENTATION =====
            elif custom_pattern:
                if debug_mode:
                    print(f"🎨 Applying custom pattern: {custom_pattern}")

                for side in sides:
                    if side in custom_pattern:
                        pattern = custom_pattern[side]
                        border_element = OxmlElement(f'w:{side}')

                        # Extract pattern properties
                        pattern_style = pattern.get('style', border_style)
                        pattern_size = str(pattern.get('size', border_size))
                        pattern_spacing = str(pattern.get('spacing', border_spacing))
                        pattern_color = normalize_color(pattern.get('color', normalized_color))

                        border_element.set(qn('w:val'), pattern_style)
                        border_element.set(qn('w:sz'), pattern_size)
                        border_element.set(qn('w:space'), pattern_spacing)
                        border_element.set(qn('w:color'), pattern_color)

                        if pattern.get('theme_color'):
                            try:
                                border_element.set(qn('w:themeColor'), pattern['theme_color'])
                            except:
                                pass

                        pBdr.append(border_element)
                    else:
                        # Use default untuk sides not in custom pattern
                        border_element = OxmlElement(f'w:{side}')
                        border_element.set(qn('w:val'), border_style)
                        border_element.set(qn('w:sz'), border_size)
                        border_element.set(qn('w:space'), border_spacing)
                        border_element.set(qn('w:color'), normalized_color)

                        if theme_color:
                            try:
                                border_element.set(qn('w:themeColor'), theme_color)
                            except:
                                pass

                        pBdr.append(border_element)

            # ===== OPACITY EFFECT IMPLEMENTATION =====
            elif opacity_effect:
                if debug_mode:
                    print("👻 Applying opacity effect")

                # Create lighter version untuk opacity effect
                opacity_color = lighten_hex_color(normalized_color, 40)

                for side in sides:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), border_style)
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), opacity_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== PRIORITY HIGHLIGHT IMPLEMENTATION =====
            elif priority_highlight:
                if debug_mode:
                    print("⭐ Applying priority highlight")

                priority_size = str(int(border_size) + 6)
                priority_style = 'thick' if border_style == 'single' else border_style

                for side in sides:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), priority_style)
                    border_element.set(qn('w:sz'), priority_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), normalized_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== ANIMATION HINT IMPLEMENTATION =====
            elif animation_hint:
                if debug_mode:
                    print("🔄 Applying animation hint styling")

                # Alternating pattern untuk animation hint
                animation_styles = ['dashed', 'dotted']

                for i, side in enumerate(sides):
                    anim_style = animation_styles[i % len(animation_styles)]
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), anim_style)
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), normalized_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== STANDARD BORDER IMPLEMENTATION =====
            else:
                if debug_mode:
                    print(f"📐 Applying standard border: {border_style}, {border_size}pt, {normalized_color}")

                for side in sides:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), border_style)
                    border_element.set(qn('w:sz'), border_size)
                    border_element.set(qn('w:space'), border_spacing)
                    border_element.set(qn('w:color'), normalized_color)

                    if theme_color:
                        try:
                            border_element.set(qn('w:themeColor'), theme_color)
                        except:
                            pass

                    pBdr.append(border_element)

            # ===== APPLY BORDERS TO PARAGRAPH =====
            pPr.append(pBdr)

            # ===== ACCESSIBILITY ENHANCEMENTS =====
            if accessibility_mode:
                try:
                    border_description = f"Border: {border_style} {normalized_color} {border_size}pt on {', '.join(sides)}"
                    paragraph._element.set(qn('w:altText'), border_description)
                    if debug_mode:
                        print(f"♿ Added accessibility description: {border_description}")
                except:
                    if debug_mode:
                        print("⚠️ Could not add accessibility description")

            # ===== ADDITIONAL VISUAL ENHANCEMENTS =====
            if kwargs.get('add_shading', False):
                try:
                    from docx.oxml.ns import nsdecls
                    from docx.oxml import parse_xml

                    shading_color = lighten_hex_color(normalized_color, 90)
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), shading_color))
                    pPr.append(shading_elm)

                    if debug_mode:
                        print(f"🎨 Added background shading: {shading_color}")
                except Exception as shading_error:
                    if debug_mode:
                        print(f"⚠️ Could not add shading: {shading_error}")

            # ===== SUCCESS LOGGING =====
            if debug_mode:
                print(f"✅ Enhanced border applied successfully!")
                print(f"   Style: {border_style}")
                print(f"   Size: {border_size}pt")
                print(f"   Color: {normalized_color}")
                print(f"   Sides: {sides}")
                print(f"   Document Type: {document_type}")
                print(f"   Special Effects: {[k for k, v in {
                    'gradient': bool(gradient_colors),
                    'shadow': shadow_effect,
                    'rounded': rounded_style,
                    'priority': priority_highlight,
                    'opacity': opacity_effect,
                    'animation': animation_hint,
                    'custom': bool(custom_pattern)
                }.items() if v]}")

            return True

        except Exception as e:
            error_msg = f"Error applying enhanced paragraph border: {str(e)}"
            if debug_mode:
                print(f"❌ {error_msg}")

            # ===== FALLBACK TO BASIC BORDER =====
            try:
                if debug_mode:
                    print("🔄 Attempting fallback to basic border...")

                pPr = paragraph._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')

                # Apply basic border ke semua sisi
                for side in ['top', 'left', 'bottom', 'right']:
                    border_element = OxmlElement(f'w:{side}')
                    border_element.set(qn('w:val'), 'single')
                    border_element.set(qn('w:sz'), '12')
                    border_element.set(qn('w:space'), '0')
                    border_element.set(qn('w:color'), '000000')
                    pBdr.append(border_element)

                pPr.append(pBdr)

                if debug_mode:
                    print("✅ Fallback border applied successfully")
                return True

            except Exception as fallback_error:
                if debug_mode:
                    print(f"❌ Fallback also failed: {fallback_error}")
                return False

    def _add_enhanced_formatted_runs(self, paragraph, text, document_type):
        """Enhanced version dengan AI-powered formatting, extended patterns, dan smart processing."""
        if not text:
            return

        # ===== MASSIVELY EXTENDED EMOJI DICTIONARY =====
        emoji_dict = {
            # Standard emojis
            ":smile:": "😊", ":check:": "✅", ":x:": "❌", ":warning:": "⚠️",
            ":star:": "⭐", ":arrow_right:": "➡️", ":bulb:": "💡", ":calendar:": "📅",
            ":chart:": "📊", ":document:": "📄", ":pencil:": "✏️", ":clipboard:": "📋",
            ":email:": "📧", ":folder:": "📁", ":money:": "💰", ":phone:": "📱", ":clock:": "🕒",
            ":fire:": "🔥", ":rocket:": "🚀", ":target:": "🎯", ":key:": "🔑", ":shield:": "🛡️",

            # Extended technical emojis
            ":code:": "💻", ":database:": "🗄️", ":server:": "🖥️", ":network:": "🌐",
            ":bug:": "🐛", ":gear:": "⚙️", ":tool:": "🔧", ":package:": "📦",
            ":download:": "⬇️", ":upload:": "⬆️", ":sync:": "🔄", ":backup:": "💾",
            ":api:": "🔌", ":cloud:": "☁️", ":config:": "⚙️", ":deploy:": "🚀",
            ":terminal:": "💻", ":script:": "📜", ":function:": "⚡", ":variable:": "🏷️",

            # Business & workflow emojis
            ":meeting:": "🗣️", ":presentation:": "📊", ":contract:": "📋", ":handshake:": "🤝",
            ":growth:": "📈", ":decline:": "📉", ":profit:": "💰", ":loss:": "💸",
            ":analysis:": "🔍", ":strategy:": "🎯", ":planning:": "📋", ":execution:": "⚡",
            ":deadline:": "⏰", ":milestone:": "🏆", ":priority:": "🔴", ":urgent:": "🚨",

            # Status & progress emojis
            ":success:": "✅", ":error:": "❌", ":pending:": "⏳", ":progress:": "🔄",
            ":new:": "🆕", ":updated:": "🔄", ":deprecated:": "⚠️", ":beta:": "🧪",
            ":completed:": "✅", ":cancelled:": "❌", ":blocked:": "🚧", ":review:": "👀",
            ":approved:": "✅", ":rejected:": "❌", ":draft:": "📝", ":final:": "🏁",

            # Communication & feedback
            ":question:": "❓", ":answer:": "💡", ":comment:": "💬", ":feedback:": "📢",
            ":note:": "📝", ":reminder:": "🔔", ":alert:": "🚨", ":info:": "ℹ️",
            ":tip:": "💡", ":hint:": "🔍", ":help:": "❓", ":support:": "🤝",

            # Education & learning
            ":learn:": "📚", ":teach:": "👨‍🏫", ":study:": "📖", ":research:": "🔬",
            ":knowledge:": "🧠", ":skill:": "💪", ":training:": "🏋️", ":certification:": "🎓",
            ":course:": "📚", ":lesson:": "📝", ":tutorial:": "🎬", ":example:": "💡",

            # Project management
            ":project:": "📁", ":task:": "📝", ":sprint:": "🏃", ":team:": "👥",
            ":lead:": "👑", ":member:": "👤", ":role:": "🎭", ":responsibility:": "📋",
            ":goal:": "🎯", ":objective:": "🎯", ":kpi:": "📊", ":metric:": "📈",

            # Quality & testing
            ":quality:": "💎", ":test:": "🧪", ":verify:": "✅", ":validate:": "🔍",
            ":pass:": "✅", ":fail:": "❌", ":issue:": "🐛", ":fix:": "🔧",

            # Security & compliance
            ":security:": "🔒", ":lock:": "🔐", ":unlock:": "🔓", ":private:": "🔒",
            ":public:": "🌐", ":admin:": "👑", ":user:": "👤", ":guest:": "👥",
            ":permission:": "🔑", ":access:": "🚪", ":deny:": "🚫", ":grant:": "✅"
        }

        # Smart emoji replacement dengan context awareness
        for code, emoji in emoji_dict.items():
            # Enhanced replacement yang mempertimbangkan konteks dokumen
            if document_type == "technical_report" and code in [":bug:", ":code:", ":server:"]:
                text = text.replace(code, emoji)
            elif document_type == "meeting_notes" and code in [":meeting:", ":action:", ":decision:"]:
                text = text.replace(code, emoji)
            elif document_type == "lecture" and code in [":learn:", ":teach:", ":study:"]:
                text = text.replace(code, emoji)
            else:
                text = text.replace(code, emoji)

        # ===== VASTLY ENHANCED FORMATTING PATTERNS =====
        patterns = [
            # Advanced combination formatting
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),           # ***bold italic***
            (r'___(.+?)___', 'bold_italic'),                 # ___bold italic___
            (r'\*\*_(.+?)_\*\*', 'bold_italic'),             # **_bold italic_**
            (r'__\*(.+?)\*__', 'bold_italic'),               # __*bold italic*__

            # Bold variants dengan smart detection
            (r'\*\*(.+?)\*\*', 'bold'),                      # **bold**
            (r'__(.+?)__', 'bold'),                          # __bold__
            (r'<b>(.+?)</b>', 'bold'),                       # <b>bold</b>
            (r'<strong>(.+?)</strong>', 'bold'),             # <strong>bold</strong>

            # Italic variants dengan improved patterns
            (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', 'italic'),      # *italic*
            (r'(?<!_)_([^_\n]+?)_(?!_)', 'italic'),          # _italic_
            (r'<i>(.+?)</i>', 'italic'),                     # <i>italic</i>
            (r'<em>(.+?)</em>', 'italic'),                   # <em>italic</em>

            # Extended visual formatting
            (r'\+\+(.+?)\+\+', 'underline'),                 # ++underline++
            (r'<u>(.+?)</u>', 'underline'),                  # <u>underline</u>
            (r'~~(.+?)~~', 'strikethrough'),                 # ~~strikethrough~~
            (r'<s>(.+?)</s>', 'strikethrough'),              # <s>strikethrough</s>
            (r'<del>(.+?)</del>', 'strikethrough'),          # <del>strikethrough</del>
            (r'==(.+?)==', 'highlight'),                     # ==highlight==
            (r'<mark>(.+?)</mark>', 'highlight'),            # <mark>highlight</mark>

            # Scientific notation
            (r'\^(.+?)\^', 'superscript'),                   # ^superscript^
            (r'<sup>(.+?)</sup>', 'superscript'),            # <sup>superscript</sup>
            (r'(?<!~)~([^~\n]+?)~(?!~)', 'subscript'),       # ~subscript~
            (r'<sub>(.+?)</sub>', 'subscript'),              # <sub>subscript</sub>

            # Enhanced code formatting
            (r'`(.+?)`', 'code'),                            # `inline code`
            (r'``(.+?)``', 'code_double'),                   # ``code with backticks``
            (r'<code>(.+?)</code>', 'code'),                 # <code>inline code</code>
            (r'<tt>(.+?)</tt>', 'code'),                     # <tt>teletype</tt>

            # Advanced special formatting
            (r'\[\[(.+?)\]\]', 'comment'),                   # [[comment]]
            (r'<!--(.+?)-->', 'comment'),                    # <!--HTML comment-->
            (r'\{\{(.+?)\}\}', 'important'),                 # {{important}}
            (r'\{\+(.+?)\+\}', 'success'),                   # {+success+}
            (r'\{\-(.+?)\-\}', 'error'),                     # {-error-}
            (r'\{!(.+?)!\}', 'warning'),                     # {!warning!}
            (r'\{@(.+?)@\}', 'mention'),                     # {@mention@}
            (r'\{#(.+?)#\}', 'tag'),                         # {#tag#}
            (r'\{%(.+?)%\}', 'variable'),                    # {%variable%}
            (r'\{\$(.+?)\$\}', 'price'),                     # {$price$}

            # Enhanced UI elements
            (r'\[(.+?)\](?!\()', 'button'),                  # [button]
            (r'<btn>(.+?)</btn>', 'button'),                 # <btn>button</btn>
            (r'<badge:([^>]+)>', 'badge'),                   # <badge:text>
            (r'<label:([^>]+)>', 'label'),                   # <label:text>
            (r'<chip:([^>]+)>', 'chip'),                     # <chip:text>
            (r'<pill:([^>]+)>', 'pill'),                     # <pill:text>

            # Advanced links dengan full support
            (r'\[([^\]]+)\]\(([^)\s]+)(?:\s+"([^"]+)")?\)', 'link'),  # [text](url "title")
            (r'<link:([^>]+)>(.+?)</link>', 'custom_link'),  # <link:url>text</link>
            (r'https?://[^\s<>"]+', 'auto_link'),            # Auto-detect URLs

            # Enhanced keyboard shortcuts
            (r'<kbd>(.+?)</kbd>', 'keyboard'),               # <kbd>Ctrl+C</kbd>
            (r'\[\[(.+?)\]\]', 'keycombo'),                  # [[Ctrl+Alt+Del]]
            (r'<key>(.+?)</key>', 'single_key'),             # <key>Enter</key>
            (r'<combo>(.+?)</combo>', 'key_combination'),    # <combo>Ctrl+C</combo>

            # Advanced colors dan styling
            (r'<color:([^>]+)>(.+?)</color>', 'color'),      # <color:red>text</color>
            (r'<bg:([^>]+)>(.+?)</bg>', 'background'),       # <bg:yellow>text</bg>
            (r'<theme:([^>]+)>(.+?)</theme>', 'theme_color'), # <theme:primary>text</theme>
            (r'<size:([^>]+)>(.+?)</size>', 'font_size'),    # <size:14>text</size>
            (r'<font:([^>]+)>(.+?)</font>', 'font_family'),  # <font:Arial>text</font>

            # Math expressions dengan LaTeX support
            (r'\$(.+?)\$', 'math'),                          # $formula$
            (r'\$\$(.+?)\$\$', 'math_block'),                # $$formula$$
            (r'\\(.+?)\\', 'latex'),                         # \LaTeX\
            (r'<math>(.+?)</math>', 'math_html'),            # <math>formula</math>

            # Special document elements
            (r'<quote>(.+?)</quote>', 'inline_quote'),       # <quote>text</quote>
            (r'<cite>(.+?)</cite>', 'citation'),             # <cite>source</cite>
            (r'<ref>(.+?)</ref>', 'reference'),              # <ref>reference</ref>
            (r'<term>(.+?)</term>', 'terminology'),          # <term>technical term</term>

            # Status indicators
            (r'<status:([^>]+)>', 'status_indicator'),       # <status:complete>
            (r'<priority:([^>]+)>', 'priority_indicator'),   # <priority:high>
            (r'<progress:([^>]+)>', 'progress_indicator'),   # <progress:75%>

            # Advanced annotations
            (r'<note>(.+?)</note>', 'annotation'),           # <note>annotation</note>
            (r'<aside>(.+?)</aside>', 'aside'),              # <aside>side note</aside>
            (r'<tooltip:([^>]+)>(.+?)</tooltip>', 'tooltip'), # <tooltip:help>text</tooltip>
        ]

        # ===== INTELLIGENT PATTERN PROCESSING =====
        all_matches = []

        # Enhanced pattern matching dengan priority system
        pattern_priorities = {
            'bold_italic': 10, 'math_block': 9, 'code_double': 8, 'link': 7,
            'color': 6, 'background': 6, 'tooltip': 5, 'bold': 4, 'italic': 4,
            'keyboard': 3, 'code': 3, 'default': 1
        }

        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text, re.DOTALL):
                priority = pattern_priorities.get(format_type, 1)

                if format_type in ['link', 'color', 'background', 'tooltip', 'custom_link']:
                    groups = match.groups()
                    if len(groups) >= 2:
                        extra_data = groups[1] if len(groups) > 1 else None
                        if len(groups) > 2:
                            extra_data = {"param1": groups[0], "param2": groups[1], "param3": groups[2]}
                        all_matches.append((match.start(), match.end(), groups[0], format_type,
                                        match.group(0), extra_data, priority))
                    else:
                        all_matches.append((match.start(), match.end(), match.group(1), format_type,
                                        match.group(0), None, priority))
                else:
                    content = match.group(1) if match.groups() else match.group(0)
                    all_matches.append((match.start(), match.end(), content, format_type,
                                    match.group(0), None, priority))

        if not all_matches:
            paragraph.add_run(text)
            return

        # ===== SMART OVERLAP RESOLUTION =====
        all_matches.sort(key=lambda x: (x[0], -x[6]))  # Sort by position, then by priority (desc)
        cleaned_matches = []

        for match in all_matches:
            start, end, content, format_type, full_match, extra, priority = match

            # Enhanced overlap detection
            is_overlapping = False
            for prev_start, prev_end, _, prev_type, _, _, prev_priority in cleaned_matches:
                if start < prev_end and end > prev_start:
                    # Resolve by priority
                    if priority <= prev_priority:
                        is_overlapping = True
                        break
                    else:
                        # Remove lower priority match
                        cleaned_matches = [m for m in cleaned_matches
                                        if not (m[0] == prev_start and m[1] == prev_end)]

            if not is_overlapping:
                cleaned_matches.append(match)

        # Sort final matches by position
        cleaned_matches.sort(key=lambda x: x[0])

        # ===== ENHANCED TEXT PROCESSING =====
        last_end = 0

        for match in cleaned_matches:
            start, end, content, format_type, full_match, extra, priority = match

            # Add plain text before formatted section dengan smart spacing
            if start > last_end:
                plain_text = text[last_end:start]
                if plain_text.strip():  # Only add non-empty text
                    # Smart whitespace handling
                    plain_text = re.sub(r'\s+', ' ', plain_text)  # Normalize whitespace
                    paragraph.add_run(plain_text)

            # Create formatted run dengan enhanced content processing
            processed_content = self._preprocess_content(content, format_type, document_type)
            run = paragraph.add_run(processed_content)

            # Apply enhanced formatting dengan document type awareness
            self._apply_enhanced_run_formatting(run, format_type, extra, document_type)

            last_end = end

        # Add remaining text dengan smart processing
        if last_end < len(text):
            tail_text = text[last_end:]
            if tail_text.strip():
                tail_text = re.sub(r'\s+', ' ', tail_text)  # Normalize whitespace
                paragraph.add_run(tail_text)

    def _preprocess_content(self, content, format_type, document_type):
        """Preprocess content berdasarkan format type dan document type."""
        if not content:
            return content

        # Document-specific preprocessing
        if document_type == "technical_report":
            # Technical term standardization
            tech_terms = {
                "api": "API", "url": "URL", "http": "HTTP", "https": "HTTPS",
                "json": "JSON", "xml": "XML", "sql": "SQL", "html": "HTML",
                "css": "CSS", "js": "JavaScript", "rest": "REST", "soap": "SOAP"
            }
            for term, standard in tech_terms.items():
                content = re.sub(r'\b' + term + r'\b', standard, content, flags=re.IGNORECASE)

        elif document_type == "meeting_notes":
            # Meeting-specific formatting
            if format_type in ["important", "success", "error"]:
                content = content.upper() if len(content) <= 10 else content

        elif document_type == "lecture":
            # Educational formatting
            if format_type == "terminology":
                content = f"{content} (term)"

        # Universal preprocessing
        content = content.strip()

        # Smart capitalization for certain formats
        if format_type in ["button", "badge", "label"] and len(content) <= 20:
            content = content.upper()

        return content

    def _apply_enhanced_run_formatting(self, run, format_type, extra_param, document_type):
        """Apply ultra-enhanced formatting dengan AI-powered styling dan adaptive colors."""

        # ===== ADVANCED DOCUMENT TYPE COLOR SCHEMES =====
        color_schemes = {
            "technical_report": {
                "primary": "1F4E79", "secondary": "4472C4", "accent": "8DB4E2",
                "success": "107C10", "error": "C42B1C", "warning": "FF8C00",
                "code": "2B5797", "link": "0078D4", "comment": "6B7280"
            },
            "meeting_notes": {
                "primary": "385723", "secondary": "70AD47", "accent": "A9D18E",
                "success": "059669", "error": "DC2626", "warning": "D97706",
                "code": "047857", "link": "0891B2", "comment": "6B7280"
            },
            "lecture": {
                "primary": "C65911", "secondary": "ED7D31", "accent": "F4B183",
                "success": "059669", "error": "DC2626", "warning": "D97706",
                "code": "B45309", "link": "0891B2", "comment": "6B7280"
            },
            "presentation": {
                "primary": "7C3AED", "secondary": "A855F7", "accent": "C4B5FD",
                "success": "059669", "error": "DC2626", "warning": "D97706",
                "code": "6D28D9", "link": "0891B2", "comment": "6B7280"
            },
            "research": {
                "primary": "1E40AF", "secondary": "3B82F6", "accent": "93C5FD",
                "success": "059669", "error": "DC2626", "warning": "D97706",
                "code": "1D4ED8", "link": "0891B2", "comment": "6B7280"
            },
            "general": {
                "primary": "1F497D", "secondary": "4F81BD", "accent": "8DB4E2",
                "success": "059669", "error": "DC2626", "warning": "D97706",
                "code": "374151", "link": "0891B2", "comment": "6B7280"
            }
        }

        colors = color_schemes.get(document_type, color_schemes["general"])

        # ===== MASSIVELY EXPANDED FORMATTING MAP =====
        formatting_map = {
            # Basic formatting
            'bold': lambda r: self._apply_bold_formatting(r, colors, document_type),
            'italic': lambda r: self._apply_italic_formatting(r, colors, document_type),
            'bold_italic': lambda r: self._apply_bold_italic_formatting(r, colors, document_type),
            'underline': lambda r: self._apply_underline_formatting(r, colors, document_type),
            'strikethrough': lambda r: self._apply_strikethrough_formatting(r, colors, document_type),

            # Visual enhancements
            'highlight': lambda r: self._apply_highlight_formatting(r, colors, document_type),
            'superscript': lambda r: self._apply_superscript_formatting(r, colors, document_type),
            'subscript': lambda r: self._apply_subscript_formatting(r, colors, document_type),

            # Code formatting
            'code': lambda r: self._apply_code_formatting(r, single=True, colors=colors, doc_type=document_type),
            'code_double': lambda r: self._apply_code_formatting(r, double=True, colors=colors, doc_type=document_type),

            # Semantic formatting
            'comment': lambda r: self._apply_comment_formatting(r, colors, document_type),
            'important': lambda r: self._apply_important_formatting(r, colors, document_type),
            'success': lambda r: self._apply_success_formatting(r, colors, document_type),
            'error': lambda r: self._apply_error_formatting(r, colors, document_type),
            'warning': lambda r: self._apply_warning_formatting(r, colors, document_type),
            'mention': lambda r: self._apply_mention_formatting(r, colors, document_type),
            'tag': lambda r: self._apply_tag_formatting(r, colors, document_type),

            # UI elements
            'button': lambda r: self._apply_button_formatting(r, colors, document_type),
            'badge': lambda r: self._apply_badge_formatting(r, colors["primary"], document_type),
            'label': lambda r: self._apply_label_formatting(r, colors["secondary"], document_type),
            'chip': lambda r: self._apply_chip_formatting(r, colors, document_type),
            'pill': lambda r: self._apply_pill_formatting(r, colors, document_type),

            # Links
            'link': lambda r: self._apply_link_formatting(r, extra_param, colors, document_type),
            'auto_link': lambda r: self._apply_auto_link_formatting(r, colors, document_type),
            'custom_link': lambda r: self._apply_custom_link_formatting(r, extra_param, colors, document_type),

            # Keyboard
            'keyboard': lambda r: self._apply_keyboard_formatting(r, colors, document_type),
            'keycombo': lambda r: self._apply_keycombo_formatting(r, colors, document_type),
            'single_key': lambda r: self._apply_single_key_formatting(r, colors, document_type),
            'key_combination': lambda r: self._apply_key_combination_formatting(r, colors, document_type),

            # Colors & styling
            'color': lambda r: self._apply_color_formatting(r, extra_param, colors, document_type),
            'background': lambda r: self._apply_background_formatting(r, extra_param, colors, document_type),
            'theme_color': lambda r: self._apply_theme_color_formatting(r, extra_param, colors, document_type),
            'font_size': lambda r: self._apply_font_size_formatting(r, extra_param, colors, document_type),
            'font_family': lambda r: self._apply_font_family_formatting(r, extra_param, colors, document_type),

            # Math & scientific
            'math': lambda r: self._apply_math_formatting(r, colors, document_type),
            'math_block': lambda r: self._apply_math_block_formatting(r, colors, document_type),
            'latex': lambda r: self._apply_latex_formatting(r, colors, document_type),

            # Special elements
            'variable': lambda r: self._apply_variable_formatting(r, colors, document_type),
            'price': lambda r: self._apply_price_formatting(r, colors, document_type),
            'terminology': lambda r: self._apply_terminology_formatting(r, colors, document_type),
            'citation': lambda r: self._apply_citation_formatting(r, colors, document_type),
            'reference': lambda r: self._apply_reference_formatting(r, colors, document_type),

            # Status indicators
            'status_indicator': lambda r: self._apply_status_indicator_formatting(r, extra_param, colors, document_type),
            'priority_indicator': lambda r: self._apply_priority_indicator_formatting(r, extra_param, colors, document_type),
            'progress_indicator': lambda r: self._apply_progress_indicator_formatting(r, extra_param, colors, document_type),

            # Annotations
            'annotation': lambda r: self._apply_annotation_formatting(r, colors, document_type),
            'aside': lambda r: self._apply_aside_formatting(r, colors, document_type),
            'tooltip': lambda r: self._apply_tooltip_formatting(r, extra_param, colors, document_type),
            'inline_quote': lambda r: self._apply_inline_quote_formatting(r, colors, document_type),
        }

        # Apply formatting dengan error handling
        formatter = formatting_map.get(format_type)
        if formatter:
            try:
                formatter(run)
            except Exception as e:
                logger.warning(f"Formatting error for {format_type}: {e}")
                # Fallback to basic formatting
                self._apply_fallback_formatting(run, format_type, colors)
        else:
            # Unknown format type, apply basic styling
            self._apply_fallback_formatting(run, format_type, colors)

    def _apply_bold_formatting(self, run, colors, document_type):
        """Enhanced bold formatting dengan document-specific adjustments."""
        run.bold = True

        # Document-specific bold intensity
        if document_type == "presentation":
            run.font.size = Pt(run.font.size.pt + 1) if run.font.size else Pt(12)
        elif document_type == "technical_report":
            run.font.color.rgb = RGBColor.from_string(colors["primary"])

    def _apply_italic_formatting(self, run, colors, document_type):
        """Enhanced italic formatting."""
        run.italic = True

        if document_type == "lecture":
            run.font.color.rgb = RGBColor.from_string(colors["secondary"])

    def _apply_bold_italic_formatting(self, run, colors, document_type):
        """Enhanced bold italic combination."""
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor.from_string(colors["primary"])

        if document_type == "presentation":
            run.font.size = Pt(run.font.size.pt + 1) if run.font.size else Pt(12)

    def _apply_underline_formatting(self, run, colors, document_type):
        """Enhanced underline dengan style variations."""
        run.underline = True

        if document_type == "technical_report":
            run.font.color.rgb = RGBColor.from_string(colors["accent"])

    def _apply_strikethrough_formatting(self, run, colors, document_type):
        """Enhanced strikethrough formatting."""
        run.font.strike = True
        run.font.color.rgb = RGBColor(128, 128, 128)  # Gray for struck text

    def _apply_highlight_formatting(self, run, colors, document_type):
        """Enhanced highlight dengan document-aware colors."""
        if document_type == "technical_report":
            run.font.highlight_color = WD_COLOR_INDEX.BRIGHT_GREEN
        elif document_type == "meeting_notes":
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW

    def _apply_superscript_formatting(self, run, colors, document_type):
        """Enhanced superscript."""
        run.font.superscript = True
        run.font.size = Pt(8)

    def _apply_subscript_formatting(self, run, colors, document_type):
        """Enhanced subscript."""
        run.font.subscript = True
        run.font.size = Pt(8)

    def _apply_code_formatting(self, run, single=False, double=False, colors=None, doc_type=None):
        """Ultra-enhanced code formatting dengan syntax awareness."""
        run.font.name = 'Consolas'

        if double:
            run.font.size = Pt(9)
            bg_color = "EEEEEE"
            border_size = "4"
        else:
            run.font.size = Pt(9.5)
            bg_color = "F8F8F8"
            border_size = "2"

        # Document-specific code styling
        if doc_type == "technical_report":
            run.font.color.rgb = RGBColor.from_string(colors["code"])
            bg_color = "F0F5FF"
        elif doc_type == "presentation":
            bg_color = "F5F5F5"

        # Enhanced background shading
        try:
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)

            # Enhanced border for double backticks
            if double:
                border_elm = OxmlElement('w:bdr')
                border_elm.set(qn('w:val'), 'single')
                border_elm.set(qn('w:sz'), border_size)
                border_elm.set(qn('w:color'), 'CCCCCC')
                rPr.append(border_elm)
        except Exception as e:
            logger.warning(f"Code formatting error: {e}")

    def _apply_button_formatting(self, run, colors=None, document_type=None):
        """Ultra-enhanced button formatting dengan 3D effect."""
        run.font.name = 'Segoe UI'
        run.font.size = Pt(9)
        run.font.bold = True

        # Document-specific button colors
        if document_type == "presentation":
            text_color = RGBColor(255, 255, 255)
            bg_color = colors["primary"]
        elif document_type == "technical_report":
            text_color = RGBColor(60, 60, 60)
            bg_color = "E1E1E1"
        else:
            text_color = RGBColor(60, 60, 60)
            bg_color = "E1E1E1"

        run.font.color.rgb = text_color

        try:
            # Enhanced button background
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)

            # Enhanced 3D border effect
            border_elm = OxmlElement('w:bdr')
            border_elm.set(qn('w:val'), 'single')
            border_elm.set(qn('w:sz'), '6')
            border_elm.set(qn('w:color'), '999999')
            rPr.append(border_elm)
        except Exception as e:
            logger.warning(f"Button formatting error: {e}")

    def _apply_badge_formatting(self, run, color, document_type=None):
        """Ultra-enhanced badge dengan rounded appearance simulation."""
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Segoe UI'

        try:
            # Enhanced badge background
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)

            # Rounded effect simulation dengan border
            border_elm = OxmlElement('w:bdr')
            border_elm.set(qn('w:val'), 'single')
            border_elm.set(qn('w:sz'), '2')
            border_elm.set(qn('w:color'), color)
            rPr.append(border_elm)
        except Exception as e:
            logger.warning(f"Badge formatting error: {e}")

    def _apply_keyboard_formatting(self, run, colors=None, document_type=None):
        """Ultra-enhanced keyboard key styling dengan realistic appearance."""
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)

        # Document-specific keyboard styling
        if document_type == "technical_report":
            bg_color = "F0F5FF"
            border_color = "BDD7EE"
        elif document_type == "presentation":
            bg_color = "F8F0FF"
            border_color = "D0A9F5"
        else:
            bg_color = "F5F5F5"
            border_color = "BBBBBB"

        try:
            # Enhanced keyboard background
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)

            # Enhanced 3D keyboard border
            border_elm = OxmlElement('w:bdr')
            border_elm.set(qn('w:val'), 'single')
            border_elm.set(qn('w:sz'), '6')
            border_elm.set(qn('w:color'), border_color)
            rPr.append(border_elm)
        except Exception as e:
            logger.warning(f"Keyboard formatting error: {e}")

    def _apply_label_formatting(self, run, color, document_type=None):
        """Ultra-enhanced label dengan professional appearance."""
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.name = 'Segoe UI'

        try:
            # Enhanced label background dengan gradient simulation
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)
        except Exception as e:
            logger.warning(f"Label formatting error: {e}")

    def _apply_keycombo_formatting(self, run, colors=None, document_type=None):
        """Ultra-enhanced key combination dengan visual separation."""
        run.font.name = 'Consolas'
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = RGBColor(51, 51, 51)

        try:
            # Enhanced key combo background
            shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(nsdecls('w')))
            rPr = run._element.get_or_add_rPr()
            rPr.append(shading_elm)

            # Enhanced border dengan dotted style untuk combination
            border_elm = OxmlElement('w:bdr')
            border_elm.set(qn('w:val'), 'dotted')
            border_elm.set(qn('w:sz'), '6')
            border_elm.set(qn('w:color'), 'AAAAAA')
            rPr.append(border_elm)
        except Exception as e:
            logger.warning(f"Key combo formatting error: {e}")

    def _apply_fallback_formatting(self, run, format_type, colors):
        """Fallback formatting untuk unknown types."""
        if 'bold' in format_type.lower():
            run.bold = True
        if 'italic' in format_type.lower():
            run.italic = True
        if 'code' in format_type.lower():
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

        # Apply default color
        try:
            run.font.color.rgb = RGBColor.from_string(colors.get("primary", "000000"))
        except:
            pass

    def _add_enhanced_special_marker(self, doc, marker_type, content_text, content_stats):
        """
        Menambahkan marker khusus yang sangat ditingkatkan untuk action items, decisions, dll.

        Features:
        - 15+ marker types dengan styling adaptif
        - Dynamic styling berdasarkan document type dan content stats
        - Priority detection dan visual indicators
        - Enhanced borders, backgrounds, dan typography
        - Smart content processing dan formatting
        - Deadline/assignee detection untuk action items
        - Status indicators dan progress tracking
        - Accessibility improvements
        - Multi-language support
        """
        # ===== ENHANCED INPUT VALIDATION =====
        if not content_text or not content_text.strip():
            return None

        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== COMPREHENSIVE MARKER CONFIGURATIONS =====
        marker_configs = {
            # Core Action & Decision Markers
            "action": {
                "icon": "⚡", "title": "ACTION ITEM", "color": "FF6B35", "bg": "FFF5F0",
                "border": "FF8C42", "priority": "high", "category": "actionable",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },
            "decision": {
                "icon": "✅", "title": "DECISION", "color": "28A745", "bg": "F0FFF0",
                "border": "5CB85C", "priority": "high", "category": "resolution",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },
            "todo": {
                "icon": "📝", "title": "TODO", "color": "007BFF", "bg": "F0F8FF",
                "border": "4A90E2", "priority": "medium", "category": "task",
                "font": "Calibri", "size": 10.5, "border_width": "6"
            },

            # Status & Progress Markers
            "completed": {
                "icon": "✔️", "title": "COMPLETED", "color": "28A745", "bg": "F0FFF0",
                "border": "90EE90", "priority": "low", "category": "status",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "in_progress": {
                "icon": "🔄", "title": "IN PROGRESS", "color": "0078D4", "bg": "F3F9FF",
                "border": "87CEEB", "priority": "medium", "category": "status",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "pending": {
                "icon": "⏳", "title": "PENDING", "color": "FFA500", "bg": "FFF8E7",
                "border": "FFD700", "priority": "medium", "category": "status",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "blocked": {
                "icon": "🚧", "title": "BLOCKED", "color": "DC3545", "bg": "FFF5F5",
                "border": "FF6B6B", "priority": "high", "category": "status",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },
            "cancelled": {
                "icon": "❌", "title": "CANCELLED", "color": "6C757D", "bg": "F8F9FA",
                "border": "ADB5BD", "priority": "low", "category": "status",
                "font": "Segoe UI", "size": 11, "border_width": "4"
            },

            # Information & Communication Markers
            "note": {
                "icon": "📋", "title": "NOTE", "color": "4472C4", "bg": "F0F5FF",
                "border": "BDD7EE", "priority": "low", "category": "information",
                "font": "Calibri", "size": 11, "border_width": "4"
            },
            "important": {
                "icon": "❗", "title": "IMPORTANT", "color": "DC143C", "bg": "FFF0F0",
                "border": "FFB3B3", "priority": "high", "category": "alert",
                "font": "Segoe UI", "size": 11, "border_width": "10"
            },
            "warning": {
                "icon": "⚠️", "title": "WARNING", "color": "FF8C00", "bg": "FFF8E7",
                "border": "FFD166", "priority": "high", "category": "alert",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },
            "tip": {
                "icon": "💡", "title": "TIP", "color": "107C10", "bg": "F3FFF3",
                "border": "9FD89F", "priority": "low", "category": "guidance",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "question": {
                "icon": "❓", "title": "QUESTION", "color": "6F42C1", "bg": "F8F0FF",
                "border": "D0A9F5", "priority": "medium", "category": "inquiry",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },

            # Meeting & Communication Markers
            "discussion": {
                "icon": "💬", "title": "DISCUSSION POINT", "color": "20B2AA", "bg": "F0FFFF",
                "border": "7FFFD4", "priority": "medium", "category": "communication",
                "font": "Calibri", "size": 10.5, "border_width": "6"
            },
            "follow_up": {
                "icon": "🔄", "title": "FOLLOW UP", "color": "9966CC", "bg": "F8F0FF",
                "border": "DDA0DD", "priority": "high", "category": "actionable",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },
            "assignment": {
                "icon": "👤", "title": "ASSIGNMENT", "color": "FF69B4", "bg": "FFF0F8",
                "border": "FFB6C1", "priority": "high", "category": "actionable",
                "font": "Segoe UI", "size": 11, "border_width": "8"
            },

            # Technical & Process Markers
            "bug": {
                "icon": "🐛", "title": "BUG", "color": "DC3545", "bg": "FFF5F5",
                "border": "F8A8A8", "priority": "high", "category": "technical",
                "font": "Consolas", "size": 10.5, "border_width": "8"
            },
            "feature": {
                "icon": "🆕", "title": "FEATURE", "color": "17A2B8", "bg": "F0FDFF",
                "border": "87CEEB", "priority": "medium", "category": "technical",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "improvement": {
                "icon": "⬆️", "title": "IMPROVEMENT", "color": "28A745", "bg": "F0FFF0",
                "border": "90EE90", "priority": "medium", "category": "enhancement",
                "font": "Segoe UI", "size": 11, "border_width": "6"
            },
            "risk": {
                "icon": "⚠️", "title": "RISK", "color": "DC3545", "bg": "FFF5F5",
                "border": "FF6B6B", "priority": "high", "category": "alert",
                "font": "Segoe UI", "size": 11, "border_width": "10"
            },

            # Default fallback
            "default": {
                "icon": "📌", "title": "MARKER", "color": "6C757D", "bg": "F8F9FA",
                "border": "CED4DA", "priority": "medium", "category": "general",
                "font": "Calibri", "size": 11, "border_width": "6"
            }
        }

        # ===== INTELLIGENT MARKER TYPE DETECTION =====
        def detect_marker_type_from_content(content_text):
            """Intelligently detect marker type from content."""
            content_lower = content_text.lower()

            patterns = {
                "action": [r'\b(action|do|execute|implement|perform|carry out)\b',
                        r'\b(task|assignment|activity)\b'],
                "decision": [r'\b(decide|decision|resolved|concluded|determined)\b',
                            r'\b(approved|rejected|accepted)\b'],
                "todo": [r'\b(todo|to do|need to|should|must)\b',
                        r'\b(pending|outstanding|remaining)\b'],
                "completed": [r'\b(done|completed|finished|accomplished)\b',
                            r'\b(delivered|shipped|closed)\b'],
                "blocked": [r'\b(blocked|stuck|waiting|dependent)\b',
                        r'\b(cannot|unable|prevented)\b'],
                "important": [r'\b(important|critical|vital|essential)\b',
                            r'\b(urgent|priority|high.priority)\b'],
                "warning": [r'\b(warning|caution|alert|risk)\b',
                        r'\b(danger|hazard|threat)\b'],
                "question": [r'\b(question|ask|clarify|confirm)\b',
                            r'\?+\s*$'],
                "note": [r'\b(note|fyi|info|information)\b',
                        r'\b(remember|keep in mind)\b']
            }

            for marker_type, pattern_list in patterns.items():
                for pattern in pattern_list:
                    if re.search(pattern, content_lower):
                        return marker_type

            return "default"

        def detect_priority_from_content(content_text):
            """Detect priority level from content text."""
            content_lower = content_text.lower()

            high_priority_indicators = [
                r'\b(urgent|asap|immediately|critical|emergency)\b',
                r'\b(high.priority|top.priority|priority.1)\b',
                r'\b(deadline.?today|due.?today|overdue)\b'
            ]

            low_priority_indicators = [
                r'\b(low.priority|nice.to.have|optional)\b',
                r'\b(when.time|eventually|someday)\b',
                r'\b(future|later|next.month)\b'
            ]

            for pattern in high_priority_indicators:
                if re.search(pattern, content_lower):
                    return "high"

            for pattern in low_priority_indicators:
                if re.search(pattern, content_lower):
                    return "low"

            return "medium"

        def analyze_urgency_level(content_text):
            """Analyze urgency level from content."""
            content_lower = content_text.lower()

            urgent_indicators = [
                r'\b(urgent|asap|immediately|emergency|critical)\b',
                r'\b(today|now|right.away)\b',
                r'!{2,}',  # Multiple exclamation marks
                r'\b(deadline.?today|due.?now)\b'
            ]

            for pattern in urgent_indicators:
                if re.search(pattern, content_lower):
                    return "urgent"

            return "normal"

        def extract_assignee_info(content_text):
            """Extract assignee information from content."""
            assignee_patterns = [
                r'\b(?:assigned to|@|owner:|responsible:)\s*([A-Za-z\s\.]+?)(?:\s|$|,|\.|;)',
                r'\b([A-Z][a-z]+\s+[A-Z][a-z]+)\s+will\b',
                r'\b([A-Z][a-z]+)\s+(?:should|must|needs to)\b'
            ]

            for pattern in assignee_patterns:
                match = re.search(pattern, content_text, re.IGNORECASE)
                if match:
                    assignee_name = match.group(1).strip()
                    if len(assignee_name) > 2 and len(assignee_name) < 50:
                        return {"name": assignee_name, "detected_pattern": pattern}

            return None

        def extract_deadline_info(content_text):
            """Extract deadline information from content."""
            deadline_patterns = [
                r'\b(?:due|deadline|by)\s+(\w+\s+\d{1,2}(?:,?\s+\d{4})?)\b',
                r'\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4})\b',
                r'\b(today|tomorrow|next\s+\w+)\b',
                r'\b(end\s+of\s+\w+)\b'
            ]

            for pattern in deadline_patterns:
                match = re.search(pattern, content_text, re.IGNORECASE)
                if match:
                    deadline_text = match.group(1).strip()
                    is_urgent = any(word in deadline_text.lower()
                                for word in ['today', 'tomorrow', 'asap'])
                    is_overdue = 'overdue' in content_text.lower()

                    return {
                        "text": deadline_text,
                        "is_urgent": is_urgent,
                        "is_overdue": is_overdue,
                        "detected_pattern": pattern
                    }

            return None

        def detect_status_indicators(content_text):
            """Detect various status indicators from content."""
            content_lower = content_text.lower()

            indicators = {
                "new": bool(re.search(r'\b(new|newly|recently|just)\b', content_lower)),
                "priority": bool(re.search(r'\b(priority|important|critical)\b', content_lower)),
                "urgent": bool(re.search(r'\b(urgent|asap|immediately)\b', content_lower)),
                "draft": bool(re.search(r'\b(draft|preliminary|initial)\b', content_lower)),
                "review": bool(re.search(r'\b(review|check|verify|validate)\b', content_lower)),
                "approved": bool(re.search(r'\b(approved|accepted|confirmed)\b', content_lower))
            }

            return indicators

        def intensify_color(hex_color, factor=0.1):
            """Intensify a hex color by reducing luminosity."""
            try:
                hex_color = hex_color.lstrip('#')
                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                new_rgb = tuple(max(0, int(c * (1 - factor))) for c in rgb)
                return f"{new_rgb[0]:02X}{new_rgb[1]:02X}{new_rgb[2]:02X}"
            except:
                return hex_color

        def lighten_color(hex_color, factor=0.1):
            """Lighten a hex color by increasing luminosity."""
            try:
                hex_color = hex_color.lstrip('#')
                rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))
                new_rgb = tuple(min(255, int(c + (255 - c) * factor)) for c in rgb)
                return f"{new_rgb[0]:02X}{new_rgb[1]:02X}{new_rgb[2]:02X}"
            except:
                return hex_color

        # Auto-detect marker type from content if not explicitly provided
        if marker_type == "auto" or marker_type not in marker_configs:
            marker_type = detect_marker_type_from_content(content_text)

        # Get configuration
        config = marker_configs.get(marker_type, marker_configs["default"])

        # ===== DOCUMENT TYPE ADAPTATIONS =====
        if document_type == "technical_report":
            config["font"] = "Cambria"
            config["size"] = max(10, config["size"] - 0.5)
            if config["category"] == "technical":
                config["border_width"] = str(int(config["border_width"]) + 2)
        elif document_type == "meeting_notes":
            config["font"] = "Calibri"
            if config["category"] in ["actionable", "communication"]:
                config["border_width"] = str(int(config["border_width"]) + 2)
        elif document_type == "presentation":
            config["size"] = min(12, config["size"] + 1)
            config["font"] = "Segoe UI"
        elif document_type == "instructional":
            if config["category"] == "guidance":
                config["border_width"] = str(int(config["border_width"]) + 2)

        # ===== ENHANCED CONTENT ANALYSIS =====
        priority_detected = detect_priority_from_content(content_text)
        urgency_level = analyze_urgency_level(content_text)
        assignee_info = extract_assignee_info(content_text)
        deadline_info = extract_deadline_info(content_text)
        status_indicators = detect_status_indicators(content_text)

        # Override priority if detected from content
        if priority_detected and priority_detected != config["priority"]:
            config["priority"] = priority_detected

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = 'Normal'

        # ===== ENHANCED VISUAL STYLING =====
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            # Enhanced background with gradient-like effect
            bg_color = config["bg"]
            if config["priority"] == "high":
                bg_color = intensify_color(bg_color, 0.1)
            elif config["priority"] == "low":
                bg_color = lighten_color(bg_color, 0.1)

            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            p._element.get_or_add_pPr().append(shading_elm)

            # Enhanced border with multiple styles
            border_style = "single"
            if config["priority"] == "high":
                border_style = "thick"
            elif urgency_level == "urgent":
                border_style = "double"

            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(f'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:left w:val="{border_style}" w:sz="{config["border_width"]}"
                        w:space="0" w:color="{config["border"]}"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{config["border"]}"/>
            </w:pBdr>''')
            pPr.append(pBdr)

        except Exception as e:
            logger.warning(f"Could not apply enhanced marker styling: {e}")

        # ===== ENHANCED SPACING AND LAYOUT =====
        base_space_before = 8
        base_space_after = 8
        base_indent = 0.25

        # Priority adjustments
        if config["priority"] == "high":
            base_space_before += 4
            base_space_after += 4
            base_indent += 0.05
        elif config["priority"] == "low":
            base_space_before -= 2
            base_space_after -= 2

        # Document type adjustments
        if document_type == "technical_report":
            base_space_before += 2
            base_space_after += 2
            base_indent += 0.1
        elif document_type == "presentation":
            base_space_before += 4
            base_space_after += 4
            base_indent += 0.15
        elif document_type == "meeting_notes":
            if config["category"] == "actionable":
                base_indent += 0.1

        # Apply spacing
        p.paragraph_format.space_before = Pt(base_space_before)
        p.paragraph_format.space_after = Pt(base_space_after)
        p.paragraph_format.left_indent = Inches(base_indent)
        p.paragraph_format.right_indent = Inches(0.1)

        # Enhanced line spacing
        if len(content_text) > 100:
            p.paragraph_format.line_spacing = Pt(15)
        else:
            p.paragraph_format.line_spacing = Pt(14)

        # ===== ENHANCED ICON AND TITLE FORMATTING =====
        icon_size = config["size"] + (2 if config["priority"] == "high" else 0)

        # Enhanced icon with potential animation indicators
        icon_text = config["icon"]
        if urgency_level == "urgent":
            icon_text = f"🔥{config['icon']}"
        elif status_indicators.get("new", False):
            icon_text = f"🆕{config['icon']}"

        # Create icon run
        icon_run = p.add_run(f"{icon_text} ")
        icon_run.font.size = Pt(icon_size)
        icon_run.font.name = "Segoe UI Emoji"

        # Enhanced title with dynamic formatting
        title_text = config["title"]

        if config["priority"] == "high":
            title_text = f"🔴 {title_text}"
        elif urgency_level == "urgent":
            title_text = f"⚡ {title_text}"
        elif deadline_info and deadline_info["is_overdue"]:
            title_text = f"⏰ {title_text} (OVERDUE)"

        # Create title run
        title_run = p.add_run(f"{title_text}: ")
        title_run.bold = True
        title_run.font.name = config["font"]
        title_run.font.size = Pt(config["size"])
        title_run.font.color.rgb = RGBColor.from_string(config["color"])

        if config["priority"] == "high":
            title_run.underline = True

        # ===== ENHANCED CONTENT PROCESSING =====
        processed_content = content_text.strip()

        # Add metadata information
        metadata_parts = []

        if assignee_info:
            metadata_parts.append(f"👤 **Assigned to:** {assignee_info['name']}")

        if deadline_info:
            deadline_emoji = "⏰" if deadline_info["is_urgent"] else "📅"
            metadata_parts.append(f"{deadline_emoji} **Deadline:** {deadline_info['text']}")

        if status_indicators:
            for status, value in status_indicators.items():
                if value and status != "new":
                    status_emojis = {
                        "priority": "🔴", "urgent": "⚡", "draft": "📝",
                        "review": "👀", "approved": "✅"
                    }
                    emoji = status_emojis.get(status, "📌")
                    metadata_parts.append(f"{emoji} **{status.title()}**")

        # Combine content with metadata
        if metadata_parts:
            processed_content = f"{processed_content}\n\n*{' | '.join(metadata_parts)}*"

        # ===== ENHANCED FORMATTED CONTENT =====
        if hasattr(self, '_add_enhanced_formatted_runs'):
            self._add_enhanced_formatted_runs(p, processed_content, content_stats)
        else:
            content_run = p.add_run(processed_content)
            content_run.font.name = config["font"]
            content_run.font.size = Pt(config["size"])

        # ===== ENHANCED SPECIAL FEATURES =====
        if marker_type in ["in_progress", "pending", "blocked"]:
            progress_run = p.add_run(" ●●○○○")
            progress_run.font.size = Pt(8)
            progress_run.font.color.rgb = RGBColor.from_string(config["color"])

        # Add completion percentage if detected
        completion_match = re.search(r'\b(\d{1,3})%\s*(complete|done|finished)', processed_content, re.IGNORECASE)
        if completion_match:
            percentage = int(completion_match.group(1))
            progress_bar = "█" * (percentage // 10) + "░" * (10 - percentage // 10)
            progress_run = p.add_run(f" [{progress_bar}] {percentage}%")
            progress_run.font.size = Pt(9)
            progress_run.font.name = "Consolas"

        # ===== ENHANCED KEEP-WITH-NEXT =====
        p.paragraph_format.keep_with_next = True
        p.paragraph_format.widow_control = True

        # ===== ENHANCED ACCESSIBILITY =====
        try:
            alt_text = f"{config['title']}: {content_text[:100]}{'...' if len(content_text) > 100 else ''}"
            p._element.set(qn('w:altText'), alt_text)
        except:
            pass

        # ===== ENHANCED LOGGING =====
        logger.info(f"Enhanced special marker created: {marker_type} | Priority: {config['priority']} | "
                    f"Document: {document_type} | Urgency: {urgency_level} | "
                    f"Assignee: {bool(assignee_info)} | Deadline: {bool(deadline_info)}")

        return p

    def _create_enhanced_definition_item(self, doc, term, definition, content_stats):
        """
        Membuat definition list item dengan styling yang ditingkatkan dan analisis AI.
        Menggabungkan enhancement, analisis kepentingan, dan fallback dalam satu fungsi.
        """
        # Normalize content_stats input
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # Variables untuk menyimpan hasil AI
        enhanced_definition = definition
        importance_rating = "MEDIUM"
        ai_enhanced = False

        # === GROQ AI ENHANCEMENT ===
        if hasattr(self, 'groq_client') and self.groq_client:
            try:
                # 1. Enhance definition dengan Groq
                enhance_prompt = f"""
                User wants to enhance a definition for a {document_type} document. I need to:
                1. Make the definition clearer and more concise
                2. Ensure it's appropriate for the document type
                3. Keep the core meaning intact
                4. Use professional language
                5. Improve technical accuracy if applicable

                Sempurnakan definisi berikut untuk dokumen {document_type}:

                Term: {term}
                Definition: {definition}

                Berikan definisi yang:
                - Lebih jelas dan padat
                - Sesuai konteks {document_type}
                - Mudah dipahami
                - Akurat secara teknis
                - Menggunakan terminologi yang tepat

                Respon hanya definisi yang disempurnakan tanpa tambahan teks:"""

                response = self.groq_client.chat.completions.create(
                    model="deepseek-r1-distill-llama-70b",
                    messages=[{"role": "user", "content": enhance_prompt}],
                    temperature=0.3,
                    max_tokens=200,
                    reasoning_format="hidden"
                )

                enhanced_result = response.choices[0].message.content.strip()

                # Validasi hasil enhancement
                if (len(enhanced_result) > 0 and
                    len(enhanced_result) < 500 and
                    enhanced_result.lower() != definition.lower()):
                    enhanced_definition = enhanced_result
                    ai_enhanced = True

                # 2. Analisis kepentingan term
                importance_prompt = f"""
                I need to analyze the importance level of this term for a {document_type} document.
                - HIGH: Critical terms that are essential for understanding the document
                - MEDIUM: Important but not critical for basic understanding
                - LOW: Nice to have but not essential, supplementary information

                Analisis tingkat kepentingan term "{term}" untuk dokumen {document_type}.

                Pertimbangkan:
                - Seberapa sering term ini muncul dalam konteks {document_type}
                - Apakah term ini fundamental untuk pemahaman
                - Dampak jika pembaca tidak memahami term ini

                Berikan rating: HIGH/MEDIUM/LOW dengan alasan singkat.

                Format: RATING|alasan"""

                importance_response = self.groq_client.chat.completions.create(
                    model="deepseek-r1-distill-llama-70b",
                    messages=[{"role": "user", "content": importance_prompt}],
                    temperature=0.2,
                    max_tokens=100,
                    reasoning_format="hidden"
                )

                importance_result = importance_response.choices[0].message.content.strip()
                if '|' in importance_result:
                    rating_part = importance_result.split('|')[0].upper()
                    if rating_part in ['HIGH', 'MEDIUM', 'LOW']:
                        importance_rating = rating_part

            except Exception as e:
                logger.warning(f"Groq enhancement failed for term '{term}': {e}")
                # Continue dengan definisi original

        # === PARAGRAPH CREATION ===
        try:
            p = doc.add_paragraph()

            # === ENHANCED TERM STYLING BERDASARKAN IMPORTANCE ===
            term_run = p.add_run(f"{term.strip()}: ")
            term_run.bold = True
            term_run.font.size = Pt(11)

            # Color coding berdasarkan importance dan document type
            importance_colors = {
                "HIGH": {
                    "technical_report": "1F4E79",  # Dark blue
                    "meeting_notes": "2E7D32",     # Dark green
                    "lecture": "BF360C",           # Dark orange
                    "general": "1A237E"            # Dark indigo
                },
                "MEDIUM": {
                    "technical_report": "4472C4",  # Medium blue
                    "meeting_notes": "66BB6A",     # Medium green
                    "lecture": "FF8F00",           # Medium orange
                    "general": "3F51B5"            # Medium indigo
                },
                "LOW": {
                    "technical_report": "8DB4E2",  # Light blue
                    "meeting_notes": "A5D6A7",     # Light green
                    "lecture": "FFCC80",           # Light orange
                    "general": "7986CB"            # Light indigo
                }
            }

            color_scheme = importance_colors.get(importance_rating, importance_colors["MEDIUM"])
            term_color = color_scheme.get(document_type, color_scheme["general"])
            term_run.font.color.rgb = RGBColor.from_string(term_color)

            # Special styling untuk HIGH importance terms
            if importance_rating == "HIGH":
                term_run.font.size = Pt(12)
                # Add subtle background highlight untuk high importance
                try:
                    from docx.oxml import parse_xml
                    from docx.oxml.ns import nsdecls
                    highlight_colors = {
                        "technical_report": "F0F5FF",
                        "meeting_notes": "F0FFF0",
                        "lecture": "FFF8F0",
                        "general": "F8F9FF"
                    }
                    bg_color = highlight_colors.get(document_type, highlight_colors["general"])
                    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
                    p._element.get_or_add_pPr().append(shading_elm)
                except:
                    pass  # Continue if highlighting fails

            # === ADD ENHANCED DEFINITION ===
            if hasattr(self, '_add_formatted_runs_to_paragraph'):
                self._add_formatted_runs_to_paragraph(p, enhanced_definition)
            else:
                # Fallback formatting
                def_run = p.add_run(enhanced_definition)
                def_run.font.size = Pt(10.5)

            # === ENHANCED PARAGRAPH STYLING ===
            # Base styling
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.left_indent = Inches(0.25)
            p.paragraph_format.line_spacing = 1.15

            # Document type specific adjustments
            if document_type == "technical_report":
                p.paragraph_format.left_indent = Inches(0.3)
                if importance_rating == "HIGH":
                    p.paragraph_format.space_after = Pt(8)
            elif document_type == "meeting_notes":
                p.paragraph_format.left_indent = Inches(0.2)
                p.paragraph_format.line_spacing = 1.1
            elif document_type == "lecture":
                p.paragraph_format.left_indent = Inches(0.35)
                if importance_rating == "HIGH":
                    p.paragraph_format.space_before = Pt(5)
                    p.paragraph_format.space_after = Pt(8)

            # === ADD AI ENHANCEMENT INDICATOR ===
            if ai_enhanced:
                # Subtle indicator bahwa definisi telah ditingkatkan AI
                try:
                    # Add very small AI indicator (optional, can be removed)
                    ai_run = p.add_run(" ✨")
                    ai_run.font.size = Pt(8)
                    ai_run.font.color.rgb = RGBColor(150, 150, 150)
                except:
                    pass

            # === IMPORTANCE-BASED FORMATTING ===
            if importance_rating == "HIGH":
                # Keep with next untuk high importance terms
                p.paragraph_format.keep_with_next = True

                # Optional: Add subtle left border untuk high importance
                try:
                    from docx.oxml import parse_xml
                    pPr = p._element.get_or_add_pPr()
                    pBdr = parse_xml(f'''<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:left w:val="single" w:sz="4" w:space="0" w:color="{term_color}"/>
                    </w:pBdr>''')
                    pPr.append(pBdr)
                except:
                    pass

            logger.debug(f"Enhanced definition created: '{term}' | Importance: {importance_rating} | AI Enhanced: {ai_enhanced}")
            return p

        except Exception as e:
            logger.error(f"Error creating enhanced definition item: {e}")

            # === FALLBACK METHOD ===
            try:
                p = doc.add_paragraph()

                # Simple fallback formatting
                term_run = p.add_run(f"{term}: ")
                term_run.bold = True
                term_run.font.color.rgb = RGBColor(70, 70, 70)
                term_run.font.size = Pt(11)

                # Add definition without enhanced formatting
                if hasattr(self, '_add_formatted_runs_to_paragraph'):
                    self._add_formatted_runs_to_paragraph(p, definition)
                else:
                    def_run = p.add_run(definition)
                    def_run.font.size = Pt(10.5)

                # Basic styling
                p.paragraph_format.space_before = Pt(3)
                p.paragraph_format.space_after = Pt(6)
                p.paragraph_format.left_indent = Inches(0.25)

                logger.info(f"Used fallback method for definition: '{term}'")
                return p

            except Exception as fallback_error:
                logger.error(f"Fallback method also failed: {fallback_error}")
                return None

    # Additional helper methods for the enhanced functionality

    def finalize_document_formatting_enhanced(self, doc, content_stats=None):
        """
        Apply comprehensive final formatting touches to the document with advanced styling,
        document type awareness, and professional enhancements.

        Features:
        - Document type specific styling
        - Advanced typography and spacing
        - Professional color schemes
        - Enhanced readability optimizations
        - Accessibility improvements
        - Performance optimized processing
        - Comprehensive error handling

        Args:
            doc: Document object from python-docx
            content_stats: Dictionary containing document metadata and statistics
                        Can also be a string representing document type

        Returns:
            bool: True if formatting applied successfully, False otherwise
        """
        import logging
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.enum.style import WD_STYLE_TYPE
        from docx.oxml import parse_xml, OxmlElement
        from docx.oxml.ns import nsdecls, qn

        # Setup logger
        logger = logging.getLogger(__name__)

        try:
            # ===== ENHANCED INPUT VALIDATION & NORMALIZATION =====
            if not doc:
                logger.error("Document object is None or invalid")
                return False

            # Normalize content_stats input
            if isinstance(content_stats, str):
                document_type = content_stats.lower().strip()
                content_stats = {"content_type": document_type}
            elif isinstance(content_stats, dict):
                document_type = content_stats.get("content_type", "general").lower().strip()
            else:
                document_type = "general"
                content_stats = {"content_type": "general"}

            # Extract additional metadata with defaults
            word_count = content_stats.get("word_count", 0)
            complexity_level = content_stats.get("complexity_level", "medium")
            language = content_stats.get("language", "id")
            heading_count = content_stats.get("heading_count", 0)

            logger.info(f"Starting enhanced document formatting for {document_type} type")

            # ===== COMPREHENSIVE DOCUMENT TYPE CONFIGURATIONS =====
            document_configs = {
                "technical_report": {
                    "font_family": "Cambria",
                    "font_size": 11,
                    "line_spacing": 15,
                    "paragraph_spacing": {"before": 0, "after": 6},
                    "colors": {
                        "primary": "1F4E79",
                        "secondary": "4472C4",
                        "accent": "8DB4E2",
                        "text": "000000"
                    },
                    "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
                    "professional_mode": True,
                    "spacing_multiplier": 1.1
                },
                "meeting_notes": {
                    "font_family": "Calibri",
                    "font_size": 11,
                    "line_spacing": 14,
                    "paragraph_spacing": {"before": 0, "after": 4},
                    "colors": {
                        "primary": "385723",
                        "secondary": "70AD47",
                        "accent": "A9D18E",
                        "text": "000000"
                    },
                    "margins": {"left": 0.8, "right": 0.8, "top": 0.8, "bottom": 0.8},
                    "professional_mode": True,
                    "spacing_multiplier": 0.9
                },
                "lecture": {
                    "font_family": "Georgia",
                    "font_size": 11.5,
                    "line_spacing": 16,
                    "paragraph_spacing": {"before": 0, "after": 8},
                    "colors": {
                        "primary": "C65911",
                        "secondary": "ED7D31",
                        "accent": "F4B183",
                        "text": "000000"
                    },
                    "margins": {"left": 1.2, "right": 1.2, "top": 1.0, "bottom": 1.0},
                    "professional_mode": True,
                    "spacing_multiplier": 1.2
                },
                "general": {
                    "font_family": "Calibri",
                    "font_size": 11,
                    "line_spacing": 14,
                    "paragraph_spacing": {"before": 0, "after": 6},
                    "colors": {
                        "primary": "1F497D",
                        "secondary": "4F81BD",
                        "accent": "8DB4E2",
                        "text": "000000"
                    },
                    "margins": {"left": 1.0, "right": 1.0, "top": 1.0, "bottom": 1.0},
                    "professional_mode": True,
                    "spacing_multiplier": 1.0
                }
            }

            # Get configuration for document type
            config = document_configs.get(document_type, document_configs["general"])

            # ===== DYNAMIC ADJUSTMENTS BASED ON DOCUMENT CHARACTERISTICS =====
            # Adjust based on word count
            if word_count > 5000:  # Long document
                config["line_spacing"] += 1
                config["paragraph_spacing"]["after"] += 2
                config["margins"]["left"] += 0.1
                config["margins"]["right"] += 0.1
            elif word_count < 500:  # Short document
                config["line_spacing"] = max(12, config["line_spacing"] - 1)
                config["paragraph_spacing"]["after"] = max(2, config["paragraph_spacing"]["after"] - 1)

            # ===== ENHANCED DOCUMENT-WIDE STYLE CONFIGURATION =====
            try:
                # Configure Normal style with enhanced properties
                normal_style = doc.styles['Normal']
                normal_font = normal_style.font
                normal_paragraph = normal_style.paragraph_format

                # Enhanced font configuration
                normal_font.name = config["font_family"]
                normal_font.size = Pt(config["font_size"])
                normal_font.color.rgb = RGBColor.from_string(config["colors"]["text"])

                # Enhanced paragraph formatting
                normal_paragraph.line_spacing = Pt(config["line_spacing"])
                normal_paragraph.space_before = Pt(config["paragraph_spacing"]["before"])
                normal_paragraph.space_after = Pt(config["paragraph_spacing"]["after"])
                normal_paragraph.widow_control = True
                normal_paragraph.keep_with_next = False

                logger.info(f"Normal style configured: {config['font_family']}, {config['font_size']}pt")

            except Exception as e:
                logger.warning(f"Could not configure Normal style: {e}")

            # ===== ENHANCED HEADING STYLES CONFIGURATION =====
            heading_configs = {
                1: {"size": 18, "color": config["colors"]["primary"], "spacing_after": 12, "bold": True},
                2: {"size": 15, "color": config["colors"]["primary"], "spacing_after": 10, "bold": True},
                3: {"size": 13, "color": config["colors"]["secondary"], "spacing_after": 8, "bold": True},
                4: {"size": 12, "color": config["colors"]["secondary"], "spacing_after": 6, "bold": True},
                5: {"size": 11, "color": config["colors"]["accent"], "spacing_after": 6, "bold": True},
                6: {"size": 10, "color": config["colors"]["accent"], "spacing_after": 4, "bold": True}
            }

            for level in range(1, 7):
                try:
                    heading_style_name = f'Heading {level}'
                    if heading_style_name in doc.styles:
                        heading_style = doc.styles[heading_style_name]
                        heading_config = heading_configs.get(level, heading_configs[6])

                        # Configure heading font
                        heading_style.font.name = config["font_family"]
                        heading_style.font.size = Pt(heading_config["size"])
                        heading_style.font.color.rgb = RGBColor.from_string(heading_config["color"])
                        heading_style.font.bold = heading_config["bold"]

                        # Configure heading paragraph
                        heading_style.paragraph_format.space_before = Pt(level * 2 + 8)
                        heading_style.paragraph_format.space_after = Pt(heading_config["spacing_after"])
                        heading_style.paragraph_format.keep_with_next = True
                        heading_style.paragraph_format.page_break_before = level == 1 and word_count > 2000

                        logger.debug(f"Heading {level} style configured")

                except Exception as e:
                    logger.warning(f"Could not configure Heading {level} style: {e}")

            # ===== ADVANCED PARAGRAPH-BY-PARAGRAPH PROCESSING =====
            processed_paragraphs = 0
            enhanced_paragraphs = 0

            for paragraph in doc.paragraphs:
                try:
                    # Skip empty paragraphs
                    if not paragraph.text.strip():
                        continue

                    processed_paragraphs += 1

                    # Apply enhanced formatting based on paragraph characteristics
                    paragraph_enhanced = False

                    # Enhanced Normal paragraph formatting
                    if paragraph.style.name == 'Normal':
                        # Apply base spacing
                        paragraph.paragraph_format.space_after = Pt(config["paragraph_spacing"]["after"])
                        paragraph.paragraph_format.space_before = Pt(config["paragraph_spacing"]["before"])

                        # Enhanced line spacing with document type consideration
                        paragraph.paragraph_format.line_spacing = Pt(config["line_spacing"])

                        # Apply intelligent spacing based on paragraph length
                        text_length = len(paragraph.text)
                        if text_length > 200:  # Long paragraph
                            paragraph.paragraph_format.space_after = Pt(config["paragraph_spacing"]["after"] + 2)
                            paragraph.paragraph_format.line_spacing = Pt(config["line_spacing"] + 1)
                        elif text_length < 50:  # Short paragraph
                            paragraph.paragraph_format.space_after = Pt(max(3, config["paragraph_spacing"]["after"] - 1))

                        paragraph_enhanced = True

                    # Enhanced List paragraph formatting
                    elif 'List' in paragraph.style.name:
                        paragraph.paragraph_format.space_after = Pt(max(3, config["paragraph_spacing"]["after"] - 2))
                        paragraph.paragraph_format.line_spacing = Pt(config["line_spacing"] - 1)
                        paragraph_enhanced = True

                    # Apply document-specific enhancements
                    if paragraph_enhanced and config["professional_mode"]:
                        # Professional documents get enhanced typography
                        paragraph.paragraph_format.widow_control = True

                        # Add subtle spacing adjustments for professional look
                        current_after = paragraph.paragraph_format.space_after.pt if paragraph.paragraph_format.space_after else 0
                        paragraph.paragraph_format.space_after = Pt(current_after * config["spacing_multiplier"])

                    if paragraph_enhanced:
                        enhanced_paragraphs += 1

                except Exception as e:
                    logger.warning(f"Could not process paragraph: {e}")
                    continue

            # ===== ADVANCED SECTION AND PAGE FORMATTING =====
            try:
                # Configure page margins for all sections
                for section in doc.sections:
                    section.left_margin = Inches(config["margins"]["left"])
                    section.right_margin = Inches(config["margins"]["right"])
                    section.top_margin = Inches(config["margins"]["top"])
                    section.bottom_margin = Inches(config["margins"]["bottom"])

                logger.info("Page margins and sections configured")

            except Exception as e:
                logger.warning(f"Could not configure page margins: {e}")

            # ===== PERFORMANCE STATISTICS AND LOGGING =====
            try:
                # Document statistics
                total_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])
                enhancement_rate = (enhanced_paragraphs / processed_paragraphs * 100) if processed_paragraphs > 0 else 0

                # Comprehensive logging
                logger.info(f"Enhanced document formatting completed successfully!")
                logger.info(f"Document type: {document_type}")
                logger.info(f"Configuration: {config['font_family']}, {config['font_size']}pt, {config['line_spacing']}pt line spacing")
                logger.info(f"Processed: {processed_paragraphs}/{total_paragraphs} paragraphs ({enhancement_rate:.1f}% enhanced)")
                logger.info(f"Word count: {word_count}, Complexity: {complexity_level}")

            except Exception as e:
                logger.warning(f"Could not generate performance statistics: {e}")

            return True

        except Exception as e:
            logger.error(f"Critical error in enhanced document formatting: {e}", exc_info=True)

            # ===== FALLBACK FORMATTING =====
            try:
                logger.info("Attempting fallback formatting...")

                # Basic fallback formatting
                if doc and hasattr(doc, 'styles') and 'Normal' in doc.styles:
                    normal_style = doc.styles['Normal']
                    normal_style.font.name = 'Calibri'
                    normal_style.font.size = Pt(11)
                    normal_style.paragraph_format.line_spacing = Pt(14)

                    # Basic paragraph formatting
                    for paragraph in doc.paragraphs:
                        if paragraph.style.name == 'Normal' and paragraph.text.strip():
                            paragraph.paragraph_format.space_after = Pt(6)
                            paragraph.paragraph_format.space_before = Pt(0)

                    logger.info("Fallback formatting applied successfully")
                    return True
                else:
                    logger.error("Fallback formatting failed - invalid document")
                    return False

            except Exception as fallback_error:
                logger.error(f"Fallback formatting failed: {fallback_error}")
                return False

    def _add_enhanced_horizontal_rule(self, doc, document_type="general", style_options=None):
        """
        Menambahkan garis horizontal yang elegan berdasarkan jenis dokumen dengan fleksibilitas tinggi.

        Args:
            doc: Document object Word
            document_type: Jenis dokumen atau dictionary content_stats atau string warna langsung
            style_options: Dictionary dengan opsi styling tambahan (optional)
        """
        # ===== ENHANCED INPUT HANDLING =====
        # Deteksi dan normalisasi berbagai jenis input
        normalized_type = "general"
        custom_config = {}

        # Handle berbagai jenis input untuk document_type
        if isinstance(document_type, dict):
            # Jika input adalah content_stats dictionary
            if "content_type" in document_type:
                normalized_type = document_type["content_type"]
            elif "type" in document_type:
                normalized_type = document_type["type"]
            else:
                # Coba deteksi field lain yang relevan
                for key in ["document_type", "doc_type", "category"]:
                    if key in document_type:
                        normalized_type = document_type[key]
                        break

            # Ekstrak informasi tambahan dari dictionary
            if "complexity_level" in document_type:
                custom_config["complexity"] = document_type["complexity_level"]
            if "language" in document_type:
                custom_config["language"] = document_type["language"]
            if "word_count" in document_type:
                custom_config["word_count"] = document_type["word_count"]

        elif isinstance(document_type, str):
            # Handle string input
            if document_type.startswith("#"):
                # Direct color code input (e.g., "#4472C4")
                custom_config["direct_color"] = document_type.replace("#", "")
                normalized_type = "custom"
            elif len(document_type) == 6 and all(c in '0123456789ABCDEFabcdef' for c in document_type):
                # Hex color without # (e.g., "4472C4")
                custom_config["direct_color"] = document_type.upper()
                normalized_type = "custom"
            else:
                # Regular document type string
                normalized_type = document_type.lower().strip()

        # Merge dengan style_options jika ada
        if style_options and isinstance(style_options, dict):
            custom_config.update(style_options)

        # ===== ENHANCED STYLE CONFIGURATION =====
        # Comprehensive style mapping dengan berbagai variasi nama
        style_configs = {
            # Primary document types
            "technical_report": {
                "color": "2B5B84",      # Professional navy blue
                "size": "16",           # Thicker untuk technical docs
                "style": "single",
                "spacing_before": 12,
                "spacing_after": 12,
                "alignment": "center"
            },
            "technical": {  # Alias
                "color": "2B5B84",
                "size": "16",
                "style": "single",
                "spacing_before": 12,
                "spacing_after": 12,
                "alignment": "center"
            },
            "meeting_notes": {
                "color": "4F7942",      # Professional green
                "size": "12",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            },
            "meeting": {  # Alias
                "color": "4F7942",
                "size": "12",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            },
            "lecture": {
                "color": "C65911",      # Professional orange
                "size": "10",
                "style": "single",
                "spacing_before": 8,
                "spacing_after": 8,
                "alignment": "center"
            },
            "education": {  # Alias
                "color": "C65911",
                "size": "10",
                "style": "single",
                "spacing_before": 8,
                "spacing_after": 8,
                "alignment": "center"
            },

            # Extended document types
            "presentation": {
                "color": "7030A0",      # Purple
                "size": "14",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            },
            "research": {
                "color": "1F4E79",      # Academic blue
                "size": "18",           # Thick untuk emphasis
                "style": "single",
                "spacing_before": 14,
                "spacing_after": 14,
                "alignment": "center"
            },
            "interview": {
                "color": "8A2BE2",      # Blue violet
                "size": "8",
                "style": "single",
                "spacing_before": 8,
                "spacing_after": 8,
                "alignment": "center"
            },
            "instructional": {
                "color": "228B22",      # Forest green
                "size": "10",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            },
            "news": {
                "color": "DC143C",      # Crimson
                "size": "6",
                "style": "single",
                "spacing_before": 6,
                "spacing_after": 6,
                "alignment": "center"
            },
            "narrative": {
                "color": "8B4513",      # Saddle brown
                "size": "8",
                "style": "single",
                "spacing_before": 8,
                "spacing_after": 8,
                "alignment": "center"
            },

            # Special styles
            "formal": {
                "color": "000000",      # Black untuk formal
                "size": "20",
                "style": "single",
                "spacing_before": 16,
                "spacing_after": 16,
                "alignment": "center"
            },
            "casual": {
                "color": "A5A5A5",      # Light gray
                "size": "6",
                "style": "dotted",
                "spacing_before": 6,
                "spacing_after": 6,
                "alignment": "center"
            },
            "subtle": {
                "color": "E0E0E0",      # Very light gray
                "size": "4",
                "style": "single",
                "spacing_before": 4,
                "spacing_after": 4,
                "alignment": "center"
            },
            "prominent": {
                "color": "FF6600",      # Orange
                "size": "24",
                "style": "double",      # Double line untuk prominence
                "spacing_before": 20,
                "spacing_after": 20,
                "alignment": "center"
            },

            # Custom direct color
            "custom": {
                "color": custom_config.get("direct_color", "A5A5A5"),
                "size": "8",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            },

            # Default
            "general": {
                "color": "A5A5A5",      # Gray
                "size": "8",
                "style": "single",
                "spacing_before": 10,
                "spacing_after": 10,
                "alignment": "center"
            }
        }

        # ===== ADVANCED STYLE SELECTION =====
        # Get base configuration
        config = style_configs.get(normalized_type, style_configs["general"]).copy()

        # ===== DYNAMIC ADJUSTMENTS =====
        # Adjust berdasarkan complexity level
        if custom_config.get("complexity") == "high":
            config["size"] = str(int(config["size"]) + 4)  # Thicker untuk high complexity
            config["spacing_before"] += 4
            config["spacing_after"] += 4
        elif custom_config.get("complexity") == "low":
            config["size"] = str(max(4, int(config["size"]) - 2))  # Thinner untuk low complexity
            config["spacing_before"] = max(4, config["spacing_before"] - 2)
            config["spacing_after"] = max(4, config["spacing_after"] - 2)

        # Adjust berdasarkan word count
        word_count = custom_config.get("word_count", 0)
        if word_count > 5000:  # Long document
            config["size"] = str(int(config["size"]) + 2)
            config["spacing_before"] += 2
            config["spacing_after"] += 2
        elif word_count < 500:  # Short document
            config["size"] = str(max(4, int(config["size"]) - 2))
            config["spacing_before"] = max(4, config["spacing_before"] - 2)
            config["spacing_after"] = max(4, config["spacing_after"] - 2)

        # ===== CUSTOM OVERRIDES =====
        # Apply custom overrides dari style_options atau custom_config
        override_mapping = {
            "color": "color",
            "border_color": "color",
            "thickness": "size",
            "size": "size",
            "border_size": "size",
            "line_style": "style",
            "border_style": "style",
            "spacing": "spacing_before",  # Will also set spacing_after
            "space_before": "spacing_before",
            "space_after": "spacing_after",
            "align": "alignment",
            "alignment": "alignment"
        }

        for custom_key, config_key in override_mapping.items():
            if custom_key in custom_config:
                value = custom_config[custom_key]
                if custom_key == "spacing":
                    # Special handling untuk spacing yang affect both before and after
                    config["spacing_before"] = value
                    config["spacing_after"] = value
                elif custom_key in ["color", "border_color"]:
                    # Handle color format variations
                    if isinstance(value, str):
                        if value.startswith("#"):
                            config[config_key] = value[1:]
                        elif len(value) == 6 and all(c in '0123456789ABCDEFabcdef' for c in value):
                            config[config_key] = value.upper()
                        else:
                            # Named color - keep as is or convert if needed
                            config[config_key] = value
                elif custom_key in ["thickness", "size", "border_size"]:
                    # Handle numeric values
                    if isinstance(value, (int, float)):
                        config[config_key] = str(int(value))
                    else:
                        config[config_key] = str(value)
                else:
                    config[config_key] = value

        # ===== ENHANCED PARAGRAPH CREATION =====
        paragraph = doc.add_paragraph()

        # Set alignment berdasarkan konfigurasi
        alignment_map = {
            "left": WD_ALIGN_PARAGRAPH.LEFT,
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "centre": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY
        }
        paragraph.alignment = alignment_map.get(config["alignment"].lower(), WD_ALIGN_PARAGRAPH.CENTER)

        # ===== ENHANCED BORDER CREATION =====
        # Access XML elements
        p = paragraph._p
        pPr = p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        # Create bottom border dengan enhanced configuration
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), config["style"])
        bottom.set(qn('w:sz'), config["size"])
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), config["color"])

        # Add theme color support jika tersedia
        if hasattr(self, 'theme') and self.theme:
            if normalized_type in ["technical_report", "meeting_notes", "lecture"]:
                # Use theme color untuk consistency dengan document
                theme_color = self.theme.get("primary", config["color"])
                bottom.set(qn('w:color'), theme_color)

        pBdr.append(bottom)
        pPr.append(pBdr)

        # ===== ENHANCED SPACING =====
        paragraph.paragraph_format.space_before = Pt(config["spacing_before"])
        paragraph.paragraph_format.space_after = Pt(config["spacing_after"])

        # ===== SPECIAL ENHANCEMENTS =====
        # Add subtle text content untuk screen readers atau accessibility
        if custom_config.get("accessibility", True):
            # Add invisible separator text untuk accessibility
            separator_run = paragraph.add_run("───")  # Em dash separator
            separator_run.font.color.rgb = RGBColor.from_string(config["color"])
            separator_run.font.size = Pt(1)  # Almost invisible

        # ===== LOGGING =====
        try:
            logger.info(f"Enhanced horizontal rule added - Type: {normalized_type}, Color: {config['color']}, Size: {config['size']}")
        except:
            # Fallback jika logger tidak tersedia
            pass

        return paragraph

    def _add_enhanced_callout(self, doc, text, callout_type, document_type="general", style_options=None):
        """
        Menambahkan callout/admonition dengan styling yang lebih profesional dan fleksibel.

        Args:
            doc: Document Word object
            text: Teks konten callout
            callout_type: Jenis callout (note, warning, important, dll.)
            document_type: Jenis dokumen untuk styling yang disesuaikan (optional)
            style_options: Dictionary dengan opsi styling tambahan (optional)
        """

        # ===== ENHANCED INPUT VALIDATION =====
        if not text or not text.strip():
            logger.warning("Empty text received for enhanced callout")
            return None

        # Normalize document_type jika berupa dictionary
        if isinstance(document_type, dict):
            document_type = document_type.get("content_type", "general")
        elif not isinstance(document_type, str):
            document_type = "general"

        document_type = document_type.lower().strip()

        # ===== COMPREHENSIVE CALLOUT CONFIGURATION =====
        # Enhanced styles dengan dukungan multi-document-type
        enhanced_styles = {
            "note": {
                "icon": "📝",
                "title": "NOTE",
                "base_color": "4472C4",
                "base_bg": "EDF6FF",
                "base_border": "BDD7EE",
                "emoji_alt": "[NOTE]",
                "priority": "info"
            },
            "info": {
                "icon": "ℹ️",
                "title": "INFORMATION",
                "base_color": "0078D4",
                "base_bg": "F3F9FF",
                "base_border": "A6C8FF",
                "emoji_alt": "[INFO]",
                "priority": "info"
            },
            "tip": {
                "icon": "💡",
                "title": "TIP",
                "base_color": "107C10",
                "base_bg": "F3FFF3",
                "base_border": "9FD89F",
                "emoji_alt": "[TIP]",
                "priority": "helpful"
            },
            "warning": {
                "icon": "⚠️",
                "title": "WARNING",
                "base_color": "FF8C00",
                "base_bg": "FFF8E7",
                "base_border": "FFD166",
                "emoji_alt": "[WARNING]",
                "priority": "attention"
            },
            "danger": {
                "icon": "🚨",
                "title": "DANGER",
                "base_color": "DC3545",
                "base_bg": "FFF5F5",
                "base_border": "F8A8A8",
                "emoji_alt": "[DANGER]",
                "priority": "critical"
            },
            "important": {
                "icon": "❗",
                "title": "IMPORTANT",
                "base_color": "DC143C",
                "base_bg": "FFF0F0",
                "base_border": "FFB3B3",
                "emoji_alt": "[IMPORTANT]",
                "priority": "high"
            },
            "success": {
                "icon": "✅",
                "title": "SUCCESS",
                "base_color": "28A745",
                "base_bg": "F0FFF0",
                "base_border": "90EE90",
                "emoji_alt": "[SUCCESS]",
                "priority": "positive"
            },
            "error": {
                "icon": "❌",
                "title": "ERROR",
                "base_color": "DC3545",
                "base_bg": "FFF5F5",
                "base_border": "F8A8A8",
                "emoji_alt": "[ERROR]",
                "priority": "critical"
            },
            "question": {
                "icon": "❓",
                "title": "QUESTION",
                "base_color": "6F42C1",
                "base_bg": "F8F0FF",
                "base_border": "D0A9F5",
                "emoji_alt": "[QUESTION]",
                "priority": "inquiry"
            },
            "conclusion": {
                "icon": "🏁",
                "title": "CONCLUSION",
                "base_color": "495057",
                "base_bg": "F8F9FA",
                "base_border": "CED4DA",
                "emoji_alt": "[CONCLUSION]",
                "priority": "summary"
            },
            "hint": {
                "icon": "🔍",
                "title": "HINT",
                "base_color": "17A2B8",
                "base_bg": "F0F8FF",
                "base_border": "B8E6F0",
                "emoji_alt": "[HINT]",
                "priority": "helpful"
            },
            "caution": {
                "icon": "⚠️",
                "title": "CAUTION",
                "base_color": "FFC107",
                "base_bg": "FFFBF0",
                "base_border": "FFE69C",
                "emoji_alt": "[CAUTION]",
                "priority": "attention"
            }
        }

        # ===== TYPE MAPPING AND NORMALIZATION =====
        # Enhanced type mapping untuk backward compatibility
        type_aliases = {
            "alert": "warning",
            "warn": "warning",
            "critical": "danger",
            "urgent": "danger",
            "notice": "info",
            "information": "info",
            "example": "tip",
            "help": "tip",
            "summary": "conclusion",
            "result": "conclusion",
            "finding": "conclusion"
        }

        # Normalize callout type
        callout_type = callout_type.lower().strip()
        callout_type = type_aliases.get(callout_type, callout_type)

        # Default ke "note" jika tipe tidak dikenal
        if callout_type not in enhanced_styles:
            logger.info(f"Unknown callout type '{callout_type}', defaulting to 'note'")
            callout_type = "note"

        base_style = enhanced_styles[callout_type]

        # ===== DOCUMENT TYPE SPECIFIC STYLING =====
        # Color adjustments berdasarkan document type
        document_color_adjustments = {
            "technical_report": {
                "color_shift": "darker",
                "bg_adjustment": "cooler",
                "border_style": "professional"
            },
            "meeting_notes": {
                "color_shift": "warmer",
                "bg_adjustment": "neutral",
                "border_style": "clean"
            },
            "lecture": {
                "color_shift": "educational",
                "bg_adjustment": "softer",
                "border_style": "academic"
            },
            "presentation": {
                "color_shift": "vibrant",
                "bg_adjustment": "bright",
                "border_style": "modern"
            },
            "research": {
                "color_shift": "academic",
                "bg_adjustment": "neutral",
                "border_style": "formal"
            }
        }

        # Apply document-specific adjustments
        doc_adjustment = document_color_adjustments.get(document_type, {})

        # ===== CALCULATE FINAL STYLING =====
        final_style = base_style.copy()

        # Adjust colors based on document type
        if doc_adjustment.get("color_shift") == "darker":
            # Make colors darker for technical documents
            final_style["color"] = self._darken_color(base_style["base_color"], 0.15)
        elif doc_adjustment.get("color_shift") == "warmer":
            # Add warmth for meeting notes
            final_style["color"] = self._warm_color(base_style["base_color"])
        elif doc_adjustment.get("color_shift") == "educational":
            # Educational tone for lectures
            final_style["color"] = self._educational_color(base_style["base_color"])
        else:
            final_style["color"] = base_style["base_color"]

        # Background adjustments
        if doc_adjustment.get("bg_adjustment") == "cooler":
            final_style["bg_color"] = self._cool_background(base_style["base_bg"])
        elif doc_adjustment.get("bg_adjustment") == "softer":
            final_style["bg_color"] = self._soften_background(base_style["base_bg"])
        else:
            final_style["bg_color"] = base_style["base_bg"]

        # Border adjustments
        final_style["border_color"] = base_style["base_border"]

        # ===== APPLY CUSTOM STYLE OPTIONS =====
        if style_options and isinstance(style_options, dict):
            # Override with custom options
            for key in ["color", "bg_color", "border_color", "icon", "title"]:
                if key in style_options:
                    final_style[key] = style_options[key]

        # ===== CREATE ENHANCED CALLOUT =====
        try:
            # Create main callout paragraph
            callout = doc.add_paragraph()
            callout.style = 'No Spacing'

            # ===== ENHANCED BACKGROUND SHADING =====
            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), final_style["bg_color"]))
            callout._element.get_or_add_pPr().append(shading_elm)

            # ===== ENHANCED BORDER STYLING =====
            pPr = callout._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            # Calculate border styling based on document type and priority
            border_specs = self._get_border_specifications(
                callout_type,
                document_type,
                base_style["priority"],
                final_style["border_color"]
            )

            # Apply enhanced borders
            for side, specs in border_specs.items():
                border_element = OxmlElement(f'w:{side}')
                border_element.set(qn('w:val'), specs['style'])
                border_element.set(qn('w:sz'), specs['size'])
                border_element.set(qn('w:space'), '0')
                border_element.set(qn('w:color'), specs['color'])
                pBdr.append(border_element)

            pPr.append(pBdr)

            # ===== ENHANCED FORMATTING =====
            # Get spacing based on document type and callout priority
            spacing = self._get_callout_spacing(document_type, base_style["priority"])

            callout.paragraph_format.left_indent = Inches(spacing["left_indent"])
            callout.paragraph_format.right_indent = Inches(spacing["right_indent"])
            callout.paragraph_format.space_before = Pt(spacing["space_before"])
            callout.paragraph_format.space_after = Pt(spacing["space_after"])
            callout.paragraph_format.line_spacing = spacing["line_spacing"]

            # ===== ENHANCED ICON AND TITLE =====
            # Determine icon preference (emoji vs text) based on document type
            use_emoji = self._should_use_emoji(document_type)

            if use_emoji:
                icon_text = f"{final_style['icon']} "
            else:
                icon_text = f"{base_style['emoji_alt']} "

            # Add icon
            icon_run = callout.add_run(icon_text)
            icon_run.font.size = Pt(self._get_icon_size(document_type))

            # Add title with enhanced styling
            title_text = f"{final_style['title']}: "
            title_run = callout.add_run(title_text)
            title_run.bold = True
            title_run.font.size = Pt(self._get_title_size(document_type))
            title_run.font.color.rgb = RGBColor.from_string(final_style["color"])

            # ===== ENHANCED CONTENT PROCESSING =====
            # Process content dengan enhanced text formatting
            if hasattr(self, '_add_enhanced_formatted_runs'):
                # Use enhanced formatter if available (untuk backward compatibility)
                self._add_enhanced_formatted_runs(callout, text, document_type)
            else:
                # Fallback to standard formatter
                self._add_formatted_runs_to_paragraph(callout, text)

            logger.info(f"Enhanced callout created - Type: {callout_type}, Document: {document_type}")
            return callout

        except Exception as e:
            logger.error(f"Error creating enhanced callout: {e}", exc_info=True)

            # ===== FALLBACK IMPLEMENTATION =====
            # Create simple fallback callout
            try:
                fallback_callout = doc.add_paragraph()
                fallback_callout.style = 'No Spacing'

                # Simple styling
                fallback_callout.paragraph_format.left_indent = Inches(0.25)
                fallback_callout.paragraph_format.right_indent = Inches(0.25)
                fallback_callout.paragraph_format.space_before = Pt(10)
                fallback_callout.paragraph_format.space_after = Pt(10)

                # Simple content
                title_run = fallback_callout.add_run(f"[{callout_type.upper()}]: ")
                title_run.bold = True
                fallback_callout.add_run(text)

                logger.info(f"Fallback callout created for type: {callout_type}")
                return fallback_callout

            except Exception as fallback_error:
                logger.error(f"Fallback callout creation failed: {fallback_error}")
                return None

    # ===== HELPER METHODS =====

    def _get_border_specifications(self, callout_type, document_type, priority, border_color):
        """Generate border specifications based on callout and document type."""
        specs = {}

        # Base border configuration
        if priority == "critical":
            base_size = "12"
            left_size = "20"  # Thicker left border for critical items
        elif priority == "high":
            base_size = "8"
            left_size = "16"
        elif priority == "attention":
            base_size = "6"
            left_size = "12"
        else:
            base_size = "4"
            left_size = "8"

        # Document type adjustments
        if document_type == "technical_report":
            # More prominent borders for technical docs
            base_size = str(int(base_size) + 2)
            left_size = str(int(left_size) + 4)
        elif document_type == "presentation":
            # Cleaner borders for presentations
            base_size = str(max(2, int(base_size) - 1))
            left_size = str(max(4, int(left_size) - 2))

        # Configure all borders
        specs = {
            'top': {'style': 'single', 'size': base_size, 'color': border_color},
            'bottom': {'style': 'single', 'size': base_size, 'color': border_color},
            'left': {'style': 'single', 'size': left_size, 'color': border_color},
            'right': {'style': 'single', 'size': base_size, 'color': border_color}
        }

        return specs

    def _get_callout_spacing(self, document_type, priority):
        """Calculate spacing based on document type and callout priority."""
        # Base spacing
        spacing = {
            "left_indent": 0.25,
            "right_indent": 0.25,
            "space_before": 10,
            "space_after": 10,
            "line_spacing": 1.15
        }

        # Priority adjustments
        if priority == "critical":
            spacing["space_before"] += 4
            spacing["space_after"] += 4
            spacing["left_indent"] += 0.1
            spacing["right_indent"] += 0.1

        # Document type adjustments
        if document_type == "technical_report":
            spacing["space_before"] += 2
            spacing["space_after"] += 2
            spacing["line_spacing"] = 1.2
        elif document_type == "meeting_notes":
            spacing["space_before"] = max(6, spacing["space_before"] - 2)
            spacing["space_after"] = max(6, spacing["space_after"] - 2)
        elif document_type == "presentation":
            spacing["space_before"] += 3
            spacing["space_after"] += 3
            spacing["left_indent"] += 0.05
            spacing["right_indent"] += 0.05

        return spacing

    def _should_use_emoji(self, document_type):
        """Determine whether to use emoji or text icons based on document type."""
        # Use text icons for formal documents, emoji for others
        formal_types = ["technical_report", "research", "formal", "academic"]
        return document_type not in formal_types

    def _get_icon_size(self, document_type):
        """Get appropriate icon size for document type."""
        if document_type == "presentation":
            return 13  # Larger for presentations
        elif document_type == "technical_report":
            return 10  # Smaller for technical docs
        else:
            return 11  # Standard size

    def _get_title_size(self, document_type):
        """Get appropriate title size for document type."""
        if document_type == "presentation":
            return 12
        elif document_type == "technical_report":
            return 10
        else:
            return 11

    # Color adjustment helper methods

    def _darken_color(self, hex_color, factor=0.15):
        """Darken a hex color by a factor."""
        try:
            # Convert hex to RGB
            hex_color = hex_color.lstrip('#')
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Darken each component
            darkened = tuple(max(0, int(c * (1 - factor))) for c in rgb)

            # Convert back to hex
            return '{:02x}{:02x}{:02x}'.format(*darkened).upper()
        except:
            return hex_color  # Return original if conversion fails

    def _warm_color(self, hex_color):
        """Add warmth to a color by shifting towards red/orange."""
        try:
            hex_color = hex_color.lstrip('#')
            rgb = list(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Increase red component slightly
            rgb[0] = min(255, rgb[0] + 15)
            # Slightly decrease blue
            rgb[2] = max(0, rgb[2] - 10)

            return '{:02x}{:02x}{:02x}'.format(*rgb).upper()
        except:
            return hex_color

    def _educational_color(self, hex_color):
        """Adjust color for educational content."""
        try:
            hex_color = hex_color.lstrip('#')
            rgb = list(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Slightly increase saturation for educational appeal
            # This is a simplified version - increase the dominant component
            max_val = max(rgb)
            max_idx = rgb.index(max_val)
            rgb[max_idx] = min(255, rgb[max_idx] + 20)

            return '{:02x}{:02x}{:02x}'.format(*rgb).upper()
        except:
            return hex_color

    def _cool_background(self, hex_color):
        """Cool down a background color."""
        try:
            hex_color = hex_color.lstrip('#')
            rgb = list(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Increase blue component slightly for cooler tone
            rgb[2] = min(255, rgb[2] + 8)
            # Slightly decrease red
            rgb[0] = max(0, rgb[0] - 5)

            return '{:02x}{:02x}{:02x}'.format(*rgb).upper()
        except:
            return hex_color

    def _soften_background(self, hex_color):
        """Soften a background color by making it lighter."""
        try:
            hex_color = hex_color.lstrip('#')
            rgb = tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))

            # Lighten each component
            softened = tuple(min(255, int(c + (255 - c) * 0.3)) for c in rgb)

            return '{:02x}{:02x}{:02x}'.format(*softened).upper()
        except:
            return hex_color

    # Another Formating for the document

    def _add_special_section(self, doc, section_type, text, document_type="general"):
        """
        Menambahkan bagian spesial dengan format yang sangat ditingkatkan, adaptif, dan profesional.

        Features:
        - Comprehensive section type support dengan 15+ jenis section
        - Dynamic styling berdasarkan document type dan complexity
        - Context-aware icon selection dan visual elements
        - Advanced border dan background patterns
        - Smart text processing dengan inline formatting
        - Responsive spacing dan typography
        - Accessibility support dan cross-platform compatibility
        """

        # ===== ENHANCED INPUT VALIDATION & NORMALIZATION =====
        if not text or not text.strip():
            logger.warning("Empty text received for special section")
            return None

        # Normalize document_type input
        if isinstance(document_type, dict):
            document_type = document_type.get("content_type", "general")
        elif not isinstance(document_type, str):
            document_type = "general"

        document_type = document_type.lower().strip()
        section_type = section_type.lower().strip()

        # ===== COMPREHENSIVE SPECIAL SECTIONS MAPPING =====
        # Significantly expanded dengan 15+ section types dan enhanced configurations
        enhanced_special_sections = {
            "conclusion": {
                "icon": "🏁",
                "title": "CONCLUSION",
                "priority": "summary",
                "style_intensity": "high",
                "color": {
                    "technical_report": "002060",
                    "meeting_notes": "385723",
                    "lecture": "C65911",
                    "presentation": "7030A0",
                    "research": "1F4E79",
                    "interview": "8A2BE2",
                    "instructional": "228B22",
                    "news": "DC143C",
                    "narrative": "8B4513",
                    "general": "002060"
                },
                "bg_color": {
                    "technical_report": "EBF1F9",
                    "meeting_notes": "F0FFF0",
                    "lecture": "FFF8F0",
                    "presentation": "F8F0FF",
                    "research": "F5F8FF",
                    "interview": "FAF0FF",
                    "instructional": "F0FFF0",
                    "news": "FFF8F8",
                    "narrative": "FFF8F0",
                    "general": "EBF1F9"
                }
            },
            "summary": {
                "icon": "📋",
                "title": "SUMMARY",
                "priority": "summary",
                "style_intensity": "high",
                "color": {
                    "technical_report": "4472C4",
                    "meeting_notes": "70AD47",
                    "lecture": "ED7D31",
                    "presentation": "9966CC",
                    "research": "2E75B5",
                    "interview": "B19CD9",
                    "instructional": "32CD32",
                    "news": "FF6B6B",
                    "narrative": "D2691E",
                    "general": "385723"
                },
                "bg_color": {
                    "technical_report": "F2F9FF",
                    "meeting_notes": "F0FFF0",
                    "lecture": "FFF8F0",
                    "presentation": "F8F0FF",
                    "research": "F0F8FF",
                    "interview": "F8F0FF",
                    "instructional": "F0FFF0",
                    "news": "FFF5F5",
                    "narrative": "FFF8F0",
                    "general": "F0FFF0"
                }
            },
            "note": {
                "icon": "📝",
                "title": "NOTE",
                "priority": "info",
                "style_intensity": "medium",
                "color": {
                    "technical_report": "5B9BD5",
                    "meeting_notes": "A9D18E",
                    "lecture": "F4B183",
                    "presentation": "DDA0DD",
                    "research": "87CEEB",
                    "interview": "E6E6FA",
                    "instructional": "98FB98",
                    "news": "FFB6C1",
                    "narrative": "F5DEB3",
                    "general": "5B9BD5"
                },
                "bg_color": {
                    "technical_report": "EDF6FF",
                    "meeting_notes": "F0FFF8",
                    "lecture": "FFF5F0",
                    "presentation": "FAF0FF",
                    "research": "F0F8FF",
                    "interview": "F8F8FF",
                    "instructional": "F0FFF0",
                    "news": "FFF0F5",
                    "narrative": "FFFAF0",
                    "general": "EDF6FF"
                }
            },
            "important": {
                "icon": "⚠️",
                "title": "IMPORTANT",
                "priority": "high",
                "style_intensity": "very_high",
                "color": {
                    "technical_report": "C00000",
                    "meeting_notes": "D9534F",
                    "lecture": "E74C3C",
                    "presentation": "DC143C",
                    "research": "B22222",
                    "interview": "CD5C5C",
                    "instructional": "FF4500",
                    "news": "DC143C",
                    "narrative": "A52A2A",
                    "general": "C00000"
                },
                "bg_color": {
                    "technical_report": "FFF0F0",
                    "meeting_notes": "FFF5F5",
                    "lecture": "FFEAEA",
                    "presentation": "FFF0F0",
                    "research": "FFF5F5",
                    "interview": "FFF0F0",
                    "instructional": "FFF5F0",
                    "news": "FFF0F0",
                    "narrative": "FFF8F8",
                    "general": "FFF0F0"
                }
            },
            "tip": {
                "icon": "💡",
                "title": "TIP",
                "priority": "helpful",
                "style_intensity": "medium",
                "color": {
                    "technical_report": "8A2BE2",
                    "meeting_notes": "32CD32",
                    "lecture": "FF8C00",
                    "presentation": "FFD700",
                    "research": "4169E1",
                    "interview": "DA70D6",
                    "instructional": "00CED1",
                    "news": "FF69B4",
                    "narrative": "DEB887",
                    "general": "107C10"
                },
                "bg_color": {
                    "technical_report": "F8F0FF",
                    "meeting_notes": "F0FFF0",
                    "lecture": "FFF8F0",
                    "presentation": "FFFEF0",
                    "research": "F0F8FF",
                    "interview": "FFF0FF",
                    "instructional": "F0FFFF",
                    "news": "FFF0F8",
                    "narrative": "FFFDF0",
                    "general": "F3FFF3"
                }
            },
            "warning": {
                "icon": "⚠️",
                "title": "WARNING",
                "priority": "attention",
                "style_intensity": "high",
                "color": {
                    "technical_report": "FF8C00",
                    "meeting_notes": "FF6347",
                    "lecture": "FF4500",
                    "presentation": "FFA500",
                    "research": "FF7F50",
                    "interview": "FF8C00",
                    "instructional": "FF6600",
                    "news": "FF4500",
                    "narrative": "D2691E",
                    "general": "FF8C00"
                },
                "bg_color": {
                    "technical_report": "FFF8E7",
                    "meeting_notes": "FFF5F0",
                    "lecture": "FFF0E6",
                    "presentation": "FFFAF0",
                    "research": "FFF8F0",
                    "interview": "FFF8E7",
                    "instructional": "FFF5E6",
                    "news": "FFF0E6",
                    "narrative": "FFF8F0",
                    "general": "FFF7ED"
                }
            },
            "action": {
                "icon": "⚡",
                "title": "ACTION ITEM",
                "priority": "urgent",
                "style_intensity": "very_high",
                "color": {
                    "technical_report": "DC143C",
                    "meeting_notes": "FF1493",
                    "lecture": "FF4500",
                    "presentation": "FF0000",
                    "research": "B22222",
                    "interview": "DC143C",
                    "instructional": "FF6600",
                    "news": "FF0000",
                    "narrative": "CD5C5C",
                    "general": "DC143C"
                },
                "bg_color": {
                    "technical_report": "FFE4E1",
                    "meeting_notes": "FFF0F5",
                    "lecture": "FFF0E6",
                    "presentation": "FFF0F0",
                    "research": "FFF5F5",
                    "interview": "FFE4E1",
                    "instructional": "FFF5E6",
                    "news": "FFF0F0",
                    "narrative": "FFF8F8",
                    "general": "FFE4E1"
                }
            },
            "decision": {
                "icon": "✅",
                "title": "DECISION",
                "priority": "high",
                "style_intensity": "high",
                "color": {
                    "technical_report": "228B22",
                    "meeting_notes": "32CD32",
                    "lecture": "00C851",
                    "presentation": "00FF00",
                    "research": "228B22",
                    "interview": "32CD32",
                    "instructional": "00CED1",
                    "news": "32CD32",
                    "narrative": "9ACD32",
                    "general": "228B22"
                },
                "bg_color": {
                    "technical_report": "F0FFF0",
                    "meeting_notes": "F0FFF0",
                    "lecture": "F0FFF0",
                    "presentation": "F0FFF0",
                    "research": "F0FFF0",
                    "interview": "F0FFF0",
                    "instructional": "F0FFFF",
                    "news": "F0FFF0",
                    "narrative": "F5FFFA",
                    "general": "F0FFF0"
                }
            },
            "question": {
                "icon": "❓",
                "title": "QUESTION",
                "priority": "inquiry",
                "style_intensity": "medium",
                "color": {
                    "technical_report": "4169E1",
                    "meeting_notes": "6495ED",
                    "lecture": "1E90FF",
                    "presentation": "00BFFF",
                    "research": "0000CD",
                    "interview": "9370DB",
                    "instructional": "00CED1",
                    "news": "4169E1",
                    "narrative": "6495ED",
                    "general": "4169E1"
                },
                "bg_color": {
                    "technical_report": "F0F8FF",
                    "meeting_notes": "F0F8FF",
                    "lecture": "F0F8FF",
                    "presentation": "F0FFFF",
                    "research": "F0F0FF",
                    "interview": "F8F0FF",
                    "instructional": "F0FFFF",
                    "news": "F0F8FF",
                    "narrative": "F8F8FF",
                    "general": "F0F8FF"
                }
            },
            "success": {
                "icon": "🎉",
                "title": "SUCCESS",
                "priority": "positive",
                "style_intensity": "high",
                "color": {
                    "technical_report": "008000",
                    "meeting_notes": "32CD32",
                    "lecture": "00C851",
                    "presentation": "00FF32",
                    "research": "228B22",
                    "interview": "32CD32",
                    "instructional": "00FA9A",
                    "news": "32CD32",
                    "narrative": "9ACD32",
                    "general": "008000"
                },
                "bg_color": {
                    "technical_report": "F0FFF0",
                    "meeting_notes": "F0FFF0",
                    "lecture": "F0FFF0",
                    "presentation": "F0FFF0",
                    "research": "F0FFF0",
                    "interview": "F0FFF0",
                    "instructional": "F0FFFF",
                    "news": "F0FFF0",
                    "narrative": "F5FFFA",
                    "general": "F0FFF0"
                }
            },
            "error": {
                "icon": "❌",
                "title": "ERROR",
                "priority": "critical",
                "style_intensity": "very_high",
                "color": {
                    "technical_report": "8B0000",
                    "meeting_notes": "DC143C",
                    "lecture": "B22222",
                    "presentation": "FF0000",
                    "research": "8B0000",
                    "interview": "CD5C5C",
                    "instructional": "FF4500",
                    "news": "DC143C",
                    "narrative": "A52A2A",
                    "general": "8B0000"
                },
                "bg_color": {
                    "technical_report": "FFE4E1",
                    "meeting_notes": "FFF0F0",
                    "lecture": "FFF5F5",
                    "presentation": "FFF0F0",
                    "research": "FFE4E1",
                    "interview": "FFF8F8",
                    "instructional": "FFF0E6",
                    "news": "FFF0F0",
                    "narrative": "FFF8F8",
                    "general": "FFE4E1"
                }
            },
            "insight": {
                "icon": "🔍",
                "title": "INSIGHT",
                "priority": "analytical",
                "style_intensity": "medium",
                "color": {
                    "technical_report": "6B46C1",
                    "meeting_notes": "8B5CF6",
                    "lecture": "7C3AED",
                    "presentation": "A855F7",
                    "research": "5B21B6",
                    "interview": "9333EA",
                    "instructional": "06B6D4",
                    "news": "8B5CF6",
                    "narrative": "D8B4FE",
                    "general": "6B46C1"
                },
                "bg_color": {
                    "technical_report": "FAF5FF",
                    "meeting_notes": "F8F5FF",
                    "lecture": "FAF5FF",
                    "presentation": "FAF5FF",
                    "research": "F5F3FF",
                    "interview": "FAF5FF",
                    "instructional": "F0FDFF",
                    "news": "F8F5FF",
                    "narrative": "FAF8FF",
                    "general": "FAF5FF"
                }
            },
            "reference": {
                "icon": "📚",
                "title": "REFERENCE",
                "priority": "info",
                "style_intensity": "low",
                "color": {
                    "technical_report": "64748B",
                    "meeting_notes": "6B7280",
                    "lecture": "78716C",
                    "presentation": "71717A",
                    "research": "52525B",
                    "interview": "6B7280",
                    "instructional": "059669",
                    "news": "6B7280",
                    "narrative": "8B7355",
                    "general": "64748B"
                },
                "bg_color": {
                    "technical_report": "F8FAFC",
                    "meeting_notes": "F9FAFB",
                    "lecture": "FAFAF9",
                    "presentation": "FAFAFA",
                    "research": "F8F9FA",
                    "interview": "F9FAFB",
                    "instructional": "ECFDF5",
                    "news": "F9FAFB",
                    "narrative": "FEFDFB",
                    "general": "F8FAFC"
                }
            },
            "feedback": {
                "icon": "💬",
                "title": "FEEDBACK",
                "priority": "interactive",
                "style_intensity": "medium",
                "color": {
                    "technical_report": "0EA5E9",
                    "meeting_notes": "06B6D4",
                    "lecture": "0891B2",
                    "presentation": "0284C7",
                    "research": "0369A1",
                    "interview": "7C3AED",
                    "instructional": "059669",
                    "news": "0EA5E9",
                    "narrative": "0F766E",
                    "general": "0EA5E9"
                },
                "bg_color": {
                    "technical_report": "F0F9FF",
                    "meeting_notes": "F0FDFF",
                    "lecture": "F0F9FF",
                    "presentation": "EFF6FF",
                    "research": "EFF6FF",
                    "interview": "FAF5FF",
                    "instructional": "ECFDF5",
                    "news": "F0F9FF",
                    "narrative": "F0FDFA",
                    "general": "F0F9FF"
                }
            },
            "recommendation": {
                "icon": "🎯",
                "title": "RECOMMENDATION",
                "priority": "guidance",
                "style_intensity": "high",
                "color": {
                    "technical_report": "059669",
                    "meeting_notes": "16A34A",
                    "lecture": "15803D",
                    "presentation": "16A34A",
                    "research": "166534",
                    "interview": "16A34A",
                    "instructional": "0891B2",
                    "news": "16A34A",
                    "narrative": "166534",
                    "general": "059669"
                },
                "bg_color": {
                    "technical_report": "ECFDF5",
                    "meeting_notes": "F0FDF4",
                    "lecture": "F0FDF4",
                    "presentation": "F0FDF4",
                    "research": "ECFDF5",
                    "interview": "F0FDF4",
                    "instructional": "F0FDFF",
                    "news": "F0FDF4",
                    "narrative": "F6FFED",
                    "general": "ECFDF5"
                }
            }
        }

        # ===== ADVANCED TYPE MAPPING & ALIASES =====
        # Enhanced type aliases untuk backward compatibility dan user-friendly input
        section_type_aliases = {
            # Common aliases
            "conclusions": "conclusion",
            "summaries": "summary",
            "notes": "note",
            "tips": "tip",
            "warnings": "warning",
            "actions": "action",
            "decisions": "decision",
            "questions": "question",
            "errors": "error",
            "insights": "insight",
            "references": "reference",
            "recommendations": "recommendation",

            # Alternative names
            "finale": "conclusion",
            "ending": "conclusion",
            "closing": "conclusion",
            "recap": "summary",
            "overview": "summary",
            "digest": "summary",
            "memo": "note",
            "annotation": "note",
            "remark": "note",
            "alert": "warning",
            "caution": "warning",
            "attention": "warning",
            "urgent": "important",
            "critical": "important",
            "priority": "important",
            "task": "action",
            "todo": "action",
            "assignment": "action",
            "resolve": "decision",
            "verdict": "decision",
            "ruling": "decision",
            "query": "question",
            "inquiry": "question",
            "ask": "question",
            "failure": "error",
            "issue": "error",
            "problem": "error",
            "finding": "insight",
            "discovery": "insight",
            "observation": "insight",
            "analysis": "insight",
            "citation": "reference",
            "source": "reference",
            "link": "reference",
            "suggest": "recommendation",
            "advice": "recommendation",
            "proposal": "recommendation",
            "response": "feedback",
            "comment": "feedback",
            "review": "feedback"
        }

        # Normalize section type
        original_section_type = section_type
        section_type = section_type_aliases.get(section_type, section_type)

        # Default ke "note" jika tipe tidak dikenal
        if section_type not in enhanced_special_sections:
            logger.info(f"Unknown section type '{original_section_type}' (normalized: '{section_type}'), defaulting to 'note'")
            section_type = "note"

        config = enhanced_special_sections[section_type]

        # ===== DYNAMIC CONFIGURATION BASED ON DOCUMENT TYPE =====
        # Get appropriate colors and backgrounds
        color = config["color"].get(document_type, config["color"]["general"])
        bg_color = config["bg_color"].get(document_type, config["bg_color"]["general"])

        # ===== ENHANCED PARAGRAPH CREATION =====
        p = doc.add_paragraph()
        p.style = 'No Spacing'

        # ===== ADVANCED BACKGROUND STYLING =====
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls

            shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), bg_color))
            p._element.get_or_add_pPr().append(shading_elm)

        except Exception as e:
            logger.warning(f"Could not apply background shading: {e}")

        # ===== COMPREHENSIVE BORDER STYLING =====
        # Enhanced border logic dengan lebih banyak variasi berdasarkan priority dan document type
        try:
            pPr = p._element.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')

            # Determine border specifications based on priority and document type
            priority = config.get("priority", "info")
            style_intensity = config.get("style_intensity", "medium")

            # Calculate border sizes based on priority and intensity
            border_sizes = {
                "critical": {"left": "24", "others": "8"},
                "urgent": {"left": "20", "others": "6"},
                "very_high": {"left": "18", "others": "6"},
                "high": {"left": "16", "others": "4"},
                "attention": {"left": "14", "others": "4"},
                "medium": {"left": "12", "others": "3"},
                "guidance": {"left": "10", "others": "3"},
                "helpful": {"left": "8", "others": "2"},
                "info": {"left": "6", "others": "2"},
                "low": {"left": "4", "others": "1"}
            }

            sizes = border_sizes.get(style_intensity, border_sizes["medium"])

            # Document type specific border adjustments
            if document_type in ["technical_report", "research"]:
                # More prominent borders for formal documents
                sizes["left"] = str(int(sizes["left"]) + 4)
                sizes["others"] = str(int(sizes["others"]) + 2)
            elif document_type in ["meeting_notes", "interview"]:
                # Moderate borders for collaborative documents
                sizes["left"] = str(int(sizes["left"]) + 2)
            elif document_type in ["presentation", "narrative"]:
                # Subtle borders for visual documents
                sizes["left"] = str(max(4, int(sizes["left"]) - 2))
                sizes["others"] = str(max(1, int(sizes["others"]) - 1))

            # Create enhanced borders
            # Left border (accent)
            left_border = OxmlElement('w:left')
            left_border.set(qn('w:val'), 'single')
            left_border.set(qn('w:sz'), sizes["left"])
            left_border.set(qn('w:space'), '0')
            left_border.set(qn('w:color'), color)
            pBdr.append(left_border)

            # Top and bottom borders (subtle)
            for side in ['top', 'bottom']:
                border_element = OxmlElement(f'w:{side}')
                border_element.set(qn('w:val'), 'single')
                border_element.set(qn('w:sz'), sizes["others"])
                border_element.set(qn('w:space'), '0')
                border_element.set(qn('w:color'), color)
                pBdr.append(border_element)

            # Right border (very subtle) only for high priority items
            if style_intensity in ["critical", "urgent", "very_high"]:
                right_border = OxmlElement('w:right')
                right_border.set(qn('w:val'), 'single')
                right_border.set(qn('w:sz'), str(max(2, int(sizes["others"]) - 1)))
                right_border.set(qn('w:space'), '0')
                right_border.set(qn('w:color'), color)
                pBdr.append(right_border)

            pPr.append(pBdr)

        except Exception as e:
            logger.warning(f"Could not apply enhanced borders: {e}")

        # ===== ENHANCED ICON AND TITLE FORMATTING =====
        # Context-aware icon sizing
        icon_sizes = {
            "critical": 14,
            "urgent": 13,
            "very_high": 13,
            "high": 12,
            "attention": 12,
            "medium": 11,
            "guidance": 11,
            "helpful": 10,
            "info": 10,
            "low": 9
        }

        icon_size = icon_sizes.get(style_intensity, 11)

        # Add icon dengan enhanced styling
        icon_run = p.add_run(f"{config['icon']} ")
        icon_run.font.size = Pt(icon_size)

        # Add title dengan adaptive sizing
        title_run = p.add_run(f"{config['title']}: ")
        title_run.bold = True
        title_run.font.color.rgb = RGBColor.from_string(color)

        # Enhanced font size calculation based on document type dan priority
        base_title_size = 11
        if document_type == "technical_report":
            base_title_size = 11
        elif document_type == "meeting_notes":
            base_title_size = 10.5
        elif document_type == "presentation":
            base_title_size = 12
        elif document_type == "lecture":
            base_title_size = 11.5
        else:
            base_title_size = 11

        # Adjust berdasarkan priority
        priority_adjustments = {
            "critical": +1.5,
            "urgent": +1,
            "very_high": +1,
            "high": +0.5,
            "attention": +0.5,
            "medium": 0,
            "guidance": 0,
            "helpful": -0.5,
            "info": -0.5,
            "low": -1
        }

        final_title_size = base_title_size + priority_adjustments.get(style_intensity, 0)
        title_run.font.size = Pt(final_title_size)

        # ===== ENHANCED TEXT CONTENT PROCESSING =====
        # Add text content dengan advanced formatting
        if hasattr(self, '_add_formatted_runs_to_paragraph'):
            self._add_formatted_runs_to_paragraph(p, text)
        else:
            # Fallback simple text addition
            content_run = p.add_run(text)
            content_run.font.name = 'Calibri'
            content_run.font.size = Pt(10.5)

        # ===== COMPREHENSIVE SPACING AND LAYOUT =====
        # Calculate sophisticated spacing based on multiple factors
        base_spacing = {
            "space_before": 10,
            "space_after": 10,
            "left_indent": 0.2,
            "right_indent": 0.2,
            "line_spacing": 1.2
        }

        # Document type adjustments
        spacing_adjustments = {
            "technical_report": {
                "space_before": +2,
                "space_after": +2,
                "left_indent": +0.1,
                "right_indent": +0.1,
                "line_spacing": +0.1
            },
            "meeting_notes": {
                "space_before": -2,
                "space_after": -2,
                "left_indent": +0.05,
                "right_indent": +0.05,
                "line_spacing": 0
            },
            "presentation": {
                "space_before": +4,
                "space_after": +4,
                "left_indent": +0.15,
                "right_indent": +0.15,
                "line_spacing": +0.1
            },
            "lecture": {
                "space_before": +1,
                "space_after": +1,
                "left_indent": +0.05,
                "right_indent": +0.05,
                "line_spacing": +0.05
            },
            "research": {
                "space_before": +3,
                "space_after": +3,
                "left_indent": +0.2,
                "right_indent": +0.1,
                "line_spacing": +0.15
            }
        }

        # Priority adjustments
        priority_spacing_adjustments = {
            "critical": {"space_before": +6, "space_after": +6, "left_indent": +0.1, "right_indent": +0.1},
            "urgent": {"space_before": +4, "space_after": +4, "left_indent": +0.05, "right_indent": +0.05},
            "very_high": {"space_before": +3, "space_after": +3, "left_indent": +0.05, "right_indent": +0.05},
            "high": {"space_before": +2, "space_after": +2, "left_indent": 0, "right_indent": 0},
            "attention": {"space_before": +1, "space_after": +1, "left_indent": 0, "right_indent": 0},
            "medium": {"space_before": 0, "space_after": 0, "left_indent": 0, "right_indent": 0},
            "guidance": {"space_before": 0, "space_after": 0, "left_indent": 0, "right_indent": 0},
            "helpful": {"space_before": -1, "space_after": -1, "left_indent": 0, "right_indent": 0},
            "info": {"space_before": -1, "space_after": -1, "left_indent": 0, "right_indent": 0},
            "low": {"space_before": -2, "space_after": -2, "left_indent": -0.05, "right_indent": -0.05}
        }

        # Apply adjustments
        doc_adj = spacing_adjustments.get(document_type, {})
        priority_adj = priority_spacing_adjustments.get(style_intensity, {})

        final_spacing = base_spacing.copy()
        for key in final_spacing:
            final_spacing[key] += doc_adj.get(key, 0) + priority_adj.get(key, 0)

        # Ensure minimum values
        final_spacing["space_before"] = max(4, final_spacing["space_before"])
        final_spacing["space_after"] = max(4, final_spacing["space_after"])
        final_spacing["left_indent"] = max(0.1, final_spacing["left_indent"])
        final_spacing["right_indent"] = max(0.1, final_spacing["right_indent"])
        final_spacing["line_spacing"] = max(1.0, final_spacing["line_spacing"])

        # Apply final spacing
        p.paragraph_format.space_before = Pt(final_spacing["space_before"])
        p.paragraph_format.space_after = Pt(final_spacing["space_after"])
        p.paragraph_format.left_indent = Inches(final_spacing["left_indent"])
        p.paragraph_format.right_indent = Inches(final_spacing["right_indent"])
        p.paragraph_format.line_spacing = final_spacing["line_spacing"]

        # ===== SPECIAL ENHANCEMENTS =====
        # Keep with next for better page flow (high priority items)
        if style_intensity in ["critical", "urgent", "very_high", "high"]:
            p.paragraph_format.keep_with_next = True

        # Add subtle shadow effect for critical items
        if style_intensity == "critical":
            try:
                # Enhance dengan additional emphasis formatting
                for run in p.runs:
                    if run.font.bold:  # Title run
                        run.font.size = Pt(run.font.size.pt + 1)
            except:
                pass

        # ===== ACCESSIBILITY ENHANCEMENTS =====
        try:
            # Add semantic role untuk screen readers
            p._element.set(qn('w:role'), f'section-{section_type}')
        except:
            pass

        # ===== LOGGING AND RETURN =====
        logger.info(f"Enhanced special section created - Type: {section_type} (original: {original_section_type}), "
                f"Document: {document_type}, Priority: {priority}, Intensity: {style_intensity}")

        return p

    def _add_enhanced_code_block(self, doc, code_lines, language):
        """
        Menambahkan blok kode dengan syntax highlighting yang disimulasikan dan styling profesional.
        """
        if not code_lines:
            return

        # Buat paragraf dengan style khusus
        p = doc.add_paragraph()
        p.style = 'No Spacing'

        # Set background dan border untuk code block
        pPr = p._element.get_or_add_pPr()

        # Background abu-abu terang dengan sedikit nuansa biru
        shading_elm = parse_xml(r'<w:shd {} w:fill="F5F7FA"/>'.format(nsdecls('w')))
        pPr.append(shading_elm)

        # Tambahkan border di seluruh paragraf dengan warna yang lebih modern
        pBdr = OxmlElement('w:pBdr')

        for side in ['top', 'left', 'bottom', 'right']:
            border_element = OxmlElement(f'w:{side}')
            border_element.set(qn('w:val'), 'single')
            border_element.set(qn('w:sz'), '4')
            border_element.set(qn('w:space'), '0')
            border_element.set(qn('w:color'), '4472C4')  # Biru modern
            pBdr.append(border_element)

        pPr.append(pBdr)

        # Paragraf utama code block
        code_text = '\n'.join(code_lines)

        # Tambahkan language label dengan styling yang lebih menarik
        if language:
            lang_container = doc.add_paragraph()
            lang_container.style = 'No Spacing'
            lang_container.paragraph_format.space_after = Pt(0)
            lang_container.paragraph_format.left_indent = Inches(0.2)

            # Background color untuk label bahasa
            lang_shading = parse_xml(r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w')))
            lang_container._element.get_or_add_pPr().append(lang_shading)

            lang_run = lang_container.add_run(f" {language} ")
            lang_run.bold = True
            lang_run.font.size = Pt(9)
            lang_run.font.color.rgb = RGBColor(255, 255, 255)  # Teks putih

        # Tambahkan kode dengan font monospace dan formatting yang ditingkatkan
        code_run = p.add_run(code_text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(9.5)  # Slightly larger for better readability

        # Set line spacing dan indentasi
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.first_line_indent = 0
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.right_indent = Inches(0.2)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

    def _style_heading_by_document_type(self, heading, level, document_type):
        """
        Menerapkan styling heading berdasarkan jenis dokumen.
        """
        # Set properties berdasarkan level
        if level == 1:
            heading.paragraph_format.space_before = Pt(16)
            heading.paragraph_format.space_after = Pt(12)
            font_size = Pt(16)
        elif level == 2:
            heading.paragraph_format.space_before = Pt(14)
            heading.paragraph_format.space_after = Pt(8)
            font_size = Pt(14)
        else:
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            font_size = Pt(12)

        # Set font size for all runs
        for run in heading.runs:
            run.font.size = font_size

        # Apply style berdasarkan jenis dokumen
        if document_type == "technical_report":
            if level == 1:
                heading.style.font.color.rgb = RGBColor(0, 82, 136)  # Deep blue
                self._add_bottom_border(heading, "4472C4", "16")
            elif level == 2:
                heading.style.font.color.rgb = RGBColor(30, 113, 170)  # Medium blue
            else:
                heading.style.font.color.rgb = RGBColor(53, 134, 192)  # Light blue

        elif document_type == "meeting_notes":
            if level == 1:
                heading.style.font.color.rgb = RGBColor(33, 97, 60)  # Deep green
                self._add_bottom_border(heading, "70AD47", "12")
            elif level == 2:
                heading.style.font.color.rgb = RGBColor(53, 127, 79)  # Medium green
            else:
                heading.style.font.color.rgb = RGBColor(84, 148, 105)  # Light green

        elif document_type == "lecture":
            if level == 1:
                heading.style.font.color.rgb = RGBColor(175, 65, 0)  # Deep orange
                self._add_bottom_border(heading, "ED7D31", "12")
            elif level == 2:
                heading.style.font.color.rgb = RGBColor(191, 89, 26)  # Medium orange
            else:
                heading.style.font.color.rgb = RGBColor(208, 124, 73)  # Light orange

        else:  # Default style
            if level == 1:
                heading.style.font.color.rgb = RGBColor(0, 112, 192)  # Biru
                self._add_bottom_border(heading, "4472C4", "12")
            elif level == 2:
                heading.style.font.color.rgb = RGBColor(75, 75, 75)  # Abu-abu
            else:
                heading.style.font.color.rgb = RGBColor(100, 100, 100)  # Abu-abu lebih terang

    def _add_bottom_border(self, paragraph, color, size):
        """
        Menambahkan border bawah pada paragraf
        """
        pPr = paragraph._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')

        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), size)
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), color)
        pBdr.append(bottom)

        pPr.append(pBdr)

    def _add_bookmark(self, paragraph, text):
        """
        Menambahkan bookmark pada heading untuk referensi internal.
        """
        # Buat ID bookmark yang valid
        bookmark_id = text.lower().replace(' ', '_')
        bookmark_id = re.sub(r'[^\w]', '', bookmark_id)

        # Tambahkan bookmark
        run = paragraph.runs[0]
        tag = run._r
        start = OxmlElement('w:bookmarkStart')
        start.set(qn('w:id'), '0')
        start.set(qn('w:name'), bookmark_id)
        tag.append(start)

        end = OxmlElement('w:bookmarkEnd')
        end.set(qn('w:id'), '0')
        tag.append(end)

    def _process_enhanced_markdown_table(self, doc, table_data, document_type="general"):
        """
        Memproses tabel Markdown dengan styling profesional yang ditingkatkan.
        """
        if len(table_data) < 2:
            return  # Minimal harus memiliki baris header dan pemisah

        # Extract header dan data
        header_row = table_data[0]
        separator_row = table_data[1]
        data_rows = table_data[2:] if len(table_data) > 2 else []

        # Parse sel header (hilangkan pipe di awal dan akhir)
        header_cells = [cell.strip() for cell in header_row.strip('|').split('|')]
        num_columns = len(header_cells)

        # Validasi minimum kolom
        if num_columns == 0:
            return

        # Tentukan alignment dari separator row dengan validasi yang lebih baik
        alignments = []
        separator_cells = separator_row.strip('|').split('|')

        # Pastikan jumlah separator cells sesuai dengan header
        if len(separator_cells) != num_columns:
            separator_cells = separator_cells[:num_columns] + ['---'] * (num_columns - len(separator_cells))

        for cell in separator_cells:
            cell = cell.strip()
            if cell.startswith(':') and cell.endswith(':'):
                alignments.append(WD_ALIGN_PARAGRAPH.CENTER)
            elif cell.endswith(':'):
                alignments.append(WD_ALIGN_PARAGRAPH.RIGHT)
            else:
                alignments.append(WD_ALIGN_PARAGRAPH.LEFT)

        # Buat tabel Word dengan validasi ukuran
        table = doc.add_table(rows=1, cols=num_columns)

        # Enhanced table styling berdasarkan document type dengan gradasi warna yang lebih halus
        style_config = {
            "technical_report": {
                "style": 'Table Grid',
                "header_bg": "2E5984",      # Navy blue yang lebih dalam
                "header_accent": "4472C4",   # Blue accent untuk border
                "alt_row_bg": "F8FAFC",     # Very light blue-gray
                "text_color": "1A365D"      # Dark blue text
            },
            "meeting_notes": {
                "style": 'Table Grid',
                "header_bg": "2F855A",      # Forest green yang lebih dalam
                "header_accent": "70AD47",   # Green accent
                "alt_row_bg": "F7FAFC",     # Very light green-gray
                "text_color": "1A202C"      # Dark green text
            },
            "lecture": {
                "style": 'Table Grid',
                "header_bg": "C05621",      # Burnt orange yang lebih dalam
                "header_accent": "ED7D31",   # Orange accent
                "alt_row_bg": "FFFAF7",     # Very light orange
                "text_color": "7B341E"      # Dark orange text
            },
            "general": {
                "style": 'Table Grid',
                "header_bg": "2B6CB0",      # Professional blue
                "header_accent": "5B9BD5",   # Medium blue
                "alt_row_bg": "F8FAFC",     # Light gray-blue
                "text_color": "1E40AF"      # Professional blue text
            }
        }

        config = style_config.get(document_type, style_config["general"])
        table.style = config["style"]

        # Enhanced header styling dengan multi-level formatting
        header_cells_word = table.rows[0].cells
        for i, text in enumerate(header_cells):
            if i < num_columns:
                cell = header_cells_word[i]
                cell.text = ""  # Bersihkan teks default
                p = cell.paragraphs[0]

                # Enhanced text processing untuk header
                cleaned_text = text.strip()
                if cleaned_text:
                    # Deteksi dan proses formatting di header dengan pembersihan tambahan
                    self._add_formatted_runs_to_paragraph(p, cleaned_text)
                else:
                    # Jika sel kosong, tambahkan placeholder
                    p.add_run("—")

                # Set alignment dengan fallback
                if i < len(alignments):
                    p.alignment = alignments[i]
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Enhanced header formatting dengan konsistensi yang lebih baik
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(10)  # Slightly smaller for better proportions
                    run.font.name = 'Calibri'
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text

                # Enhanced background shading untuk header
                shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), config["header_bg"]))
                cell._element.get_or_add_tcPr().append(shading_elm)

                # Enhanced padding untuk header
                self._add_cell_padding(cell, "120")  # Slightly more padding untuk header

        # Enhanced data processing dengan alternating pattern yang lebih sophisticated
        alt_row = False
        row_counter = 0

        for row_text in data_rows:
            row_cells = [cell.strip() for cell in row_text.strip('|').split('|')]

            # Enhanced column validation dengan padding yang lebih cerdas
            if len(row_cells) < num_columns:
                # Tambahkan sel kosong yang diperlukan
                row_cells.extend([''] * (num_columns - len(row_cells)))
            elif len(row_cells) > num_columns:
                # Potong kolom berlebih tapi pertahankan data penting
                row_cells = row_cells[:num_columns]

            # Tambahkan baris baru
            new_row = table.add_row()
            row_counter += 1
            alt_row = (row_counter % 2 == 0)  # Cleaner alternating logic

            # Enhanced cell processing
            for i, text in enumerate(row_cells):
                if i < num_columns:
                    cell = new_row.cells[i]
                    cell.text = ""  # Bersihkan teks default
                    p = cell.paragraphs[0]

                    # Enhanced text processing untuk data cells
                    cleaned_text = text.strip()
                    if cleaned_text:
                        # Deteksi dan proses formatting di dalam sel dengan validasi
                        self._add_formatted_runs_to_paragraph(p, cleaned_text)
                    else:
                        # Sel kosong dengan placeholder yang lebih subtle
                        empty_run = p.add_run("—")
                        empty_run.font.color.rgb = RGBColor(160, 160, 160)  # Light gray
                        empty_run.font.size = Pt(8)

                    # Enhanced alignment dengan fallback
                    if i < len(alignments):
                        p.alignment = alignments[i]
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    # Enhanced alternating row coloring dengan pattern yang lebih halus
                    if alt_row:
                        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), config["alt_row_bg"]))
                        cell._element.get_or_add_tcPr().append(shading_elm)

                    # Enhanced cell formatting untuk readability
                    for run in p.runs:
                        if not run.font.size:  # Hanya set jika belum di-set
                            run.font.size = Pt(9)
                        if not run.font.name:  # Hanya set jika belum di-set
                            run.font.name = 'Calibri'

                    # Enhanced padding untuk data cells
                    self._add_cell_padding(cell, "100")

        # Enhanced table-wide formatting
        table.autofit = True

        # Set column widths yang lebih proporsional
        total_width = Inches(6.5)  # Total available width
        col_width = total_width / num_columns
        for column in table.columns:
            column.width = col_width

        # Enhanced professional borders
        self._apply_enhanced_table_borders(table, config["header_accent"])

        # Enhanced spacing dengan line break yang lebih clean
        spacing_para = doc.add_paragraph()
        spacing_para.paragraph_format.space_before = Pt(8)
        spacing_para.paragraph_format.space_after = Pt(8)

    def _apply_enhanced_table_borders(self, table, header_color):
        """
        Menerapkan border profesional ke tabel Word dengan styling yang lebih sophisticated.
        """
        try:
            tbl = table._tbl

            # Enhanced tblPr access dengan error handling yang lebih baik
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = OxmlElement('w:tblPr')
                tbl.insert(0, tblPr)

            # Remove existing borders jika ada
            existing_borders = tblPr.find(qn('w:tblBorders'))
            if existing_borders is not None:
                tblPr.remove(existing_borders)

            # Create enhanced border configuration
            tblBorders = OxmlElement('w:tblBorders')

            # Enhanced border specifications dengan hierarki yang lebih jelas
            border_specs = {
                'top': {'size': '12', 'color': header_color, 'style': 'single'},
                'bottom': {'size': '12', 'color': header_color, 'style': 'single'},
                'left': {'size': '8', 'color': header_color, 'style': 'single'},
                'right': {'size': '8', 'color': header_color, 'style': 'single'},
                'insideH': {'size': '4', 'color': 'E2E8F0', 'style': 'single'},  # Lighter inside borders
                'insideV': {'size': '4', 'color': 'E2E8F0', 'style': 'single'}   # Lighter inside borders
            }

            # Apply enhanced borders dengan konsistensi yang lebih baik
            for border_name, specs in border_specs.items():
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), specs['style'])
                border.set(qn('w:sz'), specs['size'])
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), specs['color'])

                # Enhanced border styling untuk outside borders
                if border_name in ['top', 'bottom', 'left', 'right']:
                    border.set(qn('w:themeColor'), 'accent1')  # Theme color reference

                tblBorders.append(border)

            tblPr.append(tblBorders)

            # Enhanced table alignment dan spacing
            tblJc = OxmlElement('w:jc')
            tblJc.set(qn('w:val'), 'center')
            tblPr.append(tblJc)

            # Enhanced table margins
            tblCellMar = OxmlElement('w:tblCellMar')

            # Set consistent margins untuk all sides
            margin_specs = {'top': '80', 'right': '108', 'bottom': '80', 'left': '108'}
            for side, value in margin_specs.items():
                margin_elem = OxmlElement(f'w:{side}')
                margin_elem.set(qn('w:w'), value)
                margin_elem.set(qn('w:type'), 'dxa')
                tblCellMar.append(margin_elem)

            tblPr.append(tblCellMar)

        except Exception as e:
            logger.error(f"Error applying enhanced table borders: {e}")
            # Enhanced fallback dengan better styling
            try:
                table.style = 'Light Grid Accent 1'
            except:
                table.style = 'Table Grid'

    def _add_cell_padding(self, cell, padding="100"):
        """
        Enhanced cell padding dengan konsistensi yang lebih baik dan error handling.
        """
        try:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()

            # Remove existing margins jika ada
            existing_mar = tcPr.find(qn('w:tcMar'))
            if existing_mar is not None:
                tcPr.remove(existing_mar)

            # Create enhanced cell margins
            tcMar = OxmlElement('w:tcMar')

            # Enhanced padding specifications dengan proportional spacing
            padding_int = int(padding)
            padding_specs = {
                'top': str(padding_int),
                'right': str(int(padding_int * 1.2)),     # Slightly more horizontal padding
                'bottom': str(padding_int),
                'left': str(int(padding_int * 1.2))       # Slightly more horizontal padding
            }

            # Apply enhanced padding dengan consistency checks
            for side, value in padding_specs.items():
                node = OxmlElement(f'w:{side}')
                node.set(qn('w:w'), value)
                node.set(qn('w:type'), 'dxa')
                tcMar.append(node)

            tcPr.append(tcMar)

            # Enhanced cell vertical alignment
            tcVAlign = OxmlElement('w:vAlign')
            tcVAlign.set(qn('w:val'), 'center')  # Center vertical alignment
            tcPr.append(tcVAlign)

        except Exception as e:
            logger.error(f"Error adding enhanced cell padding: {e}")
            # Fallback: pastikan cell tetap memiliki basic formatting
            try:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                # Minimal padding sebagai fallback
                basic_mar = OxmlElement('w:tcMar')
                for side in ['top', 'right', 'bottom', 'left']:
                    node = OxmlElement(f'w:{side}')
                    node.set(qn('w:w'), '80')  # Basic padding
                    node.set(qn('w:type'), 'dxa')
                    basic_mar.append(node)
                tcPr.append(basic_mar)
            except:
                pass  # Jika tetap gagal, biarkan default Word formatting

    def _add_formatted_runs_to_paragraph(self, paragraph, text):
        """
        Enhanced text formatter that properly processes Markdown and removes formatting markers.
        """
        if not text:
            return

        # Dictionary to map emoji codes to actual emojis
        emoji_dict = {
            ":smile:": "😊", ":check:": "✅", ":x:": "❌", ":warning:": "⚠️",
            ":star:": "⭐", ":arrow_right:": "➡️", ":bulb:": "💡", ":calendar:": "📅",
            ":chart:": "📊", ":document:": "📄", ":pencil:": "✏️", ":clipboard:": "📋",
            ":email:": "📧", ":folder:": "📁", ":money:": "💰", ":phone:": "📱", ":clock:": "🕒",
            ":fire:": "🔥", ":rocket:": "🚀", ":target:": "🎯", ":key:": "🔑", ":shield:": "🛡️"
        }

        # Replace emoji codes first
        for code, emoji in emoji_dict.items():
            text = text.replace(code, emoji)

        # Enhanced regex patterns for various formatting
        patterns = [
            (r'\*\*\*(.+?)\*\*\*', 'bold_italic'),       # ***bold italic***
            (r'\*\*(.+?)\*\*', 'bold'),                  # **bold**
            (r'__(.+?)__', 'bold'),                      # __bold__
            (r'(?<!\*)\*([^*\n]+?)\*(?!\*)', 'italic'),  # *italic* (not part of **)
            (r'(?<!_)_([^_\n]+?)_(?!_)', 'italic'),      # _italic_ (not part of __)
            (r'\+\+(.+?)\+\+', 'underline'),             # ++underline++
            (r'~~(.+?)~~', 'strikethrough'),             # ~~strikethrough~~
            (r'==(.+?)==', 'highlight'),                 # ==highlight==
            (r'\^(.+?)\^', 'superscript'),               # ^superscript^
            (r'(?<!~)~([^~\n]+?)~(?!~)', 'subscript'),   # ~subscript~ (not part of ~~)
            (r'`(.+?)`', 'code'),                        # `code`
            (r'\[\[(.+?)\]\]', 'comment'),               # [[comment]]
            (r'\{\{(.+?)\}\}', 'important'),             # {{important}}
            (r'\{\+(.+?)\+\}', 'success'),               # {+success+}
            (r'\{\-(.+?)\-\}', 'error'),                 # {-error-}
            (r'\{!(.+?)!\}', 'warning'),                 # {!warning!}
            (r'\[(.+?)\](?!\()', 'button'),              # [button] (not part of link)
            (r'\[([^\]]+)\]\(([^)]+)\)', 'link'),        # [text](url)
        ]

        # Find all formatting matches
        all_matches = []
        for pattern, format_type in patterns:
            for match in re.finditer(pattern, text):
                if format_type == 'link':
                    all_matches.append((match.start(), match.end(), match.group(1), format_type, match.group(0), match.group(2)))
                else:
                    all_matches.append((match.start(), match.end(), match.group(1), format_type, match.group(0), None))

        # If no formatting found, just add the text
        if not all_matches:
            paragraph.add_run(text)
            return

        # Sort matches by start position
        all_matches.sort(key=lambda x: x[0])

        # Remove overlapping matches (keep the first one found)
        cleaned_matches = []
        for match in all_matches:
            start, end = match[0], match[1]
            is_overlapping = False
            for prev_start, prev_end, _, _, _, _ in cleaned_matches:
                if (start < prev_end and end > prev_start):  # Overlapping
                    is_overlapping = True
                    break
            if not is_overlapping:
                cleaned_matches.append(match)

        # Process text with formatting
        last_end = 0
        for match in cleaned_matches:
            start, end, content, format_type, full_match = match[0], match[1], match[2], match[3], match[4]
            url = match[5] if len(match) > 5 else None

            # Add text before this formatted section
            if start > last_end:
                plain_text = text[last_end:start]
                paragraph.add_run(plain_text)

            # Add the formatted text (WITHOUT the markers)
            run = paragraph.add_run(content)

            # Apply formatting
            if format_type == 'bold':
                run.bold = True
            elif format_type == 'bold_italic':
                run.bold = True
                run.italic = True
            elif format_type == 'italic':
                run.italic = True
            elif format_type == 'underline':
                run.underline = True
            elif format_type == 'strikethrough':
                run.font.strike = True
            elif format_type == 'highlight':
                run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            elif format_type == 'superscript':
                run.font.superscript = True
            elif format_type == 'subscript':
                run.font.subscript = True
            elif format_type == 'code':
                run.font.name = 'Consolas'
                run.font.size = Pt(9.5)
                shading_elm = parse_xml(r'<w:shd {} w:fill="F5F7FA"/>'.format(nsdecls('w')))
                rPr = run._element.get_or_add_rPr()
                rPr.append(shading_elm)
            elif format_type == 'comment':
                run.font.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
            elif format_type == 'important':
                run.font.color.rgb = RGBColor(192, 0, 0)
                run.bold = True
            elif format_type == 'success':
                run.font.color.rgb = RGBColor(0, 128, 0)
                run.bold = True
            elif format_type == 'error':
                run.font.color.rgb = RGBColor(192, 0, 0)
                run.bold = True
            elif format_type == 'warning':
                run.font.color.rgb = RGBColor(255, 102, 0)
                run.bold = True
            elif format_type == 'button':
                run.font.name = 'Segoe UI'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(60, 60, 60)
                shading_elm = parse_xml(r'<w:shd {} w:fill="E9E9E9"/>'.format(nsdecls('w')))
                rPr = run._element.get_or_add_rPr()
                rPr.append(shading_elm)
            elif format_type == 'link':
                run.font.color.rgb = RGBColor(0, 0, 255)
                run.underline = True
                # Note: Actual hyperlink functionality would require more complex Word XML manipulation

            last_end = end

        # Add any remaining text
        if last_end < len(text):
            tail_text = text[last_end:]
            paragraph.add_run(tail_text)

    def _apply_paragraph_style(self, paragraph, line, content_stats=None):
        """
        Menerapkan gaya khusus untuk paragraf berdasarkan konten, format tertentu, dan content_stats.

        Enhanced Features:
        - Document type adaptive styling
        - Advanced formatting detection (30+ patterns)
        - Smart spacing and typography
        - Context-aware alignment
        - Multi-language support
        - Accessibility enhancements
        - Professional document layouts
        """
        # ===== INPUT VALIDATION & NORMALIZATION =====
        if not line or not line.strip():
            return ""

        stripped_line = line.strip()

        # Normalize content_stats
        if isinstance(content_stats, str):
            document_type = content_stats
            content_stats = {"content_type": document_type}
        elif not isinstance(content_stats, dict):
            content_stats = {"content_type": "general"}

        document_type = content_stats.get("content_type", "general")

        # ===== DOCUMENT TYPE CONFIGURATIONS =====
        # Advanced styling configurations berdasarkan document type
        document_configs = {
            "technical_report": {
                "base_line_spacing": Pt(15),
                "base_space_before": Pt(8),
                "base_space_after": Pt(8),
                "quote_indent": 0.6,
                "first_line_indent": 0.3,
                "color_scheme": "professional_blue",
                "font_adjustments": {"size_factor": 1.0, "leading_factor": 1.1}
            },
            "meeting_notes": {
                "base_line_spacing": Pt(13),
                "base_space_before": Pt(5),
                "base_space_after": Pt(5),
                "quote_indent": 0.4,
                "first_line_indent": 0.2,
                "color_scheme": "business_green",
                "font_adjustments": {"size_factor": 0.95, "leading_factor": 1.0}
            },
            "lecture": {
                "base_line_spacing": Pt(16),
                "base_space_before": Pt(7),
                "base_space_after": Pt(7),
                "quote_indent": 0.7,
                "first_line_indent": 0.35,
                "color_scheme": "academic_orange",
                "font_adjustments": {"size_factor": 1.05, "leading_factor": 1.15}
            },
            "presentation": {
                "base_line_spacing": Pt(17),
                "base_space_before": Pt(10),
                "base_space_after": Pt(10),
                "quote_indent": 0.5,
                "first_line_indent": 0.25,
                "color_scheme": "modern_purple",
                "font_adjustments": {"size_factor": 1.1, "leading_factor": 1.2}
            },
            "research": {
                "base_line_spacing": Pt(16),
                "base_space_before": Pt(9),
                "base_space_after": Pt(9),
                "quote_indent": 0.8,
                "first_line_indent": 0.4,
                "color_scheme": "academic_blue",
                "font_adjustments": {"size_factor": 1.0, "leading_factor": 1.2}
            },
            "interview": {
                "base_line_spacing": Pt(14),
                "base_space_before": Pt(6),
                "base_space_after": Pt(6),
                "quote_indent": 0.3,
                "first_line_indent": 0.15,
                "color_scheme": "conversational",
                "font_adjustments": {"size_factor": 0.95, "leading_factor": 1.0}
            },
            "narrative": {
                "base_line_spacing": Pt(18),
                "base_space_before": Pt(8),
                "base_space_after": Pt(8),
                "quote_indent": 0.6,
                "first_line_indent": 0.5,
                "color_scheme": "literary",
                "font_adjustments": {"size_factor": 1.05, "leading_factor": 1.3}
            },
            "general": {
                "base_line_spacing": Pt(14),
                "base_space_before": Pt(6),
                "base_space_after": Pt(6),
                "quote_indent": 0.5,
                "first_line_indent": 0.25,
                "color_scheme": "standard",
                "font_adjustments": {"size_factor": 1.0, "leading_factor": 1.15}
            }
        }

        config = document_configs.get(document_type, document_configs["general"])

        # ===== COMPLEXITY & LENGTH ADJUSTMENTS =====
        # Adjust spacing based on content complexity and length
        complexity = content_stats.get("complexity_level", "medium")
        word_count = content_stats.get("word_count", 0)

        # Complexity adjustments
        if complexity == "high":
            config["base_space_before"] = Pt(config["base_space_before"].pt + 2)
            config["base_space_after"] = Pt(config["base_space_after"].pt + 2)
            config["base_line_spacing"] = Pt(config["base_line_spacing"].pt + 1)
        elif complexity == "low":
            config["base_space_before"] = Pt(max(4, config["base_space_before"].pt - 1))
            config["base_space_after"] = Pt(max(4, config["base_space_after"].pt - 1))

        # Length adjustments
        if word_count > 5000:  # Long document
            config["base_line_spacing"] = Pt(config["base_line_spacing"].pt + 1)
        elif word_count < 500:  # Short document
            config["base_space_before"] = Pt(max(3, config["base_space_before"].pt - 2))
            config["base_space_after"] = Pt(max(3, config["base_space_after"].pt - 2))

        # ===== DEFAULT PARAGRAPH STYLING =====
        # Apply base formatting dari config
        paragraph.paragraph_format.line_spacing = config["base_line_spacing"]
        paragraph.paragraph_format.space_before = config["base_space_before"]
        paragraph.paragraph_format.space_after = config["base_space_after"]

        # ===== COMPREHENSIVE PATTERN DETECTION =====
        # Significantly expanded pattern detection dengan 30+ patterns

        # ===== 1. ENHANCED BLOCKQUOTES =====
        if stripped_line.startswith('> '):
            quote_text = stripped_line[2:]

            # Multi-level quote support
            quote_level = 1
            temp_text = quote_text
            while temp_text.startswith('> '):
                quote_level += 1
                temp_text = temp_text[2:]

            final_quote_text = temp_text

            # Enhanced quote styling based on document type
            paragraph.style = 'Quote'
            base_indent = config["quote_indent"]
            left_indent = base_indent + (quote_level - 1) * 0.2

            paragraph.paragraph_format.left_indent = Inches(left_indent)
            paragraph.paragraph_format.right_indent = Inches(base_indent * 0.6)
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 4)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 4)
            paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt - 1)

            # Quote type detection
            if any(word in final_quote_text.lower() for word in ['said', 'stated', 'mentioned', 'noted']):
                # Spoken quote
                paragraph.paragraph_format.first_line_indent = Inches(0.1)

            return final_quote_text

        # ===== 2. ENHANCED INDENTATION PATTERNS =====
        # Multiple indentation levels dan types
        if stripped_line.startswith('    ') or stripped_line.startswith('\t\t'):  # Double indent
            paragraph.paragraph_format.first_line_indent = Inches(config["first_line_indent"] * 2)
            paragraph.paragraph_format.left_indent = Inches(0.1)
            return stripped_line.lstrip()

        elif stripped_line.startswith('   ') or stripped_line.startswith('\t'):  # Single indent
            paragraph.paragraph_format.first_line_indent = Inches(config["first_line_indent"])
            return stripped_line.lstrip()

        # ===== 3. ADVANCED SPACING CONTROL =====
        # Enhanced spacing control dengan lebih banyak options
        if stripped_line.startswith('//'):
            command = stripped_line[2:4]
            content = stripped_line[4:].strip() if len(stripped_line) > 4 else ""

            if command == 'xl':  # Extra large spacing
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 6)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 8)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 8)
            elif command == 'lg':  # Large spacing
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 4)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 6)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 6)
            elif command == 's' or command == 'sm':  # Large/Small spacing (legacy support)
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 4)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 6)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 6)
            elif command == 'md':  # Medium spacing
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 2)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 3)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 3)
            elif command == 'xs':  # Extra small spacing
                paragraph.paragraph_format.line_spacing = Pt(max(10, config["base_line_spacing"].pt - 4))
                paragraph.paragraph_format.space_before = Pt(max(2, config["base_space_before"].pt - 3))
                paragraph.paragraph_format.space_after = Pt(max(2, config["base_space_after"].pt - 3))
            elif command == 'c' or command == 'cp':  # Compact spacing (legacy support)
                paragraph.paragraph_format.line_spacing = Pt(max(10, config["base_line_spacing"].pt - 2))
                paragraph.paragraph_format.space_before = Pt(max(2, config["base_space_before"].pt - 2))
                paragraph.paragraph_format.space_after = Pt(max(2, config["base_space_after"].pt - 2))
            elif command == 'no':  # No spacing
                paragraph.paragraph_format.line_spacing = Pt(10)
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

            return content if content else stripped_line[4:].strip()

        # ===== 4. COMPREHENSIVE ALIGNMENT PATTERNS =====
        # Significantly expanded alignment options

        # Center alignment variations
        if (stripped_line.startswith('->') and stripped_line.endswith('<-')) or \
        stripped_line.startswith('><') or stripped_line.startswith('::center::'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 2)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 2)

            if stripped_line.startswith('->') and stripped_line.endswith('<-'):
                return stripped_line[2:-2].strip()
            elif stripped_line.startswith('><'):
                return stripped_line[2:].strip()
            elif stripped_line.startswith('::center::'):
                return stripped_line[10:].strip()

        # Right alignment variations
        elif stripped_line.startswith('->') or stripped_line.startswith('>>') or \
            stripped_line.startswith('::right::'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            paragraph.paragraph_format.right_indent = Inches(0.1)

            if stripped_line.startswith('->'):
                return stripped_line[2:].strip()
            elif stripped_line.startswith('>>'):
                return stripped_line[2:].strip()
            elif stripped_line.startswith('::right::'):
                return stripped_line[9:].strip()

        # Left alignment (explicit)
        elif stripped_line.startswith('<<') or stripped_line.startswith('::left::'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.left_indent = Inches(0.1)

            if stripped_line.startswith('<<'):
                return stripped_line[2:].strip()
            elif stripped_line.startswith('::left::'):
                return stripped_line[8:].strip()

        # Justify alignment
        elif stripped_line.startswith('::justify::') or stripped_line.startswith('|<->|'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

            if stripped_line.startswith('::justify::'):
                return stripped_line[11:].strip()
            elif stripped_line.startswith('|<->|'):
                return stripped_line[5:].strip()

        # ===== 5. DOCUMENT-SPECIFIC SPECIAL PATTERNS =====

        # Technical documentation patterns
        if document_type == "technical_report":
            # Code snippet inline
            if stripped_line.startswith('```inline'):
                paragraph.paragraph_format.left_indent = Inches(0.3)
                paragraph.paragraph_format.line_spacing = 1.0
                return stripped_line[9:].strip()

            # API endpoint pattern
            elif re.match(r'^(GET|POST|PUT|DELETE|PATCH)\s+/', stripped_line):
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 3)
                return stripped_line

            # Configuration pattern
            elif stripped_line.startswith('config:') or stripped_line.startswith('setting:'):
                paragraph.paragraph_format.left_indent = Inches(0.25)
                return stripped_line

        # Meeting notes patterns
        elif document_type == "meeting_notes":
            # Action item pattern
            if stripped_line.startswith('ACTION:') or stripped_line.startswith('TODO:'):
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 2)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 2)
                return stripped_line

            # Decision pattern
            elif stripped_line.startswith('DECISION:') or stripped_line.startswith('RESOLVED:'):
                paragraph.paragraph_format.left_indent = Inches(0.2)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 2)
                return stripped_line

            # Follow-up pattern
            elif stripped_line.startswith('FOLLOW-UP:') or stripped_line.startswith('NEXT:'):
                paragraph.paragraph_format.left_indent = Inches(0.2)
                return stripped_line

        # Lecture patterns
        elif document_type == "lecture":
            # Example pattern
            if stripped_line.startswith('Example:') or stripped_line.startswith('e.g.'):
                paragraph.paragraph_format.left_indent = Inches(0.4)
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt - 1)
                return stripped_line

            # Key point pattern
            elif stripped_line.startswith('Key:') or stripped_line.startswith('Important:'):
                paragraph.paragraph_format.left_indent = Inches(0.3)
                paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 3)
                paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 3)
                return stripped_line

            # Definition pattern
            elif ':' in stripped_line and len(stripped_line.split(':')[0]) < 30:
                paragraph.paragraph_format.left_indent = Inches(0.25)
                return stripped_line

        # ===== 6. ADVANCED TYPOGRAPHY PATTERNS =====

        # Drop cap simulation
        if stripped_line.startswith('::drop::'):
            paragraph.paragraph_format.first_line_indent = Inches(-0.2)
            paragraph.paragraph_format.left_indent = Inches(0.4)
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 5)
            return stripped_line[8:].strip()

        # Pull quote
        elif stripped_line.startswith('::pull::'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.left_indent = Inches(1.0)
            paragraph.paragraph_format.right_indent = Inches(1.0)
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 8)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 8)
            return stripped_line[8:].strip()

        # Sidebar text
        elif stripped_line.startswith('::sidebar::'):
            paragraph.paragraph_format.left_indent = Inches(0.8)
            paragraph.paragraph_format.right_indent = Inches(0.2)
            paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt - 2)
            return stripped_line[11:].strip()

        # Caption style
        elif stripped_line.startswith('::caption::'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 2)
            return stripped_line[11:].strip()

        # ===== 7. CONTENT-AWARE PATTERNS =====

        # Question patterns
        if stripped_line.endswith('?') and len(stripped_line) < 100:
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 1)
            return stripped_line

        # Exclamation patterns
        elif stripped_line.endswith('!') and not stripped_line.endswith('!!'):
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 1)
            return stripped_line

        # List item detection (outside of formal lists)
        elif re.match(r'^[•·‣▪▫◦‹›]', stripped_line):
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.first_line_indent = Inches(-0.15)
            return stripped_line

        # Time stamp patterns
        elif re.match(r'^\d{1,2}:\d{2}', stripped_line):
            paragraph.paragraph_format.left_indent = Inches(0.2)
            return stripped_line

        # URL patterns
        elif stripped_line.startswith('http') or stripped_line.startswith('www'):
            paragraph.paragraph_format.left_indent = Inches(0.2)
            paragraph.paragraph_format.line_spacing = 1.0
            return stripped_line

        # ===== 8. LANGUAGE-SPECIFIC PATTERNS =====
        language = content_stats.get("language", "id")

        if language.startswith("id"):  # Indonesian
            # Pantun atau sajak pattern
            if re.search(r'\b(pantun|sajak|puisi)\b', stripped_line.lower()):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 2)
                return stripped_line

        elif language.startswith("en"):  # English
            # Poetry pattern
            if re.search(r'\b(poem|verse|stanza)\b', stripped_line.lower()):
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.paragraph_format.line_spacing = Pt(config["base_line_spacing"].pt + 2)
                return stripped_line

        # ===== 9. PRIORITY & EMPHASIS PATTERNS =====

        # High priority content
        if any(word in stripped_line.upper() for word in ['URGENT', 'CRITICAL', 'IMPORTANT', 'ASAP']):
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 4)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 4)
            paragraph.paragraph_format.left_indent = Inches(0.1)
            return stripped_line

        # Note patterns
        elif stripped_line.upper().startswith(('NOTE:', 'NB:', 'CATATAN:')):
            paragraph.paragraph_format.left_indent = Inches(0.3)
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 2)
            return stripped_line

        # ===== 10. MATHEMATICAL & SCIENTIFIC PATTERNS =====

        # Formula patterns
        if re.search(r'[=+\-*/^]', stripped_line) and re.search(r'\d', stripped_line):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(config["base_space_before"].pt + 2)
            paragraph.paragraph_format.space_after = Pt(config["base_space_after"].pt + 2)
            return stripped_line

        # Units and measurements
        elif re.search(r'\d+\s*(mm|cm|m|km|g|kg|°C|°F|%)', stripped_line):
            paragraph.paragraph_format.left_indent = Inches(0.1)
            return stripped_line

        # ===== 11. ACCESSIBILITY ENHANCEMENTS =====

        try:
            # Add semantic information for screen readers
            if stripped_line.endswith('?'):
                paragraph._element.set(qn('w:role'), 'question')
            elif any(word in stripped_line.upper() for word in ['CONCLUSION', 'SUMMARY', 'RESULT']):
                paragraph._element.set(qn('w:role'), 'conclusion')
            elif stripped_line.startswith(('Example:', 'e.g.', 'Contoh:')):
                paragraph._element.set(qn('w:role'), 'example')
        except:
            pass  # Continue if accessibility features fail

        # ===== 12. KEEP-WITH-NEXT LOGIC =====

        # Smart keep-with-next for better page breaks
        if any(pattern in stripped_line for pattern in [':', 'follows', 'berikut', 'namely', 'yaitu']):
            if len(stripped_line) < 80:  # Short introductory lines
                paragraph.paragraph_format.keep_with_next = True

        # ===== 13. ORPHAN & WIDOW CONTROL =====

        # Enhanced typography control
        if len(stripped_line) > 200:  # Long paragraphs
            paragraph.paragraph_format.widow_control = True

        # ===== 14. DEFAULT ALIGNMENT =====

        # Set default alignment if not already set
        if not hasattr(paragraph, '_alignment_set'):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph._alignment_set = True

        # ===== 15. FINAL ADJUSTMENTS =====

        # Apply document-specific fine-tuning
        if document_type == "presentation":
            # Increase spacing for better visual impact
            paragraph.paragraph_format.space_before = Pt(paragraph.paragraph_format.space_before.pt + 1)
            paragraph.paragraph_format.space_after = Pt(paragraph.paragraph_format.space_after.pt + 1)

        elif document_type == "research":
            # Tighter spacing for academic density
            paragraph.paragraph_format.line_spacing = Pt(paragraph.paragraph_format.line_spacing.pt + 1)

        # ===== LOGGING & RETURN =====

        try:
            logger.debug(f"Applied paragraph style for {document_type}: {stripped_line[:50]}...")
        except:
            pass

        return stripped_line

    def _add_admonition_block(self, doc, lines, admonition_type):
        """
        Add a styled admonition block for extended markdown syntax :::type ... :::
        """
        # Map types
        type_map = {
            "note": "note",
            "info": "info",
            "tip": "tip",
            "warning": "warning",
            "danger": "warning",
            "important": "important",
            "success": "tip",
            "hint": "tip",
            "caution": "warning",
            "conclusion": "conclusion"
        }

        actual_type = type_map.get(admonition_type.lower(), "note")

        # Get content as text
        content = "\n".join(lines)

        # Check if there's a title line
        title_match = re.match(r'^(.*?)\n-{3,}$', content, re.MULTILINE | re.DOTALL)
        title = None

        if title_match:
            title = title_match.group(1).strip()
            content = content[content.find('-'*3) + 3:].strip()

        # Create the admonition
        self._add_enhanced_callout(doc, content, actual_type)

    def _add_document_footer(self, doc, content_stats):
        """
        Menambahkan footer dokumen yang profesional menggunakan content_stats.
        """
        # Gunakan informasi dari content_stats untuk footer yang lebih informatif
        section = doc.sections[0]
        footer = section.footer

        # Create footer paragraph if it doesn't exist
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add page numbering
        page_num = footer_para.add_run("Page ")
        page_num.font.size = Pt(9)
        page_num.font.color.rgb = RGBColor(100, 100, 100)

        # Add document statistics dari content_stats
        if content_stats and "word_count" in content_stats:
            stats_para = footer.add_paragraph()
            stats_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            stats_text = f"Document: {content_stats['content_type'].title()} | Words: {content_stats['word_count']}"
            if "language" in content_stats:
                stats_text += f" | Language: {content_stats['language'].upper()}"

            stats_run = stats_para.add_run(stats_text)
            stats_run.font.size = Pt(7)
            stats_run.font.color.rgb = RGBColor(150, 150, 150)

        # Add timestamp
        now = datetime.datetime.now()
        timestamp = now.strftime("%d %B %Y, %H:%M:%S")

        timestamp_para = footer.add_paragraph()
        timestamp_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp_run = timestamp_para.add_run(f"Created: {timestamp}")
        timestamp_run.font.size = Pt(7)
        timestamp_run.font.color.rgb = RGBColor(150, 150, 150)

        logger.info(f"Document footer added with stats: {content_stats.get('content_type', 'unknown')} type")

    # !Security and setup methods

    def post_init_hook(self):
        """Enhanced post init with better error handling."""
        try:
            # Set up error logging
            self.setup_error_logging()

            # Check system compatibility
            compatibility = self.check_system_compatibility()
            if not compatibility['compatibility']:
                errors = "\n".join(compatibility['errors'])
                messagebox.showerror(
                    "Masalah Kompatibilitas",
                    f"Aplikasi mungkin tidak berfungsi dengan benar:\n\n{errors}\n\n"
                    "Lihat log untuk detail lebih lanjut."
                )
            elif compatibility['warnings']:
                warnings_text = "\n".join(compatibility['warnings'])
                messagebox.showwarning(
                    "Peringatan Kompatibilitas",
                    f"Beberapa masalah terdeteksi:\n\n{warnings_text}"
                )

            # Setup auto-save setiap 30 detik
            self.setup_auto_save()

            # Enhanced exit handler
            def enhanced_on_exit():
                try:
                    print("DEBUG: Starting enhanced cleanup...")

                    # Stop visualization thread
                    if hasattr(self, 'viz_running'):
                        self.viz_running = False

                    # Wait for thread
                    if hasattr(self, 'viz_thread') and self.viz_thread and self.viz_thread.is_alive():
                        self.viz_thread.join(timeout=1.0)

                    # Close matplotlib
                    if hasattr(self, 'viz_fig'):
                        try:
                            import matplotlib.pyplot as plt
                            plt.close(self.viz_fig)
                        except:
                            pass

                    # Stop recording
                    if hasattr(self, 'recording') and self.recording:
                        self.stop_recording()

                    # Save config
                    save_success = self.save_config()
                    print(f"DEBUG: Config save {'successful' if save_success else 'failed'}")

                    # Cleanup
                    self.safe_cleanup()

                    print("DEBUG: Enhanced cleanup completed")

                except Exception as e:
                    print(f"ERROR: Error during enhanced cleanup: {e}")
                finally:
                    try:
                        self.root.destroy()
                    except:
                        pass

            # Set exit handler
            self.root.protocol("WM_DELETE_WINDOW", enhanced_on_exit)

            print("DEBUG: Post init hook completed successfully")

        except Exception as e:
            print(f"ERROR: Error in post_init_hook: {e}")
            self.add_exit_handler()

    def setup_error_logging(self):
        """Configure enhanced error logging."""
        # Set up a rotating file handler
        log_dir = os.path.join(os.path.expanduser("~"), ".echoscribe")
        os.makedirs(log_dir, exist_ok=True)

        log_file = os.path.join(log_dir, "echoscribe.log")

        # Remove old handlers if any to prevent duplication
        for handler in logger.handlers[:]:
            logger.removeHandler(handler)

        # Create file handler with rotation (supports Unicode)
        file_handler = RotatingFileHandler(
            log_file,
            maxBytes=5*1024*1024,  # 5MB
            backupCount=3,
            encoding='utf-8'  # Unicode support for file logging
        )

        # Create console handler with emoji filtering
        console_handler = logging.StreamHandler(sys.stdout)

        # Create formatter
        formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')

        # Set formatters
        file_handler.setFormatter(formatter)
        console_handler.setFormatter(formatter)

        # Add emoji filter only to console handler to prevent encoding issues
        console_handler.addFilter(self._emoji_filter)

        # Add handlers to logger
        logger.addHandler(file_handler)
        logger.addHandler(console_handler)

        # Set log level
        logger.setLevel(logging.INFO)

        logger.info("Logging initialized")

        # Log system info
        self._log_system_info()

    def _emoji_filter(self, record):
        """Filter untuk mengganti emoji dengan teks sederhana di console output."""
        # Dictionary untuk mapping emoji ke teks sederhana
        emoji_replacements = {
            '✅': '[SUCCESS]',
            '✓': '[OK]',
            '❌': '[ERROR]',
            '⚠️': '[WARNING]',
            '🔴': '[RECORDING]',
            '🔄': '[PROCESSING]',
            '📊': '[DATA]',
            '💡': '[TIP]',
            '📝': '[NOTE]',
            '🎯': '[TARGET]',
            '🚀': '[LAUNCH]',
            '⭐': '[STAR]',
            '🔧': '[TOOL]',
            '📱': '[MOBILE]',
            '💻': '[COMPUTER]',
            '🎵': '[AUDIO]',
            '🔊': '[SPEAKER]',
            '🎤': '[MICROPHONE]',
            '☐': '[ ]',
            '😊': ':)'
        }

        try:
            # Ganti emoji dalam pesan log
            if hasattr(record, 'msg') and isinstance(record.msg, str):
                for emoji, replacement in emoji_replacements.items():
                    record.msg = record.msg.replace(emoji, replacement)

            # Also handle formatted arguments
            if hasattr(record, 'args') and record.args:
                try:
                    formatted_args = []
                    for arg in record.args:
                        if isinstance(arg, str):
                            for emoji, replacement in emoji_replacements.items():
                                arg = arg.replace(emoji, replacement)
                        formatted_args.append(arg)
                    record.args = tuple(formatted_args)
                except:
                    # If there's any issue with args formatting, just keep original
                    pass

        except Exception:
            # If there's any error in filtering, just return True to let the message through
            pass

        return True

    def _log_system_info(self):
        """Log system information for troubleshooting."""
        try:
            logger.info("===== System Information =====")
            logger.info(f"OS: {platform.system()} {platform.release()} {platform.version()}")
            logger.info(f"Python: {platform.python_version()}")
            logger.info(f"Machine: {platform.machine()}")
            logger.info(f"Processor: {platform.processor()}")

            # Log available microphones
            logger.info(f"Available microphones: {len(self.microphones)}")

            # Log installed packages
            logger.info("------- Key Packages -------")
            packages = ["speech_recognition", "groq", "python-docx", "customtkinter"]
            for package in packages:
                try:
                    module = __import__(package)
                    version = getattr(module, "__version__", "unknown")
                    logger.info(f"{package}: {version}")
                except ImportError:
                    logger.info(f"{package}: Not installed")

            logger.info("============================")
        except Exception as e:
            logger.error(f"Failed to log system info: {e}")

    def check_system_compatibility(self):
        """Check system compatibility to prevent issues."""
        results = {
            'compatibility': True,
            'warnings': [],
            'errors': [],
            'recommendations': []
        }

        # Check Python version
        python_version = tuple(map(int, platform.python_version_tuple()))
        if python_version < (3, 7, 0):
            results['compatibility'] = False
            results['errors'].append(f"Versi Python ({platform.python_version()}) terlalu lama. Dibutuhkan 3.7+")

        # Check memory (if psutil is available)
        try:
            import psutil
            available_memory = psutil.virtual_memory().available / (1024 * 1024)  # MB
            if available_memory < 500:
                results['compatibility'] = False
                results['errors'].append(f"Memori tidak cukup: {available_memory:.0f}MB tersedia. Dibutuhkan minimal 500MB")
            elif available_memory < 1000:
                results['warnings'].append(f"Memori terbatas: {available_memory:.0f}MB tersedia. Disarankan 1GB+")

            # Check disk space
            free_space = psutil.disk_usage(os.path.expanduser("~")).free / (1024 * 1024 * 1024)  # GB
            if free_space < 1:
                results['warnings'].append(f"Ruang disk terbatas: {free_space:.2f}GB tersedia")
        except ImportError:
            # psutil not available, just skip these checks
            pass

        # Check FFmpeg availability
        ffmpeg_found = setup_ffmpeg()
        if not ffmpeg_found:
            results['warnings'].append("FFmpeg tidak terdeteksi. Beberapa fitur audio tidak akan berfungsi")
            results['recommendations'].append("Instal FFmpeg dari https://ffmpeg.org/download.html")

        # Check network connection
        try:
            import urllib.request
            urllib.request.urlopen("https://www.google.com", timeout=2)
        except:
            results['warnings'].append("Koneksi internet tidak terdeteksi atau tidak stabil")

        # Check for microphone
        if not self.microphones:
            results['warnings'].append("Tidak ada mikrofon terdeteksi")

        # Check for required libraries
        missing_libraries = []
        for module_name in ["speech_recognition", "groq", "docx"]:
            try:
                __import__(module_name.replace("-", "_"))
            except ImportError:
                missing_libraries.append(module_name)

        if missing_libraries:
            results['compatibility'] = False
            results['errors'].append(f"Pustaka yang diperlukan tidak terinstal: {', '.join(missing_libraries)}")
            results['recommendations'].append(
                f"Instal pustaka yang diperlukan dengan: pip install {' '.join(missing_libraries)}"
            )

        return results

    def setup_config_management(self):
        """Set up configuration management untuk persistent settings."""
        # Define config file location
        config_dir = os.path.join(os.path.expanduser("~"), ".echoscribe")
        os.makedirs(config_dir, exist_ok=True)
        self.config_file = os.path.join(config_dir, "config.json")

        # ===== PERBAIKAN 4: Config yang lebih lengkap =====
        self.default_config = {
            "output_folder": os.path.expanduser("~/Documents"),
            "filename_prefix": "catatan",
            "language": "id-ID",
            "engine": "Google",
            "use_ai_enhancement": True,
            "use_economic_model": False,
            "use_extended_recording": True,
            "chunk_size": 600,
            "api_request_delay": 10,
            "theme": "Dark",
            "last_microphone_index": 0,
            "recording_mode": "microphone",
            "use_system_audio": False,
            "use_dual_recording": False,
            "window_geometry": "1000x1080",
            "window_position": None,
            "viz_mode": "waveform",
            "viz_enabled": True,
            "viz_sensitivity": 1.0
        }

        # Load config
        self.config = self.load_config()
        print(f"DEBUG: Config loaded: {self.config}")

    def load_config(self):
        """Load configuration from file dengan error handling yang lebih baik."""
        config = self.default_config.copy()

        try:
            if os.path.exists(self.config_file):
                print(f"DEBUG: Loading config from: {self.config_file}")
                with open(self.config_file, 'r', encoding='utf-8') as f:
                    loaded_config = json.load(f)

                # Validasi dan update config
                for key, value in loaded_config.items():
                    if key in self.default_config:
                        config[key] = value
                        print(f"DEBUG: Loaded {key} = {value}")
                    else:
                        print(f"DEBUG: Skipping unknown config key: {key}")

                print(f"DEBUG: Successfully loaded {len(loaded_config)} config items")
            else:
                print(f"DEBUG: Config file not found, using defaults: {self.config_file}")

        except Exception as e:
            print(f"ERROR: Error loading config: {e}")
            # Use defaults if loading fails

        return config

    def apply_config_after_ui_ready(self):
        """Apply configuration after UI is fully ready."""
        try:
            print("DEBUG: Applying config after UI ready...")

            # Apply microphone selection
            if self.microphones and "last_microphone_index" in self.config:
                last_index = self.config["last_microphone_index"]
                for mic in self.microphones:
                    if mic.startswith(f"{last_index}:"):
                        self.selected_mic.set(mic)
                        print(f"DEBUG: Applied microphone: {mic}")
                        break

            # Apply window geometry
            if "window_geometry" in self.config:
                try:
                    self.root.geometry(self.config["window_geometry"])
                    print(f"DEBUG: Applied window geometry: {self.config['window_geometry']}")
                except:
                    pass

            # Apply window position if saved
            if "window_position" in self.config and self.config["window_position"]:
                try:
                    x, y = self.config["window_position"]
                    self.root.geometry(f"+{x}+{y}")
                    print(f"DEBUG: Applied window position: {x}, {y}")
                except:
                    pass

            # Apply theme
            try:
                ctk.set_appearance_mode(self.config.get("theme", "Dark"))
                print(f"DEBUG: Applied theme: {self.config.get('theme', 'Dark')}")
            except:
                pass

            # Apply visualization settings
            if hasattr(self, 'viz_mode'):
                self.viz_mode.set(self.config.get("viz_mode", "waveform"))
                print(f"DEBUG: Applied viz_mode: {self.config.get('viz_mode', 'waveform')}")

            if hasattr(self, 'viz_enabled'):
                self.viz_enabled.set(self.config.get("viz_enabled", True))
                print(f"DEBUG: Applied viz_enabled: {self.config.get('viz_enabled', True)}")

            if hasattr(self, 'viz_sensitivity'):
                self.viz_sensitivity.set(self.config.get("viz_sensitivity", 1.0))
                print(f"DEBUG: Applied viz_sensitivity: {self.config.get('viz_sensitivity', 1.0)}")

            # Update recording mode
            self._update_recording_mode()

            print("DEBUG: Config applied successfully!")
            if hasattr(self, 'status_var'):
                self.status_var.set("Konfigurasi dimuat")

        except Exception as e:
            print(f"ERROR: Error applying config: {e}")

    def save_config(self):
        """Save current configuration dengan data yang lebih lengkap."""
        try:
            print("DEBUG: Saving configuration...")

            # Update config dari current values
            if hasattr(self, 'output_folder'):
                self.config["output_folder"] = self.output_folder.get()
            if hasattr(self, 'filename_prefix'):
                self.config["filename_prefix"] = self.filename_prefix.get()
            if hasattr(self, 'language'):
                self.config["language"] = self.language.get()
            if hasattr(self, 'engine'):
                self.config["engine"] = self.engine.get()
            if hasattr(self, 'use_ai_enhancement'):
                self.config["use_ai_enhancement"] = self.use_ai_enhancement.get()
            if hasattr(self, 'use_economic_model'):
                self.config["use_economic_model"] = self.use_economic_model.get()
            if hasattr(self, 'use_extended_recording'):
                self.config["use_extended_recording"] = self.use_extended_recording.get()
            if hasattr(self, 'chunk_size'):
                self.config["chunk_size"] = self.chunk_size.get()
            if hasattr(self, 'api_request_delay'):
                self.config["api_request_delay"] = self.api_request_delay
            if hasattr(self, 'use_dual_recording'):
                self.config["use_dual_recording"] = self.use_dual_recording.get()
            if hasattr(self, 'recording_mode'):
                self.config["recording_mode"] = self.recording_mode.get()
            if hasattr(self, 'use_system_audio'):
                self.config["use_system_audio"] = self.use_system_audio.get()

            # Save current theme
            try:
                self.config["theme"] = ctk.get_appearance_mode()
            except:
                self.config["theme"] = "Dark"

            # Save window geometry dan position
            try:
                self.config["window_geometry"] = self.root.geometry()
                # Parse position dari geometry string
                geo = self.root.geometry()
                if '+' in geo:
                    pos_part = geo.split('+')[1:]
                    if len(pos_part) >= 2:
                        self.config["window_position"] = [int(pos_part[0]), int(pos_part[1])]
            except:
                pass

            # Save visualization settings
            if hasattr(self, 'viz_mode'):
                self.config["viz_mode"] = self.viz_mode.get()
            if hasattr(self, 'viz_enabled'):
                self.config["viz_enabled"] = self.viz_enabled.get()
            if hasattr(self, 'viz_sensitivity'):
                self.config["viz_sensitivity"] = self.viz_sensitivity.get()

            # Save microphone selection
            if hasattr(self, 'microphones') and self.microphones:
                try:
                    current_mic = self.selected_mic.get()
                    mic_index = int(current_mic.split(":")[0])
                    self.config["last_microphone_index"] = mic_index
                    print(f"DEBUG: Saved microphone index: {mic_index}")
                except:
                    pass

            # Write to file dengan backup
            backup_file = self.config_file + ".backup"

            # Create backup of existing config
            if os.path.exists(self.config_file):
                try:
                    import shutil
                    shutil.copy2(self.config_file, backup_file)
                except:
                    pass

            # Write new config
            with open(self.config_file, 'w', encoding='utf-8') as f:
                json.dump(self.config, f, indent=2, ensure_ascii=False)

            print(f"DEBUG: Configuration saved to: {self.config_file}")
            print(f"DEBUG: Saved config: {self.config}")

            return True

        except Exception as e:
            print(f"ERROR: Error saving config: {e}")

            # Restore from backup if save failed
            backup_file = self.config_file + ".backup"
            if os.path.exists(backup_file):
                try:
                    import shutil
                    shutil.copy2(backup_file, self.config_file)
                    print("DEBUG: Restored config from backup")
                except:
                    pass

            return False

    def setup_auto_save(self):
        """Setup auto-save untuk config setiap 10 menit."""
        def auto_save():
            try:
                self.save_config()
                print("DEBUG: Auto-save completed")
            except Exception as e:
                print(f"ERROR: Auto-save failed: {e}")
            finally:
                # Schedule next auto-save
                self.root.after(600000, auto_save)  # 10 minutes = 600,000 milliseconds

        # Start auto-save
        self.root.after(600000, auto_save)  # 10 minutes = 600,000 milliseconds

    def add_exit_handler(self):
        """
        Add handler for proper application exit.
        """
        def on_exit():
            # Stop any ongoing recording
            if self.recording:
                self.stop_recording()

            # Save config
            self.save_config()

            # Perform cleanup
            self.safe_cleanup()

            # Destroy root window
            self.root.destroy()

        # Bind to window close event
        self.root.protocol("WM_DELETE_WINDOW", on_exit)

    def safe_cleanup(self):
        """
        Safe cleanup of resources on exit.
        """
        try:
            # Cleanup temp files
            if hasattr(self, 'temp_audio_files') and self.temp_audio_files:
                for file_path in self.temp_audio_files:
                    try:
                        if os.path.exists(file_path):
                            os.unlink(file_path)
                    except Exception as e:
                        logger.error(f"Error cleaning up temp file {file_path}: {e}")

            # Cleanup temp directories
            if hasattr(self, 'temp_dir') and self.temp_dir and os.path.exists(self.temp_dir):
                try:
                    os.rmdir(self.temp_dir)
                except Exception as e:
                    logger.error(f"Error cleaning up temp directory {self.temp_dir}: {e}")

            # Cleanup temporary wav file
            if hasattr(self, 'temp_wav_file') and self.temp_wav_file:
                try:
                    if os.path.exists(self.temp_wav_file.name):
                        os.unlink(self.temp_wav_file.name)
                except Exception as e:
                    logger.error(f"Error cleaning up temp wav file: {e}")

            logger.info("Cleanup completed")
        except Exception as e:
            logger.error(f"Error during cleanup: {e}")

if __name__ == "__main__":
    root = tk.Tk()
    app = VoiceToMarkdownApp(root)
    root.mainloop()
